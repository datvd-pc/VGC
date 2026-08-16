"""
Master compiler script to generate both:
1. The-Art-of-Monetization-English-Polished.md
2. The-Art-of-Monetization-English-Polished.docx

With full cover page, text justification, and per-part/per-chapter page breaks.
"""

import os, sys
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

from data_sec0 import SEC0_FRONT_MATTER
from data_sec1 import SEC1_PART1
from data_sec2 import SEC2_PART2
from data_sec3 import SEC3_PART3
from data_sec4 import SEC4_PART4
from data_sec5 import SEC5_PART5
from data_sec6 import SEC6_PART6
from data_sec7 import SEC7_PART7
from data_sec8_10 import SEC8_PART8, SEC9_PART9, SEC10_CLOSING

OUTPUT_DIR = r"D:\CODE\VGC\Long form"
COVER_PATH = os.path.join(OUTPUT_DIR, "assets", "covers", "the-art-of-monetization-cover-option-6-money-lock-final.jpg")
MD_PATH = os.path.join(OUTPUT_DIR, "The-Art-of-Monetization-English-Polished.md")
DOCX_PATH = os.path.join(OUTPUT_DIR, "The-Art-of-Monetization-English-Polished.docx")


def generate_markdown():
    lines = []
    
    # Cover Image
    lines.append(f"![The Art of Monetization Cover](assets/covers/the-art-of-monetization-cover-option-6-money-lock-final.jpg)\n")
    lines.append("\n<div style=\"page-break-after: always;\"></div>\n\n---\n")
    
    # Title & Subtitle
    lines.append(f"# {SEC0_FRONT_MATTER['title']}")
    lines.append(f"### *{SEC0_FRONT_MATTER['subtitle']}*")
    lines.append("\n<div style=\"page-break-after: always;\"></div>\n\n---\n")
    
    # Research Note
    lines.append(f"## {SEC0_FRONT_MATTER['research_note']['title']}\n")
    for p in SEC0_FRONT_MATTER['research_note']['paragraphs']:
        lines.append(f"{p}\n")
    lines.append("\n<div style=\"page-break-after: always;\"></div>\n\n---\n")
    
    # Note to Reader
    lines.append(f"## {SEC0_FRONT_MATTER['note_to_reader']['title']}\n")
    for p in SEC0_FRONT_MATTER['note_to_reader']['paragraphs']:
        lines.append(f"{p}\n")
    lines.append("\n<div style=\"page-break-after: always;\"></div>\n\n---\n")
    
    # Key Terms
    lines.append(f"## {SEC0_FRONT_MATTER['key_terms']['title']}\n")
    lines.append(f"{SEC0_FRONT_MATTER['key_terms']['intro']}\n")
    lines.append("| " + " | ".join(SEC0_FRONT_MATTER['key_terms']['table']['headers']) + " |")
    lines.append("| " + " | ".join(["---"] * len(SEC0_FRONT_MATTER['key_terms']['table']['headers'])) + " |")
    for row in SEC0_FRONT_MATTER['key_terms']['table']['rows']:
        lines.append("| " + " | ".join(row) + " |")
    lines.append("\n<div style=\"page-break-after: always;\"></div>\n\n---\n")
    
    # Contents
    lines.append(f"## {SEC0_FRONT_MATTER['contents']['title']}\n")
    for idx, item in enumerate(SEC0_FRONT_MATTER['contents']['items'], 1):
        lines.append(f"{idx}. {item}")
    lines.append("\n<div style=\"page-break-after: always;\"></div>\n\n---\n")
    
    # How to Read
    lines.append(f"## {SEC0_FRONT_MATTER['how_to_read']['title']}\n")
    for p in SEC0_FRONT_MATTER['how_to_read']['paragraphs']:
        lines.append(f"{p}\n")
    lines.append("\n<div style=\"page-break-after: always;\"></div>\n\n---\n")
    
    # Starts Before Store
    lines.append(f"## {SEC0_FRONT_MATTER['starts_before_store']['title']}\n")
    for p in SEC0_FRONT_MATTER['starts_before_store']['paragraphs']:
        lines.append(f"{p}\n")
    lines.append("\n<div style=\"page-break-after: always;\"></div>\n\n---\n")
    
    # Case Study Clear Garden
    lines.append(f"## {SEC0_FRONT_MATTER['case_study']['title']}\n")
    for p in SEC0_FRONT_MATTER['case_study']['paragraphs']:
        lines.append(f"{p}\n")
    lines.append("\n<div style=\"page-break-after: always;\"></div>\n\n---\n")
    
    # Helper for Decision Boards in Markdown
    def md_decision_board(db):
        res = []
        res.append(f"### {db['title']}\n")
        if 'intro' in db:
            res.append(f"*{db['intro']}*\n")
        res.append("| " + " | ".join(db['table']['headers']) + " |")
        res.append("| " + " | ".join(["---"] * len(db['table']['headers'])) + " |")
        for row in db['table']['rows']:
            clean_row = [c.replace('\n', '<br>') for c in row]
            res.append("| " + " | ".join(clean_row) + " |")
        res.append("\n<div style=\"page-break-after: always;\"></div>\n\n---\n")
        return "\n".join(res)
    
    # Helper for generic table in Markdown
    def md_table(tbl):
        res = []
        res.append("| " + " | ".join(tbl['headers']) + " |")
        res.append("| " + " | ".join(["---"] * len(tbl['headers'])) + " |")
        for row in tbl['rows']:
            clean_row = [c.replace('\n', '<br>') for c in row]
            res.append("| " + " | ".join(clean_row) + " |")
        res.append("\n")
        return "\n".join(res)

    # PART I
    lines.append(f"# {SEC1_PART1['part_title']}\n")
    for ch_key in ['chapter1', 'chapter2', 'chapter3', 'chapter4']:
        ch = SEC1_PART1[ch_key]
        lines.append(f"## {ch['title']}\n")
        for p in ch['paragraphs']:
            lines.append(f"{p}\n")
        if 'table' in ch:
            lines.append(md_table(ch['table']))
        if 'audit_callout' in ch:
            lines.append(f"> **Audit Check:** {ch['audit_callout']}\n")
        lines.append("\n<div style=\"page-break-after: always;\"></div>\n")
    lines.append(md_decision_board(SEC1_PART1['decision_board']))

    # PART II
    lines.append(f"# {SEC2_PART2['part_title']}\n")
    for ch_key in ['chapter5', 'chapter6', 'chapter7', 'chapter8']:
        ch = SEC2_PART2[ch_key]
        lines.append(f"## {ch['title']}\n")
        for p in ch['paragraphs']:
            lines.append(f"{p}\n")
        lines.append("\n<div style=\"page-break-after: always;\"></div>\n")
    lines.append(md_decision_board(SEC2_PART2['decision_board']))

    # PART III
    lines.append(f"# {SEC3_PART3['part_title']}\n")
    # Chapter 9
    ch9 = SEC3_PART3['chapter9']
    lines.append(f"## {ch9['title']}\n")
    for p in ch9['paragraphs']:
        lines.append(f"{p}\n")
    lines.append(md_table(ch9['table']))
    for p in ch9['paragraphs_after_table']:
        lines.append(f"{p}\n")
    lines.append("\n<div style=\"page-break-after: always;\"></div>\n")
    # Chapter 10 & 11
    for ch_key in ['chapter10', 'chapter11']:
        ch = SEC3_PART3[ch_key]
        lines.append(f"## {ch['title']}\n")
        for p in ch['paragraphs']:
            lines.append(f"{p}\n")
        lines.append("\n<div style=\"page-break-after: always;\"></div>\n")
    # Chapter 12
    ch12 = SEC3_PART3['chapter12']
    lines.append(f"## {ch12['title']}\n")
    for p in ch12['paragraphs']:
        lines.append(f"{p}\n")
    lines.append(md_table(ch12['table']))
    for p in ch12['paragraphs_after_table']:
        lines.append(f"{p}\n")
    lines.append("\n<div style=\"page-break-after: always;\"></div>\n")
    # Chapter 13 & 14
    for ch_key in ['chapter13', 'chapter14']:
        ch = SEC3_PART3[ch_key]
        lines.append(f"## {ch['title']}\n")
        for p in ch['paragraphs']:
            lines.append(f"{p}\n")
        lines.append("\n<div style=\"page-break-after: always;\"></div>\n")
    lines.append(md_decision_board(SEC3_PART3['decision_board']))

    # PART IV
    lines.append(f"# {SEC4_PART4['part_title']}\n")
    # Chapter 15
    ch15 = SEC4_PART4['chapter15']
    lines.append(f"## {ch15['title']}\n")
    for p in ch15['paragraphs']:
        lines.append(f"{p}\n")
    lines.append(md_table(ch15['table']))
    for p in ch15['paragraphs_after_table']:
        lines.append(f"{p}\n")
    lines.append("\n<div style=\"page-break-after: always;\"></div>\n")
    # Chapter 16 & 17
    for ch_key in ['chapter16', 'chapter17']:
        ch = SEC4_PART4[ch_key]
        lines.append(f"## {ch['title']}\n")
        for p in ch['paragraphs']:
            lines.append(f"{p}\n")
        lines.append("\n<div style=\"page-break-after: always;\"></div>\n")
    # Chapter 18
    ch18 = SEC4_PART4['chapter18']
    lines.append(f"## {ch18['title']}\n")
    for p in ch18['paragraphs']:
        lines.append(f"{p}\n")
    lines.append(md_table(ch18['table']))
    for p in ch18['paragraphs_after_table']:
        lines.append(f"{p}\n")
    lines.append("\n<div style=\"page-break-after: always;\"></div>\n")
    # Chapter 19 & 20
    for ch_key in ['chapter19', 'chapter20']:
        ch = SEC4_PART4[ch_key]
        lines.append(f"## {ch['title']}\n")
        for p in ch['paragraphs']:
            lines.append(f"{p}\n")
        lines.append("\n<div style=\"page-break-after: always;\"></div>\n")
    lines.append(md_decision_board(SEC4_PART4['decision_board']))

    # PART V
    lines.append(f"# {SEC5_PART5['part_title']}\n")
    # Chapter 21
    ch21 = SEC5_PART5['chapter21']
    lines.append(f"## {ch21['title']}\n")
    for p in ch21['paragraphs']:
        lines.append(f"{p}\n")
    lines.append(md_table(ch21['table_diagnostic']))
    for p in ch21['paragraphs_between_tables']:
        lines.append(f"{p}\n")
    lines.append(md_table(ch21['table_strategic']))
    for p in ch21['paragraphs_after_strategic']:
        lines.append(f"{p}\n")
    lines.append("\n<div style=\"page-break-after: always;\"></div>\n")
    # Chapter 22
    ch22 = SEC5_PART5['chapter22']
    lines.append(f"## {ch22['title']}\n")
    for p in ch22['paragraphs']:
        lines.append(f"{p}\n")
    lines.append(md_table(ch22['table']))
    for p in ch22['paragraphs_after_table']:
        lines.append(f"{p}\n")
    lines.append("\n<div style=\"page-break-after: always;\"></div>\n")
    # Chapter 23
    ch23 = SEC5_PART5['chapter23']
    lines.append(f"## {ch23['title']}\n")
    for p in ch23['paragraphs']:
        lines.append(f"{p}\n")
    lines.append("\n<div style=\"page-break-after: always;\"></div>\n")
    # Chapter 24
    ch24 = SEC5_PART5['chapter24']
    lines.append(f"## {ch24['title']}\n")
    for p in ch24['paragraphs']:
        lines.append(f"{p}\n")
    lines.append(md_table(ch24['table']))
    for p in ch24['paragraphs_after_table']:
        lines.append(f"{p}\n")
    lines.append("\n<div style=\"page-break-after: always;\"></div>\n")
    lines.append(md_decision_board(SEC5_PART5['decision_board']))

    # PART VI
    lines.append(f"# {SEC6_PART6['part_title']}\n")
    # Chapter 25
    ch25 = SEC6_PART6['chapter25']
    lines.append(f"## {ch25['title']}\n")
    for p in ch25['paragraphs']:
        lines.append(f"{p}\n")
    lines.append(md_table(ch25['table']))
    for p in ch25['paragraphs_after_table']:
        lines.append(f"{p}\n")
    lines.append("\n<div style=\"page-break-after: always;\"></div>\n")
    # Chapter 26
    ch26 = SEC6_PART6['chapter26']
    lines.append(f"## {ch26['title']}\n")
    for p in ch26['paragraphs']:
        lines.append(f"{p}\n")
    lines.append("\n<div style=\"page-break-after: always;\"></div>\n")
    # Chapter 27
    ch27 = SEC6_PART6['chapter27']
    lines.append(f"## {ch27['title']}\n")
    for p in ch27['paragraphs']:
        lines.append(f"{p}\n")
    lines.append(md_table(ch27['table_strategy']))
    lines.append(md_table(ch27['table_memo_template']))
    lines.append("\n<div style=\"page-break-after: always;\"></div>\n")
    # Chapter 28 & 29
    for ch_key in ['chapter28', 'chapter29']:
        ch = SEC6_PART6[ch_key]
        lines.append(f"## {ch['title']}\n")
        for p in ch['paragraphs']:
            lines.append(f"{p}\n")
        lines.append("\n<div style=\"page-break-after: always;\"></div>\n")
    lines.append(md_decision_board(SEC6_PART6['decision_board']))

    # PART VII
    lines.append(f"# {SEC7_PART7['part_title']}\n")
    # Chapter 30
    ch30 = SEC7_PART7['chapter30']
    lines.append(f"## {ch30['title']}\n")
    for p in ch30['paragraphs']:
        lines.append(f"{p}\n")
    lines.append(md_table(ch30['table']))
    for p in ch30['paragraphs_after_table']:
        lines.append(f"{p}\n")
    lines.append("\n<div style=\"page-break-after: always;\"></div>\n")
    # Chapter 31
    ch31 = SEC7_PART7['chapter31']
    lines.append(f"## {ch31['title']}\n")
    for p in ch31['paragraphs']:
        lines.append(f"{p}\n")
    lines.append("\n<div style=\"page-break-after: always;\"></div>\n")
    # Chapter 32
    ch32 = SEC7_PART7['chapter32']
    lines.append(f"## {ch32['title']}\n")
    for p in ch32['paragraphs']:
        lines.append(f"{p}\n")
    lines.append(md_table(ch32['table']))
    for p in ch32['paragraphs_after_table']:
        lines.append(f"{p}\n")
    lines.append("\n<div style=\"page-break-after: always;\"></div>\n")
    # Chapter 33
    ch33 = SEC7_PART7['chapter33']
    lines.append(f"## {ch33['title']}\n")
    for p in ch33['paragraphs']:
        lines.append(f"{p}\n")
    lines.append("\n<div style=\"page-break-after: always;\"></div>\n")
    lines.append(md_decision_board(SEC7_PART7['decision_board']))

    # PART VIII
    lines.append(f"# {SEC8_PART8['part_title']}\n")
    ch34 = SEC8_PART8['chapter34']
    lines.append(f"## {ch34['title']}\n")
    for p in ch34['paragraphs']:
        lines.append(f"{p}\n")
    lines.append(md_table(ch34['table_segment']))
    for p in ch34['paragraphs_between_tables']:
        lines.append(f"{p}\n")
    lines.append(md_table(ch34['table_mechanics']))
    for p in ch34['paragraphs_after_tables']:
        lines.append(f"{p}\n")
    lines.append("\n<div style=\"page-break-after: always;\"></div>\n")
    ch35 = SEC8_PART8['chapter35']
    lines.append(f"## {ch35['title']}\n")
    for p in ch35['paragraphs']:
        lines.append(f"{p}\n")
    lines.append("\n<div style=\"page-break-after: always;\"></div>\n")
    lines.append(md_decision_board(SEC8_PART8['decision_board']))

    # PART IX
    lines.append(f"# {SEC9_PART9['part_title']}\n")
    ch36 = SEC9_PART9['chapter36']
    lines.append(f"## {ch36['title']}\n")
    for p in ch36['paragraphs']:
        lines.append(f"{p}\n")
    lines.append(md_table(ch36['table']))
    lines.append("\n<div style=\"page-break-after: always;\"></div>\n")
    ch37 = SEC9_PART9['chapter37']
    lines.append(f"## {ch37['title']}\n")
    for p in ch37['paragraphs']:
        lines.append(f"{p}\n")
    for item in ch37['list_items']:
        lines.append(f"{item}\n")
    lines.append("\n<div style=\"page-break-after: always;\"></div>\n")
    lines.append(md_decision_board(ch37['decision_board']))

    # CLOSING & REFERENCES
    lines.append(f"# {SEC10_CLOSING['closing']['title']}\n")
    for p in SEC10_CLOSING['closing']['paragraphs']:
        lines.append(f"{p}\n")
    lines.append("\n<div style=\"page-break-after: always;\"></div>\n\n---\n")
    lines.append(f"## {SEC10_CLOSING['references']['title']}\n")
    for src in SEC10_CLOSING['references']['sources']:
        lines.append(f"• {src}\n")

    full_md = "\n".join(lines)
    with open(MD_PATH, 'w', encoding='utf-8') as f:
        f.write(full_md)
    print(f"Markdown generated successfully at: {MD_PATH}")


def set_cell_shading(cell, color_hex):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
    tcPr.append(tcMar)

def set_table_borders(table, color="CBD5E1", sz="4", val="single"):
    tblPr = table._tbl.tblPr
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'<w:top w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:bottom w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:insideH w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:left w:val="none"/>'
        f'<w:right w:val="none"/>'
        f'<w:insideV w:val="none"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)

def set_decision_board_borders(table, color="94A3B8", sz="8"):
    tblPr = table._tbl.tblPr
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'<w:top w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:bottom w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:left w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:right w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:insideH w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:insideV w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)


def generate_docx():
    doc = Document()
    
    # -------------------------------------------------------------
    # SECTION 0: COVER PAGE (Zero margins, Full Page Cover Image)
    # -------------------------------------------------------------
    cover_section = doc.sections[0]
    cover_section.page_width = Pt(595.3)   # A4 Width
    cover_section.page_height = Pt(841.9)  # A4 Height
    cover_section.top_margin = Pt(0)
    cover_section.bottom_margin = Pt(0)
    cover_section.left_margin = Pt(0)
    cover_section.right_margin = Pt(0)
    
    p_cov = doc.add_paragraph()
    p_cov.paragraph_format.space_before = Pt(0)
    p_cov.paragraph_format.space_after = Pt(0)
    p_cov.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_cov = p_cov.add_run()
    if os.path.exists(COVER_PATH):
        r_cov.add_picture(COVER_PATH, width=Pt(595.3), height=Pt(841.9))
    else:
        print(f"Warning: Cover not found at {COVER_PATH}")

    # -------------------------------------------------------------
    # SECTION 1: BODY CONTENT (Standard A4 with 51pt margins)
    # -------------------------------------------------------------
    body_section = doc.add_section()
    body_section.page_width = Pt(595.3)
    body_section.page_height = Pt(841.9)
    body_section.top_margin = Pt(51.0)
    body_section.bottom_margin = Pt(51.0)
    body_section.left_margin = Pt(51.0)
    body_section.right_margin = Pt(51.0)
    
    # Base Styles
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Calibri'
    normal_style.font.size = Pt(10.5)
    normal_style.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)
    normal_style.paragraph_format.line_spacing = 1.15
    normal_style.paragraph_format.space_after = Pt(4)
    normal_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    def add_title(text):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(text)
        r.font.name = 'Calibri'
        r.font.size = Pt(26)
        r.font.bold = True
        r.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
        return p

    def add_subtitle(text):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(24)
        r = p.add_run(text)
        r.font.name = 'Calibri'
        r.font.size = Pt(14)
        r.font.italic = True
        r.font.color.rgb = RGBColor(0x47, 0x55, 0x69)
        return p

    def add_h1(text, page_break=True):
        p = doc.add_paragraph()
        if page_break:
            p.paragraph_format.page_break_before = True
        p.paragraph_format.space_before = Pt(20)
        p.paragraph_format.space_after = Pt(8)
        p.paragraph_format.keep_with_next = True
        r = p.add_run(text)
        r.font.name = 'Calibri'
        r.font.size = Pt(17)
        r.font.bold = True
        r.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
        return p

    def add_h2(text, page_break=True):
        p = doc.add_paragraph()
        if page_break:
            p.paragraph_format.page_break_before = True
        p.paragraph_format.space_before = Pt(16)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.keep_with_next = True
        r = p.add_run(text)
        r.font.name = 'Calibri'
        r.font.size = Pt(13.5)
        r.font.bold = True
        r.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)
        return p

    def add_h3(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.keep_with_next = True
        r = p.add_run(text)
        r.font.name = 'Calibri'
        r.font.size = Pt(11.5)
        r.font.bold = True
        r.font.color.rgb = RGBColor(0x33, 0x41, 0x55)
        return p

    def add_p(text):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(5)
        p.paragraph_format.line_spacing = 1.15
        r = p.add_run(text)
        r.font.name = 'Calibri'
        r.font.size = Pt(10.5)
        r.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)
        return p

    def add_bullet(text):
        p = doc.add_paragraph(style='List Bullet')
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(2.5)
        r = p.add_run(text)
        r.font.name = 'Calibri'
        r.font.size = Pt(10.5)
        r.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)
        return p

    def add_doc_table(tbl_data):
        headers = tbl_data['headers']
        rows = tbl_data['rows']
        t = doc.add_table(rows=len(rows)+1, cols=len(headers))
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        set_table_borders(t, color="CBD5E1")
        
        # Header row
        hdr_row = t.rows[0]
        hdr_row._tr.get_or_add_trPr().append(parse_xml(f'<w:tblHeader {nsdecls("w")}/>'))
        for idx, h_text in enumerate(headers):
            cell = hdr_row.cells[idx]
            set_cell_shading(cell, "F1F5F9")
            set_cell_margins(cell, top=120, bottom=120, left=150, right=150)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(h_text)
            r.font.name = 'Calibri'
            r.font.size = Pt(10)
            r.font.bold = True
            r.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)

        # Body rows
        for r_idx, row_data in enumerate(rows):
            row = t.rows[r_idx + 1]
            bg_color = "FFFFFF" if r_idx % 2 == 0 else "F8FAFC"
            for c_idx, cell_value in enumerate(row_data):
                cell = row.cells[c_idx]
                set_cell_shading(cell, bg_color)
                set_cell_margins(cell, top=100, bottom=100, left=150, right=150)
                p = cell.paragraphs[0]
                p.paragraph_format.space_after = Pt(0)
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                r = p.add_run(cell_value)
                r.font.name = 'Calibri'
                r.font.size = Pt(9.5)
                r.font.color.rgb = RGBColor(0x33, 0x41, 0x55)
        
        doc.add_paragraph().paragraph_format.space_after = Pt(6)

    def add_doc_decision_board(db):
        add_h3(db['title'])
        if 'intro' in db:
            p_intro = doc.add_paragraph()
            p_intro.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p_intro.paragraph_format.space_after = Pt(4)
            r_intro = p_intro.add_run(db['intro'])
            r_intro.font.name = 'Calibri'
            r_intro.font.size = Pt(10)
            r_intro.font.italic = True
            r_intro.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

        headers = db['table']['headers']
        rows = db['table']['rows']
        t = doc.add_table(rows=len(rows)+1, cols=len(headers))
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        set_decision_board_borders(t, color="94A3B8", sz="6")

        # Header row
        for idx, h_text in enumerate(headers):
            cell = t.rows[0].cells[idx]
            set_cell_shading(cell, "E2E8F0")
            set_cell_margins(cell, top=140, bottom=140, left=180, right=180)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(h_text)
            r.font.name = 'Calibri'
            r.font.size = Pt(10)
            r.font.bold = True
            r.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)

        # Body rows
        for r_idx, row_data in enumerate(rows):
            row = t.rows[r_idx + 1]
            for c_idx, cell_value in enumerate(row_data):
                cell = row.cells[c_idx]
                set_cell_shading(cell, "F8FAFC")
                set_cell_margins(cell, top=120, bottom=120, left=180, right=180)
                p = cell.paragraphs[0]
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.15
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                # Parse lines inside cell
                lines_in_cell = cell_value.split('\n')
                for l_i, line in enumerate(lines_in_cell):
                    if l_i > 0:
                        p = cell.add_paragraph()
                        p.paragraph_format.space_after = Pt(2)
                        p.paragraph_format.line_spacing = 1.15
                        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    if line.startswith("•") or line.startswith("CORE TAKEAWAYS:") or line.startswith("TEAM MEETING AGENDA"):
                        r = p.add_run(line)
                        r.font.name = 'Calibri'
                        r.font.size = Pt(9.5)
                        if "TAKEAWAYS:" in line or "AGENDA" in line:
                            r.font.bold = True
                            r.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
                        else:
                            r.font.color.rgb = RGBColor(0x33, 0x41, 0x55)
                    else:
                        r = p.add_run(line)
                        r.font.name = 'Calibri'
                        r.font.size = Pt(9.5)
                        r.font.color.rgb = RGBColor(0x33, 0x41, 0x55)

        doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # 1. Title Page (Front Matter)
    add_title(SEC0_FRONT_MATTER['title'])
    add_subtitle(SEC0_FRONT_MATTER['subtitle'])

    add_h1(SEC0_FRONT_MATTER['research_note']['title'], page_break=True)
    for p in SEC0_FRONT_MATTER['research_note']['paragraphs']:
        add_p(p)

    add_h1(SEC0_FRONT_MATTER['note_to_reader']['title'], page_break=True)
    for p in SEC0_FRONT_MATTER['note_to_reader']['paragraphs']:
        add_p(p)

    add_h1(SEC0_FRONT_MATTER['key_terms']['title'], page_break=True)
    add_p(SEC0_FRONT_MATTER['key_terms']['intro'])
    add_doc_table(SEC0_FRONT_MATTER['key_terms']['table'])

    add_h1(SEC0_FRONT_MATTER['contents']['title'], page_break=True)
    for idx, item in enumerate(SEC0_FRONT_MATTER['contents']['items'], 1):
        add_bullet(f"{idx}. {item}")

    add_h1(SEC0_FRONT_MATTER['how_to_read']['title'], page_break=True)
    for p in SEC0_FRONT_MATTER['how_to_read']['paragraphs']:
        add_p(p)

    add_h1(SEC0_FRONT_MATTER['starts_before_store']['title'], page_break=True)
    for p in SEC0_FRONT_MATTER['starts_before_store']['paragraphs']:
        add_p(p)

    add_h1(SEC0_FRONT_MATTER['case_study']['title'], page_break=True)
    for p in SEC0_FRONT_MATTER['case_study']['paragraphs']:
        add_p(p)

    # 2. PART I
    add_h1(SEC1_PART1['part_title'], page_break=True)
    for ch_key in ['chapter1', 'chapter2', 'chapter3', 'chapter4']:
        ch = SEC1_PART1[ch_key]
        add_h2(ch['title'], page_break=True)
        for p in ch['paragraphs']:
            add_p(p)
        if 'table' in ch:
            add_doc_table(ch['table'])
        if 'audit_callout' in ch:
            p_call = doc.add_paragraph()
            p_call.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p_call.paragraph_format.space_before = Pt(4)
            p_call.paragraph_format.space_after = Pt(6)
            r_bold = p_call.add_run("Audit Check: ")
            r_bold.font.name = 'Calibri'
            r_bold.font.size = Pt(10)
            r_bold.font.bold = True
            r_bold.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)
            r_text = p_call.add_run(ch['audit_callout'])
            r_text.font.name = 'Calibri'
            r_text.font.size = Pt(10)
            r_text.font.italic = True
            r_text.font.color.rgb = RGBColor(0x33, 0x41, 0x55)
    add_doc_decision_board(SEC1_PART1['decision_board'])

    # 3. PART II
    add_h1(SEC2_PART2['part_title'], page_break=True)
    for ch_key in ['chapter5', 'chapter6', 'chapter7', 'chapter8']:
        ch = SEC2_PART2[ch_key]
        add_h2(ch['title'], page_break=True)
        for p in ch['paragraphs']:
            add_p(p)
    add_doc_decision_board(SEC2_PART2['decision_board'])

    # 4. PART III
    add_h1(SEC3_PART3['part_title'], page_break=True)
    ch9 = SEC3_PART3['chapter9']
    add_h2(ch9['title'], page_break=True)
    for p in ch9['paragraphs']:
        add_p(p)
    add_doc_table(ch9['table'])
    for p in ch9['paragraphs_after_table']:
        add_p(p)

    for ch_key in ['chapter10', 'chapter11']:
        ch = SEC3_PART3[ch_key]
        add_h2(ch['title'], page_break=True)
        for p in ch['paragraphs']:
            add_p(p)

    ch12 = SEC3_PART3['chapter12']
    add_h2(ch12['title'], page_break=True)
    for p in ch12['paragraphs']:
        add_p(p)
    add_doc_table(ch12['table'])
    for p in ch12['paragraphs_after_table']:
        add_p(p)

    for ch_key in ['chapter13', 'chapter14']:
        ch = SEC3_PART3[ch_key]
        add_h2(ch['title'], page_break=True)
        for p in ch['paragraphs']:
            add_p(p)
    add_doc_decision_board(SEC3_PART3['decision_board'])

    # 5. PART IV
    add_h1(SEC4_PART4['part_title'], page_break=True)
    ch15 = SEC4_PART4['chapter15']
    add_h2(ch15['title'], page_break=True)
    for p in ch15['paragraphs']:
        add_p(p)
    add_doc_table(ch15['table'])
    for p in ch15['paragraphs_after_table']:
        add_p(p)

    for ch_key in ['chapter16', 'chapter17']:
        ch = SEC4_PART4[ch_key]
        add_h2(ch['title'], page_break=True)
        for p in ch['paragraphs']:
            add_p(p)

    ch18 = SEC4_PART4['chapter18']
    add_h2(ch18['title'], page_break=True)
    for p in ch18['paragraphs']:
        add_p(p)
    add_doc_table(ch18['table'])
    for p in ch18['paragraphs_after_table']:
        add_p(p)

    for ch_key in ['chapter19', 'chapter20']:
        ch = SEC4_PART4[ch_key]
        add_h2(ch['title'], page_break=True)
        for p in ch['paragraphs']:
            add_p(p)
    add_doc_decision_board(SEC4_PART4['decision_board'])

    # 6. PART V
    add_h1(SEC5_PART5['part_title'], page_break=True)
    ch21 = SEC5_PART5['chapter21']
    add_h2(ch21['title'], page_break=True)
    for p in ch21['paragraphs']:
        add_p(p)
    add_doc_table(ch21['table_diagnostic'])
    for p in ch21['paragraphs_between_tables']:
        add_p(p)
    add_doc_table(ch21['table_strategic'])
    for p in ch21['paragraphs_after_strategic']:
        add_p(p)

    ch22 = SEC5_PART5['chapter22']
    add_h2(ch22['title'], page_break=True)
    for p in ch22['paragraphs']:
        add_p(p)
    add_doc_table(ch22['table'])
    for p in ch22['paragraphs_after_table']:
        add_p(p)

    ch23 = SEC5_PART5['chapter23']
    add_h2(ch23['title'], page_break=True)
    for p in ch23['paragraphs']:
        add_p(p)

    ch24 = SEC5_PART5['chapter24']
    add_h2(ch24['title'], page_break=True)
    for p in ch24['paragraphs']:
        add_p(p)
    add_doc_table(ch24['table'])
    for p in ch24['paragraphs_after_table']:
        add_p(p)
    add_doc_decision_board(SEC5_PART5['decision_board'])

    # 7. PART VI
    add_h1(SEC6_PART6['part_title'], page_break=True)
    ch25 = SEC6_PART6['chapter25']
    add_h2(ch25['title'], page_break=True)
    for p in ch25['paragraphs']:
        add_p(p)
    add_doc_table(ch25['table'])
    for p in ch25['paragraphs_after_table']:
        add_p(p)

    ch26 = SEC6_PART6['chapter26']
    add_h2(ch26['title'], page_break=True)
    for p in ch26['paragraphs']:
        add_p(p)

    ch27 = SEC6_PART6['chapter27']
    add_h2(ch27['title'], page_break=True)
    for p in ch27['paragraphs']:
        add_p(p)
    add_doc_table(ch27['table_strategy'])
    add_doc_table(ch27['table_memo_template'])

    for ch_key in ['chapter28', 'chapter29']:
        ch = SEC6_PART6[ch_key]
        add_h2(ch['title'], page_break=True)
        for p in ch['paragraphs']:
            add_p(p)
    add_doc_decision_board(SEC6_PART6['decision_board'])

    # 8. PART VII
    add_h1(SEC7_PART7['part_title'], page_break=True)
    ch30 = SEC7_PART7['chapter30']
    add_h2(ch30['title'], page_break=True)
    for p in ch30['paragraphs']:
        add_p(p)
    add_doc_table(ch30['table'])
    for p in ch30['paragraphs_after_table']:
        add_p(p)

    ch31 = SEC7_PART7['chapter31']
    add_h2(ch31['title'], page_break=True)
    for p in ch31['paragraphs']:
        add_p(p)

    ch32 = SEC7_PART7['chapter32']
    add_h2(ch32['title'], page_break=True)
    for p in ch32['paragraphs']:
        add_p(p)
    add_doc_table(ch32['table'])
    for p in ch32['paragraphs_after_table']:
        add_p(p)

    ch33 = SEC7_PART7['chapter33']
    add_h2(ch33['title'], page_break=True)
    for p in ch33['paragraphs']:
        add_p(p)
    add_doc_decision_board(SEC7_PART7['decision_board'])

    # 9. PART VIII
    add_h1(SEC8_PART8['part_title'], page_break=True)
    ch34 = SEC8_PART8['chapter34']
    add_h2(ch34['title'], page_break=True)
    for p in ch34['paragraphs']:
        add_p(p)
    add_doc_table(ch34['table_segment'])
    for p in ch34['paragraphs_between_tables']:
        add_p(p)
    add_doc_table(ch34['table_mechanics'])
    for p in ch34['paragraphs_after_tables']:
        add_p(p)

    ch35 = SEC8_PART8['chapter35']
    add_h2(ch35['title'], page_break=True)
    for p in ch35['paragraphs']:
        add_p(p)
    add_doc_decision_board(SEC8_PART8['decision_board'])

    # 10. PART IX
    add_h1(SEC9_PART9['part_title'], page_break=True)
    ch36 = SEC9_PART9['chapter36']
    add_h2(ch36['title'], page_break=True)
    for p in ch36['paragraphs']:
        add_p(p)
    add_doc_table(ch36['table'])

    ch37 = SEC9_PART9['chapter37']
    add_h2(ch37['title'], page_break=True)
    for p in ch37['paragraphs']:
        add_p(p)
    for item in ch37['list_items']:
        add_bullet(item)
    add_doc_decision_board(ch37['decision_board'])

    # 11. CLOSING & REFERENCES
    add_h1(SEC10_CLOSING['closing']['title'], page_break=True)
    for p in SEC10_CLOSING['closing']['paragraphs']:
        add_p(p)

    add_h1(SEC10_CLOSING['references']['title'], page_break=True)
    for src in SEC10_CLOSING['references']['sources']:
        add_bullet(src)

    doc.save(DOCX_PATH)
    print(f"DOCX generated successfully at: {DOCX_PATH}")


if __name__ == '__main__':
    print("Starting generation of English Polished Ebook with Cover Page & Formatted Layout...")
    generate_markdown()
    generate_docx()
    print("All tasks completed successfully!")
