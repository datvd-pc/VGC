from pathlib import Path
import re
from docx import Document

PATH = Path(r"D:\CODE\VGC\Long form\The-Art-of-Monetization-Vietnamese-Editable.docx")

def replace(document, index, expected, value):
    paragraph = document.paragraphs[index]
    compact = lambda text: re.sub(r"\s+", " ", text).strip()
    if compact(paragraph.text) != compact(expected):
        raise ValueError(f"Paragraph {index} was changed manually and was not replaced.")
    paragraph.clear()
    paragraph.add_run(value)

doc = Document(PATH)
replace(
    doc,
    100,
    "Trong hồ sơ công bố năm 2022, FTC cho biết họ cáo buộc Fortnite có những bố cục nút bấm khiến người chơi có thể bị tính tiền khi chỉ định xem một vật phẩm, khi game đang tải hoặc khi đánh thức game từ chế độ ngủ. FTC cũng cho biết Epic nhận được hơn một triệu khiếu nại về các khoản phí không mong muốn. Case này nói về payment flow, không nói thay cho mọi quyết định thiết kế level. Nhưng nó đặt ra một tiêu chuẩn rất cụ thể: người chơi phải hiểu điều gì sẽ xảy ra trước khi một thao tác trở thành giao dịch.",
    "Hãy mở một level puzzle gần đây nhất của game. Sau khi người chơi thua, họ nhìn thấy gì trước: nguyên nhân của thất bại, một nước đi khác, hay một offer? Nếu booster xuất hiện, người chơi có hiểu nó sẽ thay đổi tình huống nào không? Và nếu không mua, họ còn một cách hợp lý để tiếp tục hay không? Đây là self-audit, không phải một case thị trường: câu trả lời phải đến từ chính build và dữ liệu của team.",
)
replace(
    doc,
    104,
    "Với mỗi purchase flow, hãy kiểm tra bốn điểm: người chơi có hiểu giá và thứ nhận được không; có một bước xác nhận phù hợp không; có thể tìm thấy cách hủy hoặc hoàn tiền không; và có thể tiếp tục chơi mà không mua không? FTC không cung cấp công thức cho puzzle game, nhưng case này nhắc team rằng payment design, consent và disclosure là trách nhiệm của sản phẩm. Nguồn: FTC, 2022, “Fortnite Video Game Maker Epic Games to Pay More Than Half a Billion Dollars over FTC Allegations of Privacy Violations and Unwanted Charges.”",
    "Với mỗi purchase flow, hãy kiểm tra bốn điểm: người chơi có hiểu giá và thứ nhận được không; có một bước xác nhận phù hợp không; có thể tìm thấy cách hủy hoặc hoàn tiền không; và có thể tiếp tục chơi mà không mua không? Đây không phải là một công thức cho mọi puzzle game. Đây là tiêu chuẩn tối thiểu để một lựa chọn trả tiền vẫn là lựa chọn có hiểu biết.",
)
doc.save(PATH)
print("Refocused Chapter 3 on a puzzle self-audit.")
