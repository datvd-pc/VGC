import argparse
import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


def set_cell_shading(cell, color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color)
    tc_pr.append(shading)


def set_cell_margins(cell, top=90, start=90, bottom=90, end=90):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def add_runs(paragraph, text, bold=False):
    position = 0
    for match in re.finditer(r"(\*[^*]+\*|`[^`]+`|\[[^\]]+\]\([^\)]+\))", text):
        if match.start() > position:
            paragraph.add_run(text[position:match.start()])
        token = match.group(0)
        if token.startswith("*"):
            run = paragraph.add_run(token[1:-1])
            run.italic = True
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            run.font.name = "Consolas"
        else:
            label, url = re.match(r"\[([^\]]+)\]\(([^\)]+)\)", token).groups()
            run = paragraph.add_run(f"{label} ({url})")
            run.font.color.rgb = RGBColor(8, 123, 134)
        position = match.end()
    if position < len(text):
        paragraph.add_run(text[position:])
    for run in paragraph.runs:
        run.bold = bold


def markdown_row(line):
    return [cell.strip() for cell in line.strip().strip("|").split("|")]


def add_table(document, rows):
    parsed = [markdown_row(row) for row in rows if not all(re.fullmatch(r"[-: ]+", c) for c in markdown_row(row))]
    table = document.add_table(rows=len(parsed), cols=len(parsed[0]))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for r, row in enumerate(parsed):
        for c, value in enumerate(row):
            cell = table.cell(r, c)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            if r == 0:
                set_cell_shading(cell, "E9F0EF")
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            add_runs(p, value, bold=(r == 0))
            for run in p.runs:
                run.font.size = Pt(8.5)
    document.add_paragraph()


def add_header(section):
    header = section.header
    p = header.paragraphs[0]
    p.text = "NGHỆ THUẬT KIẾM TIỀN TRONG GAME"
    p.style = "Normal"
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        run.font.name = "Segoe UI"
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(95, 91, 85)


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--input", required=True)
    parser.add_argument("--cover", required=True)
    parser.add_argument("--output", required=True)
    args = parser.parse_args()

    document = Document()
    cover_section = document.sections[0]
    cover_section.top_margin = Cm(0)
    cover_section.bottom_margin = Cm(0)
    cover_section.left_margin = Cm(0)
    cover_section.right_margin = Cm(0)
    cover_section.page_width = Cm(21)
    cover_section.page_height = Cm(29.7)
    cover = document.add_paragraph()
    cover.paragraph_format.space_after = Pt(0)
    cover.paragraph_format.space_before = Pt(0)
    cover.add_run().add_picture(args.cover, width=Cm(21), height=Cm(29.7))

    content_section = document.add_section(WD_SECTION_START.NEW_PAGE)
    content_section.top_margin = Cm(1.8)
    content_section.bottom_margin = Cm(1.8)
    content_section.left_margin = Cm(1.8)
    content_section.right_margin = Cm(1.8)
    add_header(content_section)

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Segoe UI"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(7)
    normal.paragraph_format.line_spacing = 1.25
    for style_name, size, color in (("Title", 25, RGBColor(23, 23, 23)), ("Heading 1", 18, RGBColor(8, 123, 134)), ("Heading 2", 14, RGBColor(23, 23, 23))):
        style = styles[style_name]
        style.font.name = "Segoe UI Semibold"
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(18)
        style.paragraph_format.space_after = Pt(8)

    lines = Path(args.input).read_text(encoding="utf-8").splitlines()
    index = 0
    in_code = False
    code_lines = []
    while index < len(lines):
        line = lines[index]
        if line.startswith("```"):
            if in_code:
                p = document.add_paragraph()
                p.paragraph_format.left_indent = Cm(0.45)
                p.paragraph_format.space_after = Pt(9)
                run = p.add_run("\n".join(code_lines))
                run.font.name = "Consolas"
                run.font.size = Pt(8.7)
                code_lines = []
            in_code = not in_code
            index += 1
            continue
        if in_code:
            code_lines.append(line)
            index += 1
            continue
        if line.startswith("|") and index + 1 < len(lines) and lines[index + 1].startswith("|"):
            rows = [line]
            index += 1
            while index < len(lines) and lines[index].startswith("|"):
                rows.append(lines[index])
                index += 1
            add_table(document, rows)
            continue
        if not line.strip():
            index += 1
            continue
        if line.startswith("# "):
            p = document.add_paragraph(style="Title")
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_runs(p, line[2:].replace("<br/>", "\n"), bold=True)
        elif line.startswith("## "):
            p = document.add_paragraph(style="Heading 1")
            add_runs(p, line[3:].replace("<br/>", "\n"), bold=True)
        elif line.startswith("### "):
            p = document.add_paragraph(style="Heading 2")
            add_runs(p, line[4:].replace("<br/>", "\n"), bold=True)
        elif re.match(r"^[-*] ", line):
            p = document.add_paragraph(style="Normal")
            p.style = "List Bullet"
            add_runs(p, line[2:])
        elif re.match(r"^\d+\. ", line):
            p = document.add_paragraph(style="Normal")
            p.style = "List Number"
            add_runs(p, re.sub(r"^\d+\. ", "", line))
        else:
            p = document.add_paragraph(style="Normal")
            add_runs(p, line)
        index += 1

    output = Path(args.output)
    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(output)


if __name__ == "__main__":
    main()
