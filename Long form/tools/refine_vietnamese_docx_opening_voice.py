from pathlib import Path
import re

from docx import Document


PATH = Path(r"D:\CODE\VGC\Long form\The-Art-of-Monetization-Vietnamese-Editable.docx")


def replace_paragraph(document, index, expected, replacement):
    paragraph = document.paragraphs[index]
    normalized = lambda value: re.sub(r"\s+", " ", value).strip()
    if normalized(paragraph.text) != normalized(expected):
        raise ValueError(f"Paragraph {index} has been edited since review; it was not replaced.")
    paragraph.clear()
    paragraph.add_run(replacement)


def main():
    document = Document(PATH)
    edits = {
        4: (
            "Trong một mô hình kinh doanh game, một quyết định đúng hiếm khi đến từ riêng một cơ chế, một mẫu quảng cáo nổi bật hay một bảng số liệu. Một cơ chế phổ biến chưa tạo thành luận điểm sản phẩm. Một mẫu quảng cáo táo bạo không bù nổi economics yếu. Bảng số liệu không thay được cảm giác chơi, tay nghề thiết kế level hay quyết định khó nhất: dừng một ý tưởng trông đầy hứa hẹn.",
            "Trong kinh doanh game, rất ít quyết định có thể đứng vững nếu chỉ dựa vào một mechanic, một creative nổi bật hoặc một dashboard đẹp. Một mechanic được ưa chuộng chưa đủ để trở thành định hướng sản phẩm; một creative táo bạo cũng không thể che đi một economy thiếu bền vững. Dashboard có thể chỉ ra nơi team cần nhìn, nhưng không thay thế được cảm nhận khi chơi, năng lực thiết kế level hay quyết định khó nhất: dừng một ý tưởng có vẻ hứa hẹn.",
        ),
        5: (
            "Nghiên cứu này bắt đầu từ một Game Event do Publisher tổ chức và từ việc quan sát kỹ hơn thể loại puzzle. Cơ hội thị trường rất tiềm năng.",
            "Nghiên cứu này bắt đầu từ một Game Event do publisher tổ chức, rồi mở rộng thành việc quan sát kỹ hơn category puzzle. Thị trường có cơ hội rõ ràng, nhưng cơ hội chỉ trở nên có ý nghĩa khi một team hiểu mình đang tham gia vào phần nào của nó.",
        ),
        6: (
            "Tuy nhiên, cơ chế chơi, chủ đề và cách kiếm tiền của những game thành công được sao chép nhanh hơn khả năng giải thích nhu cầu người chơi, khối lượng sản xuất và logic kinh tế phía sau chúng. Vì vậy, Ebook này được viết để giúp mọi người tìm hiểu và nghiên cứu nhanh hơn.",
            "Cơ chế chơi, chủ đề và cách kiếm tiền của những game thành công thường được sao chép nhanh hơn tốc độ mà chúng được lý giải. Đằng sau một lựa chọn tưởng như đơn giản luôn có nhu cầu của người chơi, khối lượng sản xuất và logic economy riêng. Ebook này được viết để giúp người đọc nghiên cứu những điều đó kỹ hơn, trước khi biến một tín hiệu thị trường thành quyết định sản phẩm.",
        ),
        7: (
            "Framework trong tài liệu này gồm giả thuyết, công cụ, bằng chứng công khai và câu hỏi. Nội dung sẽ tiếp tục được cập nhật dựa trên góp ý của độc giả, kinh nghiệm phát hành game, hành vi người chơi và dữ liệu theo cohort.",
            "Tài liệu này tập hợp các giả thuyết, công cụ, bằng chứng công khai và những câu hỏi chưa có đáp án cuối cùng. Nội dung sẽ tiếp tục được cập nhật từ góp ý của độc giả, kinh nghiệm phát hành game, hành vi người chơi và dữ liệu theo cohort.",
        ),
        8: (
            "Nếu có câu hỏi, phản biện hoặc đóng góp, hãy chia sẻ các vấn đề cần giải quyết để cập nhật nội dung ebook chính xác hơn, thông qua các bằng chứng mâu thuẫn với nó, hoặc một decision tool đã giúp team tránh một quyết định xấu.",
            "Phản biện có giá trị nhất khi chỉ ra điều kiện khiến một kết luận không còn đúng, đưa ra bằng chứng mâu thuẫn hoặc chia sẻ một decision tool đã giúp team tránh được một quyết định xấu. Những đóng góp như vậy sẽ giúp ebook chính xác hơn qua từng lần cập nhật.",
        ),
        9: (
            "Mục tiêu là đóng góp một ngôn ngữ rõ ràng hơn cho các quyết định làm game và cải thiện framework này cùng những người làm việc gần sản phẩm nhất.",
            "Mục tiêu là tạo ra một ngôn ngữ rõ ràng hơn cho các quyết định làm game, rồi cải thiện ngôn ngữ đó cùng những người đang làm việc gần sản phẩm nhất.",
        ),
        11: (
            "Làm game đã khó. Kiếm tiền từ game còn khó hơn, vì thiết kế, kinh tế trong game, thu hút người chơi, sản phẩm, dữ liệu và vận hành phải gặp nhau tại cùng một điểm.",
            "Làm game vốn đã khó. Việc kiếm tiền còn phức tạp hơn, vì thiết kế, economy, thu hút người chơi, sản phẩm, dữ liệu và vận hành phải gặp nhau trong cùng một trải nghiệm.",
        ),
        12: (
            "Mỗi vai trò có kỹ năng và kinh nghiệm khác nhau, nhưng để làm việc cùng nhau cần một bộ hiểu biết và ngôn ngữ chung: founder, người phụ trách sản phẩm, game designer, analyst, người làm UA, publisher, hoặc một team nhỏ nơi mỗi người phải đảm nhận nhiều vai.",
            "Mỗi vai trò mang theo kỹ năng và kinh nghiệm riêng. Để cùng đưa một game đi xa hơn, founder, người phụ trách sản phẩm, game designer, analyst, người làm UA, publisher và những team nhỏ phải dùng chung một số khái niệm, dù mỗi người tiếp cận chúng từ một góc khác.",
        ),
        13: (
            "Mỗi framework trong sách là một checklist để kiểm tra một nguyên mẫu: mẫu quảng cáo nào cần test, level nào cần xem, quảng cáo nên đặt ở đâu, gói bán nào có lý do tồn tại, chỉ số nào cần đọc và khi nào nên dừng lại.",
            "Mỗi framework trong sách là một cách kiểm tra một nguyên mẫu đang hình thành: creative nào cần test, level nào cần xem lại, quảng cáo nên xuất hiện ở đâu, offer nào có lý do tồn tại, chỉ số nào cần được đọc cùng nhau và khi nào nên dừng một hướng đi.",
        ),
        14: (
            "Hãy mở sổ ghi chú cùng ebook này. Đặt bản game và bảng số liệu bên cạnh. Một ebook hay về kiếm tiền trong game phải khiến bạn mở game ra với ít chắc chắn hơn, nhưng nhiều câu hỏi tốt hơn.",
            "Hãy đọc cùng một cuốn sổ ghi chú, bản game và bảng số liệu. Giá trị của ebook không nằm ở việc đưa ra câu trả lời thay cho team, mà ở việc giúp team mở game ra với những câu hỏi chính xác hơn.",
        ),
    }
    for index, edit in edits.items():
        replace_paragraph(document, index, *edit)
    document.save(PATH)
    print(f"Updated {len(edits)} opening paragraphs.")


if __name__ == "__main__":
    main()
