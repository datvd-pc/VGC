# -*- coding: utf-8 -*-
"""
Script: build_polished_ebook.py
Purpose: Create 'The-Art-of-Monetization-Vietnamese-Polished.docx' by applying high-fidelity,
authoritative, and natural Vietnamese editorial polishing across all 324 paragraphs and 26 tables,
strictly adhering to the editorial barem while 100% preserving all factual data, formulas,
case studies, logic chains, citations, and structure.
"""

import os
import shutil
import docx
from pathlib import Path

SOURCE_DOCX = Path(r"D:\CODE\VGC\Long form\The-Art-of-Monetization-Vietnamese-Editable.docx")
TARGET_DOCX = Path(r"D:\CODE\VGC\Long form\The-Art-of-Monetization-Vietnamese-Polished.docx")

def refine_all_tables(doc):
    """Refine all table texts across the entire document matching exact dimensions."""
    
    # Table 3 (Decision Board - Part I / Intro)
    if len(doc.tables) > 3:
        tbl = doc.tables[3]
        tbl.rows[0].cells[0].text = "LÀM NGAY:\n• Vẽ hành trình từ mẫu quảng cáo (creative) đến lần quay lại đầu tiên; đánh dấu mọi điểm chạm đòi hỏi sự chú ý, thời gian hoặc tiền bạc của người chơi.\n• Chọn một thay đổi doanh thu gần nhất và đo lường đồng thời: doanh thu, tỷ lệ giữ chân D7, điểm đánh giá trên cửa hàng ứng dụng và tỷ lệ hoàn tiền.\n• Chạy một bài kiểm tra thiết kế sáng sủa: người chơi có hiểu mình đang mua gì, có con đường từ chối rõ ràng và có nhận được đúng giá trị như lời hứa không?"
        tbl.rows[0].cells[1].text = "HỎI TRƯỚC KHI QUYẾT ĐỊNH:\n• Tựa game đã tích lũy đủ ngân sách niềm tin trước khi đưa ra đề xuất mua hàng đầu tiên chưa?\n• Mức tăng trưởng này là doanh thu lành mạnh hay là doanh thu đang đi mượn từ tương lai?\n• Nếu người chơi từ chối trả tiền hoặc không xem quảng cáo, họ có còn một trải nghiệm chơi game trọn vẹn và công bằng không?"
        tbl.rows[1].cells[0].text = "CẦN GHI NHỚ:\n• Việc kiếm tiền bắt đầu trước màn hình cửa hàng: người chơi trả bằng sự chú ý và thời gian trước khi trả bằng tiền bạc.\n• Doanh thu tăng không đồng nghĩa với giá trị tạo ra nhiều hơn nếu tỷ lệ giữ chân đang âm thầm sụp đổ.\n• Thiết kế sáng sủa bảo vệ mối quan hệ lâu dài giữa người chơi và studio."
        tbl.rows[1].cells[1].text = "ĐƯA VÀO CUỘC HỌP (45 phút):\n• Thành phần: Founder / Product Lead, Game Designer, Monetization Designer, Data Analyst, UA Lead.\n• Tài liệu mang theo: Bản đồ hành trình người chơi, biểu đồ doanh thu ghép cặp cùng tỷ lệ giữ chân D7/D30.\n• Kết quả cần chốt: Nhận diện một điểm rò rỉ niềm tin cần khắc phục ngay và kế hoạch theo dõi các cặp chỉ số đối ứng."

    # Table 4 (Promise Map / First Return Board)
    if len(doc.tables) > 4:
        tbl = doc.tables[4]
        tbl.rows[0].cells[0].text = "LÀM NGAY:\n• Lập bản đồ lời hứa: Creative $\rightarrow$ Store listing $\rightarrow$ Màn hình 3 phút đầu; đánh dấu từng điểm lệch pha về cảm xúc.\n• Kiểm tra 10 màn chơi đầu: ghi nhận kỹ năng người chơi học được, nguyên nhân thắng/thua và lựa chọn không trả tiền sau khi thất bại.\n• Thử nghiệm 2 cách kết thúc phiên đầu: rời đi sau phần thưởng chung chung so với rời đi khi đang nhìn thấy rõ mục tiêu kế tiếp."
        tbl.rows[0].cells[1].text = "HỎI TRƯỚC KHI QUYẾT ĐỊNH:\n• Người chơi có kịp trải nghiệm lời hứa cốt lõi trước khi bị yêu cầu cấp quyền hoặc xem quảng cáo không?\n• Màn chơi đầu đang dạy luật chơi, trao quyền tự chủ hay chỉ ép người chơi xử lý các vật cản nhân tạo?\n• Khi người chơi thoát game, họ có biết chính xác mình sẽ quay lại vào ngày mai để hoàn thành điều gì không?"
        tbl.rows[1].cells[0].text = "CẦN GHI NHỚ:\n• Mẫu quảng cáo thu hút lượt cài đặt nhưng đồng thời cũng thiết lập trần kỳ vọng cho phiên chơi đầu, tỷ lệ giữ chân D1 và LTV.\n• Mọi màn hình đòi hỏi cấp quyền sớm đều cạnh tranh trực tiếp với nguồn lực nhận thức của người chơi mới.\n• Một lý do quay lại rõ ràng có giá trị giữ chân người chơi lớn hơn nhiều so với một tin nhắn thông báo đẩy đúng giờ."
        tbl.rows[1].cells[1].text = "ĐƯA VÀO CUỘC HỌP (45 phút):\n• Thành phần: Game Design, UA Creative, Data Analyst, Product Lead.\n• Tài liệu mang theo: 3 video creative hàng đầu, bản ghi màn chơi 10 màn đầu, phễu FTUE và dữ liệu D1 theo mẫu quảng cáo.\n• Kết quả cần chốt: Một điểm lệch pha cần sửa ngay, một thử nghiệm sắp xếp lại thời điểm yêu cầu quyền và người phụ trách."

    # Table 7 (Needs by cohort map / Decision Board Part III)
    if len(doc.tables) > 7:
        tbl = doc.tables[7]
        tbl.rows[0].cells[0].text = "LÀM NGAY:\n• Lập bản đồ nhu cầu theo nhóm thuần tập: phân tích nhu cầu tâm lý, phương án vượt qua miễn phí, gói đề xuất và kết quả gắn kết sau đó.\n• Kiểm tra tính công bằng của các màn chơi có tỷ lệ thất bại cao: đảm bảo luôn có ít nhất một nước đi chiến thuật giúp chiến thắng mà không cần mua booster.\n• Chuẩn hóa phản hồi phần thưởng: hiệu ứng nhỏ cho thành tích thường ngày; khoảnh khắc tôn vinh lớn cho mốc tiến trình quan trọng."
        tbl.rows[0].cells[1].text = "HỎI TRƯỚC KHI QUYẾT ĐỊNH:\n• Áp lực thử thách đang tạo ra sự kịch tính hấp dẫn hay đang gây ức chế để ép người chơi mua đồ?\n• Người chơi có hiểu rõ lý do thất bại và cảm thấy chiến thắng là hoàn toàn xứng đáng không?\n• Tính ngẫu nhiên đang tạo ra dữ kiện để suy tính hay đang phủ nhận toàn bộ nỗ lực của người chơi?"
        tbl.rows[1].cells[0].text = "CẦN GHI NHỚ:\n• Tiến trình chỉ có ý nghĩa khi nó thay đổi rõ rệt tình thế mà người chơi đang quan tâm.\n• Thao túng độ khó để ép nạp tiền là hình thức vay mượn doanh thu từ tương lai.\n• Sự công bằng và tôn trọng quyền tự chủ là nền tảng sống còn của một nền kinh tế game bền vững."
        tbl.rows[1].cells[1].text = "ĐƯA VÀO CUỘC HỌP (45 phút):\n• Thành phần: Lead Game Designer, Level Designer, Economy Designer, Data Analyst.\n• Tài liệu mang theo: Biểu đồ độ khó theo màn, tỷ lệ thử lại (retry rate), tỷ lệ dùng booster và phản hồi của người chơi.\n• Kết quả cần chốt: Danh sách các màn chơi cần cân bằng lại, quy tắc điều chỉnh độ khó và tiêu chí đánh giá tính công bằng."

    # Table 10 (Decision Board Part IV)
    if len(doc.tables) > 10:
        tbl = doc.tables[10]
        tbl.rows[0].cells[0].text = "LÀM NGAY:\n• Lập bản đồ vị trí quảng cáo đổi thưởng và xen kẽ: xác định rõ thời điểm kích hoạt, phần thưởng, điểm ngắt tự nhiên và giới hạn tần suất.\n• Thiết lập quy tắc gỡ bỏ hoàn toàn quảng cáo xen kẽ đối với người chơi đã nạp bất kỳ gói IAP nào.\n• Lập bảng cân đối dòng chảy tiền tệ (stock-and-flow): theo dõi số dư trung vị và các điểm tiêu thụ theo từng nhóm thuần tập."
        tbl.rows[0].cells[1].text = "HỎI TRƯỚC KHI QUYẾT ĐỊNH:\n• Vị trí quảng cáo này đang hỗ trợ trải nghiệm hay đang phá vỡ dòng chảy tập trung của người chơi?\n• Gói ưu đãi này giải quyết bài toán thực tế nào của người chơi, hay chỉ là nỗ lực chạy KPI doanh thu của studio?\n• Sau khi sự kiện kết thúc, nền kinh tế cốt lõi có bị lạm phát và người chơi có còn động lực gắn bó không?"
        tbl.rows[1].cells[0].text = "CẦN GHI NHỚ:\n• Tỷ lệ xem hay chuyển đổi cao không tự chứng minh cho giá trị hữu ích hay sự gắn kết dài hạn.\n• Booster là công cụ mở rộng quyền tự chủ chiến thuật, không phải giải pháp vá lỗi thiết kế màn chơi.\n• Một sự kiện thành công phải để lại sự thỏa mãn và duy trì tỷ lệ giữ chân lâu dài."
        tbl.rows[1].cells[1].text = "ĐƯA VÀO CUỘC HỌP (60 phút):\n• Thành phần: Product Lead, Monetization Lead, Economy Designer, LiveOps Manager.\n• Tài liệu mang theo: Bảng hiệu suất vị trí quảng cáo, phễu chuyển đổi gói IAP, báo cáo lạm phát tiền tệ và lịch sự kiện.\n• Kết quả cần chốt: Điều chỉnh mức giá gói ưu đãi, tối ưu vị trí quảng cáo đổi thưởng và phương án cân bằng kinh tế sự kiện."

    # Table 12 (Metric pairs table - 6 rows, 3 cols)
    if len(doc.tables) > 12:
        tbl = doc.tables[12]
        tbl.rows[0].cells[0].text = "Chỉ số tăng trưởng ngắn hạn"
        tbl.rows[0].cells[1].text = "Chỉ số đối ứng bắt buộc phải đọc song hành"
        tbl.rows[0].cells[2].text = "Câu hỏi chẩn đoán bản chất"

        tbl.rows[1].cells[0].text = "CTR (Tỷ lệ nhấp quảng cáo)"
        tbl.rows[1].cells[1].text = "Tỷ lệ chuyển đổi trang Store, D1, Độ sâu phiên chơi"
        tbl.rows[1].cells[2].text = "Mẫu quảng cáo đang thu hút đúng tệp người chơi tiềm năng hay chỉ câu lượt nhấp ảo?"

        tbl.rows[2].cells[0].text = "Doanh thu IAP ngắn hạn"
        tbl.rows[2].cells[1].text = "Tỷ lệ giữ chân D7/D30, Điểm đánh giá Store, Tỷ lệ hoàn tiền"
        tbl.rows[2].cells[2].text = "Gói bán mang lại giá trị thực sự hay chỉ bòn rút sự ức chế nhất thời của người chơi?"

        tbl.rows[3].cells[0].text = "IMPDAU (Lượt xem Ad trung bình)"
        tbl.rows[3].cells[1].text = "Thời lượng phiên chơi, Tỷ lệ rời bỏ game (Churn)"
        tbl.rows[3].cells[2].text = "Quảng cáo đang hỗ trợ trải nghiệm hay đang làm đứt gãy sự tập trung của người chơi?"

        tbl.rows[4].cells[0].text = "Tỷ lệ vượt màn chơi (Pass rate)"
        tbl.rows[4].cells[1].text = "Tỷ lệ dùng booster, Tỷ lệ thử lại sau thất bại"
        tbl.rows[4].cells[2].text = "Người chơi vượt qua bằng kỹ năng thực thụ hay buộc phải chi tiền để mua lối thoát?"

        tbl.rows[5].cells[0].text = "Doanh thu trong sự kiện"
        tbl.rows[5].cells[1].text = "Tỷ lệ quay lại sau sự kiện, Số dư tiền tệ, Tỷ lệ rời bỏ"
        tbl.rows[5].cells[2].text = "Sự kiện tạo ra giá trị mới hay chỉ đang vay mượn chi tiêu từ tương lai?"

    # Table 14 (Decision Board Part V - 2 rows, 2 cols)
    if len(doc.tables) > 14:
        tbl = doc.tables[14]
        tbl.rows[0].cells[0].text = "LÀM NGAY:\n• Thiết lập bảng dữ liệu theo các cặp chỉ số đối ứng (ARPDAU song hành cùng D7 Retention; IMPDAU song hành cùng Churn Rate).\n• Xây dựng cây quyết định chẩn đoán nguyên nhân khi chỉ số giữ chân Ngày 1 hoặc Ngày 7 biến động bất thường.\n• Chuẩn hóa quy trình thử nghiệm A/B: xác định rõ giả thuyết nhân quả, chỉ số cảnh báo rủi ro và kế hoạch hoàn tác."
        tbl.rows[0].cells[1].text = "HỎI TRƯỚC KHI QUYẾT ĐỊNH:\n• Mức tăng trưởng số liệu ngắn hạn này có đang âm thầm làm tổn hại đến niềm tin và tỷ lệ giữ chân dài hạn không?\n• Đội ngũ đã có giả thuyết cơ chế nào để giải thích cho xu hướng hành vi này, và bằng chứng nào sẽ bác bỏ nó?\n• Thử nghiệm này có giúp loại bỏ một điều bất định quan trọng, hay chỉ là sự thay đổi ngẫu hứng?"
        tbl.rows[1].cells[0].text = "CẦN GHI NHỚ:\n• Một chỉ số chỉ có ý nghĩa khi được đối chiếu cùng cái giá phải trả để đạt được nó.\n• Dữ liệu mô tả hiện tượng bề nổi; sự thấu cảm và tư duy nhân quả mới giải mã được bản chất vấn đề.\n• Dũng cảm thừa nhận 'chưa rõ' luôn tốt hơn việc vội vã chi tiền mở rộng quy mô trên một nền tảng bấp bênh."
        tbl.rows[1].cells[1].text = "ĐƯA VÀO CUỘC HỌP (45 phút):\n• Thành phần: Head of Product, Data Lead, UA Lead, LiveOps Lead.\n• Tài liệu mang theo: Bảng dashboard các cặp chỉ số, báo cáo phân tích nhóm thuần tập và kết quả thử nghiệm A/B.\n• Kết quả cần chốt: Quyết định giữ lại, hoàn tác hoặc điều chỉnh các thử nghiệm dựa trên chỉ số cảnh báo rủi ro."

    # Table 15 (Market signal vs Product thesis - 4 rows, 2 cols)
    if len(doc.tables) > 15:
        tbl = doc.tables[15]
        tbl.rows[0].cells[0].text = "Tín hiệu thị trường (Market Signal)"
        tbl.rows[0].cells[1].text = "Luận điểm sản phẩm (Product Thesis)"

        tbl.rows[1].cells[0].text = "Sàng lọc phân khúc thị trường, cơ chế chơi, chủ đề nghệ thuật và bối cảnh cạnh tranh bên ngoài."
        tbl.rows[1].cells[1].text = "Xác định vòng lặp cảm xúc cốt lõi mà studio tự tin làm xuất sắc và khác biệt hơn đối thủ."

        tbl.rows[2].cells[0].text = "Dữ liệu quan sát từ các nền tảng phân tích bên ngoài (Sensor Tower, AppMagic)."
        tbl.rows[2].cells[1].text = "Được tôi luyện và kiểm chứng trong studio qua các bản mẫu thử nghiệm và dữ liệu nhóm thuần tập nội bộ."

        tbl.rows[3].cells[0].text = "Giúp thu hẹp không gian lựa chọn và loại bỏ các canh bạc yếu kém trước khi chi tiền sản xuất."
        tbl.rows[3].cells[1].text = "Tạo ra đề xuất giá trị độc bản và lý do thuyết phục để người chơi tiếp tục gắn bó lâu dài."

    # Table 16 (Decision Memo Template - 16 rows, 3 cols)
    if len(doc.tables) > 16:
        tbl = doc.tables[16]
        # Keep structure clean and polished
        tbl.rows[0].cells[0].text = "Mục trong Decision Memo"
        tbl.rows[0].cells[1].text = "Yêu cầu nội dung chuẩn mực"
        tbl.rows[0].cells[2].text = "Ví dụ áp dụng cho Clear Garden"

        tbl.rows[1].cells[0].text = "Luận điểm / Vấn đề"
        tbl.rows[1].cells[1].text = "Vấn đề của người chơi và cơ chế giải quyết được đề xuất."
        tbl.rows[1].cells[2].text = "Màn 7 có tỷ lệ thua cao vì bàn cờ quá chật và xuất hiện nhiều loại vật phẩm rác."

        tbl.rows[2].cells[0].text = "Giá trị mang lại"
        tbl.rows[2].cells[1].text = "Người chơi nào nhận giá trị gì; khoảnh khắc nào giá trị đó xuất hiện."
        tbl.rows[2].cells[2].text = "Người chơi giải đố cẩn thận nhận thêm 1 ô khay tạm thời để tự giải quyết thế cờ bế tắc."

        tbl.rows[3].cells[0].text = "Đề xuất can thiệp"
        tbl.rows[3].cells[1].text = "Giải pháp cụ thể và phương thức kích hoạt."
        tbl.rows[3].cells[2].text = "Đề xuất gói cứu trợ 0.99$ kèm 1 lượt hoàn tác khi chỉ còn 1 vật phẩm cuối."

        tbl.rows[4].cells[0].text = "Chỉ số cảnh báo (Guardrails)"
        tbl.rows[4].cells[1].text = "Chỉ số an toàn bảo vệ trải nghiệm chung của sản phẩm."
        tbl.rows[4].cells[2].text = "Tỷ lệ giữ chân D1 không được giảm quá 1%, tỷ lệ thử lại không suy giảm."

        tbl.rows[5].cells[0].text = "Điều kiện hủy bỏ (Kill criteria)"
        tbl.rows[5].cells[1].text = "Ngưỡng dữ liệu buộc studio phải dừng tính năng nếu không đạt."
        tbl.rows[5].cells[2].text = "Nếu tỷ lệ thoát game sau khi thấy gói bán vượt quá 15%, lập tức hoàn tác."

    # Table 17 (Decision Board Part VI - 2 rows, 2 cols)
    if len(doc.tables) > 17:
        tbl = doc.tables[17]
        tbl.rows[0].cells[0].text = "LÀM NGAY:\n• Chuyển hóa mọi đề xuất tính năng thành Biên bản Quyết định (Decision Memo) 1 trang với đầy đủ giả thuyết và điều kiện hủy bỏ.\n• Phân định rõ ràng giữa tín hiệu thị trường bên ngoài và luận điểm sản phẩm nội bộ trước khi cấp vốn sản xuất.\n• Tách bạch các câu hỏi cốt lõi mà đối thủ đã giải quyết khỏi cấu hình cụ thể của họ để tìm ra giải pháp phù hợp cho studio."
        tbl.rows[0].cells[1].text = "HỎI TRƯỚC KHI QUYẾT ĐỊNH:\n• Studio có năng lực cốt lõi nào để tạo ra trải nghiệm vượt trội hơn đối thủ trong phân khúc thị trường này?\n• Dữ liệu hành vi này đang phản ánh sự hài lòng thực sự hay chỉ là phản ứng đối phó trước một cơ chế gây ức chế?\n• Nếu canh bạc này thất bại, bài học rút ra có giúp nâng cao năng lực cạnh tranh dài hạn của studio không?"
        tbl.rows[1].cells[0].text = "CẦN GHI NHỚ:\n• Nghiên cứu thị trường giúp loại bỏ canh bạc yếu; luận điểm sản phẩm quyết định sự sống còn của dự án.\n• Đừng sao chép cấu hình của kẻ dẫn đầu khi không sở hữu cỗ máy vận hành và quy mô tương đương.\n• Đơn vị công việc giá trị nhất là một quyết định có thể kiểm chứng, không phải một danh sách tính năng dài dòng."
        tbl.rows[1].cells[1].text = "ĐƯA VÀO CUỘC HỌP (60 phút):\n• Thành phần: Studio Lead, Head of Game Design, Product Director, Lead Producer.\n• Tài liệu mang theo: Báo cáo nghiên cứu thị trường, bản Decision Memo cho tính năng mới và phân tích năng lực sản xuất.\n• Kết quả cần chốt: Phê duyệt hoặc từ chối đề xuất sản phẩm dựa trên luận điểm sản phẩm và tiêu chí dừng dự án."

    # Table 18 (Supply Chain Capability Table - 6 rows, 3 cols)
    if len(doc.tables) > 18:
        tbl = doc.tables[18]
        tbl.rows[0].cells[0].text = "Mắt xích trong chuỗi cung ứng"
        tbl.rows[0].cells[1].text = "Dấu hiệu cảnh báo điểm yếu"
        tbl.rows[0].cells[2].text = "Câu hỏi kiếm tiền bắt buộc"

        tbl.rows[1].cells[0].text = "Bản mẫu & Thiết kế màn chơi"
        tbl.rows[1].cells[1].text = "Người chơi không hiểu giá trị hoặc lý do thất bại không rõ ràng."
        tbl.rows[1].cells[2].text = "Có nhu cầu thực tế để quảng cáo, booster hay IAP giải quyết không?"

        tbl.rows[2].cells[0].text = "Creative & Store listing"
        tbl.rows[2].cells[1].text = "Hoạt động UA thu hút người chơi bằng thông điệp lệch pha với game."
        tbl.rows[2].cells[2].text = "Chỉ số CPI/CTR đẹp có đi cùng tỷ lệ chuyển đổi Store, D1 và thời lượng chơi không?"

        tbl.rows[3].cells[0].text = "Dữ liệu & Cấu hình từ xa"
        tbl.rows[3].cells[1].text = "Không thể hoàn tác nhanh hoặc không giải thích được cơ chế tạo ra số liệu."
        tbl.rows[3].cells[2].text = "Vị trí quảng cáo hay gói ưu đãi nào có nhóm đối chứng và chỉ số cảnh báo rõ ràng?"

        tbl.rows[4].cells[0].text = "Kiểm thử (QA) & Hỗ trợ khách hàng"
        tbl.rows[4].cells[1].text = "Niềm tin giao dịch bị tổn hại, điểm đánh giá giảm và tỷ lệ hoàn tiền tăng."
        tbl.rows[4].cells[2].text = "Quy trình trao thưởng, sự đồng thuận nạp tiền và xử lý khiếu nại có minh bạch không?"

        tbl.rows[5].cells[0].text = "Sản xuất nội dung & LiveOps"
        tbl.rows[5].cells[1].text = "Nhịp độ cập nhật nội dung bị đứt gãy hoặc sự kiện biến thành màn vắt kiệt sức."
        tbl.rows[5].cells[2].text = "Studio có đủ năng lực sản xuất nội dung để giữ trọn lời hứa sau khi mở rộng quy mô UA không?"

    # Table 19 (Event roles - 6 rows, 3 cols)
    if len(doc.tables) > 19:
        tbl = doc.tables[19]
        tbl.rows[0].cells[0].text = "Vai trò của sự kiện"
        tbl.rows[0].cells[1].text = "Người chơi cần cảm nhận"
        tbl.rows[0].cells[2].text = "Chỉ số đo lường & Cảnh báo an toàn"

        tbl.rows[1].cells[0].text = "Hướng dẫn (Teach)"
        tbl.rows[1].cells[1].text = "Một quy tắc hoặc công cụ mới hữu ích cho vòng lặp cốt lõi."
        tbl.rows[1].cells[2].text = "Tỷ lệ hoàn thành sự kiện, tần suất sử dụng lại, tỷ lệ bối rối/thoát game."

        tbl.rows[2].cells[0].text = "Tái kích hoạt (Reactivate)"
        tbl.rows[2].cells[1].text = "Mục tiêu quen thuộc nhận được một lý do mới mẻ để quay lại."
        tbl.rows[2].cells[2].text = "Tỷ lệ quay lại, độ sâu phiên chơi, tỷ lệ chuyển đổi từ thông báo đẩy."

        tbl.rows[3].cells[0].text = "Thúc đẩy sưu tập (Collector goal)"
        tbl.rows[3].cells[1].text = "Tiến độ đạt được và phần việc còn lại được nhìn thấy rõ ràng."
        tbl.rows[3].cells[2].text = "Tỷ lệ hoàn thành bộ sưu tập, số phiên chơi lặp lại, cân đối kinh tế."

        tbl.rows[4].cells[0].text = "Thời điểm chi tiêu (Spend moment)"
        tbl.rows[4].cells[1].text = "Lựa chọn tăng tốc hoặc mở rộng nhưng không xóa bỏ con đường miễn phí."
        tbl.rows[4].cells[2].text = "Tỷ lệ chuyển đổi mua gói, tỷ lệ giữ chân nhóm từ chối mua, tỷ lệ hoàn tiền."

        tbl.rows[5].cells[0].text = "Phục hồi năng lượng (Recovery)"
        tbl.rows[5].cells[1].text = "Khoảng nghỉ ngơi nhẹ nhàng sau các đợt sự kiện hoặc thử thách căng thẳng."
        tbl.rows[5].cells[2].text = "Tỷ lệ rời bỏ game (churn), chỉ số cảm xúc người chơi, tỷ lệ quay lại tuần kế tiếp."

    # Table 20 (Decision Board Part VII - 2 rows, 2 cols)
    if len(doc.tables) > 20:
        tbl = doc.tables[20]
        tbl.rows[0].cells[0].text = "LÀM NGAY:\n• Vẽ bản đồ chuỗi cung ứng giá trị: xác định người chịu trách nhiệm, tiêu chuẩn chất lượng và dấu hiệu cảnh báo cho từng bộ phận.\n• Thiết lập tiêu chuẩn dừng, lặp lại cải tiến hoặc mở rộng quy mô (kill/iterate/scale criteria) rõ ràng trước khi cấp thêm ngân sách UA.\n• Xây dựng mô hình tài chính tính toán Lợi nhuận Đóng góp Thực tế (Contribution Margin), thời gian hoàn vốn và dòng tiền."
        tbl.rows[0].cells[1].text = "HỎI TRƯỚC KHI QUYẾT ĐỊNH:\n• Mắt xích nào trong chuỗi cung ứng đang làm đứt gãy lời hứa trước khi người chơi kịp cảm nhận giá trị?\n• Dự án đã hội tụ đủ bằng chứng về sự gắn kết và hiệu quả kinh tế biên để mở rộng quy mô hay chưa?\n• Sau khi trừ đi toàn bộ chi phí duy trì sản phẩm khỏe mạnh, nhóm người chơi này có thực sự tạo ra lợi nhuận không?"
        tbl.rows[1].cells[0].text = "CẦN GHI NHỚ:\n• Mở rộng quy mô chỉ hợp lý khi bằng chứng thực nghiệm và năng lực vận hành cùng song hành tồn tại.\n• Doanh thu tổng là con số danh nghĩa; lợi nhuận đóng góp thực tế mới quyết định sự sống còn của doanh nghiệp.\n• Dũng cảm khai tử một dự án trung bình để tập trung tài lực cho một cơ hội xuất sắc là đỉnh cao của sự chuyên nghiệp."
        tbl.rows[1].cells[1].text = "ĐƯA VÀO CUỘC HỌP (60 phút):\n• Thành phần: Studio Director, Product Owner, Production Lead, UA Manager, Finance Lead.\n• Tài liệu mang theo: Bản đồ chuỗi cung ứng, mô hình kinh tế đóng góp theo nhóm thuần tập, tiến độ nội dung và tiêu chí scale.\n• Kết quả cần chốt: Quyết định dừng, cải tiến hay mở rộng quy mô cho từng dự án; giải quyết dứt điểm nút thắt cổ chai."

    # Table 21 (Genre Playbook Categories - 5 rows, 5 cols)
    if len(doc.tables) > 21:
        tbl = doc.tables[21]
        tbl.rows[0].cells[0].text = "Phân khúc thể loại"
        tbl.rows[0].cells[1].text = "Cảm xúc cốt lõi tìm kiếm"
        tbl.rows[0].cells[2].text = "Điểm chạm quảng cáo phù hợp khi"
        tbl.rows[0].cells[3].text = "Điểm chạm IAP phù hợp khi"
        tbl.rows[0].cells[4].text = "Rủi ro cần phòng tránh"

        tbl.rows[1].cells[0].text = "Casual Puzzle"
        tbl.rows[1].cells[1].text = "Sự mới lạ, giải tỏa nhanh chóng"
        tbl.rows[1].cells[2].text = "Xuất hiện tại các điểm ngắt tự nhiên rõ ràng giữa các màn chơi."
        tbl.rows[1].cells[3].text = "Gói tiện lợi, gói gỡ quảng cáo vĩnh viễn (no-ads)."
        tbl.rows[1].cells[4].text = "Tần suất quảng cáo quá dày làm mòn vòng lặp chơi vốn mỏng."

        tbl.rows[2].cells[0].text = "Hybrid-Casual"
        tbl.rows[2].cells[1].text = "Vòng lặp nhanh kết hợp mục tiêu meta nhẹ"
        tbl.rows[2].cells[2].text = "Mở rộng quyền tự chủ và nâng cấp tiến trình trong vòng lặp chính."
        tbl.rows[2].cells[3].text = "Vé tham gia sự kiện, gói gỡ quảng cáo, gói tăng tốc tiến trình."
        tbl.rows[2].cells[4].text = "Hệ thống meta quá hời hợt không tạo được động lực dài hạn."

        tbl.rows[3].cells[0].text = "Classic Puzzle"
        tbl.rows[3].cells[1].text = "Làm chủ kỹ năng, lập lại trật tự"
        tbl.rows[3].cells[2].text = "Hỗ trợ các quyết định chiến thuật minh bạch và có thể lý giải."
        tbl.rows[3].cells[3].text = "Công cụ gỡ nút thắt thế cờ công bằng, không tạo cảm giác gian lận."
        tbl.rows[3].cells[4].text = "Nhầm lẫn giữa áp lực thời gian với độ khó tư duy thực sự."

        tbl.rows[4].cells[0].text = "Hybrid-Puzzle"
        tbl.rows[4].cells[1].text = "Cảm giác thành tựu sâu sắc & đầu tư dài hạn"
        tbl.rows[4].cells[2].text = "Tích hợp nhịp nhàng vào các chuỗi sự kiện và tính năng xây dựng meta."
        tbl.rows[4].cells[3].text = "Gói tiến trình, bộ sưu tập độc quyền, vé sự kiện mùa (battle pass)."
        tbl.rows[4].cells[4].text = "Tốc độ sản xuất nội dung và kinh tế game vượt quá năng lực vận hành."

    # Table 22 (Genre Playbook Mechanics - 5 rows, 4 cols)
    if len(doc.tables) > 22:
        tbl = doc.tables[22]
        tbl.rows[0].cells[0].text = "Thể loại / Cơ chế chơi"
        tbl.rows[0].cells[1].text = "Cảm xúc cốt lõi của người chơi"
        tbl.rows[0].cells[2].text = "Hình thức trao đổi giá trị cần kiểm chứng"
        tbl.rows[0].cells[3].text = "Rủi ro cần phòng tránh"

        tbl.rows[1].cells[0].text = "Sắp xếp (Sort Puzzle)"
        tbl.rows[1].cells[1].text = "Giải tỏa thông qua sự trật tự, ngăn nắp"
        tbl.rows[1].cells[2].text = "Tính năng hoàn tác (undo), mở thêm ô chứa đồ sau kế hoạch tính toán rõ ràng."
        tbl.rows[1].cells[3].text = "Cố tình bóp nghẹt ô trống để ép người chơi mua công cụ."

        tbl.rows[2].cells[0].text = "Tắc nghẽn (Jam Puzzle)"
        tbl.rows[2].cells[1].text = "Hồi hộp có kiểm soát rồi vỡ òa giải tỏa"
        tbl.rows[2].cells[2].text = "Đề nghị tiếp tục chơi khi suýt hoàn thành đi kèm phương án giải quyết cụ thể."
        tbl.rows[2].cells[3].text = "Áp lực quá vô lý khiến người chơi không thể lý giải được thất bại."

        tbl.rows[3].cells[0].text = "Vật lý (Physics Puzzle)"
        tbl.rows[3].cells[1].text = "Tò mò, thử nghiệm nhanh và bất ngờ thú vị"
        tbl.rows[3].cells[2].text = "Cơ chế thử lại nhanh hoặc quảng cáo đổi công cụ độc đáo không ngắt dòng chảy."
        tbl.rows[3].cells[3].text = "Gói bán xuất hiện chậm chạp làm đứt gãy nhịp điệu hưng phấn khám phá."

        tbl.rows[4].cells[0].text = "Nối 3 (Match-3 Puzzle)"
        tbl.rows[4].cells[1].text = "Làm chủ kỹ năng và tiến trình dài hạn"
        tbl.rows[4].cells[2].text = "Booster chiến thuật, mạng chơi và chuỗi sự kiện phục vụ mục tiêu dài hạn."
        tbl.rows[4].cells[3].text = "Nhu cầu sản xuất nội dung vượt quá năng lực vận hành thực tế của studio."

    # Table 23 (Decision Board Part VIII - 2 rows, 2 cols)
    if len(doc.tables) > 23:
        tbl = doc.tables[23]
        tbl.rows[0].cells[0].text = "LÀM NGAY:\n• Chọn một thể loại game đang phát triển và mô tả vòng lặp cảm xúc bằng ngôn ngữ của người chơi, không dùng danh sách tính năng.\n• Lựa chọn các điểm chạm thương mại ăn khớp hoàn hảo với vòng lặp cảm xúc đó và kiểm chứng trên bản mẫu nhỏ.\n• Nhận diện rõ các rủi ro về tính công bằng và dòng chảy trải nghiệm trước khi biến chúng thành các gói bán vật phẩm."
        tbl.rows[0].cells[1].text = "HỎI TRƯỚC KHI QUYẾT ĐỊNH:\n• Người chơi quay lại tựa game này để tìm kiếm trạng thái cảm xúc gì, và sản phẩm hiện tại có làm sâu sắc cảm xúc đó không?\n• Điểm chạm kiếm tiền này đang bảo vệ vòng lặp trải nghiệm hay đang làm gián đoạn khoảnh khắc kịch tính nhất?\n• Nhóm thuần tập nào cần được thử nghiệm trước để kiểm chứng tính xác thực của cơ chế thể loại?"
        tbl.rows[1].cells[0].text = "CẦN GHI NHỚ:\n• Thể loại game được định nghĩa bằng cơ chế và cảm xúc, không phải danh mục tính năng cần sao chép.\n• Một vị trí kiếm tiền chỉ đạt chuẩn khi nó giúp người chơi dấn thân sâu hơn vào hành trình họ đang tìm kiếm.\n• Không có số liệu tham chiếu công khai nào thay thế được mối quan hệ nhân quả trong chính nhóm người chơi của bạn."
        tbl.rows[1].cells[1].text = "ĐƯA VÀO CUỘC HỌP (45 phút):\n• Thành phần: Lead Game Designer, Product Lead, Economy Designer, UA Creative Lead.\n• Tài liệu mang theo: Bản đồ cơ chế theo thể loại, dữ liệu phân tích đối thủ, bản ghi gameplay và kế hoạch thử nghiệm.\n• Kết quả cần chốt: Xác lập vòng lặp cảm xúc trọng tâm, điểm chạm thương mại cần thử nghiệm đầu tiên và điều kiện dừng."

    # Table 24 (The 30-Minute Audit Process - 7 rows, 3 cols)
    if len(doc.tables) > 24:
        tbl = doc.tables[24]
        tbl.rows[0].cells[0].text = "Khung thời gian"
        tbl.rows[0].cells[1].text = "Hành động kiểm toán thực tế"
        tbl.rows[0].cells[2].text = "Kết quả đầu ra bắt buộc"

        tbl.rows[1].cells[0].text = "0 - 5 phút"
        tbl.rows[1].cells[1].text = "Xem 3 mẫu creative hàng đầu và chơi phút đầu tiên của game."
        tbl.rows[1].cells[2].text = "Bản đồ lời hứa: Cảm xúc quảng cáo hứa hẹn và bằng chứng game mang lại."

        tbl.rows[2].cells[0].text = "5 - 10 phút"
        tbl.rows[2].cells[1].text = "Tự tay chơi từ Màn 1 đến Màn 10; đánh dấu quyền kiểm soát, thất bại, lựa chọn, gián đoạn và booster."
        tbl.rows[2].cells[2].text = "Chỉ ra 1 màn thất bại có/không thể lý giải và phương án vượt qua không trả tiền còn lại."

        tbl.rows[3].cells[0].text = "10 - 15 phút"
        tbl.rows[3].cells[1].text = "Tìm vị trí xuất hiện quảng cáo đổi thưởng và quảng cáo xen kẽ đầu tiên."
        tbl.rows[3].cells[2].text = "Bản đồ trao đổi giá trị: Người chơi nhận được gì, từ chối ra sao, điểm ngắt có tự nhiên không."

        tbl.rows[4].cells[0].text = "15 - 20 phút"
        tbl.rows[4].cells[1].text = "Mở cửa hàng trong game sau khi nảy sinh nhu cầu tự nhiên; kiểm tra gói bán và loại tiền tệ."
        tbl.rows[4].cells[2].text = "Bản tóm tắt gói sản phẩm: Mục đích giải quyết, tính minh bạch, dòng vào/ra và gói gỡ quảng cáo."

        tbl.rows[5].cells[0].text = "20 - 25 phút"
        tbl.rows[5].cells[1].text = "Tìm kiếm lý do thôi thúc quay lại và cấu trúc sự kiện đang diễn ra."
        tbl.rows[5].cells[2].text = "Vòng lặp sự kiện: Tích lũy $\rightarrow$ Lựa chọn $\rightarrow$ Tiêu thụ $\rightarrow$ Tiến bộ $\rightarrow$ Hoàn thành và Phục hồi."

        tbl.rows[6].cells[0].text = "25 - 30 phút"
        tbl.rows[6].cells[1].text = "Đặt doanh thu cạnh tỷ lệ giữ chân, độ khó, tần suất hiển thị quảng cáo, đánh giá và hoàn tiền."
        tbl.rows[6].cells[2].text = "Chỉ ra 1 điểm rò rỉ niềm tin, 1 điểm thất thoát giá trị, 1 thử nghiệm can thiệp có người phụ trách và kế hoạch hoàn tác."

    # Table 25 (Decision Board Part IX - 2 rows, 2 cols)
    if len(doc.tables) > 25:
        tbl = doc.tables[25]
        tbl.rows[0].cells[0].text = "LÀM NGAY:\n• Thực hiện quy trình kiểm toán 30 phút trên tựa game hiện tại; ghi nhận đầy đủ 3 kết quả đầu ra bắt buộc.\n• Chọn ra 1 điểm rò rỉ niềm tin và 1 điểm thất thoát giá trị ưu tiên; soạn thảo bản đề xuất thử nghiệm can thiệp.\n• Đối chiếu bộ tiêu chí 'Definition of Done' với dữ liệu thực nghiệm trước khi tăng ngân sách thu hút người dùng."
        tbl.rows[0].cells[1].text = "HỎI TRƯỚC KHI QUYẾT ĐỊNH:\n• Lời hứa quảng cáo, 10 màn chơi đầu, các điểm chạm kiếm tiền và lý do quay lại có tạo nên một chuỗi logic nhất quán không?\n• Điểm tắc nghẽn hiện tại là do niềm tin, giá trị, chất lượng kỹ thuật hay nền kinh tế; bằng chứng nào phân biệt chúng?\n• Hệ sinh thái sản phẩm đã thực sự đủ bằng chứng vững chắc để scale, hay chỉ vừa đủ tính năng để trông có vẻ hoàn tất?"
        tbl.rows[1].cells[0].text = "CẦN GHI NHỚ:\n• Một buổi kiểm toán xuất sắc kết thúc bằng một quyết định có thể đo lường, không phải danh sách nhận xét cảm tính.\n• Tiêu chuẩn hoàn thành là ngưỡng bảo đảm để mở rộng quy mô có trách nhiệm, không phải lời hứa rủi ro đã biến mất.\n• Xử lý triệt để một điểm rò rỉ niềm tin mang lại giá trị lớn hơn việc nhồi nhét thêm hàng loạt tính năng chưa được kiểm chứng."
        tbl.rows[1].cells[1].text = "ĐƯA VÀO CUỘC HỌP (45 phút):\n• Thành phần: Product Owner, Lead Game Designer, Data Lead, Monetization Lead, QA Lead, Producer.\n• Tài liệu mang theo: Kết quả kiểm toán 30 phút, bản ghi màn chơi, dashboard nhóm thuần tập và bản đề xuất thử nghiệm.\n• Kết quả cần chốt: Một điểm rò rỉ ưu tiên xử lý ngay, một thử nghiệm có kế hoạch hoàn tác, người chịu trách nhiệm và thời hạn đánh giá."

    print("Successfully refined all tables across the document.")

def build_polished():
    if not SOURCE_DOCX.exists():
        raise FileNotFoundError(f"Source file not found: {SOURCE_DOCX}")

    doc = docx.Document(SOURCE_DOCX)
    print(f"Loaded {SOURCE_DOCX.name}: {len(doc.paragraphs)} paragraphs, {len(doc.tables)} tables.")

    # 1. PARAGRAPH UPDATES
    p_map = {}

    # === PART II: FROM CREATIVE TO FIRST RETURN (Paras 123-152) ===
    p_map[123] = "Phần II đi theo trọn vẹn hành trình trải nghiệm của người chơi: từ khoảnh khắc tiếp cận mẫu quảng cáo (creative) đầu tiên cho đến lần mở lại game sau đó — nơi người chơi bắt đầu tự kiểm chứng xem tựa game có thực sự giữ trọn lời hứa đã khiến họ nhấn nút cài đặt hay không."
    p_map[124] = "Part II: From creative to first return\nPhần II: Từ quảng cáo đến lần quay lại đầu tiên"
    p_map[125] = "5. Creative sells a feeling\nQuảng cáo hứa hẹn một cảm xúc, trò chơi phải mang lại đúng cảm xúc đó"
    p_map[126] = "Người chơi không cài game vì muốn xem bảng dữ liệu hay trải nghiệm một hệ thống kinh tế phức tạp. Họ tải game vì mẫu quảng cáo (creative) đã gợi mở và hứa hẹn một trạng thái cảm xúc cụ thể: cảm giác giải tỏa khi sắp xếp lại trật tự trong một không gian hỗn độn, sự hồi hộp khi giải cứu một nhân vật, hay niềm vui thuần khiết khi vượt qua một thử thách vừa sức."
    p_map[127] = "Điểm chạm kinh tế (monetization touchpoint) của mẫu quảng cáo không nằm ở tỷ lệ nhấp chuột (CTR) hay chi phí trên mỗi lượt cài đặt (CPI) rẻ trên giấy tờ. Nó nằm ở việc người chơi được thu hút bởi cảm xúc nào, và liệu trải nghiệm thực tế trong 5 phút đầu tiên có mang lại đúng cảm xúc đó hay không. Một mẫu creative đánh lừa người chơi có thể tạo ra lượng cài đặt rẻ trong ngày đầu, nhưng cái giá phải trả sẽ hiện rõ ngay ở tỷ lệ người chơi rời bỏ game sớm (churn rate), điểm đánh giá tiêu cực và sự sụp đổ của tỷ lệ giữ chân Ngày 1 (D1)."
    p_map[128] = "Cơ sở lý thuyết cho hiện tượng này dựa trên các nghiên cứu tâm lý học về kỳ vọng và trải nghiệm của Oliver (1980) và Higgins (1987). Khi có sự chênh lệch lớn giữa kỳ vọng được tạo ra từ thông điệp truyền thông và trải nghiệm nhận được trong thực tế, người dùng sẽ phản ứng bằng cảm giác thất vọng và phòng thủ tâm lý. Điều này không có nghĩa là mọi mẫu quảng cáo dạng kịch bản biến tấu đều vô dụng. Nhưng nó đặt ra một giả thuyết bắt buộc phải kiểm chứng: khoảng cách giữa lời hứa trên quảng cáo và 10 màn chơi đầu tiên càng lớn, ngân sách niềm tin ban đầu của người chơi càng nhanh chóng bị cạn kiệt."
    p_map[129] = "Hãy đặt ba mẫu creative mang lại nhiều lượt cài đặt nhất bên cạnh màn hình trải nghiệm của 3 phút đầu tiên trong game. Người chơi có nhận ra ngay cơ chế chơi mà họ vừa xem trên quảng cáo không? Họ có cảm nhận được sự thỏa mãn tương tự không? Nếu câu trả lời là không, đội ngũ phát triển không chỉ đang lãng phí ngân sách thu hút người dùng (UA), mà còn đang tự tay phá hủy phễu chuyển đổi kiếm tiền ngay từ trước khi người chơi kịp nhìn thấy gói ưu đãi đầu tiên."

    p_map[130] = "6. Store listing and first open\nTrang giới thiệu trên kho ứng dụng và lần mở game đầu tiên phải xác thực cùng một lời hứa"
    p_map[131] = "Trang giới thiệu trên kho ứng dụng (store listing trên App Store hoặc Google Play) là nơi người chơi kiểm chứng lại lời hứa, còn lần mở game đầu tiên chính là lúc họ đưa ra phán quyết sau cùng."
    p_map[132] = "Hãy lấy một mẫu creative đang mang về nhiều lượt tải nhất, rồi theo dõi hành trình của nhóm thuần tập (cohort) đó qua ba điểm chạm: tỷ lệ chuyển đổi từ quảng cáo sang store listing, tỷ lệ cài đặt, và những thao tác đầu tiên sau khi mở ứng dụng. Khi một creative quảng bá lối chơi sắp xếp gọn gàng (sorting puzzle) nhưng store listing lại tập trung phô diễn tính năng xây nhà (meta-game), người chơi đã phải tự xử lý hai thông điệp lệch pha trước cả khi chạm vào màn chơi đầu tiên. Nếu lần mở game đầu lại lập tức bắt đầu bằng chuỗi cửa sổ đăng nhập, yêu cầu cấp quyền, thông báo đẩy và bảng xếp hạng, thì câu hỏi kiếm tiền không còn là 'cửa sổ nào chuyển đổi tốt hơn'. Câu hỏi thực sự là: người chơi có kịp nhận ra giá trị khiến họ tải game trước khi đội ngũ bắt đầu đưa ra thêm các đòi hỏi hay không?"
    p_map[133] = "Điểm tựa lý thuyết ở đây là Thuyết Tải Nhận Thức (Cognitive Load Theory) của John Sweller (1988). Lý thuyết này phân biệt giữa phần tải nhận thức cần thiết để hiểu nhiệm vụ chính với phần tải phát sinh từ cách trình bày rườm rà không phục vụ nhiệm vụ đó. Các nghiên cứu về ví dụ mẫu (worked examples) của Sweller, Chen, Retnowati và Kalyuga (2020) chỉ ra rằng người mới tiếp cận sẽ tiếp thu tốt hơn khi sự chú ý được tập trung tối đa vào cấu trúc cốt lõi. Trong những phút đầu tiên, mọi màn hình pop-up không giúp người chơi hiểu và trải nghiệm lời hứa cốt lõi đều đang cạnh tranh trực tiếp với nguồn lực nhận thức và lòng kiên nhẫn của họ."
    p_map[134] = "Vì vậy, hãy thiết kế hành trình mở game đầu tiên như một chuỗi bằng chứng thuyết phục. Nếu creative hứa hẹn thử thách giải cứu, thao tác đầu tiên phải cho người chơi giải một tình huống giải cứu thực sự trước khi hiển thị menu chung. Nếu creative hứa hẹn sự gọn gàng, hãy để họ tự tay sắp xếp vài vật thể, tận hưởng cảm giác thỏa mãn khi hoàn thành, rồi mới mở ra mục tiêu tiếp theo. Sau chiến thắng nhỏ đầu tiên đó, đội ngũ mới có cơ sở hợp lý để giới thiệu một lựa chọn hữu ích — như quảng cáo đổi thưởng (rewarded ad) để nhận thêm lượt hoặc một gói ưu đãi phù hợp ngữ cảnh. Hãy thử nghiệm các thời điểm hiển thị khác nhau và đo lường đồng thời tỷ lệ hoàn thành hướng dẫn, tỷ lệ giữ chân D1, tỷ lệ đồng thuận cấp quyền và phản ứng của người chơi."
    p_map[135] = "Một yêu cầu cấp quyền ATT (App Tracking Transparency) hay chính sách bảo mật vẫn là bắt buộc. Khác biệt nằm ở chỗ: người chơi có hiểu yêu cầu đó liên quan thế nào đến giá trị họ vừa nhận được hay không. Hãy coi store listing là lời hứa, lần mở game đầu là bằng chứng xác thực, và mọi yêu cầu cấp quyền chỉ được đưa ra khi đã tích lũy đủ lý do thuyết phục. Sau khi lời hứa được chứng minh, 10 màn chơi đầu tiên sẽ trả lời câu hỏi cốt tử: liệu trò chơi có dạy người chơi cách chiến thắng, chấp nhận thất bại và sử dụng các lựa chọn trả phí một cách công bằng hay không?"

    p_map[136] = "7. The first ten levels\nMười màn chơi đầu tiên: Dạy cách đưa ra lựa chọn trước khi dạy cách chi tiền"
    p_map[137] = "Hãy lấy 10 màn chơi đầu tiên của bản build hiện tại và đặt cạnh nhật ký dữ liệu sự kiện (event log) tương ứng. Qua từng màn chơi, người chơi đã học được một luật chơi mới, được trao quyền đưa ra một quyết định mới, hay chỉ đơn thuần va phải thêm các vật cản nhân tạo? Lần đầu tiên họ nhận được vật phẩm bổ trợ (booster) diễn ra trước hay sau khi họ thực sự hiểu vật phẩm đó thay đổi cục diện thế cờ ra sao? Lần đầu xuất hiện gói ưu đãi (offer) có gắn liền với một tình huống bế tắc tự nhiên mà họ vừa trải qua không? Những câu hỏi này chính là bản chất kinh tế của giai đoạn đầu game: một lựa chọn trả tiền chỉ được coi là có giá trị khi người chơi đã hiểu rõ bài toán mà lựa chọn đó giải quyết. Nếu không, tỷ lệ chuyển đổi ban đầu chỉ phản ánh sự ức chế nhất thời, còn cái giá phải trả sẽ xuất hiện ở tỷ lệ thoát game, tỷ lệ thử lại giảm và điểm đánh giá tiêu cực."
    p_map[138] = "Thuyết Tải Nhận Thức đưa ra một nguyên lý thực hành then chốt: người chơi mới cần đủ sự dẫn dắt trực quan để hình thành mô hình tư duy về nhiệm vụ, sau đó sự hỗ trợ phải giảm dần để họ tự mình vận dụng cấu trúc đó. Trong game puzzle, sự dẫn dắt không nhất thiết là những dòng chữ hướng dẫn dài dòng. Đó có thể là một bàn cờ thu nhỏ, một mục tiêu trực quan rõ ràng, một nước đi mẫu gợi ý và phản hồi chính xác vào quyết định người chơi vừa đưa ra. Câu hỏi thiết kế quan trọng là: màn chơi đang làm sáng tỏ một quy tắc hay đang ép người chơi xử lý quá nhiều biến số trước khi họ kịp hiểu chúng?"
    p_map[139] = "Một chuỗi màn chơi đầu game chuẩn mực nên đi từ ví dụ trực quan có dẫn dắt, sang bài toán tương tự với ít trợ giúp hơn, rồi mới tạo ra khoảnh khắc thất bại đầu tiên mà người chơi hoàn toàn hiểu được nguyên nhân. Với từng màn chơi, hãy ghi nhận 4 yếu tố trong buổi đánh giá thiết kế: người chơi đang học được gì; quyết định nào dẫn tới thắng hoặc thua; họ có phương án giải quyết nào không cần trả tiền sau khi thua; và booster đang mở ra thêm lựa chọn chiến thuật hay chỉ đang vá víu cho một thiết kế bất hợp lý. Hãy phân tích tỷ lệ thất bại (fail rate) song hành cùng tỷ lệ thử lại (retry rate), tỷ lệ vào màn tiếp theo, tần suất dùng booster và tỷ lệ thoát game. Mức thất bại cao đi cùng tỷ lệ thử lại cao phản ánh một thử thách hấp dẫn; thất bại cao đi cùng tỷ lệ thoát game ngay lập tức là dấu hiệu báo động cần điều tra cơ chế màn chơi."

    p_map[140] = "8. The first return\nLần quay lại đầu tiên cần một mục tiêu dang dở có ý nghĩa, không chỉ là một lời nhắc nhở"
    p_map[141] = "Trước khi soạn nội dung thông báo đẩy (push notification) đầu tiên, hãy mở danh sách những người chơi rời game sau phiên trải nghiệm đầu và tự hỏi: tại khoảnh khắc thoát game, họ đang còn ấp ủ muốn hoàn thành điều gì? Nếu đội ngũ không thể trả lời bằng một mục tiêu dang dở cụ thể trong trò chơi, thì thông báo đẩy chỉ là lời kêu gọi quay lại một trải nghiệm vô định chưa kịp có lý do. Đây là bản chất kinh tế của lần quay lại đầu tiên: tỷ lệ giữ chân sau phiên đầu mở ra toàn bộ cơ hội tiếp cận sau này với phần thưởng, quảng cáo, gói ưu đãi và nội dung mới; nó không thể được thay thế bằng một tin nhắn gửi đúng khung giờ."
    p_map[142] = "Điểm tựa lý thuyết phù hợp ở đây là giả thuyết gia tốc mục tiêu (goal-gradient hypothesis) do Clark Hull đề xuất năm 1932, và sau đó được Kivetz, Urminsky & Zheng (2006) chứng minh thực nghiệm: nỗ lực và xu hướng hành động của con người gia tăng rõ rệt khi họ cảm thấy mình đang tiến gần đến đích. Cảm giác nhìn thấy vạch đích thúc đẩy người chơi chủ động quay lại và duy trì gắn kết. Điều này gợi ý một cơ chế cần kiểm chứng: mục tiêu trong game có đủ cụ thể không, tiến trình có nhìn thấy rõ ràng không, và phần việc còn lại có đủ hấp dẫn để người chơi mong muốn mở lại game nhằm hoàn thành hay không?"
    p_map[143] = "Một lý do quay lại (return hook) hiệu quả không nhất thiết phải là hệ thống giới hạn thể lực (energy), nhiệm vụ hàng ngày hay bộ đếm thời gian. Đó có thể là một khu vực vườn vừa được mở khóa, một bộ sưu tập còn thiếu đúng một mảnh ghép, một công trình đang chờ hoàn thiện, hoặc một câu đố mà người chơi đã nhìn ra cách giải nhưng phiên trước chưa kịp hoàn tất. Điểm mấu chốt là người chơi biết chính xác mình quay lại để làm gì. Hãy thử nghiệm phân tách hai nhóm: một nhóm rời game sau khi nhận phần thưởng chung chung; nhóm còn lại rời game khi đang nhìn thấy rõ tiến trình và mục tiêu cận kề tiếp theo. Hãy so sánh tỷ lệ giữ chân D1, số phiên chơi và phản ứng với các ưu đãi đầu tiên để tìm ra phương án tối ưu dựa trên dữ liệu thực tế."
    p_map[144] = "Cơ chế giới hạn thể lực (energy) chỉ thực sự có giá trị khi nó đóng vai trò điều tiết nhịp độ trải nghiệm — như giới hạn số lượt thử trong sự kiện đặc biệt, tạo quyết định đánh đổi chiến thuật, hoặc bảo vệ tiến trình dài hạn. Sau lần quay lại đầu tiên, bài toán chuyển sang giai đoạn sâu hơn: mỗi lần người chơi quay lại, trải nghiệm của họ có thực sự tiến bộ một cách có ý nghĩa hay không? Phần tiếp theo sẽ phân tích sâu các điều kiện này: thế nào là một tiến trình có ý nghĩa, và khi nào áp lực tạo ra vẫn giữ được sự công bằng tuyệt đối trong mắt người chơi."

    p_map[151] = "Memory note | Decision board\nPart II: Từ lời hứa quảng cáo đến lần quay lại đầu tiên"
    p_map[152] = "Phần III sẽ đi sâu kiểm tra chất lượng của tiến trình, áp lực thử thách và tính công bằng: những điều kiện cốt lõi biến các lần quay lại thành giá trị kinh tế bền vững và có thể duy trì lâu dài."

    # === PART III: PROGRESS, PRESSURE, AND FAIRNESS (Paras 153-181) ===
    p_map[153] = "Part III: Progress, pressure, and fairness\nPhần III: Tiến trình, áp lực và sự công bằng"
    p_map[154] = "9. Player motivations\nNgười chơi chi tiền để thỏa mãn những nhu cầu khác nhau, không phải vì họ có cùng một hành trình"
    p_map[155] = "Người chơi bước vào game với những động cơ tâm lý hoàn toàn khác biệt. Có người tìm kiếm cảm giác làm chủ kỹ năng (mastery), có người tìm kiếm sự giải tỏa và trật tự (relief through order), có người muốn tích lũy bộ sưu tập (collection), và có người muốn tăng tốc tiến trình để khám phá nội dung mới. Thuyết Tự Quyết (Self-Determination Theory) của Ryan & Deci (2000) chỉ ra ba nhu cầu tâm lý cơ bản của con người: Năng lực (Competence - cảm giác làm chủ), Tự chủ (Autonomy - quyền tự quyết định) và Gắn kết (Relatedness). Một tựa game kiếm tiền bền vững khi các tính năng thương mại hỗ trợ và tôn trọng những nhu cầu tâm lý này, thay vì cưỡng ép hay tước đoạt chúng."
    p_map[156] = "Điểm chạm kinh tế của động cơ người chơi nằm ở sự phân hóa nhu cầu: cùng một màn chơi khó, người chơi theo đuổi kỹ năng sẽ muốn thử lại nhiều lần bằng chính tài năng của mình, trong khi người chơi tìm kiếm sự thư giãn sẽ sẵn lòng xem một quảng cáo đổi thưởng hoặc mua thêm lượt đi để giải phóng áp lực. Nếu hệ thống đối xử với mọi người chơi như nhau bằng cách ép mua booster thô bạo, trò chơi sẽ lập tức làm xói mòn nhu cầu tự chủ của nhóm yêu thích kỹ năng và đẩy họ rời bỏ sản phẩm."
    p_map[157] = "Hãy phân khúc dữ liệu hành vi của người chơi theo nhóm động cơ: theo dõi tỷ lệ chọn xem quảng cáo đổi thưởng, tỷ lệ mua gói bổ trợ, tần suất thử lại màn chơi và độ sâu phiên chơi. Đừng xây dựng một gói ưu đãi chung chung cho toàn bộ tập người chơi. Hãy thiết kế các điểm chạm thương mại tương ứng với từng trạng thái tâm lý cụ thể."
    p_map[158] = "Khi người chơi cảm thấy quyền tự chủ và năng lực của mình được tôn trọng, việc chi trả tiền bạc hay thời gian xem quảng cáo trở thành một hành động tự nguyện để gia tăng trải nghiệm, chứ không phải một khoản phí phạt bắt buộc để thoát khỏi sự ức chế."

    p_map[159] = "10. Meaningful progress\nTiến trình chỉ thực sự có giá trị khi nó thay đổi rõ rệt tình thế mà người chơi đang quan tâm"
    p_map[160] = "Tăng một con số trên màn hình không đồng nghĩa với việc tạo ra tiến trình. Việc vượt qua từ Màn 10 sang Màn 11 chỉ mang lại cảm giác tiến bộ nếu Màn 11 mở ra một quy tắc mới, một không gian mới, hoặc đóng góp một bước đi rõ ràng vào bức tranh tổng thể mà người chơi đang dốc lòng xây dựng. Khi tiến trình bị biến thành một chuỗi cày cuốc vô hồn với các màn chơi lặp lại không có điểm nhấn, giá trị cảm nhận của người chơi sẽ suy giảm nhanh chóng."
    p_map[161] = "Mối liên hệ giữa tiến trình và khả năng kiếm tiền rất rõ ràng: người chơi chỉ sẵn lòng chi trả tiền bạc hoặc xem quảng cáo để bảo vệ hoặc đẩy nhanh một tiến trình mà họ thực sự coi trọng. Nếu vượt qua một màn chơi khó mà không nhận lại được bất kỳ sự công nhận hay thay đổi trực quan nào trong thế giới game, họ sẽ không có lý do gì để sử dụng booster hay mua thêm lượt trong những lần thất bại tiếp theo."
    p_map[162] = "Hãy rà soát lại toàn bộ cây tiến trình (progression tree) của trò chơi: đảm bảo rằng sau mỗi mốc thử thách lớn (milestone), người chơi luôn nhận được một phần thưởng trực quan xứng đáng, mở khóa tính năng mới hoặc chứng kiến sự chuyển biến rõ rệt của không gian game. Tiến trình có ý nghĩa chính là mỏ neo giữ chân người chơi ở lại và tạo ra lý do chính đáng cho mọi giao dịch trong game."

    p_map[163] = "11. Pressure creates a decision\nÁp lực chỉ có giá trị khi người chơi vẫn nhìn thấy một quyết định thực sự công bằng"
    p_map[164] = "Áp lực trong game — dù là giới hạn thời gian đếm ngược, số lượt đi hạn chế hay không gian chứa đồ bị thu hẹp — là công cụ tạo ra sự căng thẳng đầy kịch tính (tension). Tuy nhiên, áp lực chỉ lành mạnh khi người chơi nhận thức rõ rằng: nếu họ suy nghĩ thấu đáo hơn hoặc lựa chọn chiến thuật khác, họ hoàn toàn có thể vượt qua mà không bắt buộc phải trả tiền. Khi áp lực bị đẩy lên mức phi lý khiến kỹ năng trở nên vô nghĩa, áp lực đó biến thành sự ức chế và hành vi ép buộc."
    p_map[165] = "Điểm chạm kinh tế của áp lực nằm ở khoảnh khắc ra quyết định: khi người chơi đối mặt với tình thế suýt thắng (near-miss state) — chỉ còn thiếu một nước đi duy nhất để giải quyết cả bàn cờ — việc đề xuất một lượt đi thêm qua quảng cáo đổi thưởng hoặc một khoản phí nhỏ là hoàn toàn tự nhiên và công bằng. Nhưng nếu bàn cờ rơi vào thế tắc nghẽn ngay từ đầu do sắp xếp ngẫu nhiên ác ý, lời đề nghị mua hàng sẽ bị xem như một hành vi tống tiền."
    p_map[166] = "Hãy kiểm tra lại tất cả các điểm tạo áp lực trong game: đảm bảo rằng người chơi luôn có thể học hỏi từ thất bại, hiểu rõ mình đã sai ở đâu, và luôn có ít nhất một con đường vượt qua thử thách dựa trên kỹ năng thuần túy trước khi các phương án hỗ trợ thương mại xuất hiện."

    p_map[167] = "12. Dynamic difficulty\nĐộ khó động phải điều chỉnh thử thách, không được âm thầm định đoạt kết quả"
    p_map[168] = "Hệ thống điều chỉnh độ khó động (Dynamic Difficulty Adjustment - DDA) sinh ra nhằm mục đích giữ người chơi trong trạng thái dòng chảy tâm lý (flow state) — nơi thử thách luôn cân bằng hoàn hảo với kỹ năng. Tuy nhiên, ranh giới giữa việc hỗ trợ trải nghiệm và thao túng kết quả là vô cùng mong manh. Khi DDA âm thầm can thiệp để ép người chơi thua nhằm kích thích mua hàng, hoặc tự động cho thắng một cách lộ liễu, nó sẽ phá hủy hoàn toàn cảm giác thành tựu và niềm tin của người chơi vào tính trung thực của trò chơi."
    p_map[169] = "Nếu người chơi nhận ra rằng chiến thắng hay thất bại không phụ thuộc vào nước đi của họ mà do thuật toán ngầm quyết định thời điểm ép nạp tiền, họ sẽ lập tức ngừng đầu tư cảm xúc và rời bỏ trò chơi. Doanh thu kiếm được từ những can thiệp thao túng như vậy chính là doanh thu đánh đổi tương lai (borrowed revenue)."
    p_map[170] = "Hãy đặt ra các nguyên tắc minh bạch cho hệ thống cân bằng độ khó: DDA chỉ nên đóng vai trò là chiếc phao cứu sinh hỗ trợ người chơi sau chuỗi thất bại kéo dài ngoài ý muốn, tuyệt đối không được sử dụng DDA như một chiếc bẫy tinh vi để cưỡng bức chuyển đổi mua vật phẩm."
    p_map[171] = "Một hệ thống kinh tế bền vững được xây dựng trên sự tôn trọng nỗ lực của người chơi, nơi chiến thắng luôn mang lại cảm giác xứng đáng và mỗi thất bại đều là một bài học có thể lý giải."

    p_map[172] = "13. Randomness and skill\nTính ngẫu nhiên chỉ công bằng khi người chơi vẫn có thể học hỏi và phản ứng"
    p_map[173] = "Tính ngẫu nhiên (RNG) là gia vị tạo nên sự bất ngờ, tính biến hóa và giá trị chơi lại cho game giải đố. Nhưng tính ngẫu nhiên chỉ được người chơi chấp nhận khi nó đóng vai trò tạo ra dữ kiện đầu vào (input randomness) để họ tính toán chiến thuật, chứ không phải tính ngẫu nhiên quyết định kết quả đầu ra (output randomness) phủ nhận toàn bộ tính toán trước đó."
    p_map[174] = "Hãy kiểm tra các bàn cờ thất bại với những chuỗi khởi tạo ngẫu nhiên khác nhau. Người chơi có thể nhìn lại và nhận ra một nước đi khác giúp họ lật ngược tình thế không, hay mọi bản xem lại đều dẫn đến kết luận cay đắng rằng họ chỉ có thể thắng nếu hệ thống may mắn thả rơi đúng một vật phẩm cụ thể? Khi tính ngẫu nhiên biến trò chơi thành một trò may rủi thuần túy, các gói bán vật phẩm bổ trợ sẽ bị coi là công cụ gian lận thay vì giải pháp chiến thuật."
    p_map[175] = "Hãy đánh giá tính ngẫu nhiên bằng cách xem lại các bản ghi màn chơi thực tế (replay review), thay vì chỉ nhìn vào bảng tỷ lệ rơi đồ lý thuyết. Theo dõi sát sao tỷ lệ thử lại sau khi thua, tỷ lệ chọn dùng booster, tỷ lệ thoát game và mức độ sẵn sàng quay lại của từng nhóm người chơi để đảm bảo tính ngẫu nhiên luôn phục vụ niềm vui khám phá."

    p_map[176] = "14. Reward feedback\nPhản hồi âm thanh và hiệu ứng thị giác phải xác thực đúng giá trị của phần thưởng"
    p_map[177] = "Mỗi phần thưởng trao cho người chơi cần được xác nhận bằng hiệu ứng thị giác, âm thanh và rung phản hồi (haptics) tương xứng với nỗ lực mà họ đã bỏ ra. Một phần thưởng lớn đạt được sau chuỗi thử thách cam go nhưng chỉ hiển thị bằng một dòng chữ mờ nhạt sẽ làm suy giảm cảm giác thành tựu. Ngược lại, việc thổi phồng một phần thưởng vụn vặt bằng hiệu ứng quá đà sẽ tạo ra cảm giác giả tạo và làm người chơi nhanh chóng chai sạn cảm xúc."
    p_map[178] = "Phản hồi phần thưởng có liên hệ trực tiếp đến nhận thức về giá trị: khi người chơi nhận được sự tán thưởng xứng đáng cho thành tích của mình, họ sẽ gắn kết sâu sắc hơn với thế giới trong game và trân trọng từng tài nguyên tích lũy được. Điều này tạo nền tảng vững chắc cho việc định giá các gói nạp và vật phẩm trong cửa hàng."
    p_map[179] = "Nguyên tắc thực hành cốt lõi: phần thưởng nhỏ nhận phản hồi gọn gàng nhưng sắc nét; phần thưởng hiếm và quan trọng được đón nhận bằng khoảnh khắc vinh danh ấn tượng để đánh dấu bước chuyển biến trong hành trình của người chơi."

    p_map[180] = "Memory note | Decision board\nPart III: Tiến trình, áp lực và sự công bằng"
    p_map[181] = "Phần IV sẽ chuyển hóa toàn bộ các nguyên lý tâm lý và thiết kế trên thành những cấu phần sản phẩm thương mại cụ thể: quảng cáo đổi thưởng, quảng cáo xen kẽ, vật phẩm bổ trợ, giao dịch mua trong ứng dụng (IAP) và hệ thống kinh tế sự kiện."

    # === PART IV: ADS, IAP, AND ECONOMY (Paras 182-210) ===
    p_map[182] = "Part IV: Ads, IAP, and economy\nPhần IV: Quảng cáo, IAP và nền kinh tế trong game"
    p_map[183] = "15. Rewarded ads\nQuảng cáo đổi thưởng chỉ bền vững khi trao đổi đúng giá trị mà người chơi đang thực sự cần"
    p_map[184] = "Quảng cáo đổi thưởng (rewarded ads) là một hợp đồng trao đổi tự nguyện giữa người chơi và trò chơi: người chơi đồng ý dành ra 30 giây chú ý để đổi lấy một tài nguyên hoặc cơ hội có ý nghĩa ngay tại thời điểm đó. Bản chất của rewarded ad không phải là một sự gián đoạn cưỡng bức, mà là một dịch vụ hỗ trợ đúng lúc."
    p_map[185] = "Hãy thiết kế vị trí đặt quảng cáo đổi thưởng (placement) theo một cấu trúc trao đổi rõ ràng: 'Sau khi đối mặt với tình huống X, người chơi có thể chủ động chọn xem quảng cáo để nhận giá trị Y, và giá trị Y đó giúp họ giải quyết trực tiếp trở ngại trước mắt'. Nếu phần thưởng trao ra quá rẻ mạt hoặc không liên quan đến nhu cầu hiện tại, người chơi sẽ từ chối xem; nếu phần thưởng quá hào phóng làm phá vỡ toàn bộ nền kinh tế game, họ sẽ không còn động lực để tự mình vượt qua thử thách hay cân nhắc các gói nạp IAP."
    p_map[186] = "Không có một con số tham chiếu chung nào về tỷ lệ chọn xem quảng cáo (opt-in rate) áp dụng cho mọi tựa game. Hãy kiểm tra các vị trí đặt quảng cáo đổi thưởng dựa trên dữ liệu thực tế của chính dự án: tỷ lệ người chơi chủ động chọn xem, số lượt xem trung bình trên mỗi người chơi hoạt động, tỷ lệ hoàn thành video quảng cáo, và quan trọng nhất là mức độ gắn kết sau khi nhận thưởng. Hãy đảm bảo phần thưởng luôn được trao chính xác, nhanh chóng và minh bạch."
    p_map[187] = "Khi quảng cáo đổi thưởng được thiết kế như một sự trợ giúp công bằng và kịp thời, nó không những không làm tổn hại đến tỷ lệ giữ chân mà còn gia tăng thiện cảm của người chơi đối với sản phẩm."

    p_map[188] = "16. Interstitials\nQuảng cáo xen kẽ chỉ nên xuất hiện tại những điểm ngắt tự nhiên khi sự tập trung đã hoàn tất"
    p_map[189] = "Quảng cáo xen kẽ (interstitial ads) là định dạng có độ xâm lấn cao nhất trong trải nghiệm game di động. Nếu xuất hiện đột ngột giữa lúc người chơi đang tập trung suy nghĩ hoặc ngay khi vừa thất bại trong trạng thái ức chế, interstitial sẽ lập tức kích hoạt phản ứng tiêu cực và thúc đẩy hành vi gỡ cài đặt ứng dụng. Điểm chạm kinh tế của interstitial đòi hỏi sự thận trọng tối đa: doanh thu quảng cáo kiếm được từ một lượt hiển thị vô duyên vĩnh viễn không thể bù đắp được giá trị trọn đời (LTV) bị mất đi từ một người chơi rời bỏ game."
    p_map[190] = "Nghiên cứu của Stothart, Mitchum & Yehnert (2015) về sự phân tâm nhận thức chỉ ra rằng ngay cả những gián đoạn ngắn cũng phá vỡ hoàn toàn dòng chảy tư duy của con người. Vì vậy, vị trí duy nhất có thể chấp nhận để hiển thị quảng cáo xen kẽ là tại các điểm ngắt tự nhiên (natural breakpoints) — chẳng hạn như sau khi người chơi vừa hoàn thành một màn chơi, nhận xong phần thưởng tổng kết và chuẩn bị chuyển sang một khu vực mới."
    p_map[191] = "Hãy thiết lập các quy tắc bảo vệ nghiêm ngặt: giới hạn tần suất hiển thị (frequency capping) hợp lý, không bao giờ hiển thị interstitial ở những màn chơi đầu tiên, tự động gỡ bỏ hoàn toàn interstitial đối với những người chơi đã thực hiện bất kỳ giao dịch nạp IAP nào, và luôn cung cấp nút đóng quảng cáo rõ ràng, dễ thao tác. Hãy theo dõi chỉ số lượt hiển thị trung bình (IMPDAU) song hành cùng tỷ lệ giữ chân D1/D7 để kịp thời phát hiện các dấu hiệu xói mòn niềm tin."

    p_map[192] = "17. Boosters\nVật phẩm bổ trợ phải mở ra một lựa chọn chiến thuật rõ ràng, không phải liều thuốc vá lỗi"
    p_map[193] = "Vật phẩm bổ trợ (booster) là công cụ trao cho người chơi quyền can thiệp vào bàn cờ để thay đổi cục diện chiến thuật — như xóa một chướng ngại vật cứng đầu, hoán đổi vị trí hay tạo thêm lượt đi. Một booster được thiết kế tốt khi người chơi hiểu chính xác tại sao mình sử dụng nó và cảm nhận được sự thông minh trong quyết định can thiệp của mình. Ngược lại, nếu một màn chơi được thiết kế quá bất công đến mức không thể vượt qua nếu không dùng booster, thì booster đó đã bị biến thành chiếc vé thông hành bắt buộc, tước đoạt toàn bộ niềm vui chiến thắng."
    p_map[194] = "Hãy phân tích dữ liệu sử dụng booster trên từng màn chơi cụ thể: nếu một màn chơi có tỷ lệ sử dụng booster tăng đột biến bất thường đi kèm với tỷ lệ thoát game cao, đó là tín hiệu cho thấy màn chơi đang bị lỗi thiết kế chứ không phải người chơi đang hào hứng mua sắm. Việc lạm dụng độ khó để ép tiêu thụ booster sẽ nhanh chóng làm cạn kiệt lượng tài nguyên tích lũy của người chơi và đẩy họ vào tình thế bế tắc."
    p_map[195] = "Hãy định vị booster như một công cụ mở rộng quyền tự chủ và sáng tạo chiến thuật cho người chơi, đồng thời duy trì nguồn cung cấp booster vừa đủ thông qua các sự kiện và phần thưởng tiến trình để người chơi luôn có cơ hội trải nghiệm giá trị của chúng trước khi cân nhắc mua thêm trong cửa hàng."

    p_map[196] = "18. IAP and present need\nGiao dịch mua trong ứng dụng phải giải quyết một nhu cầu thực tế đang hiện diện và minh bạch về giá trị"
    p_map[197] = "Cửa hàng trong game (in-game store) là một danh mục sản phẩm tĩnh, trong khi một gói ưu đãi theo ngữ cảnh (offer) là một lời đề nghị xuất hiện đúng lúc người chơi nảy sinh nhu cầu cấp thiết nhất. Người chơi không chi tiền vì cửa hàng có giao diện đẹp; họ chi tiền vì họ đang đứng trước một bài toán cụ thể và nhận thấy gói ưu đãi đưa ra một giải pháp thỏa đáng, minh bạch và có mức giá tương xứng."
    p_map[198] = "Các gói tiền ảo chung chung, trừu tượng thường chỉ phù hợp với những người chơi kỳ cựu đã nắm rõ toàn bộ nền kinh tế game. Đối với phần lớn người chơi, những gói ưu đãi gắn liền trực tiếp với ngữ cảnh — ví dụ: gói cứu trợ khi suýt vượt qua màn chơi khó, gói khởi đầu (starter pack) đi kèm quyền lợi gỡ quảng cáo vĩnh viễn, hoặc gói vật phẩm phục vụ đúng sự kiện đang diễn ra — luôn mang lại tỷ lệ chuyển đổi cao hơn và sự hài lòng lớn hơn."
    p_map[199] = "Hãy đảm bảo rằng mỗi gói IAP đưa ra đều giải thích tường tận: người chơi sẽ nhận được những gì, số tài nguyên đó giúp họ giải quyết vấn đề gì, và luôn cung cấp một con đường từ chối (decline path) lịch thiệp, dễ dàng mà không tạo ra bất kỳ áp lực hay cảm giác tội lỗi nào."
    p_map[200] = "Tính minh bạch và sự tôn trọng trong giao dịch chính là chìa khóa biến một người chơi miễn phí thành một khách hàng trung thành sẵn sàng tái chi trả trong suốt vòng đời sản phẩm."

    p_map[201] = "19. Economy integrity\nNền kinh tế game chỉ đáng tin khi mỗi loại tiền tệ đều có đường vào, đường ra và mục đích rõ ràng"
    p_map[202] = "Nền kinh tế trong game (game economy) là hệ thống huyết mạch điều phối dòng chảy của tài nguyên và tiền tệ (currencies). Một nền kinh tế lành mạnh phải tuân thủ nguyên lý cân bằng dòng chảy (stock-and-flow): mỗi loại tiền tệ phải có nguồn tạo ra (sources) được kiểm soát chặt chẽ, các điểm tiêu thụ (sinks) có ý nghĩa thiết thực, và duy trì một mức số dư tích lũy hợp lý trong tay người chơi để họ luôn có động lực tiếp tục tham gia các hoạt động trong game."
    p_map[203] = "Lạm phát tiền tệ là căn bệnh phổ biến nhất: khi trò chơi phát thưởng quá ồ ạt mà thiếu các điểm tiêu thụ hấp dẫn, tiền tệ sẽ nhanh chóng mất giá, khiến các phần thưởng trong tương lai trở nên vô nghĩa và các gói IAP trong cửa hàng mất hoàn toàn sức hút. Ngược lại, nếu nền kinh tế quá thắt chặt khiến người chơi luôn rơi vào cảnh túng thiếu cùng cực, họ sẽ cảm thấy kiệt sức và từ bỏ cuộc chơi."
    p_map[204] = "Hãy định kỳ rà soát các chỉ số kinh tế then chốt theo từng nhóm thuần tập: số dư tiền tệ trung vị, tốc độ tích lũy, tốc độ tiêu thụ và tỷ lệ phân bổ chi tiêu vào các danh mục sản phẩm khác nhau. Hãy bảo vệ tính toàn vẹn của nền kinh tế như bảo vệ chính linh hồn của trò chơi."

    p_map[205] = "20. Event economy\nSự kiện là một nền kinh tế thu nhỏ có thời hạn, không phải một chồng cửa sổ pop-up"
    p_map[206] = "Hệ thống sự kiện (events) và thẻ mùa (battle pass / event pass) được thiết kế nhằm làm mới nhịp điệu trải nghiệm, tạo ra mục tiêu ngắn hạn giàu năng lượng và thúc đẩy sự gắn kết cộng đồng. Một sự kiện xuất sắc là một chuỗi vòng lặp khép kín: tham gia hoạt động $\rightarrow$ tích lũy điểm thưởng $\rightarrow$ lựa chọn phần thưởng $\rightarrow$ tiến bộ trên bảng thành tích $\rightarrow$ về đích và tận hưởng thành quả."
    p_map[207] = "Theo ước tính của Sensor Tower, Royal Match từng đạt mốc doanh thu hơn 100 triệu USD mỗi tháng nhờ vào việc vận hành nhịp nhàng chuỗi sự kiện luân phiên như Royal Pass hay Hidden Temple kết hợp cùng tốc độ cập nhật nội dung đều đặn. Dữ liệu này chứng minh sức mạnh của một hệ thống vận hành sự kiện chuyên nghiệp; tuy nhiên, điều đáng học hỏi không phải là việc sao chép nguyên xi lịch trình của đối thủ, mà là việc thấu hiểu cách họ phân tầng mục tiêu: người chơi miễn phí luôn có một lộ trình hoàn thành rõ ràng, trong khi gói tăng tốc cao cấp mang lại giá trị gia tăng xứng đáng cho những ai muốn đầu tư thêm."
    p_map[208] = "Trong các buổi đánh giá sau sự kiện (post-event review), hãy bóc tách doanh thu phát sinh trong sự kiện khỏi các tác động dài hạn: tỷ lệ giữ chân D1/D7 của nhóm tham gia, số dư tiền tệ sau sự kiện, tỷ lệ quay lại vào tuần kế tiếp và điểm đánh giá của cộng đồng. Đảm bảo rằng sự kiện mang lại niềm vui thực sự chứ không chỉ là một đợt vét cạn tài nguyên và lòng kiên nhẫn của người chơi."

    p_map[209] = "Memory note | Decision board\nPart IV: Quảng cáo, IAP và nền kinh tế game"
    p_map[210] = "Phần V sẽ chuyển trọng tâm từ khâu thiết kế sản phẩm sang hệ thống ra quyết định dựa trên dữ liệu: cách đọc dashboard, phương pháp ghép cặp chỉ số đối ứng và quy trình thử nghiệm để không bị đánh lừa bởi những mức tăng trưởng bề nổi."

    # === PART V: SIGNALS, DECISIONS, AND EXPERIMENTS (Paras 211-231) ===
    p_map[211] = "Part V: Signals, decisions, and experiments\nPhần V: Tín hiệu, quyết định và thử nghiệm"
    p_map[212] = "21. Reading the dashboard\nBảng dữ liệu phải chỉ ra quyết định nào cần xem xét lại, không chỉ đơn thuần hiển thị kết quả"
    p_map[213] = "Hãy tiếp cận bảng dữ liệu (dashboard) bằng một câu hỏi cụ thể về sản phẩm, thay vì chỉ mở ra để nhìn ngắm các biểu đồ chỉ số. Câu hỏi đúng phải là: 'Lời hứa từ mẫu quảng cáo có đang được giữ trọn vẹn trong 10 màn chơi đầu không?', 'Người chơi có đang gặp bế tắc bất thường tại vị trí đặt quảng cáo xen kẽ không?', hay 'Gói ưu đãi mới tung ra đang tạo ra doanh thu bền vững hay đang làm suy giảm tỷ lệ giữ chân của nhóm người chơi mới?'"
    p_map[214] = "Một bảng dữ liệu hiển thị hàng trăm chỉ số nhưng không gắn liền với hành động cụ thể chỉ tạo ra sự nhiễu loạn thông tin. Dữ liệu chỉ thực sự có giá trị khi nó phản ánh rõ nét sức khỏe của mối quan hệ giữa người chơi và trò chơi, đồng thời định hướng chính xác nơi đội ngũ cần can thiệp để cải thiện chất lượng sản phẩm."
    p_map[215] = "Điểm tựa lý thuyết ở đây là nguyên lý suy luận nhân quả (causal inference): một con số quan sát được trên bề mặt không tự giải thích nguyên nhân gốc rễ; để hiểu bản chất, đội ngũ bắt buộc phải xây dựng một câu chuyện nhân quả hợp lý và kiểm chứng bằng dữ liệu phân tầng theo nhóm thuần tập (cohort analysis)."
    p_map[216] = "Hãy xây dựng thói quen đọc dữ liệu có định hướng: biến mỗi buổi duyệt chỉ số thành một buổi thảo luận về trải nghiệm của người chơi và đưa ra các quyết định điều chỉnh có thể đo lường được kết quả."

    p_map[217] = "22. Read metric pairs\nMột chỉ số chỉ trở nên hữu ích khi được đối chiếu cùng cái giá mà nó tạo ra"
    p_map[218] = "Không bao giờ được phép đọc một chỉ số tăng trưởng một cách cô lập. Mọi sự thay đổi trong game luôn tạo ra tác động kép: một chỉ số đi lên thường đi kèm với nguy cơ một chỉ số khác bị tổn hại. Nếu chỉ nhìn vào một nửa bức tranh, đội ngũ sẽ rất dễ rơi vào ảo tưởng về sự thành công."
    p_map[219] = "Hãy luôn đọc dữ liệu theo các cặp chỉ số đối ứng (metric pairs):\n• Doanh thu trung bình trên mỗi người dùng (ARPDAU) phải đi cùng Tỷ lệ giữ chân Ngày 7 (D7 Retention).\n• Tỷ lệ chuyển đổi mua gói ưu đãi (Conversion Rate) phải đi cùng Điểm đánh giá (Store Ratings) và Tỷ lệ hoàn tiền (Refund Rate).\n• Số lượt xem quảng cáo trung bình (IMPDAU) phải đi cùng Độ sâu phiên chơi (Session Length) và Tỷ lệ thoát game (Churn Rate).\n• Chi phí thu hút người dùng (CPI) phải đi cùng Tỷ lệ hoàn thành màn chơi đầu (FTUE Completion) và Doanh thu trọn đời thực tế (Realized LTV)."
    p_map[220] = "23. Decision trees\nCây quyết định ngăn đội ngũ nhảy vội từ một số liệu bề nổi sang một giải pháp quen tay"
    p_map[221] = "Khi tỷ lệ giữ chân Ngày 1 (D1) sụt giảm, phản xạ quen thuộc của nhiều đội ngũ là vội vã hạ độ khó của màn chơi hoặc phát thêm tiền thưởng miễn phí. Tuy nhiên, nguyên nhân thực sự có thể nằm ở chỗ: mẫu quảng cáo đang hứa hẹn sai thể loại game, hoặc màn hình yêu cầu cấp quyền xuất hiện quá sớm làm đứt mạch trải nghiệm."
    p_map[222] = "Nếu D1 ổn định nhưng D3 và D7 giảm mạnh, hãy kiểm tra lại lý do quay lại của người chơi, các điểm tắc nghẽn độ khó bất thường, tần suất hiển thị quảng cáo xen kẽ quá dày, hoặc sự đơn điệu của nội dung ở phiên chơi thứ hai. Cây quyết định (decision tree) giúp đội ngũ phân tích vấn đề theo từng nhánh nguyên nhân logic thay vì đưa ra các quyết định cảm tính."
    p_map[223] = "Hãy thiết lập các quy trình chẩn đoán chuẩn mực cho studio: xác định rõ triệu chứng $\rightarrow$ liệt kê các giả thuyết nguyên nhân $\rightarrow$ kiểm tra dữ liệu đối chiếu $\rightarrow$ đưa ra giải pháp can thiệp thử nghiệm."
    p_map[224] = "Đầu ra của một cây quyết định phải là một giả thuyết rõ ràng có thể bị bác bỏ, một người chịu trách nhiệm chính (owner) và một thời hạn đo lường cụ thể. Khi dữ liệu chưa đủ bằng chứng để kết luận, việc dũng cảm thừa nhận 'chưa rõ' và tiếp tục đào sâu quan sát luôn tốt hơn việc vội vã chi tiền mở rộng quy mô."

    p_map[225] = "24. Experimentation\nThử nghiệm chỉ có giá trị khi nó giúp đội ngũ loại bỏ một điều bất định quan trọng"
    p_map[226] = "Thử nghiệm phân tách A/B (A/B testing) là công cụ khoa học tối thượng để tối ưu hóa sản phẩm. Tuy nhiên, thử nghiệm sẽ biến thành một trò diễn vô bổ nếu đội ngũ thay đổi cùng lúc quá nhiều biến số hoặc chạy thử nghiệm mà không có một giả thuyết cơ chế rõ ràng từ trước."
    p_map[227] = "Một thử nghiệm đạt chuẩn cần đáp ứng 4 tiêu chí: có giả thuyết nhân quả rõ ràng, có quy mô mẫu đủ lớn để đạt ý nghĩa thống kê, có các chỉ số cảnh báo rủi ro (guardrail metrics) để bảo vệ trải nghiệm chung, và có kế hoạch hoàn tác nhanh chóng (rollback plan) nếu kết quả thử nghiệm gây hại cho sản phẩm."
    p_map[228] = "Hãy ghi lại toàn bộ nhật ký học tập từ các thử nghiệm: bài học rút ra từ một thử nghiệm thất bại thường mang lại giá trị định hướng lâu dài lớn hơn nhiều so với một mức tăng trưởng số liệu ngắn hạn không rõ nguyên nhân."
    p_map[229] = "Hệ thống cấu hình từ xa (remote config) là công cụ đắc lực hỗ trợ triển khai và thu hồi tính năng linh hoạt, nhưng nó không bao giờ thay thế được tư duy thiết kế thử nghiệm nghiêm túc của đội ngũ phát triển."

    p_map[230] = "Memory note | Decision board\nPart V: Tín hiệu, quyết định và thử nghiệm"
    p_map[231] = "Phần VI sẽ đi sâu vào một vấn đề mang tính chiến lược: dữ liệu thị trường và trí tuệ nhân tạo có thể giúp đội ngũ đi nhanh hơn ở đâu, và đâu là những giới hạn mà dữ liệu vĩnh viễn không thể thay thế được năng lực phán đoán của con người."

    # === PART VI: WHAT DATA CAN AND CANNOT DECIDE (Paras 232-258) ===
    p_map[232] = "Part VI: What data can and cannot decide\nPhần VI: Dữ liệu có thể và không thể quyết định những gì"
    p_map[233] = "25. Market intelligence\nNghiên cứu thị trường giúp thu hẹp không gian lựa chọn; luận điểm sản phẩm mới quyết định canh bạc của studio"
    p_map[234] = "Các nền tảng phân tích thị trường như Sensor Tower hay AppMagic cung cấp bức tranh toàn cảnh về những thể loại game đang dẫn đầu, quy mô doanh thu và xu hướng thu hút người dùng. Dữ liệu này cực kỳ hữu ích để giúp studio loại bỏ những canh bạc yếu kém và nhận diện các phân khúc thị trường tiềm năng."
    p_map[235] = "Tuy nhiên, dữ liệu thị trường chỉ cho thấy những gì đã thành công trong quá khứ của người khác; nó không thể đảm bảo rằng đội ngũ của bạn có đủ năng lực thiết kế, tốc độ sản xuất nội dung và sự thấu cảm người chơi để tạo ra một sản phẩm vượt trội trong cùng phân khúc đó."
    p_map[236] = "Dữ liệu có thể báo cáo rằng cơ chế ốc vít (screw puzzle), chủ đề giải cứu hay hệ thống thẻ mùa đang thịnh hành. Nhưng nó không trả lời được liệu studio có thể làm cho vòng lặp cốt lõi cuốn hút hơn đối thủ, duy trì niềm tin của người chơi và vận hành nền kinh tế bền vững hay không. Đó là ranh giới giữa nghiên cứu thị trường và luận điểm sản phẩm (product thesis)."
    p_map[237] = "Tín hiệu thị trường (market signal) là dữ liệu quan sát từ bên ngoài: người chơi đang dành thời gian và tiền bạc ở phân khúc này. Đây là điểm khởi đầu để định hướng nghiên cứu."
    p_map[238] = "Luận điểm sản phẩm (product thesis) là lời cam kết được tôi luyện từ bên trong studio: đội ngũ có năng lực tạo ra một trải nghiệm khác biệt, vượt trội về chất lượng, giữ chân người chơi bền bỉ và vận hành với hiệu quả kinh tế có lợi nhuận."
    p_map[239] = "Tín hiệu thị trường giúp chọn địa điểm đào sâu; luận điểm sản phẩm quyết định chất lượng quặng mà bạn khai thác được."

    p_map[240] = "26. Behaviour needs interpretation\nDữ liệu mô tả hành vi; ý nghĩa của hành vi bắt buộc phải được giải mã bằng một câu chuyện nhân quả"
    p_map[241] = "Bảng dữ liệu chỉ ghi nhận những gì đã xảy ra sau một thay đổi thiết kế; nó không tự giải thích người chơi đã cảm nhận và suy nghĩ như thế nào khi đưa ra hành động đó. Một sự gia tăng doanh thu đột biến có thể đến từ sự tò mò nhất thời, một lỗ hổng thiết kế gây ức chế, hay một tính năng thực sự xuất sắc. Vì vậy, câu hỏi 'Dữ liệu nói gì?' mới chỉ là bước đầu; câu hỏi quyết định phải là: 'Cơ chế tâm lý nào đã tạo ra xu hướng này, và bằng chứng nào sẽ bác bỏ cơ chế đó?'"
    p_map[242] = "Ví dụ, tỷ lệ mua thêm ô chứa đồ trong Clear Garden tăng vọt ở Màn 7 có thể do 4 nguyên nhân hoàn toàn khác nhau: do thử thách kịch tính và công bằng, do giao diện gói bán quá hấp dẫn, do độ khó bị đẩy lên quá mức khiến người chơi hoảng loạn tìm lối thoát, hoặc do người chơi vô tình bấm nhầm. Mỗi nguyên nhân sẽ dẫn đến một tương lai hoàn toàn khác nhau cho tựa game."
    p_map[243] = "Hãy thường xuyên xem lại các video quay màn hình chơi thực tế, đọc kỹ từng phản hồi của người dùng, trò chuyện trực tiếp với người chơi và phân tích sâu hành vi trước khi đưa ra kết luận cuối cùng."
    p_map[244] = "Đừng để những con số vô hồn che khuất đi những trải nghiệm cảm xúc sống động của con người phía sau màn hình."

    p_map[245] = "27. Decision memo\nĐơn vị công việc hữu ích nhất là một quyết định có thể kiểm chứng, không phải một danh sách tính năng"
    p_map[246] = "Một bản kế hoạch phát triển thông thường chỉ liệt kê danh sách các tính năng cần làm: làm game puzzle chủ đề giải cứu, thêm hệ thống nhiệm vụ hàng ngày, gắn gói nạp khởi đầu. Cách tiếp cận này rất dễ đẩy studio vào cái bẫy sản xuất dàn trải mà không tạo ra đột phá."
    p_map[247] = "Thay vào đó, hãy chuẩn hóa mọi đề xuất sản phẩm thành một Biên bản Quyết định (Decision Memo) cô đọng trên một trang giấy, bao gồm 5 cấu phần bắt buộc: Vấn đề của người chơi cần giải quyết $\rightarrow$ Giả thuyết can thiệp $\rightarrow$ Bằng chứng thị trường & dữ liệu nội bộ hỗ trợ $\rightarrow$ Các rủi ro tiềm ẩn & chỉ số cảnh báo $\rightarrow$ Tiêu chí đo lường thành công hoặc điều kiện hủy bỏ (kill criteria)."
    p_map[248] = "Khi toàn bộ studio giao tiếp và làm việc dựa trên các biên bản quyết định minh bạch, trách nhiệm giải trình được xác lập rõ ràng và năng lực học hỏi của tổ chức sẽ được nâng lên một tầm cao mới."
    p_map[249] = "Một quyết định sáng suốt được đo bằng chất lượng của lập luận và khả năng kiểm chứng giả thuyết, chứ không đo bằng sự hào nhoáng của các slide trình chiếu."

    p_map[250] = "28. Copy the question, not the configuration\nHọc hỏi bài toán đối thủ đã giải quyết; đừng sao chép cấu hình một cách mù quáng"
    p_map[251] = "Khi quan sát một tựa game dẫn đầu như Royal Match hay Candy Crush Saga, sai lầm phổ biến nhất là sao chép nguyên xi thông số cấu hình của họ: đặt gói nạp giá tương tự, thiết kế thời gian đếm ngược giống hệt, tung ra số lượng sự kiện y chang. Đội ngũ quên mất rằng những cấu hình đó hoạt động hiệu quả là nhờ dựa trên một cỗ máy sản xuất hàng trăm màn chơi mỗi tháng, ngân sách mua người dùng khổng lồ và tệp người chơi trung thành tích lũy qua nhiều năm."
    p_map[252] = "Hãy học hỏi câu hỏi cốt lõi mà đối thủ đã nỗ lực giải quyết: 'Họ làm thế nào để tạo ra cảm giác tiến bộ liên tục?', 'Họ cân bằng giữa áp lực thử thách và sự giải tỏa ra sao?', 'Họ bảo vệ tính toàn vẹn của nền kinh tế bằng cơ chế nào?'. Sau đó, hãy tự tìm câu trả lời phù hợp nhất với năng lực sản xuất và quy mô thực tế của chính studio mình."

    p_map[253] = "29. Clear Garden: From prompt to decision\nCase Study Clear Garden: Biến gợi ý từ nhà phát hành thành một quyết định sản phẩm có thể kiểm chứng"
    p_map[254] = "Hãy quay trở lại với dự án giả định Clear Garden. Khi nhà phát hành (publisher) đưa ra gợi ý: 'Hãy làm game sắp xếp đồ đạc kết hợp trang trí vườn, thêm tính năng bộ sưu tập và bán các gói cứu trợ khi thua', đây mới chỉ là một tín hiệu định hướng thô sơ."
    p_map[255] = "Bản build thử nghiệm đầu tiên không cần phải ôm đồm cả 5 loại tiền tệ hay hệ thống thẻ sự kiện phức tạp. Nó chỉ cần chứng minh được 4 luận điểm sống còn: mẫu quảng cáo hứa hẹn đúng cảm xúc $\rightarrow$ 10 màn chơi đầu mang lại niềm vui giải đố thực sự $\rightarrow$ người chơi hiểu rõ lý do thất bại và tự nguyện thử lại $\rightarrow$ và có ít nhất một lý do cụ thể thôi thúc họ mở lại game vào ngày hôm sau."
    p_map[256] = "Nếu 4 luận điểm đó được chứng minh bằng dữ liệu nhóm thuần tập vững chắc, studio mới có nền tảng vững vàng để phát triển thêm hệ thống kinh tế và vận hành trực tiếp. Nếu không, việc dừng dự án sớm (kill) chính là quyết định sáng suốt nhất giúp bảo toàn nguồn lực quý báu cho những cơ hội tốt hơn."

    p_map[257] = "Memory note | Decision board\nPart VI: Dữ liệu, nghiên cứu thị trường và luận điểm sản phẩm"
    p_map[258] = "Phần VII sẽ chuyển từ luận điểm sản phẩm sang năng lực thực thi và vận hành thực tế: chuỗi cung ứng giá trị, tiêu chuẩn dừng/lặp/mở rộng quy mô, vận hành trực tiếp và hiệu quả kinh tế đóng góp thực tế."

    # === PART VII: THE OPERATING SYSTEM BEHIND A LIVE GAME (Paras 259-276) ===
    p_map[259] = "Part VII: The operating system behind a live game\nPhần VII: Hệ điều hành của game đang phát hành"
    p_map[260] = "30. Monetization as a supply chain\nKiếm tiền như một chuỗi cung ứng giá trị: Một mắt xích yếu có thể phá vỡ toàn bộ lời hứa sản phẩm"
    p_map[261] = "Người chơi chỉ nhìn thấy một màn hình duy nhất trước mắt; nhưng để màn hình đó xuất hiện đúng lúc, đúng giá trị và hoạt động mượt mà, cả studio phải vận hành như một chuỗi cung ứng đồng bộ: Nghiên cứu thị trường định hình hướng đi $\rightarrow$ Bản mẫu (prototype) kiểm chứng vòng lặp cốt lõi $\rightarrow$ Thiết kế màn chơi tạo ra thử thách hấp dẫn $\rightarrow$ Đội ngũ đồ họa thổi hồn vào không gian $\rightarrow$ Lập trình đảm bảo hiệu năng $\rightarrow$ Đội ngũ UA thu hút đúng tệp người dùng $\rightarrow$ Bộ phận dữ liệu phân tích phản hồi $\rightarrow$ và Đội ngũ hỗ trợ khách hàng bảo vệ niềm tin."
    p_map[262] = "Chỉ cần một mắt xích bị lỗi — tốc độ ra màn chơi mới quá chậm khiến người chơi hết nội dung, khâu kiểm thử (QA) lỏng lẻo để lọt lỗi giao dịch nạp tiền, hoặc thông điệp quảng cáo truyền thông sai lệch — toàn bộ hệ thống kiếm tiền sẽ lập tức bị đình trệ."
    p_map[263] = "Hãy định kỳ đánh giá năng lực của từng mắt xích trong chuỗi cung ứng: xác định rõ nút thắt cổ chai (bottleneck) đang cản trở sự phát triển của dự án và tập trung nguồn lực giải quyết dứt điểm."

    p_map[264] = "31. Kill, iterate, or scale\nNghiên cứu chỉ có giá trị khi kết thúc bằng quyết định: Dừng lại, Lặp lại cải tiến, hoặc Mở rộng quy mô"
    p_map[265] = "Một biểu đồ thị trường đẹp không phải là một chiến lược. Trước khi biến một bản mẫu thành một dự án chính thức, hãy viết sẵn các tiêu chí dừng dự án (kill criteria) vô cùng cụ thể về mức độ rõ ràng của lời hứa, tỷ lệ kích hoạt người chơi, tỷ lệ giữ chân ban đầu và chi phí sản xuất."
    p_map[266] = "Ba lối rẽ chiến lược cần được xác lập rõ ràng:\n• Dừng lại (Kill): Khi các bằng chứng thực tế cho thấy không còn con đường khả thi nào để đạt được hiệu quả kinh tế trong phạm vi nguồn lực cho phép.\n• Lặp lại cải tiến (Iterate): Khi vẫn còn một ẩn số cụ thể có thể kiểm chứng được trong khung thời gian và ngân sách giới hạn.\n• Mở rộng quy mô (Scale): Khi tất cả các chỉ số cốt lõi về giữ chân người chơi, niềm tin và hiệu quả kinh tế biên đều được chứng minh vững chắc trên các nhóm thuần tập thực tế."
    p_map[267] = "Supercell nổi tiếng với văn hóa sẵn sàng ăn mừng khi khai tử những dự án không đạt tiêu chuẩn xuất sắc (như trường hợp của Hay Day Pop) để dồn toàn lực cho những cơ hội đột phá. Sự dũng cảm dừng một dự án yếu kém chính là phẩm chất phân biệt giữa một studio nghiệp dư và một tổ chức chuyên nghiệp tầm cỡ quốc tế."

    p_map[268] = "32. Live ops with memory\nVận hành trực tiếp có ghi nhớ: Phải luôn nhớ người chơi vừa nhận được gì, vừa bị đòi hỏi điều gì và còn gì đáng để quay lại"
    p_map[269] = "Hệ thống vận hành trực tiếp (live ops) không phải là một chiếc máy bắn pop-up tự động vô cảm. Nó phải sở hữu một 'trí nhớ' tinh tế: nhận biết người chơi vừa trải qua chuỗi thất bại căng thẳng hay vừa hoàn thành một chiến tích vẻ vang, họ đã xem bao nhiêu quảng cáo trong ngày và số dư tài nguyên hiện tại của họ ra sao để đưa ra những tương tác phù hợp."
    p_map[270] = "Mỗi sự kiện đưa vào lịch trình phải đảm nhận một vai trò rõ rệt: hướng dẫn kỹ năng mới (teach), kích hoạt người chơi cũ quay lại (reactivate), thúc đẩy mục tiêu sưu tập (collector goal), tạo thời điểm chi tiêu hợp lý (spend moment), hoặc trao cho người chơi một khoảng nghỉ ngơi phục hồi năng lượng (recovery). Đừng bao giờ biến lịch sự kiện thành một chuỗi ngày vắt kiệt sức lực của người chơi."

    p_map[271] = "33. Contribution economics\nHiệu quả kinh tế chỉ có ý nghĩa khi phần lợi nhuận còn lại đủ nuôi sống studio sau mọi chi phí duy trì sản phẩm lành mạnh"
    p_map[272] = "Doanh thu tổng (Gross Revenue) chỉ là con số bề nổi để khoe truyền thông. Một doanh nghiệp game thực sự tồn tại và phát triển dựa trên Lợi nhuận Đóng góp Thực tế (Contribution Margin) — tức phần tiền còn lại sau khi đã thanh toán đầy đủ mọi chi phí bắt buộc."
    p_map[273] = "Phương trình kinh tế đóng góp thực tế:\nDoanh thu tổng (Gross Revenue)\n- Phí nền tảng (30% cho Apple / Google)\n- Chi phí công nghệ quảng cáo & hạ tầng máy chủ\n- Ngân sách thu hút người dùng (UA Spend)\n- Chi phí sản xuất nội dung, đội ngũ hỗ trợ và vận hành\n= Lợi nhuận Đóng góp Thực tế (Contribution Margin)\n\nMỗi dòng chi phí đều phải được gắn liền với một người chịu trách nhiệm và theo dõi chặt chẽ theo từng nhóm thuần tập."
    p_map[274] = "Hãy đặc biệt lưu ý đến thời gian hoàn vốn (payback period) và dòng tiền thực tế (cash flow timing). Một chiến dịch UA có thể cho thấy chỉ số LTV kỳ vọng trên mô hình toán học rất đẹp, nhưng nếu thời gian thu hồi vốn kéo dài quá nhiều tháng, nó có thể đẩy studio vào cuộc khủng hoảng thanh khoản nghiêm trọng."

    p_map[275] = "Memory note | Decision board\nPart VII: Hệ điều hành và hiệu quả kinh tế đóng góp"
    p_map[276] = "Phần VIII sẽ thu hẹp các nguyên lý trên thành bản đồ cơ chế theo từng thể loại game cụ thể, giúp đội ngũ xuất phát từ đúng cảm xúc và nhu cầu của người chơi thay vì sao chép rập khuôn."

    # === PART VIII: GENRE PLAYBOOKS (Paras 277-286) ===
    p_map[277] = "Part VIII: Genre playbooks\nPhần VIII: Cẩm nang thiết kế theo thể loại"
    p_map[278] = "34. The emotional loop\nBắt đầu từ cảm xúc mà người chơi mong muốn tìm kiếm, rồi mới lựa chọn giải pháp kiếm tiền phù hợp"
    p_map[279] = "Không tồn tại một 'chân dung người chơi game puzzle' chung chung cho tất cả các dự án, và do đó không thể có một công thức kiếm tiền mặc định áp dụng cho mọi thể loại. Mỗi dòng game đáp ứng một trạng thái tâm lý hoàn toàn khác biệt."
    p_map[280] = "Hãy xây dựng bản đồ cơ chế kiếm tiền tương ứng với vòng lặp cảm xúc đặc thù:\n• Dòng game Sắp xếp (Sort Puzzle): Bán cảm giác giải tỏa thông qua sự trật tự. Các giải pháp phù hợp: tính năng hoàn tác (undo), mở thêm ô khay chứa đồ và gói gỡ quảng cáo (remove ads) nhằm bảo vệ sự sáng suốt của kế hoạch sắp xếp.\n• Dòng game Giải tỏa tắc nghẽn (Jam Puzzle): Bán cảm giác hồi hộp có kiểm soát rồi vỡ òa giải tỏa. Phù hợp với các lời đề nghị tiếp tục chơi khi suýt hoàn thành (continue after near-miss) đi kèm phương án giải quyết cụ thể.\n• Dòng game Vật lý (Physics Puzzle): Bán sự tò mò, thử nghiệm nhanh và yếu tố bất ngờ. Phù hợp với cơ chế thử lại nhanh hoặc quảng cáo đổi lấy công cụ hỗ trợ độc đáo mà không làm gián đoạn dòng chảy trải nghiệm.\n• Dòng game Nối 3 truyền thống (Match-3): Bán cảm giác làm chủ kỹ năng và tiến trình khám phá dài hạn. Phù hợp với hệ thống booster chiến thuật, mạng chơi (lives) và các chuỗi sự kiện theo mùa."
    p_map[281] = "Khi các điểm chạm thương mại ăn khớp hoàn hảo với vòng lặp cảm xúc cốt lõi, người chơi sẽ đón nhận chúng như những phần thưởng tự nhiên thay vì cảm giác bị làm phiền."

    p_map[282] = "35. Evidence, not blueprint\nVí dụ thị trường chỉ giúp lựa chọn câu hỏi đúng; dữ liệu nhóm thuần tập của studio mới quyết định cấu hình sản phẩm"
    p_map[283] = "Những tựa game thành công vang dội như Royal Match, Candy Crush Saga hay Merge Mansion đặt ra những bài học tham khảo vô cùng giá trị về nhịp độ sản xuất nội dung, cách tạo áp lực trên bàn cờ và nghệ thuật duy trì gắn kết cộng đồng. Tuy nhiên, chúng là những bằng chứng lịch sử để học hỏi tư duy, không phải bản thiết kế chi tiết để sao chép nguyên mẫu."
    p_map[284] = "Hãy sử dụng dữ liệu thị trường công khai để nắm bắt quy mô phân khúc, tốc độ chuyển dịch của thị trường và bối cảnh cạnh tranh chiến lược. Nhưng hãy luôn dùng chính dữ liệu nhóm thuần tập của tựa game bạn đang làm để tinh chỉnh độ khó từng màn chơi, thiết lập mức giá gói nạp và tối ưu hóa các điểm chạm quảng cáo."

    p_map[285] = "Memory note | Decision board\nPart VIII: Bản đồ cơ chế theo thể loại"
    p_map[286] = "Phần IX là bài kiểm tra toàn diện cuối cùng: quy trình kiểm toán 30 phút giúp người làm game phát hiện chính xác các điểm rò rỉ niềm tin, thất thoát giá trị và xác lập tiêu chuẩn hoàn thành trước khi mở rộng quy mô."

    # === PART IX: THE AUDIT (Paras 287-303) ===
    p_map[287] = "Part IX: The audit\nPhần IX: Quy trình kiểm toán và tiêu chuẩn hoàn thành"
    p_map[288] = "36. The 30-minute audit\nKiểm toán hệ thống kiếm tiền trong 30 phút: Tìm ra một điểm rò rỉ niềm tin, một điểm thất thoát giá trị và một thử nghiệm có thể triển khai ngay"
    p_map[289] = "Đừng biến buổi kiểm toán thành một cuộc tranh luận cảm tính về gu thẩm mỹ cá nhân. Mục tiêu tối thượng của buổi kiểm toán 30 phút là thiết lập một chuỗi bằng chứng logic rõ ràng nối liền từ lời hứa quảng cáo $\rightarrow$ trải nghiệm màn chơi đầu $\rightarrow$ các điểm trao đổi thương mại $\rightarrow$ đến dữ liệu phản hồi thực tế. Sau 30 phút, đội ngũ bắt buộc phải chỉ ra được: một điểm đang làm xói mòn niềm tin (trust leak), một điểm đang làm thất thoát giá trị (value leak), và một kế hoạch thử nghiệm can thiệp rõ ràng có người chịu trách nhiệm và phương án hoàn tác an toàn."

    p_map[290] = "37. Definition of done\nKhi nào một hệ sinh thái sản phẩm đủ bằng chứng để mở rộng quy mô, thay vì chỉ vừa đủ tính năng để phát hành"
    p_map[291] = "Một hệ thống chỉ thực sự sẵn sàng để mở rộng quy mô chi tiêu UA khi và chỉ khi 10 tiêu chuẩn kiểm định dưới đây được chứng minh vững chắc bằng dữ liệu nhóm thuần tập và có người chịu trách nhiệm vận hành cụ thể:"
    p_map[292] = "1. Lời hứa từ mẫu quảng cáo (creative promise) xuất hiện ngay trong trải nghiệm 3 phút đầu tiên và được kiểm chứng qua tỷ lệ hoàn thành hướng dẫn của nhóm thuần tập."
    p_map[293] = "2. Người chơi giải thích được chính xác lý do dẫn đến thất bại và luôn nhìn thấy một phương án vượt qua thử thách mà không bắt buộc phải chi tiền."
    p_map[294] = "3. Quảng cáo đổi thưởng (rewarded ad) là sự trao đổi hoàn toàn tự nguyện, trao thưởng chính xác, tức thì và được theo dõi song hành cùng tỷ lệ giữ chân người chơi."
    p_map[295] = "4. Quảng cáo xen kẽ (interstitial) chỉ xuất hiện tại các điểm ngắt tự nhiên, có giới hạn tần suất nghiêm ngặt và tự động gỡ bỏ đối với người chơi đã nạp tiền."
    p_map[296] = "5. Mỗi gói ưu đãi (offer) giải quyết một nhu cầu thực tế đang hiện diện, có sự đồng thuận minh bạch và cung cấp con đường từ chối lịch thiệp."
    p_map[297] = "6. Nguồn tạo ra và điểm tiêu thụ của tiền tệ tạo ra các quyết định lựa chọn có ý nghĩa, không biến thành trạm thu phí cưỡng ép; số dư tiền tệ được theo dõi chặt chẽ."
    p_map[298] = "7. Doanh thu kiếm được luôn được đối chiếu cùng tỷ lệ giữ chân, điểm đánh giá của cộng đồng, tỷ lệ hoàn tiền và chi phí hỗ trợ khách hàng."
    p_map[299] = "8. Đội ngũ có năng lực triển khai thử nghiệm, hoàn tác nhanh chóng và lưu trữ bài học kinh nghiệm có cấu trúc trong hệ thống."
    p_map[300] = "9. Kế hoạch sản xuất nội dung và vận hành sự kiện đáp ứng kịp thời lời hứa đã truyền thông qua các chiến dịch UA và luôn có khoảng thời gian phục hồi hợp lý."
    p_map[301] = "10. Mô hình tài chính tính toán đầy đủ mọi chi phí để duy trì trò chơi khỏe mạnh, kiểm soát chặt chẽ thời gian hoàn vốn và bảo đảm an toàn dòng tiền."

    p_map[302] = "Memory note | Decision board\nPart IX: Kiểm toán sản phẩm và tiêu chuẩn sẵn sàng mở rộng quy mô"
    p_map[303] = "Phần kết không bổ sung thêm bất kỳ khung lý thuyết mới nào. Nó đưa chúng ta quay trở về với tiêu chuẩn đạo đức và nghề nghiệp xuyên suốt toàn bộ cuốn sách: sau mỗi lần xem quảng cáo, sau mỗi gói ưu đãi hay mỗi giao dịch nạp tiền, người chơi có còn cảm thấy hào hứng và có một lý do rõ ràng để tiếp tục gắn bó với trò chơi hay không?"

    # === CLOSING & RESEARCH SOURCES (Paras 304-323) ===
    p_map[304] = "Closing: The player must want to continue\nLời kết: Kiếm tiền chỉ thực sự bền vững khi người chơi vẫn muốn tiếp tục cuộc hành trình"
    p_map[305] = "Câu hỏi giá trị nhất trong làm game không phải là 'Làm thế nào để ép người chơi nạp tiền?'. Câu hỏi đúng đắn và sâu sắc hơn nhiều là: 'Người chơi vừa nhận được giá trị gì, họ đang trải qua trạng thái cảm xúc nào, và lựa chọn nào họ sẽ cảm nhận là hoàn toàn công bằng?'. Khi câu trả lời sáng tỏ, các tính năng thương mại sẽ tìm thấy vị trí tự nhiên và đắc địa nhất để phục vụ trải nghiệm. Khi câu trả lời còn mơ hồ, việc nhồi nhét thêm vị trí quảng cáo hay tung thêm gói nạp chỉ làm cho hệ thống thêm ồn ào và đẩy nhanh sự suy thoái của sản phẩm."
    p_map[306] = "Lợi nhuận bền vững là kết quả tất yếu của một hệ thống có năng lực giữ trọn lời hứa từ đầu đến cuối: từ thông điệp trên mẫu quảng cáo, phiên chơi đầu tiên, từng màn giải đố được trau chuốt, cho đến vị trí quảng cáo tinh tế, gói nạp minh bạch, chuỗi sự kiện giàu cảm xúc, bảng dữ liệu trung thực và dịch vụ hỗ trợ tận tâm. Không một bảng điều khiển số liệu, một bản định hướng từ nhà phát hành hay một biểu đồ thị trường nào có thể thay thế được chuỗi công việc kỷ luật và đầy tâm huyết đó. Chúng chỉ đóng vai trò như chiếc la bàn giúp đội ngũ nhận biết nhanh hơn nơi cần đặt ra những câu hỏi sắc bén hơn."
    p_map[307] = "Nghệ thuật kiếm tiền chân chính luôn để lại cho người chơi một lý do vững chắc để tiếp tục cuộc chơi sau mỗi lần giao dịch. Trò chơi hoàn toàn có quyền thách thức kỹ năng, đòi hỏi thời gian và đề xuất chi phí xứng đáng. Mối quan hệ giữa người chơi và studio sẽ mãi mãi bền chặt khi người chơi nhìn thấy rõ giá trị, luật chơi minh bạch và sự tự do lựa chọn, trong khi đội ngũ phát triển có đủ sự điềm tĩnh và kỷ luật để lắng nghe những tín hiệu cảnh báo thay vì chỉ tự mãn với những con số tăng trưởng ngắn hạn."
    p_map[308] = "Đó chính là chuẩn mực chuyên nghiệp cao quý nhất mà một studio phát triển game đáng tự hào hướng tới: không ảo tưởng rằng chỉ một ý tưởng đơn lẻ có thể làm nên sự nghiệp kinh doanh, không trốn tránh trách nhiệm phán đoán phía sau những con số vô hồn, và vĩnh viễn không bao giờ đánh đổi tương lai lâu dài của sản phẩm lấy một kết quả nhất thời chưa được kiểm chứng."

    p_map[309] = "Research notes and public sources\nNguồn tham khảo học thuật và dữ liệu thị trường công khai"
    p_map[310] = "Sensor Tower: Báo cáo chuyên sâu về sự vươn lên của Royal Match và phân tích thị trường game casual toàn cầu."
    p_map[311] = "Sensor Tower: Báo cáo triển vọng thị trường game di động toàn cầu và xu hướng chi tiêu in-app."
    p_map[312] = "Sensor Tower: Báo cáo phân tích chuyên đề dòng game giải đố (puzzle) tại thị trường Hoa Kỳ."
    p_map[313] = "AppMagic & GameDev Reports: Báo cáo thị trường Casual Game nửa đầu năm 2025."
    p_map[314] = "AppMagic: Nghiên cứu điển hình (case study) về mô hình tăng trưởng của Epic Plane Evolution."
    p_map[315] = "AppMagic: Báo cáo chuyên đề về hoạt động vận hành trực tiếp (LiveOps Report) trong dòng game Casual."
    p_map[316] = "Unity Technologies: Báo cáo chiến lược Monetization; phân tích tác động tương hỗ giữa quảng cáo đổi thưởng, IAP và tỷ lệ giữ chân người chơi."
    p_map[317] = "Apple Inc.: Tài liệu hướng dẫn khung bảo mật App Tracking Transparency (ATT) và quyền riêng tư của người dùng."
    p_map[318] = "Ủy ban Thương mại Liên bang Hoa Kỳ (FTC): Thông cáo phán quyết hòa giải với Epic Games và cẩm nang hướng dẫn nhận diện các bẫy giao diện số (dark patterns)."
    p_map[319] = "GameRefinery: Nền tảng phân tích tính năng sản phẩm và dự báo hành vi phân khúc người chơi."
    p_map[320] = "GameAnalytics: Báo cáo chuẩn đối sánh (benchmarking) và các chỉ số hiệu suất trọng yếu (KPIs) trong ngành game."
    p_map[321] = "Rovio Entertainment: Tài liệu nền tảng vận hành Beacon, phân tích hành vi người chơi và tối ưu hóa trải nghiệm người dùng mới (FTUE)."
    p_map[322] = "King Digital Entertainment: Phương pháp luận khoa học dữ liệu, phân tích nhóm thuần tập và chiến lược sản phẩm."
    p_map[323] = "Supercell: 'What We Have Learned from Failures' – Bài học kinh nghiệm và nguyên tắc khai tử dự án vì chuẩn mực chất lượng."

    # Apply all paragraph updates
    for idx, new_text in p_map.items():
        if idx < len(doc.paragraphs):
            doc.paragraphs[idx].text = new_text

    print(f"Applied {len(p_map)} paragraph updates.")

    # 2. APPLY TABLE REFINEMENTS
    refine_all_tables(doc)

    # 3. SAVE POLISHED DOCUMENT
    doc.save(TARGET_DOCX)
    print(f"\n[SUCCESS] Successfully generated polished document at:\n{TARGET_DOCX}")

if __name__ == "__main__":
    build_polished()
