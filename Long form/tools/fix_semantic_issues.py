# -*- coding: utf-8 -*-
"""
Script: fix_semantic_issues.py
Purpose: Fix all identified semantic, formatting, and translation issues in
The-Art-of-Monetization-Vietnamese-Polished.docx:
1. Fix LaTeX escape artifacts ($ \r\n ightarrow $) -> clean Unicode arrows (➔).
2. Fully translate and harmonize Table 11 (Operations diagnostic table) into professional Vietnamese.
3. Fully translate and harmonize Table 16 (Clear Garden Decision Memo) into professional Vietnamese.
4. Clean up remaining raw terms in P[270] and P[273].
"""

import docx
from pathlib import Path

DOCX_PATH = Path(r"D:\CODE\VGC\Long form\The-Art-of-Monetization-Vietnamese-Polished.docx")

def fix_all_issues():
    doc = docx.Document(DOCX_PATH)
    print(f"Loaded {DOCX_PATH.name} for semantic fixes...")

    # 1. FIX PARAGRAPH ARTIFACTS
    p_fixes = {
        206: "Hệ thống sự kiện (events) và thẻ mùa (battle pass / event pass) được thiết kế nhằm làm mới nhịp điệu trải nghiệm, tạo ra mục tiêu ngắn hạn giàu năng lượng và thúc đẩy sự gắn kết cộng đồng. Một sự kiện xuất sắc là một chuỗi vòng lặp khép kín: tham gia hoạt động ➔ tích lũy điểm thưởng ➔ lựa chọn phần thưởng ➔ tiến bộ trên bảng thành tích ➔ về đích và tận hưởng thành quả.",
        223: "Hãy thiết lập các quy trình chẩn đoán chuẩn mực cho studio: xác định rõ triệu chứng ➔ liệt kê các giả thuyết nguyên nhân ➔ kiểm tra dữ liệu đối chiếu ➔ đưa ra giải pháp can thiệp thử nghiệm.",
        247: "Thay vào đó, hãy chuẩn hóa mọi đề xuất sản phẩm thành một Biên bản Quyết định (Decision Memo) cô đọng trên một trang giấy, bao gồm 5 cấu phần bắt buộc: Vấn đề của người chơi cần giải quyết ➔ Giả thuyết can thiệp ➔ Bằng chứng thị trường & dữ liệu nội bộ hỗ trợ ➔ Các rủi ro tiềm ẩn & chỉ số cảnh báo ➔ Tiêu chí đo lường thành công hoặc điều kiện hủy bỏ (kill criteria).",
        255: "Bản build thử nghiệm đầu tiên không cần phải ôm đồm cả 5 loại tiền tệ hay hệ thống thẻ sự kiện phức tạp. Nó chỉ cần chứng minh được 4 luận điểm sống còn: mẫu quảng cáo hứa hẹn đúng cảm xúc ➔ 10 màn chơi đầu mang lại niềm vui giải đố thực sự ➔ người chơi hiểu rõ lý do thất bại và tự nguyện thử lại ➔ và có ít nhất một lý do cụ thể thôi thúc họ mở lại game vào ngày hôm sau.",
        261: "Người chơi chỉ nhìn thấy một màn hình duy nhất trước mắt; nhưng để màn hình đó xuất hiện đúng lúc, đúng giá trị và hoạt động mượt mà, cả studio phải vận hành như một chuỗi cung ứng đồng bộ: Nghiên cứu thị trường định hình hướng đi ➔ Bản mẫu (prototype) kiểm chứng vòng lặp cốt lõi ➔ Thiết kế màn chơi tạo ra thử thách hấp dẫn ➔ Đội ngũ đồ họa thổi hồn vào không gian ➔ Lập trình đảm bảo hiệu năng ➔ Đội ngũ UA thu hút đúng tệp người dùng ➔ Bộ phận dữ liệu phân tích phản hồi ➔ và Đội ngũ hỗ trợ khách hàng bảo vệ niềm tin.",
        270: "Mỗi sự kiện đưa vào lịch trình phải đảm nhận một vai trò rõ rệt: hướng dẫn kỹ năng mới (teach), kích hoạt người chơi cũ quay lại (reactivate), thúc đẩy mục tiêu sưu tập (collector goal), tạo thời điểm chi tiêu hợp lý (monetization moment), hoặc trao cho người chơi một khoảng nghỉ ngơi phục hồi năng lượng (recovery). Đừng bao giờ biến lịch sự kiện thành một chuỗi ngày vắt kiệt sức lực của người chơi.",
        273: "Phương trình kinh tế đóng góp thực tế:\nDoanh thu tổng (Gross Revenue)\n- Phí nền tảng (30% cho Apple / Google)\n- Chi phí công nghệ quảng cáo & hạ tầng máy chủ\n- Ngân sách thu hút người dùng (UA Spend)\n- Chi phí sản xuất nội dung, đội ngũ hỗ trợ và vận hành\n= Lợi nhuận Đóng góp Thực tế (Contribution Margin)\n\nMỗi dòng chi phí đều phải được gắn liền với một người chịu trách nhiệm và theo dõi chặt chẽ theo từng nhóm thuần tập.",
        289: "Đừng biến buổi kiểm toán thành một cuộc tranh luận cảm tính về gu thẩm mỹ cá nhân. Mục tiêu tối thượng của buổi kiểm toán 30 phút là thiết lập một chuỗi bằng chứng logic rõ ràng nối liền từ lời hứa quảng cáo ➔ trải nghiệm màn chơi đầu ➔ các điểm trao đổi thương mại ➔ đến dữ liệu phản hồi thực tế. Sau 30 phút, đội ngũ bắt buộc phải chỉ ra được: một điểm đang làm xói mòn niềm tin (trust leak), một điểm đang làm thất thoát giá trị (value leak), và một kế hoạch thử nghiệm can thiệp rõ ràng có người chịu trách nhiệm và phương án hoàn tác an toàn."
    }

    for idx, txt in p_fixes.items():
        if idx < len(doc.paragraphs):
            doc.paragraphs[idx].text = txt

    print(f"Fixed {len(p_fixes)} paragraphs with escape artifacts.")

    # 2. FIX TABLE 4 ARTIFACTS
    if len(doc.tables) > 4:
        tbl4 = doc.tables[4]
        tbl4.rows[0].cells[0].text = "LÀM NGAY:\n• Lập bản đồ lời hứa: Creative ➔ Store listing ➔ Màn hình 3 phút đầu; đánh dấu từng điểm lệch pha về cảm xúc.\n• Kiểm tra 10 màn chơi đầu: ghi nhận kỹ năng người chơi học được, nguyên nhân thắng/thua và lựa chọn không trả tiền sau khi thất bại.\n• Thử nghiệm 2 cách kết thúc phiên đầu: rời đi sau phần thưởng chung chung so với rời đi khi đang nhìn thấy rõ mục tiêu kế tiếp."

    # 3. FIX TABLE 24 ARTIFACTS
    if len(doc.tables) > 24:
        tbl24 = doc.tables[24]
        tbl24.rows[5].cells[2].text = "Vòng lặp sự kiện: Tích lũy ➔ Lựa chọn ➔ Tiêu thụ ➔ Tiến bộ ➔ Hoàn thành và Phục hồi."

    # 4. FIX TABLE 11 (Operations & Diagnostic Table)
    if len(doc.tables) > 11:
        tbl11 = doc.tables[11]
        
        # Row 0 (Header 1)
        for c in tbl11.rows[0].cells:
            c.text = "Câu hỏi chẩn đoán"
        tbl11.rows[0].cells[0].text = "Câu hỏi chẩn đoán"
        tbl11.rows[0].cells[2].text = "Chỉ số cần đọc song hành"
        tbl11.rows[0].cells[4].text = "Hành động kiểm chứng tiếp theo"

        # Row 1
        tbl11.rows[1].cells[0].text = "Lời hứa có được giữ trọn?"
        tbl11.rows[1].cells[2].text = "CTR, Tỷ lệ chuyển đổi Store, Tỷ lệ xong hướng dẫn, D1 theo mẫu quảng cáo"
        tbl11.rows[1].cells[4].text = "So sánh video quảng cáo, trang Store và bản ghi 3 phút đầu theo nhóm thuần tập."

        # Row 2
        tbl11.rows[2].cells[0].text = "Màn thua đầu có công bằng?"
        tbl11.rows[2].cells[2].text = "Tỷ lệ thua, Tỷ lệ thử lại, Tỷ lệ thoát game, Tần suất dùng booster, Ý kiến đánh giá"
        tbl11.rows[2].cells[4].text = "Xem lại bản ghi màn chơi; phân biệt rõ giữa thử thách hấp dẫn và sự ức chế do lỗi thiết kế."

        # Row 3
        tbl11.rows[3].cells[0].text = "Quảng cáo có mang lại giá trị?"
        tbl11.rows[3].cells[2].text = "Tỷ lệ chọn xem, Lượt xem trung bình, Tỷ lệ xem hết video, Tỷ lệ thoát sau xem, Tỷ lệ quay lại"
        tbl11.rows[3].cells[4].text = "Thử nghiệm A/B vị trí hoặc thời điểm hiển thị; giữ nguyên mức phần thưởng để đo lường."

        # Row 4
        tbl11.rows[4].cells[0].text = "Gói IAP có giải quyết đúng nhu cầu?"
        tbl11.rows[4].cells[2].text = "Lượt xem gói, Tỷ lệ mua, Hành vi của nhóm từ chối mua, Tỷ lệ hoàn tiền, Giữ chân người nạp"
        tbl11.rows[4].cells[4].text = "Rà soát tính minh bạch của gói bán, giá trị thực tế và trải nghiệm khi người chơi từ chối."

        # Row 5 (Header 2)
        for c in tbl11.rows[5].cells:
            c.text = "Lĩnh vực chẩn đoán"
        tbl11.rows[5].cells[0].text = "Lĩnh vực chẩn đoán"
        tbl11.rows[5].cells[2].text = "Chỉ số trọng yếu"
        tbl11.rows[5].cells[4].text = "Câu hỏi quyết định chiến lược"

        # Rows 6-12
        r6_12_data = [
            ("Thu hút (Acquisition)", "CTR, IPM, CPI, Tỷ lệ chuyển đổi trang Store", "Mẫu quảng cáo có tiếp cận đúng tệp người chơi tiềm năng không?"),
            ("Kích hoạt (Activation)", "Tốc độ tải, Tỷ lệ crash, Hoàn thành FTUE, Tiến độ Màn 1-10", "Trò chơi có chứng minh được lời hứa cốt lõi đủ nhanh không?"),
            ("Gắn kết (Engagement)", "Số phiên chơi, Độ sâu phiên, Tỷ lệ thử lại, D1/D3/D7", "Vòng lặp chơi có tạo ra lý do xứng đáng để người chơi quay lại không?"),
            ("Độ khó (Difficulty)", "Tỷ lệ thua, Tỷ lệ thoát game, Tần suất tiêu thụ booster", "Áp lực tạo ra là động lực học hỏi hay sự ức chế bế tắc?"),
            ("Quảng cáo (Ads)", "Tỷ lệ người xem, Lượt hiển thị (IMPDAU), Tỷ lệ rời bỏ sau xem", "Hình thức trao đổi quảng cáo có thực sự mang lại sự trợ giúp hữu ích?"),
            ("Giao dịch (IAP)", "Lượt xem gói, Tỷ lệ mua, Tỷ lệ tái nạp, Tỷ lệ yêu cầu hoàn tiền", "Sản phẩm có giữ được tính minh bạch và toàn vẹn kinh tế không?"),
            ("Niềm tin (Trust)", "Điểm xếp hạng sao, Đánh giá trên Store, Khiếu nại hỗ trợ", "Mối quan hệ lâu dài giữa người chơi và studio có đang bị suy yếu không?")
        ]

        for i, (col0, col1, col2) in enumerate(r6_12_data, start=6):
            tbl11.rows[i].cells[0].text = col0
            tbl11.rows[i].cells[2].text = col1
            tbl11.rows[i].cells[4].text = col2

    # 5. FIX TABLE 16 (Clear Garden Decision Memo)
    if len(doc.tables) > 16:
        tbl16 = doc.tables[16]
        
        t16_full_data = [
            ("Mục trong Decision Memo", "Yêu cầu nội dung chuẩn mực", "Ví dụ áp dụng cho Clear Garden"),
            ("1. Luận điểm / Vấn đề", "Xác định rõ vấn đề của người chơi và cơ chế giải quyết được đề xuất.", "Màn 7 có tỷ lệ thua cao vì bàn cờ quá chật và xuất hiện nhiều loại vật phẩm rác."),
            ("2. Giá trị mang lại", "Người chơi nào nhận giá trị gì; khoảnh khắc nào giá trị đó xuất hiện.", "Người chơi giải đố cẩn thận nhận thêm 1 ô khay tạm thời để tự giải quyết thế cờ bế tắc."),
            ("3. Đề xuất can thiệp", "Giải pháp cụ thể và phương thức kích hoạt.", "Đề xuất gói cứu trợ 0.99$ kèm 1 lượt hoàn tác khi chỉ còn 1 vật phẩm cuối."),
            ("4. Chỉ số cảnh báo (Guardrails)", "Chỉ số an toàn bảo vệ trải nghiệm chung của sản phẩm.", "Tỷ lệ giữ chân D1 không được giảm quá 1%, tỷ lệ thử lại không suy giảm."),
            ("5. Điều kiện hủy bỏ (Kill criteria)", "Ngưỡng dữ liệu buộc studio phải dừng tính năng nếu không đạt.", "Nếu tỷ lệ thoát game sau khi thấy gói bán vượt quá 15%, lập tức hoàn tác."),
            ("6. Người phụ trách (Owner)", "Ai quyết định, ai vận hành và khi nào đánh giá lại.", "Lead Game Designer phối hợp Monetization Lead; đánh giá lại sau 7 ngày thử nghiệm."),
            ("7. Khung chiến lược", "Nội dung định hướng chiến lược", "Minh họa thực tế cho dự án Clear Garden"),
            ("8. Tín hiệu thị trường", "Dữ liệu nhu cầu quan sát từ bên ngoài.", "Dòng game sắp xếp dọn dẹp (Order-and-clear) đang chứng minh sức hút thương mại lớn."),
            ("9. Mục tiêu của người chơi", "Động cơ tâm lý cốt lõi người chơi tìm kiếm.", "Khôi phục trật tự ngăn nắp và tận hưởng sự đổi thay trực quan của khu vườn."),
            ("10. Thể hiện sản phẩm", "Cơ chế thể hiện qua lối chơi cụ thể.", "Số lượng khay chứa giới hạn, bàn cờ sắp xếp trực quan, tiến trình cải tạo vườn rõ nét."),
            ("11. Điểm khác biệt độc bản", "Yếu tố tạo đột phá so với đối thủ.", "Trải nghiệm cải tạo khu vườn tươi đẹp xuất hiện ngay trong phút chơi đầu tiên."),
            ("12. Nhu cầu thương mại", "Điểm chạm tạo cơ hội kiếm tiền tự nhiên.", "Đề xuất mua thêm ô khay hoặc hoàn tác khi người chơi nhận ra sai lầm do tính toán không gian."),
            ("13. Áp lực sản xuất", "Năng lực studio cần có để duy trì sản phẩm.", "Đội ngũ cần sản xuất liên tục các mẫu bàn cờ mới, tài nguyên mỹ thuật, nhiệm vụ và sự kiện mùa."),
            ("14. Bằng chứng kiểm chứng", "Dữ liệu thực nghiệm chứng minh thành công.", "Mẫu quảng cáo khớp lối chơi, tỷ lệ vượt Màn 1-10 khỏe mạnh, tỷ lệ thử lại cao, lý do quay lại rõ ràng."),
            ("15. Tiêu chuẩn dừng dự án", "Dấu hiệu cảnh báo buộc studio phải hủy bỏ.", "Thu hút được sự chú ý ban đầu nhưng không tạo ra hành vi chơi lặp lại một cách tự nhiên.")
        ]

        for r_idx, (c0, c1, c2) in enumerate(t16_full_data):
            if r_idx < len(tbl16.rows):
                tbl16.rows[r_idx].cells[0].text = c0
                if len(tbl16.rows[r_idx].cells) > 1:
                    tbl16.rows[r_idx].cells[1].text = c1
                if len(tbl16.rows[r_idx].cells) > 2:
                    tbl16.rows[r_idx].cells[2].text = c2

    # Save changes
    doc.save(DOCX_PATH)
    print(f"\n[SUCCESS] Successfully applied all semantic and formatting fixes to:\n{DOCX_PATH}")

if __name__ == "__main__":
    fix_all_issues()
