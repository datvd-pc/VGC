from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from pathlib import Path


DOCX_PATH = Path(r"D:\CODE\VGC\Long form\The-Art-of-Monetization-Vietnamese-Editable.docx")


def find_unique(doc, text):
    matches = [paragraph for paragraph in doc.paragraphs if paragraph.text == text]
    if len(matches) != 1:
        raise RuntimeError(f"Expected exactly one unchanged paragraph: {text!r}; found {len(matches)}")
    return matches[0]


def replace_text(doc, old, new):
    paragraph = find_unique(doc, old)
    paragraph.text = new
    return paragraph


def set_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def write_cell(cell, value, header=False):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(2)
    run = paragraph.add_run(value)
    run.font.size = Pt(8.5)
    run.bold = header
    if header:
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)


def add_table_after(doc, anchor, headers, rows, widths):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.autofit = False
    for column, header in enumerate(headers):
        cell = table.cell(0, column)
        cell.width = Inches(widths[column])
        write_cell(cell, header, header=True)
        set_shading(cell, "005868")
    for row_index, row in enumerate(rows, start=1):
        for column, value in enumerate(row):
            cell = table.cell(row_index, column)
            cell.width = Inches(widths[column])
            write_cell(cell, value)
            if row_index % 2 == 0:
                set_shading(cell, "F4F7F8")
    anchor._p.addnext(table._tbl)
    return table


doc = Document(DOCX_PATH)

# Chapter 25
replace_text(
    doc,
    "25. Market intelligence\nMarket intelligence thu hẹp lựa chọn, không thay team quyết định",
    "25. Market intelligence\nMarket intelligence làm hẹp không gian lựa chọn; product thesis mới quyết định team sẽ đánh cược vào đâu",
)
intelligence_intro = replace_text(
    doc,
    "Publisher lớn dùng dataset lớn. Họ track mechanic, theme, price, live event, creative, motivation, market và cohort. GameRefinery công khai nói về data trên hơn 100,000 games. Rovio từng mô tả platform có dashboard, A/B test, remote config, UA attribution và live-ops calendar. King nói data của họ shape design và strategy từ hàng tỷ gameplay events.",
    "Publisher lớn thật sự có lợi thế thông tin. GameRefinery công khai dữ liệu về hơn 100.000 game; Rovio mô tả Beacon với analytics, dashboard tùy biến, A/B testing, UA và live-ops configuration; King nói đội ngũ của họ phân tích hàng tỷ gameplay events mỗi ngày để hỗ trợ design và strategy. Một studio nhỏ không nên phủ nhận lợi thế đó. Câu hỏi thực tế hơn là: sau khi nhận được một market signal, team nào có thể tạo ra một product thesis khác biệt và kiểm chứng nó nhanh hơn?",
)
replace_text(
    doc,
    "Năng lực này tăng tốc độ và quality của decision. Nó không biến decision thành lookup.",
    "Năng lực này tăng tốc độ và chất lượng của decision, nhưng không biến decision thành một lần lookup. Dữ liệu ngoài thị trường quan sát được những thứ đã xuất hiện: mechanic, theme, price, visible event, creative direction hay kết quả ước tính. Nó không nhìn thấy toàn bộ context của từng team: chi phí sản xuất, năng lực level design, chất lượng iteration, mức tin cậy của economy hoặc lý do nội tại khiến player chọn ở lại.",
)
replace_text(
    doc,
    "Data có thể báo screw puzzle, rescue fantasy, collection mechanic hay event structure đã xuất hiện trong game thành công. Nó không trả lời team cụ thể có thể làm core loop hay hơn alternative, produce content đúng cadence, giữ trust hay học được gì nếu bet fail.",
    "Data có thể báo rằng screw puzzle, rescue fantasy, collection mechanic hay một cấu trúc event đã xuất hiện trong game thành công. Nó không trả lời liệu team cụ thể có thể làm core loop hay hơn alternative, sản xuất content đúng cadence, giữ trust hay học được gì nếu bet fail. Đó là monetization link của market intelligence: nó giúp loại bớt bet yếu trước khi chi tiền, nhưng không tạo ra value proposition mà player sẽ tiếp tục trả time và money cho.",
)
replace_text(
    doc,
    "Market signal: player đang spend time và money ở territory này.",
    "Market signal: player đang spend time và money ở territory này. Đây là dữ kiện bên ngoài có thể giúp team chọn nơi đào sâu.",
)
replace_text(
    doc,
    "Product thesis: team này có thể tạo experience distinct, durable, operate nó ở quality cần thiết và acquire audience có lợi nhuận.",
    "Product thesis: team này có thể tạo một experience khác biệt, bền và vận hành được ở quality cần thiết, đồng thời acquire đúng audience với economics có lợi nhuận. Đây là lời hứa phải được kiếm trong studio, qua prototype, cohort và năng lực vận hành.",
)
replace_text(
    doc,
    "Market signal nhìn thấy từ bên ngoài. Product thesis phải được earn trong studio.",
    "Market signal nhìn thấy từ bên ngoài. Product thesis phải được earn trong studio. Bảng dưới đây giúp giữ hai việc này không bị lẫn vào nhau.",
)
add_table_after(
    doc,
    intelligence_intro,
    ["Dữ liệu bên ngoài có thể giúp", "Dữ liệu bên ngoài không thể thay team quyết định"],
    [
        ["Sàng lọc territory, mechanic, theme, visible cadence và cạnh tranh.", "Core feeling nào team có thể làm tốt hơn và vì sao player sẽ chọn nó."],
        ["Đặt benchmark, tìm câu hỏi cần đào sâu, tránh clone lỗi thời.", "Level nào fair, offer nào đúng moment, price nào hợp cohort và economy nào còn trust."],
        ["Ước lượng quy mô/rủi ro và chọn prototype nhỏ để test.", "Khả năng sản xuất content, học từ thất bại, giữ quality và mở rộng contribution margin."],
    ],
    [3.15, 3.15],
)

# Chapter 26
replace_text(
    doc,
    "26. Behaviour needs interpretation\nData cho biết hành vi, không tự tạo ra ý nghĩa",
    "26. Behaviour needs interpretation\nData mô tả hành vi; ý nghĩa của hành vi phải được kiểm tra bằng một causal story",
)
replace_text(
    doc,
    "Rovio từng lưu ý test uplift có thể là random variation hoặc novelty effect. Dashboard ghi lại điều xảy ra sau design choice. Nó không tự giải thích player đã hiểu choice đó như thế nào.",
    "Dashboard ghi lại điều xảy ra sau một design choice. Nó không tự giải thích player đã hiểu choice đó như thế nào. Một uplift có thể đến từ novelty, source mix, random variation, friction bị đẩy sang chỗ khác hoặc một cơ chế thực sự tốt hơn. Vì vậy 'data nói gì?' chưa đủ. Câu hỏi tiếp theo phải là 'cơ chế nào có thể tạo ra pattern này, và bằng chứng nào sẽ bác bỏ cơ chế đó?'",
)
replace_text(
    doc,
    "Extra-tray offer của Clear Garden convert ở level bảy có thể vì tension satisfying, board difficult but fair, offer framing tốt, hoặc object type ẩn làm player desperate để thoát. Cùng một purchase event, bốn future khác nhau.",
    "Ví dụ, extra-tray offer của Clear Garden convert ở level bảy có thể vì tension satisfying, board difficult but fair, offer framing tốt, hoặc object type ẩn làm player desperate để thoát. Cùng một purchase event có ít nhất bốn tương lai khác nhau: player quay lại vì lựa chọn hữu ích; player mua một lần rồi rời; player từ chối và vẫn học được; hoặc player rời vì cảm thấy bị gài. Chỉ conversion không đủ để phân biệt chúng.",
)
replace_text(
    doc,
    "Watch recording, đọc review, hỏi player narrate decision, inspect board trước offer, so buyer-decliner-control theo thời gian. Mục tiêu không phải thay data bằng intuition. Mục tiêu là đưa data vào một causal story đáng test.",
    "Hãy xem recording, đọc review, yêu cầu player kể lại quyết định, kiểm tra board trước offer và so buyer, decliner, control theo thời gian. Mục tiêu không phải thay data bằng intuition. Mục tiêu là đưa data vào một causal story đáng test. Câu chuyện càng cụ thể, experiment càng nhỏ và rẻ; càng mơ hồ, team càng dễ dùng data để biện hộ cho feature đã muốn ship.",
)
replace_text(
    doc,
    "Supercell từng viết về Hay Day Pop: team ship nhanh và react với data, nhưng không có strong feeling for genre và không innovate core puzzle. Information không thiếu. Product conviction thiếu.",
    "Supercell mô tả Hay Day Pop là puzzle game do một team có track record tốt thực hiện, ship nhanh và phản ứng với data, nhưng không có cảm nhận đủ mạnh về genre để đổi mới core puzzle. Đây không phải lập luận chống data; chính Supercell dùng data sâu. Nó là bằng chứng rằng information và product conviction là hai tài sản khác nhau. Khi thiếu conviction, team dễ chạy theo metric ngắn hạn vì không có tiêu chuẩn trải nghiệm đủ rõ để phản biện chính mình. Nguồn case: Supercell, 2026.",
)

# Chapter 27
replace_text(
    doc,
    "27. Decision memo\nĐơn vị công việc hữu ích là một bản ghi quyết định",
    "27. Decision memo\nĐơn vị công việc hữu ích là một quyết định có thể kiểm tra, không phải một danh sách feature",
)
replace_text(
    doc,
    "Greenlight list thường có dạng: build screw puzzle, dùng rescue theme, thêm collection meta, test starter pack, run win streak.",
    "Greenlight list thường có dạng: build screw puzzle, dùng rescue theme, thêm collection meta, test starter pack, run win streak. Những dòng này là prompt để thảo luận, chưa phải decision. Chúng không nói rõ player nào, vấn đề nào, trade-off nào, bằng chứng nào sẽ khiến team dừng lại, hay ai chịu trách nhiệm về kết quả.",
)
replace_text(
    doc,
    "Đây là prompt, chưa phải decision.",
    "Một decision memo ngắn buộc team biến prompt thành một cam kết có thể kiểm tra. Nó không làm quyết định chắc chắn đúng; nó khiến cách team sai trở nên rẻ hơn và dễ học hơn.",
)
replace_text(
    doc,
    "Mechanic là ingredient. Product là coherent promise được giao lặp lại.",
    "Mechanic là ingredient. Product là lời hứa nhất quán được giao lặp lại qua creative, first session, progression, monetization và live ops. Khi memo không mô tả được lời hứa này, team đang chọn component trước khi hiểu product cần tạo cảm giác gì.",
)
memo_anchor = find_unique(doc, "Một decision memo ngắn buộc team biến prompt thành một cam kết có thể kiểm tra. Nó không làm quyết định chắc chắn đúng; nó khiến cách team sai trở nên rẻ hơn và dễ học hơn.")
add_table_after(
    doc,
    memo_anchor,
    ["Trường", "Nội dung cần viết"],
    [
        ["Decision", "Chọn một thay đổi hoặc một bet, không phải một danh sách feature."],
        ["Player value", "Player nào nhận giá trị gì; moment nào value đó xuất hiện."],
        ["Evidence và giả thuyết", "Điều biết từ data ngoài/inside; điều vẫn chỉ là giả thuyết."],
        ["Trade-off", "Retention, trust, economy, production hoặc UA có thể phải trả giá gì."],
        ["Test và kill condition", "Metric, guardrail, window đọc và điều kiện dừng/rollback."],
        ["Owner", "Ai quyết định, ai vận hành và khi nào review lại."],
    ],
    [1.55, 4.8],
)

# Chapter 28
replace_text(
    doc,
    "28. Copy the question, not the configuration\nHọc câu hỏi, đừng sao chép cấu hình",
    "28. Copy the question, not the configuration\nHọc vấn đề đối thủ đã giải; đừng sao chép cấu hình mà không có bối cảnh",
)
replace_text(
    doc,
    "Royal Match đặt câu hỏi về content cadence và operational depth. Merge Mansion đặt câu hỏi về board pressure, narrative và event tạo return rhythm. Hybrid-casual hit đặt câu hỏi về loop ngắn mang useful exchange mà không mất speed.",
    "Royal Match đặt câu hỏi về content cadence và operational depth. Merge Mansion đặt câu hỏi về board pressure, narrative và event tạo return rhythm. Một hybrid-casual hit đặt câu hỏi về loop ngắn có thể mang useful exchange mà không mất speed. Giá trị của competitor study là phát hiện câu hỏi có thật trên thị trường, chứ không phải sao chép event calendar, price ladder, level configuration hay economy mà không biết chúng đang phục vụ cohort nào.",
)
replace_text(
    doc,
    "Dùng public data để screen bad bets. Dùng product thesis để form narrow hypothesis. Dùng cohort để học sau novelty, source mix và operational cost. Big data mạnh ở pass một. Cuốn sách này mạnh nhất ở pass hai và ba.",
    "Dùng public data để screen bad bets. Dùng product thesis để tạo narrow hypothesis. Dùng cohort để học sau novelty, source mix và operational cost. Big data mạnh ở pass một: nhìn bản đồ và loại bớt territory. Lợi thế của một team biết đọc product nằm ở pass hai và ba: tạo giả thuyết có bản sắc, rồi học xem nó có đứng được trong build thật hay không. Đây là phần ebook có thể đóng góp ngay cả khi publisher đã gửi sẵn mechanic, theme hoặc skin.",
)

# Chapter 29
replace_text(
    doc,
    "29. Clear Garden: From prompt to decision\nClear Garden: Từ gợi ý của publisher đến quyết định sản phẩm",
    "29. Clear Garden: From prompt to decision\nClear Garden: biến gợi ý từ publisher thành một quyết định sản phẩm có thể bị bác bỏ",
)
replace_text(
    doc,
    "Nếu publisher gợi ý organizational puzzle, cozy restoration, collection layer và fail offers, gợi ý đó có thể hợp lý. Câu hỏi là team có làm nó coherent được không.",
    "Nếu publisher gợi ý organizational puzzle, cozy restoration, collection layer và fail offers, gợi ý đó có thể hợp lý theo market map. Nhưng market map chỉ nói đây là một territory đáng kiểm tra. Quyết định của studio phải cụ thể hơn: Clear Garden có thể biến việc sắp xếp thành cảm giác 'phục hồi trật tự' đủ rõ để player hiểu vì sao extra tray hoặc undo là lựa chọn hữu ích hay không?",
)
replace_text(
    doc,
    "Build đầu tiên không cần năm currency hay event pass. Nó cần prove bốn điều: creative promise xuất hiện trong play; tray constraint tạo planning thay vì panic; garden thay đổi đủ để completion tangible; extra tray hoặc undo giúp sau known mistake thay vì che hidden rule.",
    "Build đầu tiên không cần năm currency hay event pass. Nó cần chứng minh bốn điều: creative promise xuất hiện trong play; tray constraint tạo planning thay vì panic; garden thay đổi đủ để completion tangible; extra tray hoặc undo giúp sau known mistake thay vì che hidden rule. Mỗi điều cần một dạng bằng chứng: creative-to-play recording, level replay, player interview ngắn, cohort metrics và review về fairness. Đây là cách gợi ý từ publisher trở thành decision memo thay vì thành backlog.",
)
replace_text(
    doc,
    "Nếu bốn điều đúng, team có foundation. Nếu không, market map vẫn đã tiết kiệm thời gian bằng cách chỉ territory testable. Nó không xóa nhu cầu kill, revise hoặc tìm core feeling khác.",
    "Nếu bốn điều đúng, team có foundation để thêm economy và live ops. Nếu không, market map vẫn đã tiết kiệm thời gian bằng cách chỉ ra territory testable; nó không xóa nhu cầu kill, revise hoặc tìm core feeling khác. Điều này không làm studio nhỏ yếu hơn publisher. Nó làm studio có một lợi thế khác: khả năng nói 'đây là điều chưa được chứng minh' trước khi bỏ thêm months, people và UA spend vào một giả định đẹp.",
)

part_vii = next(
    p for p in doc.paragraphs
    if p.text == "Part VII: The operating system behind a live game\nPhần VII: Hệ điều hành của game đang phát hành"
)
heading = part_vii.insert_paragraph_before(
    "Memory note | Decision board\nPart VI: Data, market intelligence và product thesis"
)
heading.style = "Heading 3"
board = doc.add_table(rows=2, cols=2)
board.style = "Table Grid"
board.autofit = False
content = {
    "LÀM NGAY": [
        "Tách mọi research finding thành hai cột: market signal quan sát được và product thesis cần team tự chứng minh.",
        "Viết một decision memo cho bet đang được đề xuất bởi publisher hoặc competitor study.",
        "Chọn một gợi ý ngoài thị trường và biến nó thành prototype test có kill condition rõ.",
    ],
    "HỎI TRƯỚC KHI QUYẾT ĐỊNH": [
        "Dữ liệu này cho biết sự kiện nào đã xảy ra, và phần nào team đang tự suy diễn?",
        "Team có năng lực design, content, UA và live ops nào để giao lời hứa khác biệt này?",
        "Bằng chứng nào sẽ khiến team revise hoặc kill thesis thay vì thêm feature?",
        "Đối thủ đã giải vấn đề nào mà team đang cố copy cấu hình của họ?",
    ],
    "CẦN NHỚ": [
        "Big data làm hẹp lựa chọn và tăng tốc học; nó không tạo product conviction hay thay tay nghề triển khai.",
        "Một purchase event hoặc market chart không tự kể causal story của player.",
        "Product thesis cần được kiếm qua build, cohort, quality và năng lực vận hành, không được mượn từ brand của publisher.",
    ],
    "ĐƯA VÀO CUỘC HỌP": [
        "Thành phần: studio lead, product/design, data, production và UA/BD khi nguồn gợi ý đến từ publisher. Thời lượng: 60 phút.",
        "Mang theo: source research, decision memo, prototype/replay, capability map và kill criteria.",
        "Chốt: thesis nào được test, scope nhỏ nhất, owner, deadline và điều kiện kill/iterate/scale.",
    ],
}
slots = [(0, 0, "LÀM NGAY"), (0, 1, "HỎI TRƯỚC KHI QUYẾT ĐỊNH"), (1, 0, "CẦN NHỚ"), (1, 1, "ĐƯA VÀO CUỘC HỌP")]
for row, column, title in slots:
    cell = board.cell(row, column)
    cell.width = Inches(3.15)
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(title)
    run.bold = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x00, 0x58, 0x68)
    for item in content[title]:
        bullet = cell.add_paragraph(style="List Bullet")
        bullet.paragraph_format.space_after = Pt(2)
        bullet.add_run(item).font.size = Pt(8.5)
    set_shading(cell, "F4F7F8")
heading._p.addnext(board._tbl)
bridge = doc.add_paragraph(
    "Part VII chuyển từ thesis sang năng lực thực thi: supply chain, kill/iterate/scale, live ops và contribution economics quyết định một lời hứa có sống được sau launch hay không."
)
bridge.style = "Normal"
bridge.paragraph_format.space_before = Pt(4)
bridge.paragraph_format.space_after = Pt(8)
bridge.runs[0].italic = True
board._tbl.addnext(bridge._p)

doc.save(DOCX_PATH)
print(f"Updated: {DOCX_PATH}")
