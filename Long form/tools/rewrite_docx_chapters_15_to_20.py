from docx import Document
from pathlib import Path


DOCX_PATH = Path(r"D:\CODE\VGC\Long form\The-Art-of-Monetization-Vietnamese-Editable.docx")


def replace_exact(paragraph, expected, replacement):
    if paragraph.text != expected:
        raise RuntimeError(
            "Paragraph changed since this script was prepared. Refusing to overwrite.\n"
            f"Expected: {expected!r}\nActual: {paragraph.text!r}"
        )
    paragraph.text = replacement


def add_before(anchor, text, style="Normal"):
    paragraph = anchor.insert_paragraph_before(text)
    paragraph.style = style
    return paragraph


doc = Document(DOCX_PATH)
paragraphs = doc.paragraphs

# Chapter 15
replace_exact(
    paragraphs[179],
    "15. Rewarded ads\nQuảng cáo có thưởng là một cam kết về giá trị sử dụng",
    "15. Rewarded ads\nQuảng cáo có thưởng chỉ bền khi đổi lấy một giá trị người chơi vừa cần",
)
replace_exact(
    paragraphs[180],
    "Rewarded ads mạnh nhất khi kéo dài hành động player đã muốn làm: second chance sau readable loss, double reward sau meaningful win, refill để hoàn thành plan, undo để cứu một decision đã hiểu.",
    "Hãy chọn placement rewarded ad có opt-in cao nhất và xem lại ba mươi giây trước khi nó xuất hiện. Player đang muốn hoàn thành một kế hoạch rõ ràng, cứu một sai lầm đã hiểu, hay chỉ đang bị chặn ở một việc bình thường? Monetization link của rewarded ad là trao đổi tự nguyện: player nhận một giá trị đã nhìn thấy, còn game nhận attention. Nếu placement chỉ hoạt động vì game đã cắt ngắn normal play, opt-in cao có thể đang đo friction chứ không đo utility.",
)
replace_exact(
    paragraphs[181],
    "Reward phải rõ trước ad và grant ngay sau ad. Nếu ad fail, value phải được giữ. Dùng placement biến normal task thành waiting room là đang tiêu trust.",
    "Thiết kế placement theo một câu hoàn chỉnh: sau khi làm X, player có thể xem một ad để nhận Y, và Y giúp họ làm Z. Second chance sau một fail có thể hợp lý khi fail readable và lượt tiếp theo tạo cơ hội thật. Double reward hợp lý khi reward gốc đã có trọng lượng. Undo hợp lý khi player hiểu nước đi vừa sai. Reward phải được nêu rõ trước ad, được grant ngay sau ad, và vẫn được bảo toàn khi ad fail. Đây là điều kiện sản phẩm, không phải chi tiết của mediation SDK.",
)
replace_exact(
    paragraphs[182],
    "",
    "Không có một benchmark công khai nào đủ để kết luận opt-in bao nhiêu là tốt cho mọi puzzle game. Vì vậy, hãy chạy một experiment có control: giữ nguyên reward, thay placement hoặc moment; sau đó đọc opt-in cùng completion, exit sau exposure, next-session return, tần suất dùng placement lần hai và review language. Nếu reward được dùng lại mà retention không xấu đi, giả thuyết utility đáng được giữ. Nếu exposure tăng nhưng player thoát hoặc không quay lại, placement cần được xem lại trước khi tăng frequency cap.",
)
replace_exact(
    paragraphs[183],
    "High opt-in có thể là utility, cũng có thể là friction manufactured. Đọc retention theo exposure, exit sau level và review language.",
    "High opt-in có thể là utility, cũng có thể là friction manufactured. Đọc retention theo exposure, exit sau level, completion và review language trước khi gọi placement là thành công. Interstitial tạo ra một trao đổi khác: không có phần thưởng trực tiếp, nên chi phí attention của nó cần được đặt ở một điểm nghỉ thật sự.",
)

# Chapter 16
replace_exact(
    paragraphs[184],
    "16. Interstitials\nQuảng cáo xen kẽ cần một điểm nghỉ tự nhiên",
    "16. Interstitials\nQuảng cáo xen kẽ chỉ nên xuất hiện tại điểm mà sự tập trung đã được khép lại",
)
replace_exact(
    paragraphs[185],
    "Interstitial là thuế trên attention. Nó có thể chấp nhận được sau loop ngắn đã hoàn thành, sau reward đã bank hoặc giữa sessions. Nó đặt khi cắt concentration hoặc trước first minute thì đắt hơn nhiều.",
    "Hãy review ba vị trí interstitial hiện có bằng replay, không bằng tổng impression. Mỗi vị trí có nằm sau một loop đã hoàn thành, sau khi reward đã được ghi nhận, hoặc giữa hai session không? Hay nó cắt một lần suy nghĩ, một combo hoặc tutorial? Monetization link ở đây là attention tax: interstitial có thể tạo doanh thu ngay, nhưng nó lấy attention mà player đang dùng để học, giải hoặc quay lại. Việc chọn đúng breakpoint là quyết định sản phẩm, không phải default của ad network.",
)
replace_exact(
    paragraphs[186],
    "Frequency cap là product setting, không phải network default. Segment early và mature player, session depth, exposure, platform và purchase state. Buyer remove ads không nên phải nghi sản phẩm có hoạt động hay không.",
    "Nghiên cứu của Stothart, Mitchum và Yehnert năm 2015 cho thấy chỉ riêng notification điện thoại cũng làm gián đoạn hiệu suất trong một nhiệm vụ đòi hỏi chú ý, kể cả khi người tham gia không tương tác với điện thoại. Đây không phải nghiên cứu về interstitial hay puzzle. Nó cho một cơ chế đủ cụ thể để kiểm tra: một interruption có chi phí ngay cả khi player đóng nó nhanh. Vì vậy frequency cap cần được segment theo early/mature player, session depth, platform và trạng thái đã mua remove ads; đừng để network default quyết định thay team.",
)
header_17 = next(p for p in doc.paragraphs if p.text == "17. Boosters\nBooster phải mở rộng lựa chọn")
add_before(
    header_17,
    "Một test tối thiểu là so sánh cùng frequency cap ở hai breakpoint: sau level complete và giữa một chuỗi puzzle chưa khép lại. Đọc ad revenue cùng session completion, next-level start, return và tỷ lệ tắt game trong ba mươi giây sau exposure. Nguồn nghiên cứu: Stothart, Mitchum và Yehnert, 2015. Khi player chủ động chọn một tool để thay đổi tình huống, booster có thể tạo trao đổi rõ ràng hơn interstitial, nhưng chỉ khi họ hiểu tool đó làm gì.",
)

# Chapter 17
replace_exact(
    paragraphs[187],
    "17. Boosters\nBooster phải mở rộng lựa chọn",
    "17. Boosters\nBooster phải mở ra một lựa chọn có thể hiểu được, không vá một thất bại mơ hồ",
)
replace_exact(
    paragraphs[188],
    "Booster đáng giá khi nó cho player thể hiện strategy, recover từ known mistake hoặc đi qua bottleneck đã chọn. Nó đáng nghi khi là đáp án duy nhất cho unclear design.",
    "Mở replay của một player vừa dùng booster, rồi tua về trước thời điểm offer xuất hiện. Họ có thể chỉ ra vấn đề cụ thể cần giải không: thiếu một ô trống trong sort, cần undo sau một nước đi sai, thiếu vài lượt đi để hoàn thành board đã hiểu? Nếu có, booster đang mở một lựa chọn. Nếu không, booster dễ trở thành đáp án duy nhất cho một level mơ hồ. Monetization link là willingness to pay hoặc willingness to spend currency chỉ có ý nghĩa khi player hiểu công cụ đang đổi tình huống của họ như thế nào.",
)
replace_exact(
    paragraphs[189],
    "Giới thiệu booster trong safe context, cho player thấy value, cho ít nhất một free use, rồi đợi đến khi need đang sống. Bán tool trước khi player hiểu job chỉ là bán icon.",
    "Nghiên cứu về worked examples và cognitive load cho thấy người mới học xử lý tốt hơn khi cấu trúc của nhiệm vụ được làm rõ trước khi họ phải tự giải. Chuyển nguyên tắc này sang booster: hãy giới thiệu tool trong safe context, cho một lần dùng miễn phí, chỉ ra tình huống nó thay đổi, rồi mới để player quyết định có dùng lại hay không. Nghiên cứu không đo booster revenue. Nó chỉ bảo vệ một thứ thiết kế cần có trước doanh thu: player cần có mô hình đủ rõ về 'job' của tool trước khi được mời trả tiền cho nó.",
)
header_18 = next(p for p in doc.paragraphs if p.text == "18. IAP and present need\nIAP phải giải quyết nhu cầu đang hiện diện")
add_before(
    header_18,
    "Đo booster bằng hai lớp chỉ số. Lớp hành vi: lần dùng đầu, repeat use, loại tình huống dùng và tỷ lệ hoàn thành sau dùng. Lớp tin cậy: retry không dùng booster, exit sau từ chối, review và refund nếu có. Một booster convert cao nhưng gần như chỉ xuất hiện sau fail không thể giải thích cần được sửa ở level trước khi tối ưu offer. Nguồn nghiên cứu: Chen, Retnowati và Kalyuga, 2020. IAP rộng hơn booster vì nó thường bán một kết quả, một quyền lợi hoặc một nhịp chơi chứ không chỉ một tool.",
)

# Chapter 18
replace_exact(
    paragraphs[190],
    "18. IAP and present need\nIAP phải giải quyết nhu cầu đang hiện diện",
    "18. IAP and present need\nIAP phải giải quyết một nhu cầu đang hiện diện và nói rõ cái giá của lựa chọn đó",
)
replace_exact(
    paragraphs[191],
    "Store là catalogue. Offer là một câu nói tại một moment.",
    "In-game store là catalogue; offer là một câu nói tại một moment. Hãy mở một offer đang có và yêu cầu người không làm game trả lời bốn câu trong năm giây: offer này giải quyết việc gì, vì sao xuất hiện lúc này, họ nhận chính xác gì, và nếu không mua thì chuyện gì xảy ra? Nếu bốn câu không rõ, player đang phải tự làm phần design work mà product lẽ ra phải làm. Monetization link không phải chỉ là conversion: clarity quyết định liệu payment có được cảm nhận như một lựa chọn tự nguyện hay một cái bẫy thông tin.",
)
replace_exact(
    paragraphs[192],
    "Generic currency pack bắt player làm design work: tự suy ra need, tính conversion, tự đoán future use. Contextual product giảm burden. Offer sạch cho biết nó giải quyết gì, tại sao bây giờ, nhận gì và cái gì vẫn optional.",
    "Generic currency pack có thể phù hợp với player đã hiểu economy, nhưng nó buộc người mới tự suy ra need, tính conversion và đoán future use. Contextual product giảm phần việc đó: một remove-ads product bán trải nghiệm không bị ngắt quãng; một starter pack hỗ trợ vài ngày đầu; một offer sau fail giải một tình huống vừa xảy ra, nhưng vẫn phải cho player quyền từ chối rõ ràng. Offer sạch nói được nó giải quyết gì, tại sao bây giờ, nhận gì và lựa chọn miễn phí nào vẫn còn.",
)
replace_exact(
    paragraphs[193],
    "",
    "FTC bắt đầu gửi hơn 72 triệu USD tiền hoàn vào tháng 12/2024 trong vụ Epic Games, với 629.344 khoản thanh toán đầu tiên; thỏa thuận 245 triệu USD liên quan đến các cáo buộc về purchase flow gây ra charge không mong muốn. Đây không phải case về một gói IAP trong puzzle, và không chứng minh mọi impulse purchase đều sai. Dữ kiện công khai này đặt một guardrail rất cụ thể: purchase flow không được dựa vào nút khó hiểu, xác nhận thiếu rõ ràng hoặc đường hoàn tiền bị che giấu. Conversion không bao giờ là lý do để bỏ qua consent có hiểu biết.",
)
replace_exact(
    paragraphs[194],
    "Epic Plane Evolution cho thấy \"thêm IAP\" không phải thêm pack currency. AppMagic mô tả ticket, hard currency, remove ads và VIP-style benefits, cũng như khác biệt theo platform và region. Bài học là build product quanh use case và đọc theo cohort.",
    "Khi thử IAP, hãy tạo product brief một trang: need đang sống, cohort, nội dung sản phẩm, giá, lựa chọn từ chối, metric chính và guardrail như refund, review, payer retention. Sau đó mới quyết định pack currency, remove ads, ticket, starter pack hay benefit dài hạn. Economy là nơi các sản phẩm này gặp nhau: một offer tốt vẫn thất bại nếu currency nó bán không còn giữ giá trị hoặc không có nơi dùng hợp lý.",
)

# Chapter 19
replace_exact(
    paragraphs[195],
    "19. Economy integrity\nEconomy cần nguồn vào, nơi tiêu và hệ quả",
    "19. Economy integrity\nEconomy chỉ đáng tin khi mỗi currency có đường vào, đường ra và một quyết định còn lại",
)
replace_exact(
    paragraphs[196],
    "Mỗi currency cần source, sink, cadence và purpose. Source là play, event, ads, purchase, social reward. Sink là retry, upgrade, cosmetic, unlock, gate và acceleration. Purpose là decision còn lại nếu bỏ các con số đi.",
    "Bắt đầu bằng một currency, không phải cả economy. Vẽ phương trình đơn giản cho một cohort trong một khoảng thời gian: số dư cuối = số dư đầu + mọi source - mọi sink. Sau đó ghi vào source: play, event, rewarded ad, purchase hoặc social reward; ghi vào sink: retry, upgrade, cosmetic, unlock, gate hoặc acceleration. Monetization link nằm ở phần còn lại sau khi bỏ những con số: currency này còn tạo ra quyết định nào, hay chỉ tích lại để bảng số trông bận rộn?",
)
replace_exact(
    paragraphs[197],
    "Inflation xảy ra khi reward tăng nhanh hơn meaningful use. Artificial scarcity xảy ra khi game cắt normal play chỉ để ép sink. Một cái làm reward rẻ, một cái làm game hostile.",
    "Đây là mô hình stock-and-flow, không phải một case thị trường: currency là stock, source và sink là flow, còn cadence quyết định stock tích hay cạn theo thời gian. Inflation xảy ra khi source tăng nhanh hơn meaningful use; artificial scarcity xảy ra khi normal play bị cắt chỉ để ép một sink. Một bên làm reward rẻ đi, bên kia làm game trở nên hostile. Cả hai đều làm IAP và rewarded ad mất ý nghĩa vì player không còn tin vào giá trị thật của resource.",
)
header_20 = next(p for p in doc.paragraphs if p.text == "20. Event economy\nEvent là một nền kinh tế nhỏ có thời hạn")
add_before(
    header_20,
    "Trong economy review, chọn một currency và chốt bốn số: median balance theo cohort, tốc độ earn, tốc độ spend và tỷ lệ player không còn sink hợp lý. Sau đó mô phỏng một thay đổi, như tăng reward từ ad hoặc thêm starter pack, trước khi đưa vào live. Mô hình không dự báo chính xác doanh thu; nó giúp team thấy một source mới sẽ làm nguồn lực, decision và sink thay đổi ở đâu. Event là nơi một economy nhỏ có thời hạn cho phép thử những trao đổi này mà không phá core economy.",
)

# Chapter 20
replace_exact(
    paragraphs[198],
    "20. Event economy\nEvent là một nền kinh tế nhỏ có thời hạn",
    "20. Event economy\nEvent là một nền kinh tế nhỏ có thời hạn, không phải một chồng pop-up",
)
replace_exact(
    paragraphs[199],
    "Event tạo temporary loop: earn, choose, spend, progress, finish. Deadline tạo energy. Event currency bảo vệ core economy. Premium acceleration tạo đường spend mà không làm free participation vô nghĩa.",
    "Hãy mô tả một event bằng một vòng lặp nhỏ: earn, choose, spend, progress, finish. Deadline có thể tạo energy; event currency có thể bảo vệ core economy; premium acceleration có thể tạo một đường spend. Nhưng từng phần chỉ có ý nghĩa khi free participation vẫn cho player một mục tiêu hoàn thành được và spend chỉ rút ngắn, mở rộng hoặc cá nhân hóa hành trình đó. Monetization link của event không phải doanh thu trong tuần diễn ra, mà là chất lượng của loop tạm thời và ảnh hưởng của nó lên return, core economy và hành vi sau event.",
)
replace_exact(
    paragraphs[200],
    "Đừng run event như pile of pop-ups. Event calendar trưởng thành cần cadence, segmentation, reward policy, content supply và recovery time. Team cần biết event tăng return hay chỉ kéo future spend về tuần này.",
    "Sensor Tower ước tính Royal Match đạt 112 triệu USD doanh thu gộp và 14,6 triệu lượt tải trong tháng 7/2023, đồng thời ghi nhận nhịp phát hành khoảng 200 level mỗi tháng cùng các lớp như Royal Pass và Hidden Temple. Dữ liệu này chứng minh quy mô của một hệ vận hành content và monetization có thể đạt được; nó không cho biết event nào, pass nào hoặc reward nào tạo ra từng đồng doanh thu. Vì vậy đừng copy một calendar từ game top. Hãy dùng case để đặt tiêu chuẩn vận hành: event cần content supply, cadence, segmentation, reward policy, measurement và thời gian recovery.",
)
part_v = next(
    p for p in doc.paragraphs
    if p.text == "Part V: Signals, decisions, and experiments\nPhần V: Tín hiệu, quyết định và thử nghiệm"
)
add_before(
    part_v,
    "Trong post-event review, hãy tách doanh thu trong event khỏi thay đổi sau event: D1/D7 của cohort tham gia, return tuần kế tiếp, số dư currency, churn, review và tỷ lệ dùng lại sản phẩm. Nếu event chỉ kéo future spend về tuần này hoặc làm core loop bị bỏ trống, calendar đang tạo doanh thu mượn. Nguồn market data: Sensor Tower, tháng 8/2023. Part V sẽ biến các quan sát trên thành một hệ thống đọc dashboard, cặp chỉ số và experiment để team không tự lừa mình bằng một uplift đơn lẻ.",
)

doc.save(DOCX_PATH)
print(f"Updated: {DOCX_PATH}")
