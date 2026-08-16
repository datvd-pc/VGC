from docx import Document
from pathlib import Path


DOCX_PATH = Path(r"D:\CODE\VGC\Long form\The-Art-of-Monetization-Vietnamese-Editable.docx")


def replace_exact(paragraph, expected, replacement):
    if paragraph.text != expected:
        raise RuntimeError(
            "Paragraph changed since this script was prepared. Refusing to overwrite.\n"
            f"Expected: {expected!r}\nActual: {paragraph.text!r}"
        )
    paragraph.text = replacement


def add_before(anchor, text, style="Normal"):
    paragraph = anchor.insert_paragraph_before(text)
    paragraph.style = style
    return paragraph


doc = Document(DOCX_PATH)
paragraphs = doc.paragraphs

# Chapter 9
replace_exact(
    paragraphs[153],
    "9. Player motivations\nNgười chơi khác nhau mua những dạng giải tỏa khác nhau",
    "9. Player motivations\nNgười chơi trả tiền để giải những nhu cầu khác nhau, không phải vì họ cùng một hành trình",
)
replace_exact(
    paragraphs[154],
    "Không có universal player journey.",
    "Hãy chọn một level có tỷ lệ dùng booster hoặc rewarded ad đáng chú ý, rồi chia cohort theo hành vi trước khi gặp level đó. Có người đến đây sau nhiều lần thử vì muốn tự giải; có người vừa hoàn thành một chuỗi và muốn giữ nhịp; có người chỉ muốn vượt qua để tiếp tục câu chuyện hoặc bộ sưu tập. Nếu ba nhóm được nhìn như một player trung bình, team rất dễ đưa cùng một offer vào cùng một thời điểm rồi gọi mọi người không mua là 'không có nhu cầu'. Monetization link của chương này là segmentation: sản phẩm kiếm tiền cần giải một nhu cầu đang hiện diện, không phải chỉ xuất hiện ở một điểm có pressure.",
)
replace_exact(
    paragraphs[155],
    "",
    "Theory anchor phù hợp là Self-Determination Theory, với ba nhu cầu thường được dùng để phân tích động lực: autonomy, competence và relatedness. Một thí nghiệm trong game-learning của Sheldon và Filak năm 2008 cho thấy khi hỗ trợ các nhu cầu này thay đổi, động lực và một số kết quả trải nghiệm cũng thay đổi. Nghiên cứu không dự đoán một player casual sẽ mua gói nào. Nó cho team một cách đọc tốt hơn: player đang muốn cảm thấy có quyền chọn, muốn chứng minh mình hiểu board, hay muốn duy trì một mối liên hệ với collection, câu chuyện hoặc cộng đồng? Mỗi trạng thái cần một giải pháp khác nhau.",
)
replace_exact(
    paragraphs[156],
    "Các nhãn này là công cụ suy nghĩ, không phải những chiếc hộp kín. Người chơi thất bại không tự động là người sẽ trả tiền. Có người muốn chơi lại, có người muốn hiểu level, có người sẽ rời đi. Một hệ thống mạnh phải chừa chỗ cho cả ba phản ứng.",
    "Hãy dùng các nhãn như solver, relaxer, collector, optimizer hoặc convenience payer như giả thuyết, không phải persona cố định. Cùng một player có thể đổi động lực giữa các phiên. Trong design review, mỗi offer nên trả lời: nó giúp nhóm nào làm điều gì, bằng cách nào, và lựa chọn miễn phí nào vẫn còn? Đọc acceptance, repeat use, D1/D7 và phản hồi định tính theo cohort trước khi mở rộng placement. Nguồn lý thuyết: Sheldon và Filak, 2008. Khi nhu cầu đã rõ, câu hỏi tiếp theo là game có làm player cảm thấy họ đang thực sự tiến lên hay chỉ đang tích thêm con số?",
)

# Chapter 10
replace_exact(
    paragraphs[157],
    "10. Meaningful progress\nTiến bộ phải thay đổi tình huống của người chơi",
    "10. Meaningful progress\nTiến bộ chỉ có giá trị khi nó thay đổi tình huống người chơi đang quan tâm",
)
replace_exact(
    paragraphs[158],
    "Progress có thể là mastery, completion, collection, streak, identity, narrative movement hoặc control lớn hơn. Currency counter chỉ là evidence nếu nó đổi một trong các state đó.",
    "Lấy một reward loop đang chạy tốt trên dashboard và tắt phần hiển thị số trong bản review. Sau phiên đó, player có thể nói điều gì đã thay đổi không: họ giải puzzle tốt hơn, mở được khu vực mới, tiến gần một collection, có thêm quyền kiểm soát, hay chỉ có thêm currency? Đây là monetization link của progress. Reward chỉ tạo nhu cầu bền vững khi nó dẫn đến một lựa chọn hoặc trạng thái đáng mong muốn; nếu không, economy có thể tăng activity mà không tăng lý do để quay lại, xem ad hoặc mua hỗ trợ.",
)
replace_exact(
    paragraphs[159],
    "Bỏ reward khỏi session. Player còn biết cái gì đã tốt hơn không? Nếu không, game có thể tạo activity mà không tạo advancement.",
    "Goal-gradient research của Kivetz, Urminsky và Zheng năm 2006 cung cấp một cơ chế tham khảo. Trong các chương trình phần thưởng của người tiêu dùng, người tham gia có xu hướng tăng nỗ lực khi cảm thấy gần mốc thưởng hơn; nghiên cứu cũng quan sát hành vi quay lại nhanh hơn trong bối cảnh đó. Đây không phải bằng chứng rằng một progress bar sẽ tăng retention trong puzzle. Nó cho biết phần nào cần có thật trước khi thử: mục tiêu cụ thể, tiến độ nhìn thấy được và phần thưởng đủ ý nghĩa để khoảng cách còn lại có giá trị.",
)
replace_exact(
    paragraphs[160],
    "Reward quantity không phải reward weight. Một resource khan hiếm với source, sink và decision rõ có thể làm việc tốt hơn trang currency counter.",
    "Reward quantity không phải reward weight. Một resource ít hơn nhưng có source, sink và quyết định rõ có thể đáng giá hơn nhiều loại currency cùng tăng. Hãy thử thay đổi một phần của loop, chẳng hạn cho progress mở quyền chọn board tiếp theo thay vì chỉ cộng tiền; sau đó so sánh next-session return, mục tiêu hoàn thành, spend, rewarded-ad uptake và lượng currency bị bỏ quên. Tiến bộ có ý nghĩa không loại bỏ pressure. Nó cho pressure một nền tảng công bằng, vì player biết điều gì đang bị đe dọa và điều gì đáng để cứu.",
)

# Chapter 11
replace_exact(
    paragraphs[161],
    "11. Pressure creates a decision\nÁp lực phải tạo ra một quyết định",
    "11. Pressure creates a decision\nÁp lực chỉ có giá trị khi người chơi vẫn nhìn thấy một quyết định thật",
)
replace_exact(
    paragraphs[162],
    "Lượt đi có hạn, ô chứa chật, đồng hồ đếm ngược, thời hạn event, near miss, chuỗi thắng và tài nguyên khan hiếm đều tạo áp lực. Áp lực tốt khi người chơi hiểu tình huống và thấy ít nhất hai con đường: chơi lại, đổi chiến lược, dùng công cụ miễn phí, xem quảng cáo, dùng tiền tệ trong game, mua, hoặc tạm dừng và quay lại sau.",
    "Mở lại một fail offer và mô tả nó bằng ngôn ngữ của player, không phải ngôn ngữ của dashboard. Họ đang thiếu gì, tại sao lại thiếu, họ có thể làm gì nếu không trả tiền, và offer có giải quyết đúng trở ngại đó không? Lượt đi có hạn, board chật, đồng hồ, thời hạn event, near miss và tài nguyên khan hiếm đều có thể tạo pressure. Chúng chỉ trở thành monetization context lành mạnh khi player hiểu tình huống và nhìn thấy ít nhất hai hướng đi: thử lại, đổi chiến lược, dùng công cụ miễn phí, xem rewarded ad, dùng currency, mua, hoặc tạm dừng để quay lại sau.",
)
replace_exact(
    paragraphs[163],
    "Hỏi trong design review: \"Pressure này tạo ra quyết định gì?\" Nếu câu trả lời là \"pay hoặc suffer\", design cần làm lại.",
    "Có một case công khai cần được dùng đúng phạm vi. Tháng 12/2022, FTC công bố thỏa thuận với Epic Games tổng trị giá hơn 520 triệu USD, gồm 245 triệu USD dành cho hoàn tiền cho người tiêu dùng, liên quan đến các cáo buộc về dark patterns dẫn đến charge không mong muốn, cùng các vấn đề về quyền riêng tư trẻ em. Fortnite không phải puzzle game, và case này không chứng minh một fail offer trong casual game là bất công. Nó chứng minh một điều hẹp nhưng quan trọng: khi pressure, purchase flow và consent được thiết kế khiến người dùng khó hiểu hoặc khó từ chối, rủi ro không dừng ở conversion mà có thể thành hoàn tiền, niềm tin và pháp lý.",
)
header_12 = next(p for p in doc.paragraphs if p.text == "12. Dynamic difficulty\nĐộ khó động là cơ chế tạo hoặc phá trust")
add_before(
    header_12,
    "Trong design review, hãy hỏi: pressure này tạo ra quyết định nào, và lựa chọn không trả tiền có còn đáng tin không? Đọc conversion cùng thời gian từ fail đến exit, retry không dùng trợ giúp, tỷ lệ quay lại sau offer, refund và ngôn ngữ review. Một offer chuyển đổi tốt chỉ cho thấy nó được chọn trong bối cảnh hiện tại; nó chưa cho biết bối cảnh đó công bằng hay có thể duy trì. Nguồn case: FTC, 2022. Độ khó động là nơi team thường tạo pressure theo thời gian thực, vì vậy nó cần được kiểm tra kỹ hơn ở chương tiếp theo.",
)

# Chapter 12
replace_exact(
    paragraphs[164],
    "12. Dynamic difficulty\nĐộ khó động là cơ chế tạo hoặc phá trust",
    "12. Dynamic difficulty\nĐộ khó động phải điều chỉnh thử thách, không được âm thầm định đoạt kết quả",
)
replace_exact(
    paragraphs[165],
    "Difficulty định hình cách player giải thích game. Win là earned hay được cho? Loss có thể sửa bằng decision tốt hơn hay game đã quyết định cho player thua? Giải thích riêng tư đó ảnh hưởng retry, booster use, purchase và word of mouth.",
    "Hãy chọn một rule dynamic đang được đề xuất, như thay đổi số lượt đi, blocker density, số ống trống của sort puzzle hoặc tỷ lệ xuất hiện một vật thể. Viết trước điều player có thể quan sát, điều hệ thống được phép thay đổi, và điều tuyệt đối không được thay đổi. Monetization link là rất trực tiếp: nếu độ khó động đưa player vào những thất bại họ không thể giải thích, booster use hoặc fail-offer conversion có thể tăng; nhưng game đang đánh đổi agency, retry tự nguyện và niềm tin vào kết quả.",
)
replace_exact(
    paragraphs[166],
    "Theo dõi start rate, fail rate, retry rate, next-level start, booster use và exit rate.",
    "Một nghiên cứu DDA năm 2021 trên 50 người chơi một exergame so sánh năm cách điều chỉnh độ khó. Phương pháp có dữ liệu phong phú nhất dự đoán sở thích thay đổi độ khó chính xác hơn, nhưng nghiên cứu không tìm thấy khác biệt có ý nghĩa giữa các phương pháp ở các thước đo động lực nội tại và flow; mức khớp giữa thay đổi độ khó với sở thích người chơi lại tương quan với enjoyment và pressure. Bối cảnh là Pong vận động, không phải puzzle mobile, nên không thể mang kết quả sang monetization. Giá trị của nghiên cứu là một cảnh báo: thuật toán phức tạp hoặc nhiều dữ liệu hơn không tự động tạo trải nghiệm tốt hơn.",
)
replace_exact(
    paragraphs[167],
    "",
    "Một rule có thể bắt đầu an toàn hơn bằng cách giảm độ khó sau tín hiệu confusion rõ, hoặc mở thêm một phương án chiến lược, thay vì bí mật tạo win hay spawn một vật thể vừa đủ để player thất bại. Hãy giữ một control cohort, log rule đã can thiệp ở đâu, và đọc start rate, fail rate, retry rate, next-level start, booster use, rewarded-ad uptake và exit rate theo nhóm có và không có can thiệp. Thêm khảo sát ngắn sau thất bại nếu team có khả năng: player hiểu mình có thể làm gì khác không?",
)
replace_exact(
    paragraphs[168],
    "Fail offer convert chỉ chứng minh pressure có thể convert. Nó không chứng minh level fair hay durable. Đọc post-offer retention, review và refund trước khi gọi nó healthy.",
    "Fail-offer conversion chỉ chứng minh pressure có thể chuyển đổi trong khoảnh khắc đó. Nó không chứng minh level fair hay doanh thu có thể duy trì. Đọc post-offer retention, review, refund và hành vi của player từ chối offer trước khi kết luận. Nguồn nghiên cứu: Darzi, McCrea và Novak, 2021. Khi một phần kết quả còn phụ thuộc vào random seed, ranh giới giữa thử thách điều chỉnh được và cảm giác bị hệ thống chi phối sẽ rõ hơn nữa.",
)

# Chapter 13
replace_exact(
    paragraphs[169],
    "13. Randomness and skill\nTính ngẫu nhiên phải chừa chỗ cho kỹ năng",
    "13. Randomness and skill\nTính ngẫu nhiên chỉ công bằng khi người chơi vẫn có thể học và phản ứng",
)
replace_exact(
    paragraphs[170],
    "Randomness tạo surprise và replay value, nhưng tạo fairness risk khi player không phân biệt được luck và hidden control. Dùng nó để đổi tình huống mà không xóa agency. Player phải có thể respond, plan và learn.",
    "Hãy review một board thất bại có random seed khác nhau. Sau khi nhìn lại, player có thể nêu một nước đi khác để tăng cơ hội thắng không, hay mọi bản xem lại đều dẫn đến kết luận rằng game cần rơi đúng một vật thể khác? Tính ngẫu nhiên có thể tạo surprise và replay value, nhưng nó tạo rủi ro monetization khi player không phân biệt được luck, skill và hidden control. Nếu thất bại không thể học được, booster, rewarded ad hay IAP xuất hiện ngay sau đó dễ bị hiểu là giá để thoát khỏi một kết quả đã được định đoạt.",
)
replace_exact(
    paragraphs[171],
    "Near miss hay nói: \"Tôi thấy điều mình có thể làm khác.\" Near miss tệ nói: \"Game đã quyết định tôi phải thua.\"",
    "Nghiên cứu về near miss trong gambling là một guardrail, không phải nguồn cảm hứng để chuyển nguyên xi sang puzzle. Trong game may rủi, near miss không cung cấp thông tin hữu ích về chiến lược nhưng có thể làm người chơi muốn tiếp tục; các nghiên cứu cũng gắn hiệu ứng này với cảm giác kiểm soát ảo. Puzzle có skill phải đặt chuẩn cao hơn: một near miss tốt khiến player nhận ra lựa chọn, thứ tự hoặc cách dùng tool mà họ có thể thay đổi. Một near miss chỉ được thiết kế để gợi ý rằng chiến thắng 'sắp tới' mà không cho thông tin hành động là tín hiệu cần dừng lại.",
)
header_14 = next(p for p in doc.paragraphs if p.text == "14. Reward feel\nÂm thanh, phản hồi rung và cảm giác phần thưởng")
add_before(
    header_14,
    "Hãy test randomness bằng replay review, không chỉ bằng drop-rate table. Theo dõi tỷ lệ retry sau fail, lựa chọn booster, exit, khiếu nại về rigging và chênh lệch outcome giữa các seed. Rule có thể vẫn ngẫu nhiên nhưng cần giữ một khoảng agency: player được plan, được phản ứng và thấy vì sao kết quả đổi khác. Nguồn guardrail: Clark và cộng sự, 2009; Belisle và Dixon, 2016, đều trong bối cảnh gambling. Phản hồi âm thanh và rung có thể làm một kết quả có vẻ lớn hơn hoặc nhỏ hơn thực tế, nên chương cuối Part này sẽ kiểm tra cách game truyền đạt giá trị mà không thổi phồng nó.",
)

# Chapter 14
replace_exact(
    paragraphs[172],
    "14. Reward feel\nÂm thanh, phản hồi rung và cảm giác phần thưởng",
    "14. Reward feedback\nÂm thanh và phản hồi rung phải xác nhận đúng giá trị phần thưởng",
)
replace_exact(
    paragraphs[173],
    "Audio và haptic điều tiết expectation. Sound nhỏ làm clear feel final. Haptic restrained đánh dấu risky move. Reward sound cho biết effort đã đổi state của game.",
    "Hãy đặt cạnh nhau hai reward cùng giá trị economy nhưng có phản hồi khác nhau: một reward nhỏ dùng animation, sound và haptic như jackpot; reward còn lại phản hồi vừa đủ để player hiểu trạng thái đã đổi. Player đang học điều gì về giá trị trong hai trường hợp đó? Monetization link ở đây không phải 'làm phản hồi lớn hơn để bán nhiều hơn'. Reward feedback định hình perceived value của currency, booster và purchase; khi mọi thứ đều được trình bày như một chiến thắng lớn, game tự làm mờ thang giá trị cần thiết để player đưa ra quyết định có hiểu biết.",
)
replace_exact(
    paragraphs[174],
    "Chúng không được nói dối. Tiny reward không nên có jackpot fanfare. Failed action không nên bị che bằng âm thanh vui. Inflated feedback dạy player đến chỗ không tin senses của mình.",
    "Một nghiên cứu năm 2025 về haptic feedback trên thiết bị di động quan sát rằng rung có thể tạo reward response khác với phản hồi âm thanh hoặc hình ảnh và ảnh hưởng đến lựa chọn trong các bối cảnh mua sắm trực tuyến. Nghiên cứu này không nói về game, càng không chứng minh rung làm IAP conversion tăng trong puzzle. Nhưng nó đủ để bác bỏ giả định rằng sound và haptic chỉ là phần trang trí. Chúng là một phần của cách người dùng cảm nhận việc xác nhận, giá trị và nhịp điệu của thao tác; vì vậy phải được thử nghiệm cùng với thông điệp và economy, không tách rời khỏi chúng.",
)
replace_exact(
    paragraphs[175],
    "Puzzle là market lớn, nhưng không phải một product duy nhất. Sensor Tower báo cáo puzzle ở Mỹ khoảng $5B năm 2022, classic match-3 khoảng $1.6B. Scale không cho phép mang mechanic monetization của match-3 sang sort hay jam một cách máy móc.",
    "Nguyên tắc thực hành rất ngắn: reward nhỏ nhận phản hồi nhỏ nhưng rõ; reward hiếm được phép có khoảnh khắc lớn vì trạng thái game đã thực sự thay đổi; thất bại phải được báo đúng là thất bại; purchase thành công cần xác nhận vật phẩm, quyền lợi và cách dùng mà không giả vờ rằng một gói nhỏ đã giải quyết toàn bộ game. Hãy A/B test intensity của feedback cùng một reward, rồi đọc repeat use, retention, completion, purchase refund và phản hồi định tính, thay vì chỉ đọc tap-through. Nguồn nghiên cứu: Sinha, 2025, Journal of Consumer Research. Part IV sẽ chuyển từ các điều kiện tâm lý sang các sản phẩm monetization cụ thể: rewarded ads, interstitials, booster, IAP và economy.",
)

part_iv = next(
    p for p in doc.paragraphs
    if p.text == "Part IV: Ads, IAP, and economy\nPhần IV: Quảng cáo, IAP và economy"
)
add_before(part_iv, "Memory note | Part III\nGhi nhớ trước khi bước sang Part IV", "Heading 3")
add_before(
    part_iv,
    "Keywords: segmentation là đọc nhu cầu theo cohort thay vì một player trung bình; autonomy là cảm giác vẫn có quyền chọn; competence là cảm giác hiểu và làm chủ thử thách; meaningful progress là thay đổi trạng thái player thực sự quan tâm; pressure là điều kiện tạo quyết định chứ không phải công cụ dồn ép; agency là khả năng học, lập kế hoạch và phản ứng trước kết quả; reward feedback là tín hiệu cho player biết giá trị nào vừa được xác nhận.",
)
add_before(
    part_iv,
    "Bài học cần nhớ: một offer không thể phù hợp với mọi động lực; currency chỉ có ý nghĩa khi dẫn đến trạng thái hoặc lựa chọn đáng giá; conversion trong bối cảnh pressure chưa chứng minh fairness; dynamic difficulty cần có giới hạn quan sát được và control cohort; randomness chỉ đáng giữ khi player vẫn thấy đường học; sound và haptic có thể truyền đạt giá trị nhưng không được thổi phồng nó.",
)
add_before(
    part_iv,
    "Checklist: Offer hiện tại đang giải nhu cầu nào của cohort nào? Sau một reward, player có biết điều gì đã tốt hơn không? Khi từ chối offer, họ còn lựa chọn đáng tin nào? Rule dynamic có thể được giải thích bằng hành vi player thay vì chỉ bằng conversion không? Một near miss có chỉ ra hành động khác khả thi không? Intensity của feedback có tương xứng với giá trị economy không?",
)
add_before(
    part_iv,
    "Tiếp theo, hãy biến các điều kiện trên thành sản phẩm monetization cụ thể. Câu hỏi không còn là có nên dùng ad hay IAP, mà là mỗi sản phẩm giải nhu cầu nào, ở thời điểm nào, và với guardrail nào.",
)

doc.save(DOCX_PATH)
print(f"Updated: {DOCX_PATH}")
