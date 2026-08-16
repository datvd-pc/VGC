from pathlib import Path
import re
from docx import Document

PATH = Path(r"D:\CODE\VGC\Long form\The-Art-of-Monetization-Vietnamese-Editable.docx")

def replace(document, index, expected, value):
    paragraph = document.paragraphs[index]
    compact = lambda text: re.sub(r"\s+", " ", text).strip()
    if compact(paragraph.text) != compact(expected):
        raise ValueError(f"Paragraph {index} was changed manually and was not replaced.")
    paragraph.clear()
    paragraph.add_run(value)

doc = Document(PATH)
edits = {
    98: (
        "“Chúng tôi tăng ad load” chỉ mô tả điều team đã làm. “Chúng tôi cải thiện monetization” là kết luận cần bằng chứng. Chương tiếp theo đi sâu hơn vào điều kiện của bằng chứng đó: một thiết kế tử tế phải cho người chơi hiểu luật chơi, lựa chọn và hệ quả của việc trả tiền.",
        "“Chúng tôi tăng ad load” chỉ mô tả điều team đã làm. “Chúng tôi cải thiện monetization” là kết luận cần bằng chứng. Tiếp theo, hãy cùng đi sâu vào điều kiện của bằng chứng đó: một thiết kế tử tế phải cho người chơi hiểu luật chơi, lựa chọn và hệ quả của việc trả tiền.",
    ),
    99: (
        "3. Bright design\nThiết kế tử tế cần một cam kết rõ ràng",
        "3. Bright design\nThiết kế tử tế bắt đầu từ sự minh bạch",
    ),
    100: (
        "Puzzle cần uncertainty. Không có uncertainty, nó thành paperwork.",
        "Trong hồ sơ công bố năm 2022, FTC cho biết họ cáo buộc Fortnite có những bố cục nút bấm khiến người chơi có thể bị tính tiền khi chỉ định xem một vật phẩm, khi game đang tải hoặc khi đánh thức game từ chế độ ngủ. FTC cũng cho biết Epic nhận được hơn một triệu khiếu nại về các khoản phí không mong muốn. Case này nói về payment flow, không nói thay cho mọi quyết định thiết kế level. Nhưng nó đặt ra một tiêu chuẩn rất cụ thể: người chơi phải hiểu điều gì sẽ xảy ra trước khi một thao tác trở thành giao dịch.",
    ),
    101: (
        "Thiết kế tử tế làm luật chơi, tỉ lệ, giá và hệ quả dễ hiểu. Người chơi có lựa chọn thật.",
        "Puzzle vẫn cần bất ngờ và rủi ro. Nếu mọi kết quả đều được báo trước, game chỉ còn là một chuỗi thao tác. Sự minh bạch không có nghĩa là loại bỏ thử thách; nó có nghĩa là người chơi có đủ thông tin để hiểu lựa chọn của mình, nhận ra vì sao mình thắng hoặc thua, và biết mình đang mua điều gì.",
    ),
    102: (
        "Dark pattern che luật, tạo nhầm lẫn, giấu lối thoát, hoặc bán sự giải tỏa cho một vấn đề game cố ý tạo ra.",
        "Một thiết kế tử tế làm luật chơi, xác suất, giá bán và hệ quả của lựa chọn trở nên dễ đọc. Dark pattern làm điều ngược lại: che luật, tạo nhầm lẫn, giấu lối thoát hoặc bán sự giải tỏa cho một vấn đề mà game cố ý tạo ra. Khác biệt này không nằm ở việc có payment hay không, mà nằm ở việc payment có phải là một lựa chọn có hiểu biết hay chỉ là con đường duy nhất còn dễ thấy.",
    ),
    103: (
        "Near miss là test tốt. Nó hay khi player thấy một quyết định có thể đổi kết quả. Nó độc hại khi game che intervention, lặp lại loss, rồi đặt payment thành con đường dễ hiểu duy nhất.",
        "Một lần suýt thắng có thể là thử thách tốt khi người chơi nhìn ra quyết định nào của mình có thể thay đổi kết quả. Nó trở thành trải nghiệm độc hại khi game che sự can thiệp, lặp lại thất bại, rồi đặt payment thành lựa chọn duy nhất còn dễ hiểu. Khi review một level, hãy hỏi: sau lần thua này, người chơi có thể chỉ ra một nước đi khác, một công cụ miễn phí hoặc một lý do hợp lý để thử lại không?",
    ),
    104: (
        "FTC từng có hành động với Epic Games về unwanted charges và digital dark patterns. Bài học vượt qua một company: payment design, consent và disclosure là product responsibility, không phải legal decoration.",
        "Với mỗi purchase flow, hãy kiểm tra bốn điểm: người chơi có hiểu giá và thứ nhận được không; có một bước xác nhận phù hợp không; có thể tìm thấy cách hủy hoặc hoàn tiền không; và có thể tiếp tục chơi mà không mua không? FTC không cung cấp công thức cho puzzle game, nhưng case này nhắc team rằng payment design, consent và disclosure là trách nhiệm của sản phẩm. Nguồn: FTC, 2022, “Fortnite Video Game Maker Epic Games to Pay More Than Half a Billion Dollars over FTC Allegations of Privacy Violations and Unwanted Charges.”",
    ),
}
for index, edit in edits.items():
    replace(doc, index, *edit)
doc.save(PATH)
print("Rewrote Chapter 3 and strengthened the Chapter 2 bridge.")
