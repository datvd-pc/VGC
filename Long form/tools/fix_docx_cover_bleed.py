from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.shared import Pt


DOCUMENT_PATH = Path(r"D:\CODE\VGC\Long form\The-Art-of-Monetization-Vietnamese-Editable.docx")
COVER_PATH = Path(r"D:\CODE\VGC\Long form\assets\covers\the-art-of-monetization-cover-option-6-money-lock-final.jpg")


def main():
    document = Document(DOCUMENT_PATH)
    section = document.sections[0]
    section.top_margin = 0
    section.bottom_margin = 0
    section.left_margin = 0
    section.right_margin = 0
    section.header_distance = 0
    section.footer_distance = 0
    section.header.is_linked_to_previous = False
    header = section.header
    header.paragraphs[0].clear()
    header.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
    header.paragraphs[0].paragraph_format.space_before = Pt(0)
    header.paragraphs[0].paragraph_format.space_after = Pt(0)
    header.paragraphs[0].paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    header.paragraphs[0].paragraph_format.line_spacing = Pt(1)
    header.paragraphs[0].add_run().add_picture(
        str(COVER_PATH), width=section.page_width, height=section.page_height
    )

    cover = document.paragraphs[0]
    cover.alignment = WD_ALIGN_PARAGRAPH.LEFT
    cover.paragraph_format.space_before = Pt(0)
    cover.paragraph_format.space_after = Pt(0)
    cover.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    cover.paragraph_format.line_spacing = Pt(1)

    document.save(DOCUMENT_PATH)
    print(f"Updated cover layout in {DOCUMENT_PATH.name}")


if __name__ == "__main__":
    main()
