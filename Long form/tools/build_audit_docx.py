"""
Build script for The-Art-of-Monetization-Audit-Workflow-Checklist.docx
"""

import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

OUTPUT_DIR = r"D:\CODE\VGC\Long form"
DOCX_PATH = os.path.join(OUTPUT_DIR, "The-Art-of-Monetization-Audit-Workflow-Checklist.docx")

def set_cell_shading(cell, color_hex):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def set_cell_margins(cell, top=100, bottom=100, left=120, right=120):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
    tcPr.append(tcMar)

def set_table_borders(table, color="CBD5E1", sz="4", val="single"):
    tblPr = table._tbl.tblPr
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'<w:top w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:bottom w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:left w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:right w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:insideH w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:insideV w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)

def build_docx():
    doc = Document()
    
    section = doc.sections[0]
    section.page_width = Pt(595.3)   # A4 Width
    section.page_height = Pt(841.9)  # A4 Height
    section.top_margin = Pt(45.0)
    section.bottom_margin = Pt(45.0)
    section.left_margin = Pt(45.0)
    section.right_margin = Pt(45.0)
    
    # Styles
    normal = doc.styles['Normal']
    normal.font.name = 'Calibri'
    normal.font.size = Pt(10)
    normal.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)
    normal.paragraph_format.line_spacing = 1.15
    normal.paragraph_format.space_after = Pt(3)

    def add_title(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(text)
        r.font.name = 'Calibri'
        r.font.size = Pt(22)
        r.font.bold = True
        r.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
        return p

    def add_subtitle(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(14)
        r = p.add_run(text)
        r.font.name = 'Calibri'
        r.font.size = Pt(13)
        r.font.italic = True
        r.font.color.rgb = RGBColor(0x47, 0x55, 0x69)
        return p

    def add_h1(text, page_break=False):
        p = doc.add_paragraph()
        if page_break:
            p.paragraph_format.page_break_before = True
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.keep_with_next = True
        r = p.add_run(text)
        r.font.name = 'Calibri'
        r.font.size = Pt(15)
        r.font.bold = True
        r.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
        return p

    def add_h2(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.keep_with_next = True
        r = p.add_run(text)
        r.font.name = 'Calibri'
        r.font.size = Pt(12)
        r.font.bold = True
        r.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)
        return p

    def add_callout(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(8)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        r = p.add_run(text)
        r.font.name = 'Calibri'
        r.font.size = Pt(9.5)
        r.font.italic = True
        r.font.color.rgb = RGBColor(0x33, 0x41, 0x55)
        return p

    def add_body(text, bold_prefix=None):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        if bold_prefix and text.startswith(bold_prefix):
            r = p.add_run(bold_prefix)
            r.font.bold = True
            r.font.size = Pt(9.5)
            r.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
            text = text[len(bold_prefix):]
        r = p.add_run(text)
        r.font.size = Pt(9.5)
        r.font.color.rgb = RGBColor(0x33, 0x41, 0x55)
        return p

    def add_checklist_table(headers, rows):
        t = doc.add_table(rows=len(rows)+1, cols=len(headers))
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        set_table_borders(t, color="CBD5E1")
        
        # Header
        hdr_row = t.rows[0]
        hdr_row._tr.get_or_add_trPr().append(parse_xml(f'<w:tblHeader {nsdecls("w")}/>'))
        for idx, h_text in enumerate(headers):
            cell = hdr_row.cells[idx]
            set_cell_shading(cell, "0F172A")
            set_cell_margins(cell, top=100, bottom=100, left=100, right=100)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(h_text)
            r.font.name = 'Calibri'
            r.font.size = Pt(8.5)
            r.font.bold = True
            r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

        # Rows
        for r_idx, row_data in enumerate(rows):
            row = t.rows[r_idx + 1]
            bg_color = "FFFFFF" if r_idx % 2 == 0 else "F8FAFC"
            for c_idx, cell_value in enumerate(row_data):
                cell = row.cells[c_idx]
                set_cell_shading(cell, bg_color)
                set_cell_margins(cell, top=80, bottom=80, left=90, right=90)
                p = cell.paragraphs[0]
                p.paragraph_format.space_after = Pt(0)
                lines = cell_value.split('\n')
                for l_i, line in enumerate(lines):
                    if l_i > 0:
                        p = cell.add_paragraph()
                        p.paragraph_format.space_after = Pt(0)
                        p.paragraph_format.line_spacing = 1.1
                    r = p.add_run(line)
                    r.font.name = 'Calibri'
                    r.font.size = Pt(8.5)
                    if c_idx == 0:
                        r.font.bold = True
                        r.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
                    elif "Warning" in line or "Red Flag" in line or "•" in line and c_idx == 3:
                        r.font.color.rgb = RGBColor(0x99, 0x1B, 0x1B)
                    elif c_idx == 4:
                        r.font.bold = True
                        r.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)
                    else:
                        r.font.color.rgb = RGBColor(0x33, 0x41, 0x55)
        
        doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # Document Header
    add_title("THE ART OF MONETIZATION")
    add_subtitle("Operational Monetization Audit & QA Checklist")
    add_callout("Use this document beside the live build, source-of-truth dashboard, and store configuration. For every failed check, capture evidence, classify severity, assign one owner and a due date, then retest before closing it. Do not authorize UA scale on a critical trust, consent, purchase, or reward-delivery failure.")

    add_h1("Audit Control Sheet")
    add_body("Scope: [Prototype / Soft Launch / Live Game]     Project: [________________]     Build: [________]     Date: [________]     Audit Lead: [________________]")
    add_body("Evidence standard: a check is PASS only when the auditor has directly verified the player path and, where available, linked supporting cohort telemetry. Use N/A only when the game stage genuinely makes the check inapplicable; write the reason.")
    control_headers = ["Status", "Meaning", "Required action"]
    control_rows = [
        ["PASS", "Observed behavior meets the stated requirement; evidence is recorded.", "Keep evidence link, screenshot, or query reference."],
        ["FAIL – Critical", "Trust, consent, purchase, reward delivery, or legal/platform risk.", "Block release or UA scale; owner and rollback plan required."],
        ["FAIL – High", "Material retention, fairness, economy, or monetization harm.", "Fix before the next scale decision; retest required."],
        ["FAIL – Medium / Low", "Important optimization or presentation issue without immediate player harm.", "Create a prioritized experiment or backlog item."],
        ["N/A", "Not applicable at this product stage.", "State why; review at the next stage gate."]
    ]
    add_checklist_table(control_headers, control_rows)

    add_h2("Finding Log (complete one row for every FAIL)")
    finding_headers = ["Check ID", "Status / Severity", "Evidence", "Owner", "Due date", "Retest / Decision"]
    finding_rows = [
        ["[e.g., 3.2]", "[FAIL – High]", "[screen, video, dashboard link, cohort/date]", "[DRI]", "[YYYY-MM-DD]", "[Open / Pass / Rollback]"]
        for _ in range(5)
    ]
    add_checklist_table(finding_headers, finding_rows)

    # 1. Acquisition to First Return
    add_h1("Stage 1: Acquisition to First Return Checklist")
    headers1 = ["Touchpoint", "Diagnostic Reality Check", "Red Flag Warning Signs", "Paired Metrics", "Immediate Design Fix"]
    rows1 = [
        [
            "[ ] 1.1 Ad Creative Alignment",
            "Does the gameplay in the first 60 seconds match the exact mechanic and emotion promised in your top UA video ad?",
            "• High CTR on ad but immediate drop-off after install.\n• Mismatched mechanics (e.g. sorting ad vs building meta).",
            "Ad CTR ⟷ Store Conversion ⟷ FTUE 3-min Drop-off",
            "Align the very first playable screen to fulfill the ad hook before opening any menus."
        ],
        [
            "[ ] 1.2 Store Page & Permissions",
            "Are system permissions (ATT, notifications) delayed until after the player experiences their first satisfying win?",
            "• ATT prompt firing on splash screen.\n• Review modal popping up before Level 5.",
            "Install Rate ⟷ First Launch Bounce ⟷ ATT Opt-in Rate",
            "Move permission prompts to the end of Level 2/3, right after a victory milestone."
        ],
        [
            "[ ] 1.3 First 10 Levels Pacing",
            "Do early levels fade tutorial scaffolding and teach mastery, or do they overload players with arbitrary obstacles?",
            "• Heavy tutorial text blocking the board.\n• Zero cognitive agency in the first 5 minutes.",
            "Tutorial Completion ⟷ Levels 1–10 Drop-off Velocity",
            "Replace text walls with guided ghosted hints, compact boards, and instant feedback."
        ],
        [
            "[ ] 1.4 Fail State Transparency",
            "When a player loses for the first time, can they pinpoint the exact strategic misplay that caused defeat?",
            "• Level fail feels like a random impossible brick.\n• Instant rage-quit upon defeat screen.",
            "Fail Rate ⟷ Retry Rate ⟷ App Quit Rate on Loss",
            "Ensure every fail state has a visible alternative line of play that would have won freely."
        ],
        [
            "[ ] 1.5 Day-Two Return Hook",
            "At the exact moment of exit, does the player have a clear, unfinished, goal-gradient cliffhanger to return to?",
            "• Exiting on a generic daily reward claimed screen.\n• Zero curiosity about what opens next.",
            "D1 Retention ⟷ Session Count ⟷ Push Notification CTR",
            "End sessions on an unlocked zone, nearly finished album set, or event preview."
        ]
    ]
    add_checklist_table(headers1, rows1)

    # 2. Progression, Pressure & Fairness
    add_h1("Stage 2: Progression, Pressure & Fairness Checklist", page_break=True)
    headers2 = ["Touchpoint", "Diagnostic Reality Check", "Red Flag Warning Signs", "Paired Metrics", "Immediate Design Fix"]
    rows2 = [
        [
            "[ ] 2.1 Meaningful Progress",
            "Does beating a milestone unlock a tangible visual transformation or new capability rather than just bumping a number?",
            "• Levels feel like an endless identical conveyor belt.\n• Player feels zero attachment to game progress.",
            "Milestone Pass Rate ⟷ Session Length ⟷ D7 Retention",
            "Anchor progression to high-impact meta renovations, visual badges, or new mechanics."
        ],
        [
            "[ ] 2.2 Pressure & Near-Misses",
            "Is tension generated through thrilling, close finishes, or through artificial, hopeless bottlenecks?",
            "• Board is blocked on Turn 1 due to bad initial RNG.\n• Defeat feels like developer extortion.",
            "Near-Miss Rate (1 move left) ⟷ Booster Purchases ⟷ Retry Rate",
            "Tune board generators to guarantee at least 2 viable opening strategic combinations."
        ],
        [
            "[ ] 2.3 Dynamic Difficulty (DDA)",
            "Is DDA used strictly as an invisible safety net, or is it secretly rigging losses to force IAP conversion?",
            "• Blatantly artificial win/loss streaks.\n• Players vocalize feeling manipulated in reviews.",
            "Fail-to-Pass Transition Curves ⟷ Store Review Sentiment",
            "Cap DDA assistance to 3 consecutive organic losses; never engineer deliberate defeats."
        ],
        [
            "[ ] 2.4 Input vs Output RNG",
            "Does randomness provide initial tactical inputs (Input RNG) rather than overriding player execution (Output RNG)?",
            "• Lucky piece drops are the sole way to win.\n• Strategic planning is rendered useless by chaos.",
            "Replay Win Distribution ⟷ Player Churn on Hard Levels",
            "Convert output RNG into deterministic reactions so players can forecast outcomes."
        ],
        [
            "[ ] 2.5 Reward Feedback",
            "Are major achievements celebrated with spectacular audio/visual feedback while micro-claims remain snappy?",
            "• Major milestone ends with a flat text prompt.\n• Trivial coin claim locks UI with 10s of unskippable fanfare.",
            "Reward Claim Velocity ⟷ Player Engagement Depth",
            "Create a 2-tier feedback system: <1s snappy feedback for micro-claims; full-screen fanfare for milestones."
        ]
    ]
    add_checklist_table(headers2, rows2)

    # 3. Commercial Touchpoints
    add_h1("Stage 3: Commercial Touchpoints (Ads, IAP, Economy) Checklist", page_break=True)
    headers3 = ["Touchpoint", "Diagnostic Reality Check", "Red Flag Warning Signs", "Paired Metrics", "Immediate Design Fix"]
    rows3 = [
        [
            "[ ] 3.1 Rewarded Video Ad",
            "Does every rewarded video ad solve a specific, voluntary player need with instant, reliable payoff?",
            "• Ad reward is insulting (e.g. 5 coins when level costs 500).\n• Player feels zero incentive to opt in.",
            "Opt-in Rate ⟷ Completion Rate ⟷ Post-Reward D1 Retention",
            "Recalibrate reward value to equal ~10–15% of a paid booster's utility, delivered instantly."
        ],
        [
            "[ ] 3.2 Interstitials",
            "Do interstitials appear exclusively at natural cognitive breakpoints with strict frequency caps?",
            "• Interstitial pops up mid-puzzle or right after defeat.\n• Unskippable ads triggering immediate uninstalls.",
            "IMPDAU ⟷ Session Length ⟷ Immediate Churn Rate",
            "Enforce min 180s cooldown, disable during onboarding, and auto-remove upon any IAP."
        ],
        [
            "[ ] 3.3 Tactical Boosters",
            "Do boosters expand tactical agency, or do they exist purely as paid band-aids for broken level design?",
            "• Level cannot be solved without burning a paid bomb.\n• Booster stock depletes with zero free refills.",
            "Booster Consumption by Level ⟷ Non-Spender Pass Rate",
            "Ensure free boosters trickle through progression tracks; redesign booster-dependent levels."
        ],
        [
            "[ ] 3.4 Contextual IAP Offers",
            "Does every shop bundle solve a concrete, present problem with transparent pricing and an effortless decline path?",
            "• Pushing abstract currency packs with zero context.\n• Tiny hidden 'X' close buttons on offer dialogs.",
            "Offer Conversion Rate ⟷ Non-Buyer Retention ⟷ Refund Rate",
            "Bundle currency with tangible utility (e.g. Starter Pack = No-Ads + 3 Undos + 500 Coins) with clear close buttons."
        ],
        [
            "[ ] 3.5 Economy Stock & Flow",
            "Does every in-game currency have tightly controlled faucets, meaningful sinks, and a stable median balance?",
            "• Massive currency hyperinflation devaluing rewards.\n• Extreme currency starvation choking progress.",
            "Median Wallet Balance ⟷ Sink Consumption Velocity",
            "Introduce recurring cosmetic/meta sinks and adjust event payouts to maintain equilibrium."
        ]
    ]
    add_checklist_table(headers3, rows3)

    # 4. Live Ops, Telemetry & Experiments
    add_h1("Stage 4 & 5: Live Ops, Telemetry & Experimentation Checklist", page_break=True)
    headers4 = ["Touchpoint", "Diagnostic Reality Check", "Red Flag Warning Signs", "Paired Metrics", "Immediate Design Fix"]
    rows4 = [
        [
            "[ ] 4.1 Event Purpose Rotation",
            "Does your event calendar rotate through Teach, Reactivate, Collector, Spend, and Recovery cycles?",
            "• Community fatigue from back-to-back tournaments.\n• Churn spikes following major live events.",
            "Participation Rate ⟷ Post-Event D7 Return Rate",
            "Schedule mandatory 48-hour low-pressure 'Recovery' intervals between intense events."
        ],
        [
            "[ ] 4.2 Paired Metrics Review",
            "Is every positive growth metric balanced against its counter-metric before declaring success?",
            "• Celebrating an ARPDAU jump while D7 retention collapses.\n• High CTR masking cheap unqualified traffic.",
            "ARPDAU ⟷ D7 | IMPDAU ⟷ Churn | IAP Conv ⟷ Refunds",
            "Build automated dashboard views that physically lock paired metrics together."
        ],
        [
            "[ ] 4.3 Diagnostic Trees",
            "When D1 or D7 retention experiences volatility, does the team isolate root causes through causal logic?",
            "• Slashing difficulty across the board because D1 dipped (when UA ad was broken).",
            "Cohort Anomaly Isolation ⟷ Crash Rates ⟷ Creative D1",
            "Execute the 4-step protocol: Observe ➔ Hypothesize ➔ Verify Telemetry ➔ Controlled Fix."
        ],
        [
            "[ ] 4.4 Rigorous A/B Testing",
            "Does every live experiment have 1 falsifiable hypothesis, guardrails, and a pre-committed rollback trigger?",
            "• Testing 5 variables at once in a single variant.\n• Letting tests run indefinitely without decisions.",
            "Sample Statistical Power ⟷ Guardrail Stability",
            "Draft a 1-page Decision Memo specifying exact success thresholds and rollback triggers."
        ]
    ]
    add_checklist_table(headers4, rows4)

    # 5. Monetization Technical QA
    add_h1("Stage 5B: Monetization Technical QA", page_break=True)
    add_callout("Run these checks on every release candidate that changes ads, purchase flow, pricing, entitlements, remote config, or consent. A successful UI review does not replace a verified transaction and reward-delivery test.")
    tech_headers = ["QA Check", "Pass condition", "Failure severity", "Evidence to record", "Owner"]
    tech_rows = [
        ["[ ] T1 Rewarded-ad delivery", "Opt-in is voluntary; completion grants exactly one correct reward, immediately; cancellation, timeout, and network failure do not consume value.", "Critical if reward is lost or duplicated", "Device/video; ad callback; reward ledger", "Ads Eng / QA"],
        ["[ ] T2 IAP purchase, restore & receipt", "Purchase, cancellation, pending payment, restore, and reinstall preserve correct entitlements with no duplicate charge or missing item.", "Critical", "Store sandbox receipt; entitlement record; video", "Commerce Eng / QA"],
        ["[ ] T3 No-ads entitlement", "No-ads removes every promised forced placement across relaunch, session change, and remote-config refresh; rewarded ads remain voluntary.", "Critical if promise is broken", "Before/after placement map; purchase ID", "Ads Eng / QA"],
        ["[ ] T4 Frequency cap & breakpoints", "Interstitial cap holds across level, session, background/resume, and config refresh; no interruption during play, tutorial, or defeat recovery.", "High", "Timestamped session trace; cap config", "Product / QA"],
        ["[ ] T5 Consent, price & localization", "ATT/privacy consent is correctly timed; price, currency, tax disclosure, terms, and close/decline path are readable in every target locale.", "Critical for consent or price error", "Locale screenshots; store product config", "Product / Legal / QA"],
        ["[ ] T6 Remote config & rollback", "A bad offer, ad rule, economy value, or experiment can be disabled safely without app update; change is logged and verified in production-safe test cohort.", "High", "Config revision; rollback test; owner", "Live Ops / Eng"],
    ]
    add_checklist_table(tech_headers, tech_rows)

    # 6. Definition of Done Scorecard
    add_h1("Stage 6: The 10-Point 'Definition of Done' Scale Scorecard", page_break=True)
    add_callout("Scale decision: Minimum score of 9/10 AND zero Critical failures across this scorecard and Technical QA. A missing evidence record is not a pass. Any N/A needs a documented stage rationale and a named re-review date.")
    
    dod_items = [
        ("1. Promise Validation", "The core ad creative hook is proven within the first 3 minutes of gameplay and validated by strong FTUE completion across UA cohorts."),
        ("2. Fair Failure", "Players clearly understand why they failed any given level and always have at least one viable skill-based path forward without spending money."),
        ("3. Voluntary Rewarded Ads", "Rewarded placements are 100% opt-in, grant immediate and reliable rewards, and preserve long-term player goodwill."),
        ("4. Clean Interstitials", "Forced ads appear strictly at natural cognitive breakpoints, enforce strict frequency caps, and automatically vanish upon any IAP purchase."),
        ("5. Contextual & Transparent IAP", "Every shop offer solves an authentic, present in-game need, features crystal-clear pricing, and provides an effortless decline button."),
        ("6. Balanced Stock & Flow", "Currency sources and sinks create meaningful strategic choices rather than coercive bottlenecks; median balances are actively tracked."),
        ("7. Paired Metrics Discipline", "Monetization gains are continuously evaluated alongside retention, app store ratings, refund volume, and customer support tickets."),
        ("8. Controlled Experimentation", "The team possesses the infrastructure to deploy remote config updates, test falsifiable hypotheses, and execute instant rollbacks."),
        ("9. Sustainable Live Ops Pipeline", "The content production machine can reliably satisfy player demand following UA scale, with built-in cool-down recovery cycles."),
        ("10. Positive Contribution Economics", "Unit economics demonstrate true profitability after fully deducting platform fees (30%), ad tech/server infrastructure, UA marketing spend, and live operations overhead.")
    ]

    dod_headers = ["Score Gate", "Criterion & Operational Definition", "Evidence / Owner", "Status"]
    dod_rows = [[title, desc, "[link / cohort / DRI]", "[PASS / FAIL / N/A]"] for title, desc in dod_items]
    add_checklist_table(dod_headers, dod_rows)

    add_h2("Scale Decision Record")
    decision_headers = ["Decision", "Required sign-off", "Reason / evidence", "Next review date"]
    decision_rows = [
        ["[Scale / Hold / Iterate / Kill]", "Product Lead: [ ]  Data Lead: [ ]  Monetization Lead: [ ]  QA Lead: [ ]", "[score, critical-fail status, key cohort evidence]", "[YYYY-MM-DD]"]
    ]
    add_checklist_table(decision_headers, decision_rows)

    doc.save(DOCX_PATH)
    print(f"Audit Checklist DOCX generated at: {DOCX_PATH}")

if __name__ == '__main__':
    build_docx()
