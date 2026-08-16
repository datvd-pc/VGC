from pathlib import Path
import re
from docx import Document

PATH = Path(r"D:\CODE\VGC\Long form\The-Art-of-Monetization-Vietnamese-Editable.docx")

def compact(text):
    return re.sub(r"\s+", " ", text).strip()

def replace(document, index, expected, value):
    paragraph = document.paragraphs[index]
    if compact(paragraph.text) != compact(expected):
        raise ValueError(f"Paragraph {index} was changed manually and was not replaced.")
    paragraph.clear()
    paragraph.add_run(value)

def set_cell(cell, value):
    paragraph = cell.paragraphs[0]
    paragraph.clear()
    paragraph.add_run(value)

doc = Document(PATH)
replace(doc, 105, "4. The operating map\nBản đồ vận hành gồm sáu phần", "4. The operating map\nBản đồ vận hành của một hệ thống kiếm tiền")
replace(doc, 106, "", "Tháng 7/2023, Sensor Tower ước tính Royal Match đạt khoảng 112 triệu USD doanh thu gộp và 14,6 triệu lượt tải; 61,5% lượt tải đến từ kênh trả phí. Bài viết của Sensor Tower còn ghi nhận tốc độ bổ sung khoảng 200 level mỗi tháng, cùng Royal Pass và mini-game Hidden Temple.")
replace(doc, 107, "Royal Match là một case công khai để nhớ rằng vị trí dẫn đầu không đến từ một màn hình cửa hàng. Sensor Tower từng báo cáo game tạo khoảng $112M doanh thu gộp và 14,6M lượt tải trong tháng 7/2023, với lượt cài đặt trả phí chiếm 61,5% tổng lượt tải. Bài học không phải sao chép cấu hình. Đây là một cỗ máy kết nối UA, vòng chơi cốt lõi, năng lực sản xuất nội dung và vận hành game.", "Đây là bằng chứng về quy mô và một vài thành phần có thể quan sát, không phải bản thiết kế nội bộ của Royal Match. Nó không cho biết placement quảng cáo, funnel mua hàng hay economy của game được vận hành ra sao. Điều hữu ích hơn là dùng nó để đặt câu hỏi: một game phải nối những phần nào với nhau để biến traffic, core loop và nội dung mới thành một vòng đời dài hơn?")

anchor = doc.paragraphs[108]
for text in [
    "Operating map dưới đây là một cách trả lời câu hỏi đó. Sáu phần không phải sáu feature độc lập. Chúng là sáu điều kiện liên tiếp: người chơi hiểu lời hứa, thấy mình đang tiến bộ, gặp áp lực có thể giải thích, được mời trao đổi đúng lúc, biết tiền giải quyết vấn đề gì và có lý do để quay lại.",
    "Khi một phần yếu, phần còn lại vẫn có thể tạo doanh thu trong một khoảng thời gian. Nhưng đội ngũ sẽ phải bù bằng nhiều traffic hơn, nhiều nội dung hơn hoặc nhiều áp lực hơn. Bản đồ này giúp nhận ra chỗ đang bị bù trừ trước khi một chỉ số ngắn hạn che khuất nó.",
    "Hãy chọn một game của team và đi qua từng hàng trong bảng. Nếu không trả lời được một câu hỏi bằng hành vi người chơi, dữ liệu hoặc một màn hình cụ thể, đó là phần cần nghiên cứu trước khi thêm feature mới.",
    "Chương tiếp theo bắt đầu ở mắt xích đầu tiên của bản đồ: creative. Trước khi người chơi đánh giá level, reward hay in-game store, họ đã hình thành một kỳ vọng từ quảng cáo.",
]:
    anchor.insert_paragraph_before(text)

table = doc.tables[2]
values = [
    ("Thành phần", "Câu hỏi cần trả lời", "Lỗi thường gặp"),
    ("Lời hứa (Promise)", "Vì sao người chơi muốn bắt đầu?", "Creative và game hứa hẹn hai trải nghiệm khác nhau"),
    ("Tiến bộ (Progress)", "Vì sao thêm một phút chơi còn đáng giá?", "Reward không cho thấy người chơi đã đi xa hơn"),
    ("Áp lực (Pressure)", "Vì sao cần hành động vào lúc này?", "Ma sát tùy tiện hoặc hình phạt quá mức"),
    ("Sự cho phép (Permission)", "Vì sao quảng cáo hoặc offer hợp lý ở thời điểm này?", "Trao đổi làm gián đoạn thay vì hỗ trợ"),
    ("Thanh toán (Payment)", "Tiền giải quyết vấn đề nào?", "In-game store bán currency khi chưa có nhu cầu cụ thể"),
    ("Lý do quay lại (Persistence)", "Vì sao người chơi quay lại ngày mai?", "Nội dung và event không có mục đích rõ"),
]
for row, row_values in zip(table.rows, values):
    for cell, value in zip(row.cells, row_values): set_cell(cell, value)
doc.save(PATH)
print("Rewrote Chapter 4 and its operating map.")
