import docx

doc = docx.Document('The-Art-of-Monetization-Vietnamese-Polished.docx')
print(f"Total tables: {len(doc.tables)}")

for idx, t in enumerate(doc.tables):
    rows = len(t.rows)
    cols = len(t.rows[0].cells) if rows > 0 else 0
    hdr = [c.text.strip().replace('\n', ' ') for c in t.rows[0].cells] if rows > 0 else []
    print(f"Table {idx}: {rows} rows x {cols} cols | Style: {t.style.name} | Header: {hdr[:2]}")
