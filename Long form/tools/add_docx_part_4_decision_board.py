from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from pathlib import Path


DOCX_PATH = Path(r"D:\CODE\VGC\Long form\The-Art-of-Monetization-Vietnamese-Editable.docx")


def set_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = OxmlElement("w:tcMar")
    for side, value in (("top", 100), ("start", 130), ("bottom", 100), ("end", 130)):
        node = OxmlElement(f"w:{side}")
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")
        tc_mar.append(node)
    tc_pr.append(tc_mar)


def write_cell(cell, title, items):
    cell.text = ""
    set_cell_margins(cell)
    title_paragraph = cell.paragraphs[0]
    title_paragraph.paragraph_format.space_after = Pt(4)
    title_run = title_paragraph.add_run(title)
    title_run.bold = True
    title_run.font.size = Pt(9)
    title_run.font.color.rgb = RGBColor(0x00, 0x58, 0x68)
    for item in items:
        paragraph = cell.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.space_after = Pt(2)
        paragraph.paragraph_format.line_spacing = 1.0
        run = paragraph.add_run(item)
        run.font.size = Pt(8.5)


doc = Document(DOCX_PATH)
part_v = next(
    p for p in doc.paragraphs
    if p.text == "Part V: Signals, decisions, and experiments\nPhần V: Tín hiệu, quyết định và thử nghiệm"
)

if any(p.text.startswith("Memory note | Decision board\nPart IV:") for p in doc.paragraphs):
    raise RuntimeError("Part IV Decision board already exists. Refusing to duplicate it.")

heading = part_v.insert_paragraph_before(
    "Memory note | Decision board\nPart IV: Ads, IAP và economy"
)
heading.style = "Heading 3"

table = doc.add_table(rows=2, cols=2)
table.style = "Table Grid"
table.autofit = False
content = {
    "LÀM NGAY": [
        "Lập placement map cho rewarded ad và interstitial: trigger, reward, breakpoint, frequency cap và cohort được thấy.",
        "Viết product brief một trang cho mỗi IAP: need, cohort, nội dung, giá, lựa chọn từ chối, metric chính và guardrail.",
        "Vẽ source-sink map cho một currency và event scorecard tách doanh thu trong event khỏi kết quả sau event.",
    ],
    "HỎI TRƯỚC KHI QUYẾT ĐỊNH": [
        "Rewarded ad đang kéo dài hành động player muốn làm, hay đang sửa một friction do game tạo ra?",
        "Interstitial có nằm sau một loop hoàn thành, và doanh thu thêm có bù được exit hoặc giảm return không?",
        "Player hiểu IAP giải quyết gì, nhận gì và có thể từ chối như thế nào trước khi chạm nút mua không?",
        "Source mới hoặc event reward mới sẽ làm currency balance và sink nào thay đổi?",
    ],
    "CẦN NHỚ": [
        "Opt-in, impression hay conversion cao không tự chứng minh utility, fairness hoặc retention tốt.",
        "Booster là công cụ có job rõ; IAP là lựa chọn có consent rõ; cả hai không được bù cho level mơ hồ.",
        "Event revenue chỉ đáng tin khi cohort quay lại, core economy còn khỏe và spend không bị kéo mượn từ tương lai.",
    ],
    "ĐƯA VÀO CUỘC HỌP": [
        "Thành phần: monetization/product, game design, economy, data và UA. Thời lượng: 60 phút.",
        "Mang theo: replay placement, revenue theo placement/cohort, exit/return sau exposure, refund/review, currency balance và kết quả event trước.",
        "Chốt: placement hoặc product nào giữ/sửa/dừng; một metric chính, guardrail và thời điểm review tiếp theo.",
    ],
}
slots = [
    (0, 0, "LÀM NGAY"),
    (0, 1, "HỎI TRƯỚC KHI QUYẾT ĐỊNH"),
    (1, 0, "CẦN NHỚ"),
    (1, 1, "ĐƯA VÀO CUỘC HỌP"),
]
for row, column, title in slots:
    cell = table.cell(row, column)
    cell.width = Inches(3.15)
    set_shading(cell, "F4F7F8")
    write_cell(cell, title, content[title])

heading._p.addnext(table._tbl)
bridge = doc.add_paragraph(
    "Part V chuyển từ thiết kế sản phẩm sang hệ thống ra quyết định: đọc dashboard nào, ghép chỉ số nào và thử nghiệm thế nào để không gọi một uplift đơn lẻ là câu trả lời."
)
bridge.style = "Normal"
bridge.paragraph_format.space_before = Pt(4)
bridge.paragraph_format.space_after = Pt(8)
bridge.runs[0].italic = True
table._tbl.addnext(bridge._p)

doc.save(DOCX_PATH)
print(f"Updated: {DOCX_PATH}")
