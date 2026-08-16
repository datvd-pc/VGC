from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from pathlib import Path


DOCX_PATH = Path(r"D:\CODE\VGC\Long form\The-Art-of-Monetization-Vietnamese-Editable.docx")


def find_unique(doc, text):
    matches = [paragraph for paragraph in doc.paragraphs if paragraph.text == text]
    if len(matches) != 1:
        raise RuntimeError(f"Expected exactly one unchanged paragraph: {text!r}; found {len(matches)}")
    return matches[0]


def replace_text(doc, old, new):
    paragraph = find_unique(doc, old)
    paragraph.text = new
    return paragraph


def set_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def write_cell(cell, value, bold=False, fill=None):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(2)
    run = paragraph.add_run(value)
    run.font.size = Pt(8.5)
    run.bold = bold
    if bold:
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    if fill:
        set_shading(cell, fill)


def add_table_after(doc, anchor, headers, rows, widths):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.autofit = False
    for column, header in enumerate(headers):
        cell = table.cell(0, column)
        cell.width = Inches(widths[column])
        write_cell(cell, header, bold=True, fill="005868")
    for row_index, row in enumerate(rows, start=1):
        for column, value in enumerate(row):
            cell = table.cell(row_index, column)
            cell.width = Inches(widths[column])
            write_cell(cell, value)
            if row_index % 2 == 0:
                set_shading(cell, "F4F7F8")
    anchor._p.addnext(table._tbl)
    return table


doc = Document(DOCX_PATH)

# Chapter 21
replace_text(
    doc,
    "21. Reading the dashboard\nDashboard là bản đồ của các quyết định người chơi",
    "21. Reading the dashboard\nDashboard phải chỉ ra quyết định nào cần được xem lại, không chỉ hiển thị kết quả",
)
dashboard_intro = replace_text(
    doc,
    "Dashboard hay theo journey và gom signal theo decision mà nó soi sáng.",
    "Mở dashboard với một câu hỏi trước, không phải với một metric trước. 'Lời hứa từ creative có được giữ không?', 'player có hiểu fail đầu không?', 'placement ad có tạo utility không?' là những câu hỏi có thể dẫn đến quyết định. D1, opt-in hay revenue chỉ là tín hiệu giúp thu hẹp nơi cần nhìn. Monetization link của dashboard là tốc độ học: nếu mỗi chart không gắn với owner, dữ liệu cần mở và quyết định có thể thay đổi, team sẽ nhìn nhiều hơn nhưng học không nhanh hơn.",
)
low_d1 = find_unique(doc, "Low D1 có thể do creative mismatch, store, load, crash, tutorial hoặc level 1. Metric cho biết nơi cần nhìn, không tự chẩn đoán.")
theory = low_d1.insert_paragraph_before(
    "Theory anchor là logic của causal inference: một chỉ số quan sát được không tự nói nguyên nhân; để kết luận về một thay đổi, team cần một giả thuyết về cơ chế và một cách so sánh với điều có thể đã xảy ra nếu không thay đổi. Dashboard vì thế là bề mặt để đặt câu hỏi nhân quả, không phải cỗ máy tự chẩn đoán."
)
theory.style = "Normal"
replace_text(
    doc,
    "Low D1 có thể do creative mismatch, store, load, crash, tutorial hoặc level 1. Metric cho biết nơi cần nhìn, không tự chẩn đoán.",
    "Ví dụ, low D1 có thể đến từ creative mismatch, store listing, load, crash, tutorial hoặc level 1. Metric chỉ cho biết nơi cần nhìn tiếp; nó không tự chẩn đoán. Hãy gắn mỗi signal với một replay, một cohort cut hoặc một technical check trước khi mở experiment. Bảng dưới đây là dashboard tối thiểu cho một game puzzle hoặc hybrid puzzle đang tìm product-market fit.",
)
add_table_after(
    doc,
    dashboard_intro,
    ["Câu hỏi", "Signals cần đọc cùng nhau", "Việc cần làm tiếp"],
    [
        ["Promise có được giữ?", "CTR, store CVR, tutorial completion, D1 theo creative", "So creative, store listing và first playable recording theo cohort."],
        ["Fail đầu có công bằng?", "Fail rate, retry, exit, booster use, review language", "Xem replay/board; phân biệt confusion với challenge."],
        ["Ad có utility?", "Opt-in/impression, completion, exit sau exposure, return", "So placement hoặc moment với control; giữ reward cố định."],
        ["IAP có giải need?", "Offer view, conversion, decliner behavior, refund, payer retention", "Đọc product brief, clarity và lựa chọn khi từ chối."],
    ],
    [1.35, 2.45, 2.55],
)

# Chapter 22
replace_text(
    doc,
    "22. Read metric pairs\nĐọc các cặp chỉ số, không đọc số đơn lẻ",
    "22. Read metric pairs\nMột metric chỉ trở nên hữu ích khi được đọc cùng cái giá mà nó tạo ra",
)
metric_intro = replace_text(
    doc,
    "CTR tăng: đọc thêm store CVR, D1 và session depth. RV opt-in tăng: đọc retention by exposure và level friction. IAP conversion tăng: đọc refunds, reviews và payer retention. ARPDAU tăng: đọc D3/D7, churn sau ad và rating. Event revenue tăng: đọc post-event return và resource inflation.",
    "Đọc một metric đơn lẻ thường chỉ trả lời được điều gì vừa xảy ra, không nói game đã phải trả giá gì. Vì vậy, mỗi metric chính cần một cặp hoặc nhóm guardrail đi kèm. Không phải mọi cặp đều phải tăng cùng chiều. Điều cần kiểm tra là phần lợi ích có còn đáng giá sau khi nhìn chi phí về retention, trust, quality hoặc economy hay không.",
)
replace_text(
    doc,
    "Market data cũng vậy. AppMagic H1 2025 cho thấy casual market lớn và concentrated. Category lớn có thể rất thật, nhưng vẫn rất selective. Tách category growth khỏi leader effect trước khi biến chart thành greenlight.",
    "Nguyên tắc này cũng áp dụng cho market data: category growth không phải product thesis, và chart doanh thu của leader không thay thế cohort của team. AppMagic ước tính casual games đạt 12 tỷ USD doanh thu IAP ròng trong H1 2025, trong đó puzzle đạt 4,6 tỷ USD; cùng báo cáo nhấn mạnh thị trường tập trung cao và ít game top mới. Số liệu đủ để nói territory có tiền và cạnh tranh khắc nghiệt. Nó không phải lý do để greenlight một clone hoặc sao chép economy của leader.",
)
add_table_after(
    doc,
    metric_intro,
    ["Metric chính", "Đọc thêm", "Câu hỏi phải trả lời"],
    [
        ["CTR", "store CVR, D1, session depth", "Creative đang thu hút đúng player hay chỉ thu hút click?"],
        ["Rewarded-ad opt-in", "exit/return theo exposure, level friction", "Utility tăng hay normal play đang bị cắt?"],
        ["IAP conversion", "refund, review, payer retention, decliner behavior", "Need có rõ và purchase flow có sạch không?"],
        ["ARPDAU", "D3/D7, ad churn, rating", "Doanh thu tăng có để lại lý do tiếp tục không?"],
        ["Event revenue", "post-event return, currency balance, churn", "Event tạo giá trị mới hay mượn spend từ tương lai?"],
    ],
    [1.35, 2.3, 2.7],
)

# Chapter 23
replace_text(
    doc,
    "23. Decision trees\nCây quyết định giúp team trung thực với dữ liệu",
    "23. Decision trees\nCây quyết định ngăn team nhảy từ một số liệu sang một giải pháp quen tay",
)
replace_text(
    doc,
    "Nếu D1 yếu: so creative promise với first minute, xem store conversion theo creative, check load/crash/device, watch L1-L3 không commentary, defer pop-up và interruption sớm.",
    "Khi D1 yếu, đừng mặc định giảm difficulty hoặc thêm reward. So creative promise với first minute, xem store conversion theo creative, kiểm tra load/crash/device và quan sát L1-L3 không commentary. Nếu player chưa chạm được core action, hãy hoãn mọi pop-up và interruption trước khi thử thay economy. Đó là một nhánh điều tra, không phải một kết luận về nguyên nhân.",
)
replace_text(
    doc,
    "Nếu D1 ổn mà D3 rơi: xem return reason, difficulty spike, ad exposure, novelty session hai, next goal và segment theo source.",
    "Nếu D1 ổn nhưng D3 rơi, mở return reason, difficulty spike, ad exposure, novelty của session hai, mục tiêu kế tiếp và source mix. Nếu retention giảm chỉ ở cohort có exposure lớn, placement đáng nghi hơn core loop. Nếu giảm ở mọi cohort sau một content point, team cần xem lại progression hoặc content. Segment là điều kiện để cây quyết định không biến player trung bình thành một lời giải sai.",
)
replace_text(
    doc,
    "Nếu fail offer convert: so post-offer retention của buyer, decliner và control; review replay, feedback, refund; test level fair hơn trước khi tăng pressure.",
    "Nếu fail offer convert, so post-offer retention của buyer, decliner và control; review replay, feedback và refund; thử level fair hơn trước khi tăng pressure. Nếu in-game store open cao nhưng buy yếu, gọi tên need tại store entry, kiểm tra product clarity, price ladder, region/platform và purchase failure. Test một contextual offer trước khi thêm generic pack. Một decision tree tốt luôn có nhánh 'chưa biết': khi không đủ bằng chứng, việc đúng là mở quan sát hoặc test nhỏ hơn, không phải chọn feature to hơn.",
)
replace_text(
    doc,
    "Nếu store open cao mà buy yếu: gọi tên need lúc store entry, check product clarity, price ladder, region/platform và purchase failure. Test một contextual offer thay vì thêm generic packs.",
    "Output của decision tree phải là một hypothesis có thể bị bác bỏ, một owner và một thời điểm đọc lại. Khi không có đủ bằng chứng để đi tiếp, viết 'chưa biết' là một kết quả tốt hơn giả vờ biết rồi mở rộng spend.",
)

# Chapter 24
replace_text(
    doc,
    "24. Experimentation\nThử nghiệm là một kỷ luật sản xuất",
    "24. Experimentation\nThử nghiệm chỉ có giá trị khi nó giúp team loại bớt một bất định quan trọng",
)
experiment_intro = replace_text(
    doc,
    "A/B test có giá trị khi hạ uncertainty. Nó thành theatre khi test nhiều thay đổi phụ thuộc, ship không guardrail hoặc gọi result quá sớm.",
    "A/B test có giá trị khi nó hạ được một bất định quan trọng. Nó thành theatre khi team thay nhiều thứ phụ thuộc cùng lúc, không định nghĩa guardrail, gọi kết quả quá sớm hoặc không có quyết định đã định trước. Research về risk-aware experiment decision nhấn mạnh việc tách metric thành nhóm success, guardrail, deterioration và quality thay vì để một uplift đơn lẻ quyết định toàn bộ kết luận. Đây là nguyên tắc phương pháp, không phải bảo đảm rằng mọi test có ý nghĩa kinh doanh.",
)
replace_text(
    doc,
    "Problem observed:\nHypothesis:\nChange:\nPrimary metric:\nGuardrail metrics:\nAudience and exclusions:\nDecision window:\nRollback condition:\nOwner:",
    "Problem observed: một cohort cụ thể đang gặp vấn đề gì?\nHypothesis: cơ chế nào có thể giải thích vấn đề?\nChange: một thay đổi nhỏ nhất cần thiết để test cơ chế đó.\nPrimary metric: kết quả nào sẽ tăng hoặc giảm nếu giả thuyết đúng?\nGuardrail metrics: điều gì không được xấu đi?\nAudience and exclusions: ai được vào test, ai không?\nDecision window: khi nào đủ dữ liệu để đọc?\nRollback condition: tín hiệu nào buộc phải dừng?\nOwner: ai chịu trách nhiệm cho kết luận và thao tác rollback?",
)
replace_text(
    doc,
    "Remote config phù hợp với change cần rollback nhanh. Lưu learning repository có cohort, creative context và decision. Result không có context sẽ thành superstition trong vài quarter.",
    "Remote config phù hợp với thay đổi cần rollback nhanh, nhưng không thay thế thiết kế experiment. Lưu learning repository cùng cohort, creative context, bản build, decision và kết quả sau thời gian đủ dài. Một result không có context sẽ thành superstition trong vài quarter. Sau đây là decision board của Part V; hãy dùng nó để biến dashboard và experiment thành nhịp làm việc, không phải báo cáo cuối tuần.",
)
add_table_after(
    doc,
    experiment_intro,
    ["Trước khi ship", "Chỉ được đi tiếp khi"],
    [
        ["Giả thuyết", "Nêu được cơ chế và một thay đổi nhỏ nhất để kiểm tra."],
        ["Metric", "Có primary metric, guardrail và cách đọc theo cohort."],
        ["Quyết định", "Đã viết điều kiện giữ, sửa, rollback hoặc dừng trước khi xem kết quả."],
        ["Trí nhớ team", "Có owner và nơi lưu context để test sau không lặp lại cùng sai lầm."],
    ],
    [1.7, 4.65],
)

part_vi = next(
    p for p in doc.paragraphs
    if p.text == "Part VI: What data can and cannot decide\nPhần VI: Data có thể và không thể quyết định gì"
)
heading = part_vi.insert_paragraph_before(
    "Memory note | Decision board\nPart V: Signals, decisions và experiments"
)
heading.style = "Heading 3"
board = doc.add_table(rows=2, cols=2)
board.style = "Table Grid"
board.autofit = False
board_content = {
    "LÀM NGAY": [
        "Gắn mỗi dashboard panel với một quyết định, owner, cohort cut và hành động quan sát tiếp theo.",
        "Chọn ba metric chính hiện dùng và bổ sung guardrail cho từng metric trước review kế tiếp.",
        "Dùng experiment brief cho thay đổi monetization/difficulty gần nhất; ghi rollback condition trước khi ship.",
    ],
    "HỎI TRƯỚC KHI QUYẾT ĐỊNH": [
        "Metric này đang là bằng chứng của kết quả hay đang bị dùng nhầm như lời giải thích nguyên nhân?",
        "Cohort, source, platform hoặc exposure nào cần được tách trước khi kết luận?",
        "Nếu primary metric tốt nhưng guardrail xấu, team sẽ giữ, sửa hay rollback?",
        "Kết quả nào sẽ khiến giả thuyết hiện tại bị bác bỏ?",
    ],
    "CẦN NHỚ": [
        "Dashboard thu hẹp nơi cần nhìn; replay, technical check và experiment mới giúp kiểm tra cơ chế.",
        "Mỗi uplift có một cái giá cần đọc cùng: retention, trust, quality, economy hoặc chi phí vận hành.",
        "'Chưa biết' là kết quả hợp lệ khi chưa có đủ bằng chứng để scale.",
    ],
    "ĐƯA VÀO CUỘC HỌP": [
        "Thành phần: owner của product/monetization, data, game design và UA khi source mix liên quan. Thời lượng: 45 phút.",
        "Mang theo: decision log, cohort cut, replay/technical evidence, primary metric, guardrail và kết quả control.",
        "Chốt: hypothesis nào tiếp tục, experiment nào ship/rollback và learning nào được ghi vào repository.",
    ],
}
slots = [(0, 0, "LÀM NGAY"), (0, 1, "HỎI TRƯỚC KHI QUYẾT ĐỊNH"), (1, 0, "CẦN NHỚ"), (1, 1, "ĐƯA VÀO CUỘC HỌP")]
for row, column, title in slots:
    cell = board.cell(row, column)
    cell.width = Inches(3.15)
    cell.text = ""
    paragraph = cell.paragraphs[0]
    title_run = paragraph.add_run(title)
    title_run.bold = True
    title_run.font.size = Pt(9)
    title_run.font.color.rgb = RGBColor(0x00, 0x58, 0x68)
    for item in board_content[title]:
        item_paragraph = cell.add_paragraph(style="List Bullet")
        item_paragraph.paragraph_format.space_after = Pt(2)
        item_paragraph.add_run(item).font.size = Pt(8.5)
    set_shading(cell, "F4F7F8")
heading._p.addnext(board._tbl)
bridge = doc.add_paragraph(
    "Part VI sẽ trả lời câu hỏi khó hơn: market intelligence và big data giúp team đi nhanh hơn ở đâu, và ở đâu chúng không thể thay thế product thesis, tay nghề và trách nhiệm ra quyết định."
)
bridge.style = "Normal"
bridge.paragraph_format.space_before = Pt(4)
bridge.paragraph_format.space_after = Pt(8)
bridge.runs[0].italic = True
board._tbl.addnext(bridge._p)

doc.save(DOCX_PATH)
print(f"Updated: {DOCX_PATH}")
