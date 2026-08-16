from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from pathlib import Path


DOCX_PATH = Path(r"D:\CODE\VGC\Long form\The-Art-of-Monetization-Vietnamese-Editable.docx")


def find_unique(doc, text):
    matches = [paragraph for paragraph in doc.paragraphs if paragraph.text == text]
    if len(matches) != 1:
        raise RuntimeError(f"Expected exactly one unchanged paragraph: {text!r}; found {len(matches)}")
    return matches[0]


def replace_text(doc, old, new):
    paragraph = find_unique(doc, old)
    paragraph.text = new
    return paragraph


def remove_text(doc, text):
    paragraph = find_unique(doc, text)
    paragraph._element.getparent().remove(paragraph._element)


def set_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def write_cell(cell, value, header=False):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(2)
    run = paragraph.add_run(value)
    run.font.size = Pt(8.5)
    run.bold = header
    if header:
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)


def add_table_after(doc, anchor, headers, rows, widths):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.autofit = False
    for column, header in enumerate(headers):
        cell = table.cell(0, column)
        cell.width = Inches(widths[column])
        write_cell(cell, header, header=True)
        set_shading(cell, "005868")
    for row_index, row in enumerate(rows, start=1):
        for column, value in enumerate(row):
            cell = table.cell(row_index, column)
            cell.width = Inches(widths[column])
            write_cell(cell, value)
            if row_index % 2 == 0:
                set_shading(cell, "F4F7F8")
    anchor._p.addnext(table._tbl)
    return table


def add_decision_board(doc, before_heading_text, title, content, bridge):
    before = find_unique(doc, before_heading_text)
    heading = before.insert_paragraph_before(title)
    heading.style = "Heading 3"
    board = doc.add_table(rows=2, cols=2)
    board.style = "Table Grid"
    board.autofit = False
    slots = [(0, 0, "LÀM NGAY"), (0, 1, "HỎI TRƯỚC KHI QUYẾT ĐỊNH"), (1, 0, "CẦN NHỚ"), (1, 1, "ĐƯA VÀO CUỘC HỌP")]
    for row, column, label in slots:
        cell = board.cell(row, column)
        cell.width = Inches(3.15)
        cell.text = ""
        paragraph = cell.paragraphs[0]
        title_run = paragraph.add_run(label)
        title_run.bold = True
        title_run.font.size = Pt(9)
        title_run.font.color.rgb = RGBColor(0x00, 0x58, 0x68)
        for item in content[label]:
            bullet = cell.add_paragraph(style="List Bullet")
            bullet.paragraph_format.space_after = Pt(2)
            bullet.add_run(item).font.size = Pt(8.5)
        set_shading(cell, "F4F7F8")
    heading._p.addnext(board._tbl)
    bridge_paragraph = doc.add_paragraph(bridge)
    bridge_paragraph.style = "Normal"
    bridge_paragraph.paragraph_format.space_before = Pt(4)
    bridge_paragraph.paragraph_format.space_after = Pt(8)
    bridge_paragraph.runs[0].italic = True
    board._tbl.addnext(bridge_paragraph._p)


doc = Document(DOCX_PATH)

# Part VII
replace_text(
    doc,
    "30. Monetization as a supply chain\nMonetization là một chuỗi cung ứng",
    "30. Monetization as a supply chain\nMonetization là một chuỗi năng lực: một mắt xích yếu có thể làm hỏng lời hứa ở mắt xích khác",
)
supply_intro = replace_text(
    doc,
    "Player thấy một screen. Studio vận hành một supply chain.",
    "Player thấy một screen; studio vận hành một chuỗi năng lực để screen đó tồn tại đúng lúc, đúng giá trị và đúng chất lượng. Đây là operating model, không phải claim về một studio cụ thể: monetization không thuộc riêng in-game store hay ad mediation, vì mỗi quyết định doanh thu đều phụ thuộc vào trải nghiệm, dữ liệu và khả năng vận hành đứng phía sau.",
)
replace_text(
    doc,
    "Market intelligence định hình bet. Prototype test core loop. Level design produce reason to play. Art, UI, audio và performance quyết định feel trust. Data và mediation quyết định team có thấy và serve player được không. QA, store, policy, localization, support và live ops quyết định product sống nổi với market hay không.",
    "Market intelligence định hình bet. Prototype kiểm tra core loop. Level design tạo lý do để chơi. Art, UI, audio và performance quyết định cảm giác tin cậy. Data và mediation quyết định team có nhìn thấy và phục vụ player được không. QA, in-game store, policy, localization, support và live ops quyết định product có sống được với market hay không. Mỗi phần không cần team lớn ngay từ đầu, nhưng từng phần cần một owner, một chuẩn quality và một cách báo lỗi cho phần còn lại.",
)
replace_text(
    doc,
    "Build chậm làm experiment chậm. Event mỏng làm persistence yếu. QA kém tạo purchase distrust. Creative pipeline yếu làm UA đắt. Data không reliable biến tranh luận thành taste.",
    "Build chậm làm experiment chậm. Event mỏng làm persistence yếu. QA kém tạo purchase distrust. Creative pipeline yếu làm UA đắt. Data không reliable biến tranh luận thành taste. Khi revenue không đạt, đừng mặc định sửa offer; hãy tìm mắt xích đang làm lời hứa bị gãy trước khi player kịp thấy giá trị.",
)
add_table_after(
    doc,
    supply_intro,
    ["Năng lực", "Nếu yếu", "Câu hỏi monetization"],
    [
        ["Prototype + level design", "Player không hiểu value hoặc fail không đọc được.", "Có need thật để ad, booster hay IAP giải quyết không?"],
        ["Creative + store listing", "UA mua traffic lệch lời hứa.", "CPI/CTR đẹp có đi cùng store CVR, D1 và session depth không?"],
        ["Data + remote config", "Team không rollback hoặc không biết cơ chế nào gây kết quả.", "Placement/offer nào có control và guardrail?"],
        ["QA + support + policy", "Purchase trust, rating và refund bị tổn hại.", "Delivery, consent và complaint path có rõ không?"],
        ["Content + live ops", "Return rhythm đứt hoặc event thành rush job.", "Có đủ supply để giữ promise sau UA scale không?"],
    ],
    [1.6, 2.25, 2.45],
)

replace_text(
    doc,
    "31. Kill, iterate, or scale\nNghiên cứu phải kết thúc bằng dừng, lặp lại hoặc mở rộng",
    "31. Kill, iterate, or scale\nNghiên cứu chỉ có giá trị khi kết thúc bằng dừng, lặp lại hoặc mở rộng",
)
replace_text(
    doc,
    "Market chart không phải strategy. Trước khi prototype thành project, viết kill criteria cho promise clarity, activation, early retention, production feasibility, content cost, creative volume và monetization need đầu tiên.",
    "Market chart không phải strategy. Trước khi prototype thành project, viết kill criteria cho promise clarity, activation, early retention, production feasibility, content cost, creative volume và monetization need đầu tiên. Không cần giả vờ biết mọi threshold từ ngày đầu. Cần biết tín hiệu nào, nếu xuất hiện liên tục sau các test hợp lý, sẽ làm thesis không còn đường credible.",
)
replace_text(
    doc,
    "Kill: evidence không còn đường credible.\nIterate: một uncertainty cụ thể test được trong time/cost rõ.\nScale: cohort evidence và operational capacity cho phép tăng spend.",
    "Kill: evidence không còn đường credible trong scope đã chọn.\nIterate: còn một uncertainty cụ thể có thể test trong time và cost rõ.\nScale: cohort evidence và operational capacity cùng cho phép tăng spend.\nKhông chọn: giữ project sống bằng hy vọng, vì chưa ai dám nói giả thuyết nào đã sai.",
)
replace_text(
    doc,
    "Con đường còn lại thường là cách tránh chọn.",
    "Supercell từng mô tả việc dừng Hay Day Pop dù team đã ship nhanh và phản ứng với data. Case này không tạo kill threshold cho studio khác; nó nhắc rằng dừng một bet khi product conviction và bằng chứng không còn gặp nhau là một quyết định vận hành, không phải thất bại của dashboard. Nguồn case: Supercell, 2026.",
)

replace_text(
    doc,
    "32. Live ops with memory\nLive ops là sản xuất nội dung có trí nhớ",
    "32. Live ops with memory\nLive ops phải nhớ player vừa nhận gì, vừa bị yêu cầu gì và còn điều gì đáng quay lại",
)
liveops_intro = replace_text(
    doc,
    "Mỗi event cần một role: teach, reactivate, collector goal, spend moment hoặc recovery. Calendar cần nhớ player vừa thấy, vừa spend, vừa complete gì tuần trước.",
    "Mỗi event cần một role: teach, reactivate, collector goal, spend moment hoặc recovery. Calendar cần nhớ player vừa thấy, vừa spend và vừa complete gì tuần trước. Một event không có role thường trở thành một pop-up có deadline; nó có thể tạo activity nhưng không tạo nhịp quay lại mà product có thể duy trì.",
)
replace_text(
    doc,
    "Content team cần reusable components, economy rules, localization lead time, QA, segmentation, creative support và post-event review. Không có system, event nào cũng thành rush job và result nào cũng khó đọc.",
    "Content team cần reusable components, economy rules, localization lead time, QA, segmentation, creative support và post-event review. Không có system, event nào cũng thành rush job và result nào cũng khó đọc. Hãy review một calendar theo memory: player nào vừa bị pressure mạnh cần recovery; player nào vừa hoàn thành collection cần mục tiêu mới; player nào vừa mua pass cần thấy quyền lợi đã được giao, thay vì nhận thêm một offer chưa liên quan.",
)
add_table_after(
    doc,
    liveops_intro,
    ["Role của event", "Player cần thấy", "Metric/guardrail"],
    [
        ["Teach", "Một luật hoặc tool mới có ích cho core loop.", "Completion, repeat use, confusion/exit."],
        ["Reactivate", "Mục tiêu cũ có lý do mới để quay lại.", "Return, session depth, notification-to-play."],
        ["Collector goal", "Tiến độ và phần còn lại nhìn thấy được.", "Progress completion, repeat return, economy balance."],
        ["Spend moment", "Một lựa chọn tăng tốc/mở rộng nhưng không xóa free path.", "Conversion, decliner retention, refund/review."],
        ["Recovery", "Khoảng thở sau pressure hoặc event dày.", "Churn, sentiment, return tuần sau."],
    ],
    [1.3, 2.7, 2.3],
)

replace_text(
    doc,
    "33. Contribution economics\nHiệu quả kinh tế phải sống được khi mở rộng",
    "33. Contribution economics\nHiệu quả kinh tế chỉ có ý nghĩa khi phần còn lại sống được sau chi phí để giữ game healthy",
)
replace_text(
    doc,
    "Revenue là gross. Business sống bằng phần còn lại.",
    "Revenue là gross. Business sống bằng phần còn lại sau khi trả chi phí để acquire, vận hành, hỗ trợ và giữ lời hứa với player. Đây là contribution-margin accounting, không phải một benchmark chung: cấu trúc chi phí của từng studio và market khác nhau, nhưng mọi decision scale đều phải trả lời được phần còn lại đi đâu.",
)
replace_text(
    doc,
    "Gross revenue\n- platform fees\n- ad-tech and service costs\n- UA spend\n- content, support, vendor cost\n= contribution margin",
    "Gross revenue\n- platform fees\n- ad-tech và service costs\n- UA spend\n- content, support và vendor cost\n= contribution margin\n\nMỗi dòng phải được đọc theo cohort hoặc period có ý nghĩa, không gộp một event tốt với những tuần retention đã rơi.",
)
replace_text(
    doc,
    "Đọc payback period và cash timing theo cohort. Campaign có thể profit trên modeled LTV nhưng tạo funding problem nếu payback quá dài hoặc retention unstable.",
    "Đọc payback period và cash timing theo cohort. Một campaign có thể profit trên modeled LTV nhưng vẫn tạo funding problem nếu payback quá dài, retention unstable hoặc content cost tăng nhanh khi scale. Hãy đặt câu hỏi ngược: nếu giảm ad load để giữ trust, thêm QA để giảm refund hoặc tăng content để giữ return, contribution margin còn sống không? Đây là nơi business discipline bảo vệ product discipline, thay vì ép product lấy doanh thu bằng mọi giá.",
)

add_decision_board(
    doc,
    "Part VIII: Genre playbooks\nPhần VIII: Playbook theo thể loại",
    "Memory note | Decision board\nPart VII: Operating system và contribution economics",
    {
        "LÀM NGAY": [
            "Vẽ capability map: owner, chuẩn quality và failure signal cho prototype, data, content, QA, UA và support.",
            "Viết kill/iterate/scale criteria cho bet hiện tại trước khi mở rộng scope hoặc UA spend.",
            "Làm post-event review có cohort return, currency balance, churn và contribution cost, không chỉ doanh thu.",
        ],
        "HỎI TRƯỚC KHI QUYẾT ĐỊNH": [
            "Mắt xích nào đang làm lời hứa bị gãy trước khi player thấy value?",
            "Thesis còn uncertainty nào test được trong scope rõ, và điều gì sẽ buộc team kill?",
            "Calendar có đang giao value/recovery đúng nhịp hay chỉ tăng demand?",
            "Sau chi phí để giữ game healthy, cohort này còn contribution margin không?",
        ],
        "CẦN NHỚ": [
            "Scale chỉ hợp lý khi evidence và năng lực vận hành cùng tồn tại.",
            "Event là một sản phẩm có role; không có role thì khó đo và dễ làm mòn trust.",
            "Gross revenue không phải lợi nhuận; contribution margin phải tính cả chi phí chất lượng và giữ player.",
        ],
        "ĐƯA VÀO CUỘC HỌP": [
            "Thành phần: studio lead, product, production, UA, data, economy và live ops. Thời lượng: 60 phút.",
            "Mang theo: capability map, roadmap/event calendar, cohort economics, content cost và kill criteria.",
            "Chốt: bet nào kill/iterate/scale; bottleneck nào cần owner; cost nào phải được đưa vào contribution model.",
        ],
    },
    "Part VIII thu hẹp các nguyên tắc trên thành mechanism map theo genre, để team học từ loại cảm xúc và lựa chọn player cần thay vì sao chép một funnel có sẵn.",
)

# Part VIII
replace_text(
    doc,
    "34. The emotional loop\nBắt đầu từ vòng cảm xúc",
    "34. The emotional loop\nBắt đầu từ cảm giác player muốn quay lại tìm, rồi mới chọn sản phẩm monetization",
)
sort_paragraph = find_unique(doc, "Sort bán relief through order. Undo, extra container và remove ads có thể fit. Jam bán controlled panic rồi release; continue fit sau readable near miss. Physics bán curiosity và fast retry; đừng bury loop dưới offer chậm. Match-3 bán mastery trong progression dài; booster, lives, events và deep content calendar cần economy discipline.")
emotional_anchor = sort_paragraph.insert_paragraph_before(
    "Không có một 'puzzle player' duy nhất, nên không có một monetization product mặc định cho toàn bộ genre. Bảng dưới đây là mechanism map: nó mô tả loại cảm giác cần được giữ và những trao đổi có thể hợp lý. Đây không phải dữ liệu chứng minh conversion cho từng genre; team vẫn cần kiểm tra bằng build và cohort của mình."
)
emotional_anchor.style = "Normal"
replace_text(
    doc,
    "Sort bán relief through order. Undo, extra container và remove ads có thể fit. Jam bán controlled panic rồi release; continue fit sau readable near miss. Physics bán curiosity và fast retry; đừng bury loop dưới offer chậm. Match-3 bán mastery trong progression dài; booster, lives, events và deep content calendar cần economy discipline.",
    "Sort thường bán relief through order: undo, extra container và remove ads có thể phù hợp khi chúng bảo vệ sự rõ ràng của plan. Jam thường bán controlled panic rồi release: continue chỉ hợp lý sau near miss readable. Physics bán curiosity và fast retry: đừng chôn loop dưới offer chậm. Match-3 bán mastery trong progression dài: booster, lives, event và content calendar chỉ bền khi economy vẫn giữ giá trị. Điểm chung không phải feature; là cảm giác player đang trả time để nhận.",
)
add_table_after(
    doc,
    emotional_anchor,
    ["Genre/mechanic", "Cảm giác cốt lõi", "Trao đổi cần kiểm tra", "Rủi ro"],
    [
        ["Sort", "Relief qua trật tự", "Undo/extra container sau plan đọc được", "Thiếu ô trống bị tạo để ép tool."],
        ["Jam", "Panic có kiểm soát rồi giải tỏa", "Continue sau near miss có hành động khác", "Pressure không còn giải thích được."],
        ["Physics", "Tò mò, thử nhanh, bất ngờ", "Fast retry hoặc ad utility không ngắt flow", "Offer chậm làm mất nhịp khám phá."],
        ["Match-3", "Mastery và progression dài", "Booster/lives/event phục vụ mục tiêu dài hạn", "Content/economy vượt năng lực vận hành."],
    ],
    [1.1, 1.55, 2.35, 1.3],
)

replace_text(
    doc,
    "35. Evidence, not blueprint\nVí dụ theo thể loại là bằng chứng, không phải bản thiết kế để sao chép",
    "35. Evidence, not blueprint\nVí dụ theo thể loại chỉ giúp chọn câu hỏi; cohort của team mới quyết định cấu hình",
)
replace_text(
    doc,
    "Royal Match, Candy Crush Saga, Merge Mansion và hybrid-casual hit dạy các bài học khác nhau. Không case nào cấp phép clone economy hay funnel từ public observation.",
    "Royal Match, Candy Crush Saga, Merge Mansion và hybrid-casual hit đặt ra những câu hỏi khác nhau về content, progression, board pressure, persistence và UA. Không case nào cấp phép clone economy hay funnel từ public observation. Công việc đúng là nêu cơ chế có thể chuyển giao, giới hạn công khai của case và experiment nhỏ cần chạy trước khi team đưa nó vào roadmap.",
)
replace_text(
    doc,
    "Dùng public data cho scale, market movement, visible cadence và strategic context. Dùng data của mình cho level tuning, pricing, segmentation và causal effect. Public numbers là anchors. Cohort của bạn là decision-maker.",
    "Dùng public data cho scale, market movement, visible cadence và strategic context. Dùng data của mình cho level tuning, pricing, segmentation và causal effect. Public numbers là anchors. Cohort của bạn là decision-maker. Đây là lý do playbook không có 'best price', 'best ad frequency' hay 'best fail rate': những con số tách khỏi product, cohort và cost chỉ tạo cảm giác chắc chắn giả.",
)

add_decision_board(
    doc,
    "Part IX: The audit\nPhần IX: Audit",
    "Memory note | Decision board\nPart VIII: Genre mechanism map",
    {
        "LÀM NGAY": [
            "Chọn một genre/mechanic đang cân nhắc và viết emotional loop theo ngôn ngữ player, không theo feature list.",
            "Chọn một exchange phù hợp với loop đó và test trong build nhỏ trước khi thêm economy rộng.",
            "Ghi rõ rủi ro fairness/flow của mechanic trước khi biến nó thành offer hoặc ad placement.",
        ],
        "HỎI TRƯỚC KHI QUYẾT ĐỊNH": [
            "Player quay lại để tìm cảm giác gì, và product hiện tại có làm mạnh cảm giác đó không?",
            "Exchange này bảo vệ loop hay làm gián đoạn đúng khoảnh khắc loop cần được hoàn thành?",
            "Case công khai đang chứng minh quy mô, cơ chế hay chỉ là một bối cảnh chiến lược?",
            "Cohort nào cần test trước để giả thuyết genre không biến thành stereotype?",
        ],
        "CẦN NHỚ": [
            "Genre là cơ chế và cảm xúc, không phải danh sách feature phải sao chép.",
            "Một placement chỉ fit khi nó giúp player đi sâu hơn vào loop họ đến để tìm.",
            "Không có benchmark công khai nào thay thế được causal effect trong cohort của chính game.",
        ],
        "ĐƯA VÀO CUỘC HỌP": [
            "Thành phần: game design, product, economy, UA creative và data. Thời lượng: 45 phút.",
            "Mang theo: mechanism map, competitor observation, prototype/replay và test plan theo cohort.",
            "Chốt: emotional loop, exchange đầu tiên cần test, risk cần guardrail và điều kiện dừng.",
        ],
    },
    "Part IX là bài kiểm tra cuối: trong ba mươi phút, người đọc phải tìm được một trust leak, một value leak và một experiment có owner/rollback trong một game cụ thể.",
)

# Part IX
replace_text(
    doc,
    "36. The 30-minute audit\nAudit monetization trong 30 phút",
    "36. The 30-minute audit\nAudit monetization trong 30 phút để tìm một trust leak, một value leak và một test có thể ship",
)
audit_anchor = replace_text(
    doc,
    "0-5 phút: xem ba creative. Cảm giác nào chúng bán? First playable minute có prove nó?",
    "Đừng dùng audit để chấm điểm gu. Mục tiêu là tạo một evidence trail ngắn từ lời hứa, trải nghiệm, exchange tới dữ liệu; cuối ba mươi phút phải có một vấn đề cụ thể cần xác minh, không phải mười nhận xét rời rạc.",
)
remove_text(doc, "5-10 phút: chơi first ten levels. Đánh dấu control, success, fail, choice, interruption và booster đầu tiên. Loss có explain được không?")
remove_text(doc, "10-15 phút: tìm rewarded ad và interstitial đầu tiên. Player nhận gì? Có thể decline mà vẫn normal play không?")
remove_text(doc, "15-20 phút: mở store sau natural need. Gọi tên job của từng product. Check delivery và remove-ads scope. Tìm source và sink của mỗi currency quan trọng.")
remove_text(doc, "20-25 phút: tìm return reason và event. Player earn, choose, spend và finish gì?")
remove_text(doc, "25-30 phút: mở cohort dashboard. Đặt revenue cạnh retention, difficulty, exposure, review và refund. Chọn một trust leak và một value leak. Viết một experiment có rollback condition.")
add_table_after(
    doc,
    audit_anchor,
    ["Thời gian", "Làm gì", "Output bắt buộc"],
    [
        ["0-5 phút", "Xem ba creative và first playable minute.", "Promise map: cảm giác quảng cáo bán và bằng chứng game giữ nó."],
        ["5-10 phút", "Chơi L1-L10; đánh dấu control, fail, choice, interruption, booster.", "Một fail có/không đọc được và lựa chọn không trả tiền còn lại."],
        ["10-15 phút", "Tìm rewarded ad/interstitial đầu tiên.", "Exchange map: player nhận gì, có thể decline ra sao, breakpoint có tự nhiên không."],
        ["15-20 phút", "Mở in-game store sau natural need; xem delivery và currency.", "Product brief ngắn: job, clarity, source/sink, remove-ads scope."],
        ["20-25 phút", "Tìm return reason và event.", "Return/event loop: earn, choose, spend, progress, finish và recovery."],
        ["25-30 phút", "Đặt revenue cạnh retention, difficulty, exposure, review/refund.", "Một trust leak, một value leak, một experiment có owner và rollback."],
    ],
    [1.0, 2.85, 2.75],
)

replace_text(
    doc,
    "37. Definition of done\nKhi nào một hệ thống được xem là đủ tốt",
    "37. Definition of done\nKhi nào một hệ thống đủ bằng chứng để được scale, thay vì chỉ đủ feature để được ship",
)
replace_text(
    doc,
    "System sẵn sàng scale hơn khi các câu này đúng:",
    "Một system sẵn sàng scale hơn khi những điều dưới đây đúng trong cohort và có owner vận hành, không chỉ đúng trong một bản demo. Đây là definition of done cho quyết định tăng spend hoặc tăng scope, không phải chứng nhận rằng game sẽ không còn rủi ro.",
)
for old, new in [
    ("Creative promise xuất hiện trong early product experience.", "Creative promise xuất hiện trong early product experience và được kiểm tra theo cohort."),
    ("Player giải thích được meaningful loss và options sau nó.", "Player giải thích được meaningful loss và lựa chọn sau nó, kể cả khi không chi tiền."),
    ("Rewarded ads là voluntary exchange rõ ràng.", "Rewarded ads là voluntary exchange rõ ràng, grant đáng tin và được đọc cùng retention."),
    ("Interstitial tôn trọng rest points và purchase states.", "Interstitial tôn trọng rest points, purchase states và có frequency cap do product quyết định."),
    ("Mỗi offer giải quyết named player need.", "Mỗi offer giải quyết named player need, có consent rõ và có đường từ chối đáng tin."),
    ("Currency source/sink tạo decision, không tạo toll.", "Currency source/sink tạo decision, không tạo toll; balance được theo dõi theo cohort."),
    ("Revenue được đọc cùng retention, reviews, refunds và support.", "Revenue được đọc cùng retention, review, refund, support và contribution cost."),
    ("Team test, rollback và lưu learning ở tốc độ hữu ích.", "Team test, rollback và lưu learning có context ở tốc độ hữu ích."),
    ("Content/live-ops plan phù hợp promise đã mua bằng UA.", "Content/live-ops plan phù hợp với promise đã mua bằng UA và có recovery time."),
    ("Financial model tính cả chi phí giữ game healthy.", "Financial model tính cả chi phí giữ game healthy, payback period và cash timing."),
]:
    replace_text(doc, old, new)

add_decision_board(
    doc,
    "Closing: The player must want to continue\nKết: Người chơi phải muốn tiếp tục",
    "Memory note | Decision board\nPart IX: Audit và definition of done",
    {
        "LÀM NGAY": [
            "Chạy audit 30 phút trên game hiện tại hoặc một competitor; ghi đủ ba output bắt buộc cuối audit.",
            "Chọn một trust leak và một value leak; viết experiment brief với owner, guardrail và rollback.",
            "So definition of done với cohort evidence và capability thật trước khi tăng scope hoặc UA spend.",
        ],
        "HỎI TRƯỚC KHI QUYẾT ĐỊNH": [
            "Lời hứa, core loop, ad/IAP và return reason có đang tạo một evidence trail nhất quán không?",
            "Leak đang là vấn đề trust, value, technical quality hay economy; bằng chứng nào phân biệt chúng?",
            "Team có thể rollback và học đủ nhanh nếu scale tạo ra tác dụng phụ không?",
            "System đã đủ bằng chứng để tăng spend, hay chỉ đủ feature để trông hoàn chỉnh?",
        ],
        "CẦN NHỚ": [
            "Audit tốt kết thúc bằng một quyết định có thể kiểm tra, không phải một danh sách nhận xét.",
            "Definition of done là ngưỡng để scale có trách nhiệm, không phải lời hứa rằng rủi ro đã biến mất.",
            "Một leak được xử lý tốt thường có giá trị hơn nhiều feature mới chưa chứng minh được need.",
        ],
        "ĐƯA VÀO CUỘC HỌP": [
            "Thành phần: product owner, game design, data, UA/monetization, QA và production khi scope thay đổi. Thời lượng: 45 phút.",
            "Mang theo: audit output, replay, cohort dashboard, decision memo, experiment brief và capability/cost impact.",
            "Chốt: một leak ưu tiên, một test có rollback, owner, decision window và tiêu chuẩn kill/iterate/scale.",
        ],
    },
    "Phần kết không thêm framework mới. Nó quay lại một tiêu chuẩn đã xuyên suốt ebook: sau mỗi ad, offer hoặc purchase, player có còn lý do rõ ràng để tiếp tục không?",
)

# Closing
replace_text(
    doc,
    "Closing: The player must want to continue\nKết: Người chơi phải muốn tiếp tục",
    "Closing: The player must want to continue\nKết: Monetization chỉ bền khi người chơi vẫn muốn tiếp tục",
)
replace_text(
    doc,
    "Câu hỏi hữu ích không phải \"Làm sao để player spend?\" Mà là: player vừa earn value gì, đang cảm thấy tension gì, và choice nào họ sẽ xem là fair?",
    "Câu hỏi hữu ích không phải 'Làm sao để player spend?'. Câu hỏi hữu ích hơn là: player vừa nhận được value gì, đang ở trong tension nào, và choice nào họ có thể xem là fair? Khi câu trả lời rõ, monetization product có một việc cụ thể để làm. Khi câu trả lời mơ hồ, thêm placement hoặc pack chỉ làm hệ thống ồn hơn.",
)
replace_text(
    doc,
    "Profit đến từ các câu trả lời được lặp lại trong một system có khả năng giữ lời hứa.",
    "Profit đến từ những câu trả lời được lặp lại trong một system có khả năng giữ lời hứa: từ creative, first session và level, tới ad, IAP, event, dashboard và support. Không một dashboard, publisher brief hay market chart nào thay thế được chuỗi công việc đó. Chúng chỉ giúp team thấy nhanh hơn nơi cần đặt câu hỏi tốt hơn.",
)
replace_text(
    doc,
    "Monetization tốt để lại cho player lý do tiếp tục sau ad, offer và purchase. Game vẫn có thể challenge, đòi thời gian và charge money. Relationship vẫn còn nguyên khi player thấy được value, rule và choice.",
    "Monetization tốt để lại cho player lý do tiếp tục sau ad, offer và purchase. Game vẫn có thể challenge, đòi thời gian và charge money. Mối quan hệ vẫn còn nguyên khi player thấy được value, rule và choice, còn team có đủ kỷ luật để đọc tín hiệu xấu thay vì chỉ giữ lại uplift đẹp.",
)
replace_text(
    doc,
    "Đó là standard đáng để vận hành hướng tới.",
    "Đó là standard đáng để một studio vận hành hướng tới: không ảo tưởng rằng idea một mình tạo ra business, không dùng data để né trách nhiệm phán đoán, và không đổi tương lai của product lấy một kết quả ngắn hạn chưa được kiểm chứng.",
)

doc.save(DOCX_PATH)
print(f"Updated: {DOCX_PATH}")
