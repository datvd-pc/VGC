from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from pathlib import Path


DOCX_PATH = Path(r"D:\CODE\VGC\Long form\The-Art-of-Monetization-Vietnamese-Editable.docx")


def set_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


doc = Document(DOCX_PATH)
part_ii = next(
    p for p in doc.paragraphs
    if p.text == "Part II: From creative to first return\nPhần II: Từ quảng cáo đến lần quay lại đầu tiên"
)
if any(p.text.startswith("Memory note | Decision board\nPart I:") for p in doc.paragraphs):
    raise RuntimeError("Part I Decision board already exists. Refusing to duplicate it.")

heading = part_ii.insert_paragraph_before(
    "Memory note | Decision board\nPart I: Hệ thống phía sau monetization"
)
heading.style = "Heading 3"
board = doc.add_table(rows=2, cols=2)
board.style = "Table Grid"
board.autofit = False
content = {
    "LÀM NGAY": [
        "Vẽ journey từ creative đến first return; đánh dấu mọi touchpoint xin attention, time hoặc money.",
        "Chọn một thay đổi doanh thu gần nhất và lập bảng lợi ích tức thời so với chi phí retention, review, refund và support.",
        "Audit một payment/offer flow bằng bốn điểm: giá trị, giá, xác nhận và lựa chọn không mua.",
    ],
    "HỎI TRƯỚC KHI QUYẾT ĐỊNH": [
        "Player vừa nhận value gì trước touchpoint kiếm tiền này?",
        "Chỉ số nào sẽ cho biết trust được giữ hay bị đánh đổi sau thay đổi?",
        "Pressure đang tạo ra lựa chọn có hiểu biết, hay chỉ làm payment trở thành lối thoát dễ thấy nhất?",
        "Một uplift hiện tại có thể đang được trả bằng chi phí nào ở cohort sau?",
    ],
    "CẦN NHỚ": [
        "Monetization bắt đầu trước in-game store: player trả bằng attention và time trước khi trả money.",
        "Revenue tăng không tự động là healthy revenue; cần đọc phần chi phí đi kèm.",
        "Trust không nằm trong một metric; nó tích lũy từ clarity, fairness, delivery và support xuyên journey.",
    ],
    "ĐƯA VÀO CUỘC HỌP": [
        "Thành phần: product, game design, monetization, data, UA và QA/support khi payment flow liên quan. Thời lượng: 45 phút.",
        "Mang theo: journey map, replay touchpoint, cohort metrics, review/refund và thay đổi revenue gần nhất.",
        "Chốt: một trust leak, một value leak, owner và experiment/rollback cần chạy trước khi thêm feature mới.",
    ],
}
slots = [(0, 0, "LÀM NGAY"), (0, 1, "HỎI TRƯỚC KHI QUYẾT ĐỊNH"), (1, 0, "CẦN NHỚ"), (1, 1, "ĐƯA VÀO CUỘC HỌP")]
for row, column, title in slots:
    cell = board.cell(row, column)
    cell.width = Inches(3.15)
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(title)
    run.bold = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x00, 0x58, 0x68)
    for item in content[title]:
        bullet = cell.add_paragraph(style="List Bullet")
        bullet.paragraph_format.space_after = Pt(2)
        bullet.add_run(item).font.size = Pt(8.5)
    set_shading(cell, "F4F7F8")
heading._p.addnext(board._tbl)
bridge = doc.add_paragraph(
    "Part II đi theo journey này từ creative đến first return, nơi player bắt đầu kiểm tra liệu game có thực sự giữ lời hứa đã khiến họ cài đặt hay không."
)
bridge.style = "Normal"
bridge.paragraph_format.space_before = Pt(4)
bridge.paragraph_format.space_after = Pt(8)
bridge.runs[0].italic = True
board._tbl.addnext(bridge._p)

doc.save(DOCX_PATH)
print(f"Updated: {DOCX_PATH}")
