from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from pathlib import Path


DOCX_PATH = Path(r"D:\CODE\VGC\Long form\The-Art-of-Monetization-Vietnamese-Editable.docx")


def remove_paragraph(paragraph):
    paragraph._element.getparent().remove(paragraph._element)


def set_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=100, start=130, bottom=100, end=130):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def write_cell(cell, title, items):
    cell.text = ""
    set_cell_margins(cell)
    title_paragraph = cell.paragraphs[0]
    title_paragraph.paragraph_format.space_after = Pt(4)
    title_run = title_paragraph.add_run(title)
    title_run.bold = True
    title_run.font.size = Pt(9)
    title_run.font.color.rgb = RGBColor(0x00, 0x58, 0x68)
    for item in items:
        paragraph = cell.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.space_after = Pt(2)
        paragraph.paragraph_format.line_spacing = 1.0
        run = paragraph.add_run(item)
        run.font.size = Pt(8.5)


def insert_board_after(doc, heading, content):
    table = doc.add_table(rows=2, cols=2)
    table.style = "Table Grid"
    table.autofit = False
    for row in table.rows:
        for cell in row.cells:
            cell.width = Inches(3.15)
            set_shading(cell, "F4F7F8")

    slots = [
        (0, 0, "LÀM NGAY"),
        (0, 1, "HỎI TRƯỚC KHI QUYẾT ĐỊNH"),
        (1, 0, "CẦN NHỚ"),
        (1, 1, "ĐƯA VÀO CUỘC HỌP"),
    ]
    for row, column, title in slots:
        write_cell(table.cell(row, column), title, content[title])

    heading._p.addnext(table._tbl)
    return table


def replace_board(doc, old_heading, next_part_heading, new_heading, expected_body, content, bridge):
    heading = next(p for p in doc.paragraphs if p.text == old_heading)
    part_heading = next(p for p in doc.paragraphs if p.text == next_part_heading)

    between = []
    current = heading._p.getnext()
    while current is not None and current is not part_heading._p:
        paragraph = next((p for p in doc.paragraphs if p._p is current), None)
        if paragraph is not None:
            between.append(paragraph)
        current = current.getnext()

    actual_body = [p.text for p in between]
    if actual_body != expected_body:
        raise RuntimeError(
            "Memory note changed since this script was prepared. Refusing to overwrite.\n"
            f"Expected: {expected_body!r}\nActual: {actual_body!r}"
        )

    heading.text = new_heading
    for paragraph in between:
        remove_paragraph(paragraph)
    insert_board_after(doc, heading, content)

    bridge_paragraph = doc.add_paragraph(bridge)
    bridge_paragraph.style = "Normal"
    bridge_paragraph.paragraph_format.space_before = Pt(4)
    bridge_paragraph.paragraph_format.space_after = Pt(8)
    run = bridge_paragraph.runs[0]
    run.italic = True
    heading._p.addnext(bridge_paragraph._p)


doc = Document(DOCX_PATH)

replace_board(
    doc,
    "Memory note | Part II\nGhi nhớ trước khi bước sang Part III",
    "Part III: Progress, pressure, and fairness\nPhần III: Tiến bộ, áp lực và sự công bằng",
    "Memory note | Decision board\nPart II: Từ lời hứa đến lần quay lại đầu",
    [
        "Keywords: creative promise là kỳ vọng được tạo trước khi cài đặt; store listing là nơi lời hứa được kiểm tra trước khi tải; first open là lần chứng minh bằng thao tác thật; cognitive load là phần nguồn lực nhận thức có hạn mà game phải sử dụng có chủ đích; early-game sequence là thứ tự dạy luật, quyết định và hệ quả; return hook là mục tiêu đủ rõ để player có lý do quay lại.",
        "Bài học cần nhớ: creative không chỉ mua lượt cài đặt mà còn định hình chất lượng traffic; mọi màn hình đầu game phải giúp xác thực lời hứa hoặc có lý do rõ ràng để xuất hiện; booster và offer chỉ công bằng khi player hiểu tình huống chúng hỗ trợ; notification không tạo ra nhu cầu quay lại, nó chỉ nhắc về một nhu cầu đã có.",
        "Checklist: Creative, store listing và first playable moment có nói về cùng một trải nghiệm không? Player hiểu booster đầu tiên trước khi được mời dùng hoặc mua nó không? Fail đầu tiên cho họ một cách học hoặc thử lại không cần chi tiền không? Khi thoát phiên đầu, họ nhìn thấy mục tiêu cụ thể nào đang chờ? Team có đang đọc các tín hiệu này theo cohort thay vì chỉ nhìn tổng số không?",
        "Tiếp theo, hãy cùng đi sâu vào chất lượng của tiến bộ, áp lực và sự công bằng, vì đây là nơi game biến những lần quay lại thành giá trị kinh tế có thể duy trì.",
    ],
    {
        "LÀM NGAY": [
            "Tạo Promise map: creative -> store listing -> first playable action; đánh dấu từng chỗ lời hứa đổi hướng.",
            "Lập bảng L1-L10: luật mới, fail đầu, booster/offer/rewarded ad đầu tiên và lựa chọn không trả tiền đi kèm.",
            "Chọn một return hook xuất hiện trước khi kết thúc phiên đầu; log cohort đã nhìn thấy hook đó.",
        ],
        "HỎI TRƯỚC KHI QUYẾT ĐỊNH": [
            "Cohort đến từ creative này có hoàn thành tutorial và quay lại D1 khác cohort khác không?",
            "Player đã tự dùng hoặc hiểu booster trước khi offer xuất hiện chưa?",
            "Prompt này có theo sau một giá trị player vừa nhận, hay chỉ xuất hiện vì app vừa launch?",
            "Khi thoát phiên đầu, mục tiêu cụ thể nào còn lại trên màn hình?",
        ],
        "CẦN NHỚ": [
            "Creative mua lượt cài đặt và đồng thời đặt trần kỳ vọng cho first session, D1 và LTV.",
            "Prompt sớm cạnh tranh với việc player cần hiểu và thử lời hứa cốt lõi.",
            "Notification chỉ nhắc về lý do quay lại; nó không tạo lý do đó.",
        ],
        "ĐƯA VÀO CUỘC HỌP": [
            "Thành phần: growth, product, game design và data. Thời lượng: 45 phút.",
            "Mang theo: creative đang chạy, store listing, bản ghi L1-L10 và funnel theo creative/cohort.",
            "Chốt: một lời hứa cốt lõi, first playable moment, lần exposure monetization đầu và một A/B test ưu tiên.",
        ],
    },
    "Part III kiểm tra chất lượng của tiến bộ, áp lực và fairness: các điều kiện khiến lần quay lại tạo thành giá trị kinh tế có thể duy trì.",
)

replace_board(
    doc,
    "Memory note | Part III\nGhi nhớ trước khi bước sang Part IV",
    "Part IV: Ads, IAP, and economy\nPhần IV: Quảng cáo, IAP và economy",
    "Memory note | Decision board\nPart III: Tiến bộ, áp lực và fairness",
    [
        "Keywords: segmentation là đọc nhu cầu theo cohort thay vì một player trung bình; autonomy là cảm giác vẫn có quyền chọn; competence là cảm giác hiểu và làm chủ thử thách; meaningful progress là thay đổi trạng thái player thực sự quan tâm; pressure là điều kiện tạo quyết định chứ không phải công cụ dồn ép; agency là khả năng học, lập kế hoạch và phản ứng trước kết quả; reward feedback là tín hiệu cho player biết giá trị nào vừa được xác nhận.",
        "Bài học cần nhớ: một offer không thể phù hợp với mọi động lực; currency chỉ có ý nghĩa khi dẫn đến trạng thái hoặc lựa chọn đáng giá; conversion trong bối cảnh pressure chưa chứng minh fairness; dynamic difficulty cần có giới hạn quan sát được và control cohort; randomness chỉ đáng giữ khi player vẫn thấy đường học; sound và haptic có thể truyền đạt giá trị nhưng không được thổi phồng nó.",
        "Checklist: Offer hiện tại đang giải nhu cầu nào của cohort nào? Sau một reward, player có biết điều gì đã tốt hơn không? Khi từ chối offer, họ còn lựa chọn đáng tin nào? Rule dynamic có thể được giải thích bằng hành vi player thay vì chỉ bằng conversion không? Một near miss có chỉ ra hành động khác khả thi không? Intensity của feedback có tương xứng với giá trị economy không?",
        "Tiếp theo, hãy biến các điều kiện trên thành sản phẩm monetization cụ thể. Câu hỏi không còn là có nên dùng ad hay IAP, mà là mỗi sản phẩm giải nhu cầu nào, ở thời điểm nào, và với guardrail nào.",
    ],
    {
        "LÀM NGAY": [
            "Lập needs-by-cohort map cho level hoặc event có offer: nhu cầu, lựa chọn miễn phí, sản phẩm được đề nghị và kết quả sau khi từ chối.",
            "Đánh dấu trong economy loop mỗi reward đã đổi trạng thái player như thế nào, không chỉ đổi số currency.",
            "Lập danh sách rule dynamic và random seed có thể ảnh hưởng outcome; xác định control cohort cho từng rule.",
        ],
        "HỎI TRƯỚC KHI QUYẾT ĐỊNH": [
            "Offer này giải nhu cầu nào của cohort nào, và player có thể mô tả nhu cầu đó không?",
            "Sau khi từ chối offer, player còn một lựa chọn đáng tin để retry, học hoặc quay lại không?",
            "Rule dynamic đã can thiệp tại đâu và team đọc kết quả so với control bằng metric nào?",
            "Near miss này có chỉ ra hành động khác khả thi, hay chỉ gợi ý chiến thắng đang rất gần?",
        ],
        "CẦN NHỚ": [
            "Conversion ở điểm có pressure chưa chứng minh pressure fair hoặc doanh thu bền vững.",
            "Nhiều dữ liệu hay thuật toán phức tạp hơn không tự động tạo trải nghiệm tốt hơn.",
            "Feedback phải tương xứng với giá trị thật; thổi phồng mọi reward sẽ làm mờ thang giá trị của economy.",
        ],
        "ĐƯA VÀO CUỘC HỌP": [
            "Thành phần: product, game design, economy, data và UA/monetization. Thời lượng: 60 phút.",
            "Mang theo: replay của fail/near miss, kết quả offer theo cohort, log rule dynamic, retention sau offer và review/refund nếu có.",
            "Chốt: giữ, sửa hoặc dừng từng rule gây pressure; metric chính, guardrail và ngày đọc lại kết quả.",
        ],
    },
    "Part IV biến các điều kiện này thành sản phẩm cụ thể: rewarded ads, interstitials, boosters, IAP và economy cần được thiết kế như những trao đổi có lý do.",
)

doc.save(DOCX_PATH)
print(f"Updated: {DOCX_PATH}")
