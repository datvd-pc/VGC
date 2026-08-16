# -*- coding: utf-8 -*-
import io
import sys
import docx

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

def audit_document():
    doc = docx.Document(r"D:\CODE\VGC\Long form\The-Art-of-Monetization-Vietnamese-Polished.docx")
    
    print(f"Total paragraphs: {len(doc.paragraphs)}")
    print(f"Total tables: {len(doc.tables)}")
    
    # 1. Check every paragraph
    findings = []
    
    for i, p in enumerate(doc.paragraphs):
        txt = p.text.strip()
        style = p.style.name if p.style else "Normal"
        
        # Check empty paragraphs
        if not txt:
            findings.append((i, style, "EMPTY", "Đoạn văn trống hoàn toàn."))
            continue
            
        # Check for awkward fragments, untranslated English words
        untranslated = []
        for word in ["player", "spend", "uplift", "theory anchor", "monetization link", "worked example", "opt-in", "rest points", "purchase states", "grant đáng tin", "near miss", "frequency cap"]:
            if word in txt.lower():
                untranslated.append(word)
        if untranslated:
            findings.append((i, style, "UNTRANSLATED", f"Chứa thuật ngữ thô chưa mượt: {untranslated}"))
            
        # Check for punctuation, sentence truncation
        lines = [l.strip() for l in txt.split("\n") if l.strip()]
        for l in lines:
            if len(l) > 50 and not l.endswith(('.', ':', '?', '!', '"', "'", '…', ')', ';', '-', '—')):
                if not any(k in l for k in ["Heading", "Part", "Phần", "Chương", "THE ART", "Contents", "Ghi chú"]):
                    findings.append((i, style, "PUNCTUATION", f"Câu dài không có dấu kết thúc: '{l[-40:]}'"))

    # 2. Check every table
    table_findings = []
    for t_idx, tbl in enumerate(doc.tables):
        for r_idx, row in enumerate(tbl.rows):
            for c_idx, cell in enumerate(row.cells):
                c_txt = cell.text.strip()
                if not c_txt:
                    table_findings.append((t_idx, r_idx, c_idx, "EMPTY_CELL", "Ô bảng trống."))
                # Check for untranslated terms
                untranslated = []
                for word in ["player", "spend", "uplift", "theory anchor", "monetization link", "relief through order", "exchange map", "product brief"]:
                    if word in c_txt.lower():
                        table_findings.append((t_idx, r_idx, c_idx, "TABLE_UNTRANSLATED", f"Ô bảng chứa từ thô: {untranslated}"))

    print("\n=== PARAGRAPH FINDINGS ===")
    for f in findings:
        print(f"P[{f[0]}] ({f[1]}) [{f[2]}]: {f[3]}")
        p_txt = doc.paragraphs[f[0]].text.strip()
        if p_txt:
            print(f"   Text: {p_txt[:140]}...\n")

    print("\n=== TABLE FINDINGS ===")
    for tf in table_findings:
        print(f"Table[{tf[0]}] Row[{tf[1]}] Col[{tf[2]}] [{tf[3]}]: {tf[4]}")
        c_txt = doc.tables[tf[0]].rows[tf[1]].cells[tf[2]].text.strip()
        print(f"   Cell: {c_txt[:120]}...\n")

if __name__ == "__main__":
    audit_document()
