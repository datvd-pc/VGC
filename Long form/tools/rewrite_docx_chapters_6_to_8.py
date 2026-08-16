from docx import Document
from pathlib import Path


DOCX_PATH = Path(r"D:\CODE\VGC\Long form\The-Art-of-Monetization-Vietnamese-Editable.docx")


def replace_exact(paragraph, expected, replacement):
    if paragraph.text != expected:
        raise RuntimeError(
            "Paragraph changed since this script was prepared. Refusing to overwrite.\n"
            f"Expected: {expected!r}\nActual: {paragraph.text!r}"
        )
    paragraph.text = replacement


def add_before(anchor, text, style):
    paragraph = anchor.insert_paragraph_before(text)
    paragraph.style = style
    return paragraph


doc = Document(DOCX_PATH)
paragraphs = doc.paragraphs

replace_exact(
    paragraphs[132],
    "6. Store and first open\nCửa hàng và lần mở game đầu phải chứng minh cùng một lời hứa",
    "6. Store listing and first open\nStore listing và lần mở game đầu phải xác thực cùng một lời hứa",
)
replace_exact(
    paragraphs[133],
    "Store là verification surface. First open là verdict.",
    "Hãy lấy một creative đang mang về nhiều lượt cài đặt, rồi xem hành trình của cohort đó qua ba điểm: tỷ lệ từ ad sang store listing, tỷ lệ cài đặt, và những hành động đầu tiên sau khi mở game. Khi một creative về sorting puzzle nhưng store listing lại ưu tiên một metagame khác, người chơi đã phải tự ghép hai lời hứa trước cả khi chạm vào level đầu. Khi lần mở game đầu lại bắt đầu bằng login, thông báo, xin quyền và menu, câu hỏi monetization không còn là prompt nào chuyển đổi tốt hơn. Câu hỏi là: người chơi có kịp thấy giá trị đã khiến họ cài game trước khi team bắt đầu xin thêm một sự cho phép hay không?",
)
replace_exact(
    paragraphs[134],
    "Nếu ad hứa rescue puzzle, cho player chạm vào rescue puzzle trước menu chung. Nếu ad hứa sort, cho player sort trước notification, tracking, login, rating hay special offer. Mỗi interruption sớm đòi trust khi game chưa earn quyền để xin.",
    "Theory anchor ở đây là Cognitive Load Theory của John Sweller. Lý thuyết này phân biệt phần tải cần thiết để hiểu một nhiệm vụ với phần tải phát sinh từ cách trình bày không phục vụ nhiệm vụ đó. Các nghiên cứu về worked examples cho thấy người mới học xử lý bài toán phức tạp tốt hơn khi được hướng sự chú ý vào cấu trúc cần học. Điều này không chứng minh một privacy prompt sẽ làm D1 của puzzle game giảm bao nhiêu. Nhưng nó cho một giả thuyết rõ để kiểm tra: trong những phút đầu, mọi màn hình không giúp player hiểu và thử lời hứa cốt lõi đều cạnh tranh với nguồn lực nhận thức dành cho game.",
)
replace_exact(
    paragraphs[135],
    "Cho thấy ngay cảm giác mà quảng cáo đã hứa.\nTrao quyền điều khiển thật nhanh.\nTạo một chiến thắng nhỏ, dễ hiểu.\nMở ra mục tiêu kế tiếp.\nGiới thiệu một lựa chọn hữu ích.\nChỉ xin quyền khi người chơi đã thấy giá trị.",
    "Vì vậy, hãy viết hành trình first open như một chuỗi bằng chứng. Nếu creative hứa rescue puzzle, thao tác đầu tiên nên cho player giải một tình huống rescue thật, trước menu chung. Nếu creative hứa sorting, hãy để họ sắp xếp vài vật thể, thấy vì sao nước đi đó đúng, rồi mới mở ra mục tiêu kế tiếp. Sau chiến thắng nhỏ, team mới có cơ sở để giới thiệu một lựa chọn hữu ích, như rewarded ad để nhận thêm lượt hoặc một gói hỗ trợ có ngữ cảnh. Đừng mặc định bỏ mọi prompt sớm; hãy thử thứ tự và thời điểm khác nhau, sau đó đọc cùng lúc completion của tutorial, D1, tỷ lệ chấp nhận quyền, lần xem rewarded ad đầu tiên và đánh giá của cohort.",
)
replace_exact(
    paragraphs[136],
    "ATT hay bất kỳ privacy prompt nào cũng là trust moment. Prompt đúng lúc player hiểu value khác hoàn toàn với một nút bấm lúc launch.",
    "Một ATT hoặc privacy prompt vẫn có thể cần thiết. Khác biệt nằm ở việc player có hiểu yêu cầu đó liên quan thế nào đến giá trị họ vừa nhận hay không. Hãy coi store listing là lời hứa, first open là lần chứng minh đầu tiên, và các prompt là những yêu cầu chỉ được đưa ra sau khi có đủ lý do. Nguồn lý thuyết: Sweller, 1988; Chen, Retnowati và Kalyuga, 2020 về worked example và cognitive load. Sau khi lời hứa được xác thực, mười level đầu sẽ trả lời một câu hỏi khó hơn: game có dạy người chơi cách thắng, thất bại và sử dụng lựa chọn có trả phí một cách công bằng không?",
)

replace_exact(
    paragraphs[137],
    "7. The first ten levels\nMười level đầu là một cam kết cảm xúc",
    "7. The first ten levels\nMười level đầu dạy cách đưa ra lựa chọn, trước khi dạy cách chi tiền",
)
replace_exact(
    paragraphs[138],
    "Mười level đầu dạy luật và dạy \"đạo đức\" của game. Người chơi học xem một lần thua có thể hiểu được không; booster là một lựa chọn hay bộ vá cho level lỗi; phần thưởng có trọng lượng không; game có tôn trọng thời gian giữa các lượt chơi không.",
    "Hãy lấy mười level đầu của build hiện tại và đặt cạnh event log của chúng. Với mỗi level, player đã học một luật mới, được trao một quyết định mới, hay chỉ gặp thêm vật cản? Lần đầu họ nhận booster diễn ra trước hay sau khi họ hiểu nó thay đổi thế cờ ra sao? Lần đầu xuất hiện offer có bám vào một tình huống họ vừa trải qua không? Những câu hỏi này là monetization link của early game: một lựa chọn có trả phí chỉ có thể được đánh giá là hữu ích khi player đã hiểu vấn đề mà lựa chọn đó giải quyết. Nếu không, conversion có thể chỉ phản ánh áp lực, còn chi phí sẽ hiện ở retry, exit, đánh giá và mức sẵn sàng quay lại.",
)
replace_exact(
    paragraphs[139],
    "Mỗi early level cần trả lời: player đang học gì, quyết định nào tạo success/failure, win cần cảm thấy thế nào, sau fail có gì không cần spend, booster là choice hay rescue.",
    "Cognitive Load Theory đưa ra một nguyên tắc có ích cho sequence này: người mới cần đủ hướng dẫn để hình thành cấu trúc của nhiệm vụ, sau đó mức hỗ trợ phải giảm dần để họ tự sử dụng cấu trúc đó. Trong bối cảnh puzzle, hướng dẫn không nhất thiết là một lớp tutorial bằng chữ. Nó có thể là một board nhỏ, mục tiêu nhìn thấy được, một nước đi mẫu và kết quả phản hồi đúng vào quyết định vừa đưa ra. Nghiên cứu này không đo booster conversion hay LTV trong game. Phần có thể chuyển sang design là cách đặt câu hỏi: level đang làm rõ một quy tắc hay đang buộc player xử lý nhiều thứ trước khi họ có mô hình để hiểu chúng?",
)
replace_exact(
    paragraphs[140],
    "Không cần fail rate thấp ở mọi nơi. High fail + high retry có thể là challenge tốt. High fail + low retry thường là confusion, exhaustion hoặc suspicion.",
    "Một sequence đầu game có thể đi từ ví dụ được dẫn dắt, sang bài toán tương tự có ít trợ giúp hơn, rồi mới tạo thất bại đầu tiên có thể giải thích được. Với từng level, ghi bốn cột trong design review: player đang học gì; quyết định nào tạo thắng hoặc thua; họ có phương án nào không cần chi tiền sau khi thua; và booster đang mở thêm lựa chọn hay đang sửa một điều không thể hiểu. Đọc fail rate cùng retry rate, next-level start, booster use, rewarded-ad uptake và exit rate. High fail đi cùng high retry có thể là thử thách hấp dẫn; high fail đi cùng exit chỉ là tín hiệu cần điều tra, không phải bản án về level. Chương tiếp theo chuyển từ việc dạy lựa chọn sang điều kiện để player tự nguyện mở lại game: một mục tiêu còn đủ sức kéo họ về.",
)

replace_exact(
    paragraphs[141],
    "8. The first return\nLần quay lại đầu cần lý do, không chỉ cần lời nhắc",
    "8. The first return\nLần quay lại đầu cần một mục tiêu dang dở có ý nghĩa, không chỉ một lời nhắc",
)
replace_exact(
    paragraphs[142],
    "Notification chỉ nhắc player về lý do quay lại. Nó không tự tạo lý do.",
    "Trước khi viết notification đầu tiên, hãy mở danh sách player rời game sau phiên đầu và hỏi một câu đơn giản: tại khoảnh khắc thoát, họ đang còn muốn hoàn thành điều gì? Nếu team không trả lời được bằng một mục tiêu nhìn thấy trong game, notification chỉ là lời gọi quay lại một trải nghiệm chưa kịp có lý do. Đây là monetization link của first return: retention sau phiên đầu mở ra những lần tiếp xúc sau này với reward, ad, offer và content; nó không thể được thay bằng một thông điệp gửi đúng giờ.",
)
replace_exact(
    paragraphs[143],
    "Return hook có thể là board mới, construction timer, collection completion, daily goal, event milestone hay puzzle còn dang dở. Chọn một hook sớm và để player rời game với một việc dang dở nhỏ.",
    "Một theory anchor phù hợp là goal-gradient hypothesis. Hull đề xuất từ năm 1932 rằng nỗ lực có thể tăng khi chủ thể tiến gần mục tiêu; Kivetz, Urminsky và Zheng sau đó tìm thấy trong các chương trình phần thưởng của người tiêu dùng rằng cảm giác tiến gần mốc thưởng đi kèm với việc quay lại và duy trì nỗ lực nhanh hơn. Đây không phải bằng chứng rằng mọi daily goal hay timer trong casual game sẽ giữ được D1. Nó gợi một cơ chế để kiểm tra: mục tiêu có đủ cụ thể, tiến độ có nhìn thấy được, và phần còn lại có đáng để quay lại hoàn tất không?",
)
replace_exact(
    paragraphs[144],
    "Energy chỉ để ngăn chơi thường yếu. Energy mạnh hơn khi điều tiết một rhythm lớn hơn: event attempts, meaningful decisions, social cadence hoặc long-form progression.",
    "Một return hook tốt không nhất thiết phải là energy, daily goal hay timer. Nó có thể là một board vừa mở khóa, một bộ sưu tập còn thiếu một mảnh, một construction step đang chờ, hoặc một puzzle mà player đã hiểu cách giải nhưng chưa đủ lượt để hoàn thành. Điểm chung là player biết mình sẽ quay lại để làm gì. Vì thế, hãy thử hai cách kết thúc first session: một cohort rời game sau reward chung chung; cohort còn lại rời game khi nhìn thấy tiến độ và mục tiêu kế tiếp. So sánh D1, số phiên trong ngày đầu, thời điểm quay lại và phản ứng với rewarded ad hoặc offer đầu tiên. Kết quả là dữ liệu của game, không phải suy đoán từ lý thuyết.",
)
replace_exact(
    paragraphs[145],
    "Merge Mansion là ví dụ công khai về persistence như operational capability. AppMagic-related analysis từng nói về recovery và event cadence dày hơn. Bài học không phải mọi game cần nhiều event. Live ops là calendar, economy, narrative, segmentation và measurement.",
    "Energy chỉ đáng giữ khi nó điều tiết một nhịp có giá trị hơn là đơn thuần chặn lượt chơi: chẳng hạn số lần thử trong event, một quyết định có đánh đổi, hoặc tiến độ dài hạn mà player muốn quay lại theo dõi. Sau first return, vấn đề không còn là có bao nhiêu lý do để quay lại, mà là mỗi lần quay lại có làm tình huống của player thực sự tiến lên không. Nguồn lý thuyết: Hull, 1932; Kivetz, Urminsky và Zheng, 2006. Part tiếp theo bắt đầu từ câu hỏi đó: thế nào là tiến bộ có ý nghĩa, và khi nào áp lực vẫn còn công bằng?",
)

part_iii = next(
    p for p in doc.paragraphs
    if p.text == "Part III: Progress, pressure, and fairness\nPhần III: Tiến bộ, áp lực và sự công bằng"
)

add_before(part_iii, "Memory note | Part II\nGhi nhớ trước khi bước sang Part III", "Heading 3")
add_before(
    part_iii,
    "Keywords: creative promise là kỳ vọng được tạo trước khi cài đặt; store listing là nơi lời hứa được kiểm tra trước khi tải; first open là lần chứng minh bằng thao tác thật; cognitive load là phần nguồn lực nhận thức có hạn mà game phải sử dụng có chủ đích; early-game sequence là thứ tự dạy luật, quyết định và hệ quả; return hook là mục tiêu đủ rõ để player có lý do quay lại.",
    "Normal",
)
add_before(
    part_iii,
    "Bài học cần nhớ: creative không chỉ mua lượt cài đặt mà còn định hình chất lượng traffic; mọi màn hình đầu game phải giúp xác thực lời hứa hoặc có lý do rõ ràng để xuất hiện; booster và offer chỉ công bằng khi player hiểu tình huống chúng hỗ trợ; notification không tạo ra nhu cầu quay lại, nó chỉ nhắc về một nhu cầu đã có.",
    "Normal",
)
add_before(
    part_iii,
    "Checklist: Creative, store listing và first playable moment có nói về cùng một trải nghiệm không? Player hiểu booster đầu tiên trước khi được mời dùng hoặc mua nó không? Fail đầu tiên cho họ một cách học hoặc thử lại không cần chi tiền không? Khi thoát phiên đầu, họ nhìn thấy mục tiêu cụ thể nào đang chờ? Team có đang đọc các tín hiệu này theo cohort thay vì chỉ nhìn tổng số không?",
    "Normal",
)
add_before(
    part_iii,
    "Tiếp theo, hãy cùng đi sâu vào chất lượng của tiến bộ, áp lực và sự công bằng, vì đây là nơi game biến những lần quay lại thành giá trị kinh tế có thể duy trì.",
    "Normal",
)

doc.save(DOCX_PATH)
print(f"Updated: {DOCX_PATH}")
