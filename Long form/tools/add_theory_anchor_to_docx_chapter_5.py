from pathlib import Path
import re
from docx import Document

PATH = Path(r"D:\CODE\VGC\Long form\The-Art-of-Monetization-Vietnamese-Editable.docx")

def replace(document, index, expected, value):
    paragraph = document.paragraphs[index]
    compact = lambda text: re.sub(r"\s+", " ", text).strip()
    if compact(paragraph.text) != compact(expected):
        raise ValueError(f"Paragraph {index} was changed manually and was not replaced.")
    paragraph.clear()
    paragraph.add_run(value)

doc = Document(PATH)
replace(
    doc, 128,
    "Hãy chọn creative đã mang về nhiều lượt cài đặt nhất trong chiến dịch gần đây. Người chơi nhìn thấy tình huống gì trong ba giây đầu? Họ chờ được giải tỏa, được lập lại trật tự hay được chứng kiến một cú lội ngược dòng? Sau khi cài đặt, phiên chơi đầu có đưa họ đến cảm giác đó đủ nhanh không? Đây là monetization link của creative: lời hứa ban đầu quyết định chất lượng traffic, D1 và trần LTV mà team có thể tạo ra sau này.",
    "Hãy chọn creative đã mang về nhiều lượt cài đặt nhất trong chiến dịch gần đây. Người chơi nhìn thấy tình huống gì trong ba giây đầu? Họ chờ được giải tỏa, được lập lại trật tự hay được chứng kiến một cú lội ngược dòng? Sau khi cài đặt, phiên chơi đầu có đưa họ đến cảm giác đó đủ nhanh không? Đây là monetization link của creative: lời hứa ban đầu định hình chất lượng traffic, D1 và trần LTV mà team có thể tạo ra sau này.",
)
replace(
    doc, 129,
    "Một creative tốt không chỉ liệt kê feature. Nó gọi đúng công việc cảm xúc mà game làm cho người chơi. Puzzle sắp xếp có thể hứa khoảnh khắc mọi thứ trở về đúng chỗ. Match-3 có thể hứa chuỗi phản ứng và cảm giác tiến bộ qua nhiều level. Khi creative hứa một game khác, team có thể vẫn thấy CTR hoặc CPI đẹp hơn trong ngắn hạn. Phần chi phí thường xuất hiện sau đó: tỷ lệ chuyển đổi ở store listing yếu, D1 thấp, phiên chơi ngắn và đánh giá cho rằng quảng cáo không trung thực.",
    "Nền tảng của cách đọc này là Expectation-Disconfirmation Theory, do Richard Oliver phát triển trong nghiên cứu hành vi người tiêu dùng. Lý thuyết cho rằng mức hài lòng hình thành khi người dùng so sánh kỳ vọng trước khi dùng với trải nghiệm thực tế: trải nghiệm thấp hơn kỳ vọng tạo ra negative disconfirmation, còn trải nghiệm đáp ứng hoặc vượt kỳ vọng tạo ra satisfaction. Lý thuyết không chứng minh riêng cho D1 hay LTV của một puzzle game. Nó cung cấp cơ chế để team kiểm tra giả thuyết: creative có đang tạo ra một kỳ vọng mà first session không thể giữ được hay không?",
)
replace(
    doc, 130,
    "Hãy xây creative library theo từng tình huống, không chỉ theo asset. Với mỗi creative, lưu lại lời hứa, xung đột, khoảnh khắc giải tỏa, nhóm người chơi hướng tới, bằng chứng trong phiên chơi đầu và kết quả theo cohort. AI có thể giúp team tạo nhiều biến thể nhanh hơn. Quyết định khó hơn vẫn thuộc về team: lời hứa này có trung thực với product không, và có đưa về nhóm người chơi mà economy của game có thể phục vụ không? Chương tiếp theo kiểm tra nơi người chơi bắt đầu xác minh lời hứa đó: store listing và lần mở game đầu.",
    "Vì vậy, creative tốt không chỉ liệt kê feature. Nó gọi đúng công việc cảm xúc mà game làm cho người chơi. Puzzle sắp xếp có thể hứa khoảnh khắc mọi thứ trở về đúng chỗ; match-3 có thể hứa chuỗi phản ứng và cảm giác tiến bộ qua nhiều level. Khi creative hứa một game khác, team có thể vẫn thấy CTR hoặc CPI đẹp hơn trong ngắn hạn. Phần chi phí cần kiểm chứng thường xuất hiện sau đó: tỷ lệ chuyển đổi ở store listing, D1, độ sâu phiên chơi và ngôn ngữ trong đánh giá. Hãy lưu các tín hiệu này theo từng creative và từng cohort. Nguồn lý thuyết: Oliver, 1980; tổng quan về Expectation-Disconfirmation Theory tại PMC. Chương tiếp theo kiểm tra nơi người chơi bắt đầu xác minh lời hứa đó: store listing và lần mở game đầu.",
)
doc.save(PATH)
print("Added a theory anchor to Chapter 5.")
