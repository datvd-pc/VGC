from pathlib import Path
import re

from docx import Document


DOCUMENT_PATH = Path(r"D:\CODE\VGC\Long form\The-Art-of-Monetization-Vietnamese-Editable.docx")
HEADER_TEXT = "NGHỆ THUẬT KIẾM TIỀN TRONG GAME"


def replace_paragraph(document, index, expected, replacement):
    paragraph = document.paragraphs[index]
    actual = paragraph.text
    normalize = lambda value: re.sub(r"\s+", " ", value).strip()
    if normalize(actual) != normalize(expected):
        raise ValueError(
            f"Paragraph {index} changed since this review. Expected {expected!r}; got {actual!r}."
        )
    paragraph.clear()
    paragraph.add_run(replacement)


def main():
    document = Document(DOCUMENT_PATH)

    replacements = {
        3: (
            "Research note",
            "Research note\nGhi chú nghiên cứu",
        ),
        4: (
            "Trong mô hình kinh doanh của ngành game, một quyết định đúng hiếm khi đến từ riêng một mechanic, một creative nổi bật hay một dashboard. Một mechanic phổ biến không đại diện cho sản phẩm. Một creative táo bạo không bù nổi economics yếu. Dashboard không quyết định được cảm giác chơi, design level hay quyết định khó nhất: Dừng một ý tưởng trông đầy hứa hẹn.",
            "Trong một mô hình kinh doanh game, một quyết định đúng hiếm khi đến từ riêng một cơ chế, một mẫu quảng cáo nổi bật hay một bảng số liệu. Một cơ chế phổ biến chưa tạo thành luận điểm sản phẩm. Một mẫu quảng cáo táo bạo không bù nổi economics yếu. Bảng số liệu không thay được cảm giác chơi, tay nghề thiết kế level hay quyết định khó nhất: dừng một ý tưởng trông đầy hứa hẹn.",
        ),
        5: (
            "Ebook này được truyền cảm hứng từ một Game Event do publisher tổ chức và từ việc quan sát kỹ hơn Category puzzle game. Cơ hội thị trường rất rõ, và cách tham gia mô hình này cũng cũng rõ.",
            "Nghiên cứu này bắt đầu từ một Game Event do publisher tổ chức và từ việc quan sát kỹ hơn thể loại puzzle. Cơ hội thị trường rất rõ.",
        ),
        6: (
            "Tuy nhiên, cơ chế chơi, chủ đề và cách kiếm tiền của những game thành công được sao chép nhanh hơn khả năng giải thích nhu cầu người chơi, khối lượng sản xuất và logic kinh tế phía sau chúng.",
            "Sự ồn ào cũng rõ: cơ chế chơi, chủ đề và cách kiếm tiền của những game thành công được sao chép nhanh hơn khả năng giải thích nhu cầu người chơi, khối lượng sản xuất và logic kinh tế phía sau chúng.",
        ),
        7: (
            "Framework trong tài liệu này đang được xây dựng liên tục: bao gồm giả thuyết, công cụ, bằng chứng công khai và câu hỏi. Nội dung sẽ được cập nhật liên tục dựa trên góp ý của độc giả, kinh nghiệm phát hành game, hành vi người chơi và dữ liệu theo cohort.",
            "Framework trong tài liệu này gồm giả thuyết, công cụ, bằng chứng công khai và câu hỏi. Nội dung sẽ tiếp tục được cập nhật dựa trên góp ý của độc giả, kinh nghiệm phát hành game, hành vi người chơi và dữ liệu theo cohort.",
        ),
        8: (
            "Nếu bạn có câu hỏi, phản biện hoặc đóng góp, vui lòng chia sẻ trường hợp làm khiến framework này không có tác dụng hoặc hợp lý, các bằng chứng mâu thuẫn, hoặc chia sẻ một decision tool đã giúp team tránh một quyết định xấu. ",
            "Nếu có câu hỏi, phản biện hoặc đóng góp, hãy chia sẻ điều kiện làm framework này không còn phù hợp, bằng chứng mâu thuẫn với nó, hoặc một decision tool đã giúp team tránh một quyết định xấu.",
        ),
        9: (
            "Mục tiêu ebook là đóng góp một ngôn ngữ rõ ràng hơn cho các quyết định làm game, và cải thiện nó để tạo ra các sản phẩm game hoàn thiện nhất.",
            "Mục tiêu là đóng góp một ngôn ngữ rõ ràng hơn cho các quyết định làm game và cải thiện framework này cùng những người làm việc gần sản phẩm nhất.",
        ),
        10: (
            "A note to the reader",
            "A note to the reader\nLời nhắn tới người đọc",
        ),
        11: (
            "Làm game đã khó. Kiếm tiền từ game còn khó hơn. Vì thiết kế, kinh tế trong game,việc thu hút người chơi, sản phẩm, dữ liệu và vận hành phải gặp nhau tại cùng một điểm.",
            "Làm game đã khó. Kiếm tiền từ game còn khó hơn vì thiết kế, kinh tế trong game, thu hút người chơi, sản phẩm, dữ liệu và vận hành phải gặp nhau tại cùng một điểm.",
        ),
        12: (
            "Mỗi nhân sự, vị trí lại có những kỹ năng và kinh nghiệm khác nhau, nhưng vận hành cùng nhau lại cần cùng một bộ hiểu biết, kiến thức và ngôn ngữ chung: Từ founder, người phụ trách sản phẩm, game designer, analyst, người làm UA, publisher, hoặc một team nhỏ mà mỗi người lại phải đóng nhiều vai. Đó là lý do chính của ebook này.",
            "Mỗi vai trò có kỹ năng và kinh nghiệm khác nhau, nhưng để làm việc cùng nhau cần một bộ hiểu biết và ngôn ngữ chung: founder, người phụ trách sản phẩm, game designer, analyst, người làm UA, publisher, hoặc một team nhỏ nơi mỗi người phải đảm nhận nhiều vai.",
        ),
        13: (
            "Mỗi framework trong sách là một checklist cho nguyên mẫu đang được kiểm chứng: Creative nào cần test, level nào cần xem, quảng cáo nên đặt ở đâu, gói bán nào có lý do tồn tại, chỉ số nào cần đọc, hoặc khi nào cần dừng lại.",
            "Mỗi framework trong sách là một checklist để kiểm tra một nguyên mẫu: mẫu quảng cáo nào cần test, level nào cần xem, quảng cáo nên đặt ở đâu, gói bán nào có lý do tồn tại, chỉ số nào cần đọc và khi nào nên dừng lại.",
        ),
        14: (
            "Vậy mong bạn, hãy bật song song một bản note và ebook này. Hãy mở game và bảng số liệu bên cạnh. Một ebook hay về kiếm tiền trong game phải làm bạn mở game ra với câu hỏi tốt hơn, và tự tin hơn trong việc phát hành.",
            "Hãy mở sổ ghi chú cùng ebook này. Đặt bản game và bảng số liệu bên cạnh. Một ebook hay về kiếm tiền trong game phải khiến bạn mở game ra với ít chắc chắn hơn, nhưng nhiều câu hỏi tốt hơn.",
        ),
        23: ("Hệ thống đằng sau store", "Hệ thống phía sau màn hình cửa hàng"),
        24: ("Từ creative đến lần quay lại đầu tiên", "Từ quảng cáo đến lần quay lại đầu tiên"),
        25: ("Progress, pressure và fairness", "Tiến bộ, áp lực và sự công bằng"),
        26: ("Ads, IAP và economy", "Quảng cáo, IAP và nền kinh tế trong game"),
        27: ("Signals, decisions và experiments", "Tín hiệu, quyết định và thử nghiệm"),
        28: ("Data có thể và không thể quyết định gì", "Dữ liệu có thể và không thể quyết định gì"),
        29: ("Operating system của live game", "Hệ thống vận hành game đang phát hành"),
        30: ("Playbook theo category", "Cẩm nang theo từng thể loại game"),
        31: ("Audit", "Kiểm tra tổng thể"),
        33: (
            "How to read this ebook",
            "How to read this ebook\nCách đọc ebook này",
        ),
        34: (
            "Đừng đọc như một blog post. Nó hiệu quả nhất khi bạn đang audit game.",
            "Đừng đọc như một bài blog. Ebook hiệu quả nhất khi bạn dùng nó để kiểm tra một game cụ thể.",
        ),
        35: (
            "Mỗi chương sẽ đồng hành với bạn trong từng yếu tố: Quảng cáo, trang cửa hàng, phiên chơi đầu tiên, level, khoảnh khắc thua, quảng cáo có thưởng, gói đề nghị mua, event, bảng số liệu, review hoặc buổi họp của team.",
            "Mỗi chương đi qua một điểm chạm cụ thể: quảng cáo, trang cửa hàng, phiên chơi đầu tiên, level, khoảnh khắc thua, quảng cáo có thưởng, gói đề nghị mua, sự kiện, bảng số liệu, đánh giá hoặc buổi họp của team.",
        ),
        36: (
            "Nếu game đang phát hành, giữ bản game gần bên. Khi đọc về mười level đầu, chơi lại mười level đầu. Khi đọc về quảng cáo có thưởng, tìm vị trí hiển thị đầu tiên và hỏi người chơi cần gì ở thời điểm đó. Khi đọc về IAP, mở cửa hàng và gọi tên vấn đề chính xác mà từng gói đang giải quyết. Khi đọc về chỉ số, mở bảng số liệu và tách tín hiệu khỏi nhiễu.",
            "Với game đã phát hành, hãy giữ bản game gần bên. Khi đọc về mười level đầu, chơi lại mười level đầu. Khi đọc về quảng cáo có thưởng, tìm vị trí hiển thị đầu tiên và hỏi người chơi cần gì ở thời điểm đó. Khi đọc về mua trong app (IAP), mở cửa hàng và gọi tên vấn đề mà từng gói đang giải quyết. Khi đọc về chỉ số, mở bảng số liệu và tách tín hiệu khỏi nhiễu.",
        ),
        37: (
            "Nếu bạn còn ở giai đoạn prototype, dùng các chương như những cổng kiểm tra trước soft launch. Một game chưa giải thích được lời hứa với người chơi, phiên chơi đầu, áp lực, trao đổi quảng cáo, logic của gói bán và lý do quay lại thì chưa sẵn sàng để mở rộng.",
            "Nếu game còn ở giai đoạn prototype, hãy dùng các chương như những cổng kiểm tra trước khi phát hành thử nghiệm giới hạn (soft launch). Một game chưa giải thích được lời hứa với người chơi, phiên chơi đầu, áp lực, trao đổi quảng cáo, logic của gói bán và lý do quay lại thì chưa sẵn sàng để mở rộng.",
        ),
        38: (
            "Mục tiêu không phải đồng ý với mọi framework. Mục tiêu là rời mỗi chương với một câu hỏi sắc hơn cho game của bạn, và cách tốt hơn để thách thức câu trả lời.",
            "Mục tiêu không phải là đồng ý với mọi framework. Mục tiêu là rời mỗi chương với một câu hỏi sắc hơn cho game của bạn và cách tốt hơn để thách thức câu trả lời.",
        ),
        44: (
            "Monetization starts before the store",
            "Monetization starts before the store\nViệc kiếm tiền bắt đầu trước màn hình cửa hàng",
        ),
        45: (
            "Nhiều game không thất bại trong việc monetization ở shop. Chúng thất bại từ khâu chuẩn bị.",
            "Nhiều game không thất bại ở màn hình cửa hàng. Chúng thất bại sớm hơn, ở những điều kiện tạo ra niềm tin để người chơi muốn trả tiền.",
        ),
        47: ("Trang cửa hàng không chứng minh được quảng cáo.", "Trang cửa hàng không thể tự chứng minh lời hứa của quảng cáo."),
        48: ("Hướng dẫn quá dài.", "Hướng dẫn quá dài hoặc cướp quyền điều khiển."),
        49: (
            "Quảng cáo xen kẽ đầu tiên đến trước khi người chơi quyết định game đáng thêm một phút nữa.",
            "Quảng cáo chen màn hình đầu tiên xuất hiện trước khi người chơi quyết định game có đáng thêm một phút nữa hay không.",
        ),
        50: (
            "Level có cảm giác bất công, rồi game bán booster như thuốc chữa.",
            "Một level tạo cảm giác bất công, rồi game mới bán booster như cách chữa lỗi.",
        ),
        51: (
            "Team đọc doanh thu như dấu hiệu sức khỏe trong khi retention, review, hoàn tiền và trust đang chảy máu.",
            "Team đọc doanh thu như tín hiệu sức khỏe, trong khi tỷ lệ quay lại, đánh giá, hoàn tiền và niềm tin đang suy giảm.",
        ),
        52: (
            "Nên nhớ: Store giúp bạn thu tiền, chứ không phải là cách bày ra một cái gì đó vui vui rồi thu phí bảo kê.",
            "Màn hình cửa hàng là nơi hoàn tất trao đổi giá trị. Nó không thể cứu một trải nghiệm chưa khiến người chơi muốn ở lại.",
        ),
        54: ("Vậy game sẽ thu cái gì trước khi thu được tiền?", "Trước khi thu tiền, game cần nhận được gì từ người chơi?"),
        56: ("Người chơi trả bằng sự chú ý trước.", "Trước hết là sự chú ý."),
        57: ("Rồi đến cú nhấp, lượt cài đặt, thời gian chờ, phiên chơi đầu và lần quay lại.", "Sau đó là cú nhấp, lượt cài đặt, thời gian chờ, phiên chơi đầu và lần quay lại."),
        58: ("Tiền đến sau, nếu game giữ được đủ trust.", "Tiền đến sau, khi game giữ được đủ niềm tin."),
        60: (
            "Trong casual, hybrid-casual, puzzle và hybrid puzzle, monetization là kết quả của cả journey:",
            "Trong game casual, hybrid-casual, puzzle và hybrid puzzle, việc kiếm tiền là kết quả của cả một hành trình:",
        ),
        61: (
            "See Ad -> Click -> Store -> Install -> First Open -> First 10 Levels -> First Return -> Habit -> Ads Opt-in -> First Purchase -> Repeat Purchase -> Event/Live Ops -> Share or Recommend",
            "Thấy quảng cáo -> Nhấp -> Cửa hàng -> Cài đặt -> Mở lần đầu -> 10 level đầu -> Lần quay lại đầu tiên -> Thói quen -> Tự nguyện xem quảng cáo -> Lần mua đầu -> Mua lại -> Sự kiện/vận hành game -> Chia sẻ hoặc giới thiệu",
        ),
        62: ("Mỗi điểm chạm đều đòi hỏi một thứ.", "Mỗi điểm chạm đều yêu cầu người chơi trao cho game một thứ."),
        63: (
            "Ad đòi attention. Store đòi belief. First open đòi patience. First ten levels đòi trust. Ad đầu tiên đòi permission. Offer đầu tiên đòi tiền. Event đầu tiên đòi habit.",
            "Quảng cáo cần sự chú ý. Trang cửa hàng cần niềm tin ban đầu. Lần mở đầu cần kiên nhẫn. Mười level đầu cần xây dựng niềm tin. Quảng cáo đầu tiên cần sự cho phép. Gói bán đầu tiên cần một lý do để trả tiền. Sự kiện đầu tiên cần thói quen quay lại.",
        ),
        65: ("Đây là công thức tài chính đúng, nhưng đến quá muộn cho design.", "Đây là một thấu kính tài chính hữu ích, nhưng nó xuất hiện quá muộn để hướng dẫn thiết kế."),
        66: ("Công thức vận hành hữu ích hơn là:", "Một công thức vận hành hữu ích hơn là:"),
        67: ("Monetization = Player Need * Right Context * Trust * Execution Speed", "Monetization = Nhu cầu người chơi * Đúng ngữ cảnh * Niềm tin * Tốc độ thực thi"),
        68: (
            "Nhu cầu của người chơi có thể là CẢM GIÁC được giải tỏa, chơi lại, CẢM GIÁC giỏi hơn, đi nhanh hơn, tiến bộ, sưu tập, tiện lợi, CẢM GIÁC về địa vị hoặc quyền kiểm soát.",
            "Nhu cầu của người chơi có thể là cảm giác được giải tỏa, được chơi lại, thấy mình giỏi hơn, đi nhanh hơn, tiến bộ, sưu tập, thuận tiện, có địa vị hoặc kiểm soát được tình huống.",
        ),
        69: (
            "Đúng ngữ cảnh nghĩa là gói đề nghị xuất hiện khi nhu cầu đó đang rõ nhất, không phải khi publisher / studio cần thêm doanh thu.",
            "Đúng ngữ cảnh nghĩa là gói đề nghị xuất hiện khi nhu cầu ấy rõ nhất, không phải khi publisher hoặc studio cần thêm doanh thu.",
        ),
        70: (
            "Trust là cảm giác game còn đủ công bằng để họ dành thêm thời gian hoặc tiền. Tốc độ thực thi là khả năng team học nhanh từ quảng cáo, funnel, mức độ xem quảng cáo, tỉ lệ mua gói, retention theo cohort, review và live ops.",
            "Niềm tin là cảm giác game vẫn đủ công bằng để người chơi dành thêm thời gian hoặc tiền. Tốc độ thực thi là khả năng team học nhanh từ dữ liệu quảng cáo, phễu chuyển đổi, mức độ xem quảng cáo, tỷ lệ mua gói, tỷ lệ quay lại theo cohort, đánh giá và vận hành game đang phát hành.",
        ),
        71: ("Operating map của cuốn sách có sáu phần:", "Bản đồ vận hành của ebook có sáu phần:"),
        72: ("Promise * Progress * Pressure * Permission * Payment * Persistence", "Lời hứa * Tiến bộ * Áp lực * Sự cho phép * Thanh toán * Sự gắn bó lâu dài"),
        73: ("Khi một phần thiếu hụt, revenue vẫn có thể tăng trong một thời gian. Nhưng nó nguy hiểm.", "Khi thiếu một phần trong số này, doanh thu vẫn có thể tăng trong một thời gian. Điều đó vẫn nguy hiểm."),
        75: (
            "Vì khi IMPDAU có thể tăng trong khi D3 giảm. Fail offer có thể convert trong khi review bắt đầu gọi level là unfair. Interstitial có thể nâng ARPDAU trong khi game khó scale hơn.",
            "IMPDAU, tức số lượt hiển thị quảng cáo trung bình trên mỗi người chơi hoạt động hằng ngày, có thể tăng trong khi tỷ lệ quay lại ở ngày thứ ba (D3 retention) giảm. Một ưu đãi không phù hợp có thể vẫn chuyển đổi, trong khi phần đánh giá bắt đầu gọi level là bất công. Quảng cáo chen màn hình có thể nâng ARPDAU, nhưng khiến game khó mở rộng quy mô hơn.",
        ),
        76: (
            "Revenue healthy để lại cho player lý do tiếp tục sau ad, offer và purchase. Revenue borrowed lấy giá trị từ pressure mà game không thể bảo vệ bằng rules và value.",
            "Doanh thu lành mạnh để lại lý do cho người chơi tiếp tục sau quảng cáo, gói đề nghị mua và lần thanh toán. Doanh thu vay mượn lấy giá trị từ áp lực mà game không thể bảo vệ bằng luật chơi và giá trị trao đổi.",
        ),
        77: ("Example: Clear Garden", "Example: Clear Garden\nVí dụ: Clear Garden"),
        78: ("Hãy tưởng tượng một hybrid puzzle đang soft launch tên là Clear Garden.", "Hãy tưởng tượng một game puzzle lai đang được thử nghiệm phát hành giới hạn (soft launch) với tên Clear Garden."),
        79: (
            "Core loop yêu cầu player sort đồ vật trong khu vườn bỏ hoang vào tray có giới hạn, tạo không gian và phục hồi từng góc vườn. Creative bán cảm giác relief through order và một biến đổi có thể nhìn thấy.",
            "Vòng lặp cốt lõi yêu cầu người chơi sắp đồ vật trong khu vườn bỏ hoang vào một khay có giới hạn, từ đó tạo không gian và phục hồi từng góc vườn. Mẫu quảng cáo bán cảm giác tìm lại trật tự và một biến đổi có thể nhìn thấy.",
        ),
        80: (
            "Version đầu tiên có các lỗi quen thuộc: xin tracking permission trước khi player chạm vào puzzle; đặt interstitial sau level hai; level 7 thêm quá nhiều object type rồi show extra-tray offer ngay sau người chơi fail; starter pack bán coins mà không giải thích coins thay đổi cái gì; daily reward rộng rãi nhưng khu vườn không đưa ra lý do quay lại ngày mai.",
            "Phiên bản đầu tiên mắc những lỗi quen thuộc: xin quyền theo dõi trước khi người chơi chạm vào câu đố; đặt quảng cáo chen màn hình sau level hai; ở level 7, thêm quá nhiều loại đồ vật rồi đề nghị mua thêm ô khay ngay khi người chơi thua; bán gói khởi đầu gồm tiền tệ trong game nhưng không nói nó giải quyết vấn đề gì; trao phần thưởng hằng ngày rộng rãi nhưng khu vườn không tạo lý do để quay lại vào ngày mai.",
        ),
        81: ("Vậy thì con game này sống sót bằng cách nào? Quá khó.", "Nếu đây là một bản build thật, chỉ số nào sẽ cho thấy rủi ro này trước khi team mua thêm traffic?"),
        82: ("Ví dụ về puzzle game Clear Garden chỉ là fiction. Các quyết định thì rất thật.", "Clear Garden là một ví dụ hư cấu. Các loại quyết định được nêu ra thì rất thật."),
        83: ("Mỗi part sẽ quay lại nó để biến principle thành một khoảnh khắc cụ thể trong build.", "Mỗi phần sẽ quay lại ví dụ này để biến một nguyên tắc thành một khoảnh khắc cụ thể trong phiên bản game đang thử nghiệm."),
        84: ("Part I: The system behind the store", "Part I: The system behind the store\nPhần I: Hệ thống phía sau cửa hàng"),
        85: ("1. Trust budget - Ngân sách niềm tin", "1. Trust budget\nNgân sách niềm tin"),
        87: (
            "Trust được bổ sung bằng các chi tiết bình thường: rules dễ đọc, input đúng như hứa, reward được grant, price rõ, fail dạy được điều gì đó, support sửa purchase bị lỗi. Tổng hợp của chúng quyết định ad là fair exchange hay purchase là trap.",
            "Niềm tin được bổ sung bằng những chi tiết bình thường: luật chơi dễ đọc, thao tác đúng như lời hứa, phần thưởng được trao đúng lúc, giá rõ ràng, một lần thua dạy được điều gì đó và bộ phận hỗ trợ sửa lỗi mua hàng. Tổng hợp của chúng quyết định quảng cáo là một trao đổi công bằng hay việc mua hàng là một chiếc bẫy.",
        ),
    }

    for index, (expected, replacement) in replacements.items():
        replace_paragraph(document, index, expected, replacement)

    cover_section, content_section = document.sections
    cover_section.header.is_linked_to_previous = False
    cover_section.header.paragraphs[0].text = ""
    content_section.header.is_linked_to_previous = False
    content_section.header.paragraphs[0].text = HEADER_TEXT

    document.save(DOCUMENT_PATH)
    print(f"Updated {len(replacements)} paragraphs in {DOCUMENT_PATH.name}")


if __name__ == "__main__":
    main()
