from pathlib import Path
import re
from docx import Document

PATH = Path(r"D:\CODE\VGC\Long form\The-Art-of-Monetization-Vietnamese-Editable.docx")

def replace(document, index, expected, value):
    paragraph = document.paragraphs[index]
    compact = lambda text: re.sub(r"\s+", " ", text).strip()
    if compact(paragraph.text) != compact(expected):
        raise ValueError(f"Paragraph {index} was changed manually and was not replaced.")
    paragraph.clear()
    paragraph.add_run(value)

doc = Document(PATH)
edits = {
    94: (
        "Tháng 7/2023, Sensor Tower ước tính Royal Match đạt khoảng 112 triệu USD doanh thu gộp và 14,6 triệu lượt tải, vượt Candy Crush Saga trong cùng tháng. Con số này cho thấy quy mô của một game đang vận hành tốt. Nó không cho biết ad load, economy, level design hay các quyết định nội bộ nào đã tạo ra kết quả đó.",
        "Hãy nghĩ về thay đổi gần nhất team đã đưa vào để tăng doanh thu: thêm interstitial, đặt một offer sau khi thua, tăng giá một gói hay làm một level khó hơn. Chỉ số nào tăng ngay sau thay đổi đó? Và team đã mở những chỉ số nào để kiểm tra cái giá phải trả?",
    ),
    95: (
        "Đó là giới hạn quan trọng của market data. Một số liệu doanh thu có thể xác nhận kết quả, nhưng không thể xác nhận con đường dẫn đến kết quả ấy. Khi một team nhìn thấy một game dẫn đầu, câu hỏi có ích không phải là “họ đã đặt gì vào game?”, mà là “hệ thống nào khiến người chơi tiếp tục ở lại đủ lâu để những lựa chọn đó có ý nghĩa?”.",
        "Nếu câu trả lời chỉ dừng ở doanh thu, tỷ lệ chuyển đổi hoặc lượt hiển thị quảng cáo, bức tranh vẫn chưa đủ. Một thay đổi có thể làm một chỉ số đẹp lên trong tuần này, đồng thời làm người chơi rời sớm hơn, đánh giá thấp hơn hoặc khiến chi phí nội dung tăng ở những tuần sau.",
    ),
    96: (
        "Doanh thu đánh đổi tương lai xuất hiện khi một thay đổi làm đẹp chỉ số ngắn hạn nhưng làm yếu điều kiện tạo doanh thu sau này. Interstitial quá sớm có thể tăng doanh thu quảng cáo nhưng kéo D1 xuống. Một bức tường độ khó có thể làm offer sau khi thua chuyển đổi tốt hơn, đồng thời khiến đánh giá của người chơi gọi game là bất công. Không có thay đổi nào tự động sai; điều quan trọng là team có đọc phần chi phí đi kèm hay không.",
        "Đó là ý nghĩa của doanh thu đánh đổi tương lai: một thay đổi làm đẹp chỉ số ngắn hạn nhưng làm yếu điều kiện tạo doanh thu sau này. Interstitial quá sớm có thể tăng doanh thu quảng cáo nhưng kéo D1 xuống. Một bức tường độ khó có thể làm offer sau khi thua chuyển đổi tốt hơn, đồng thời khiến đánh giá của người chơi gọi game là bất công. Không có thay đổi nào tự động sai; điều quan trọng là team có đọc phần chi phí đi kèm hay không.",
    ),
    97: (
        "Lợi ích tức thời có thể là doanh thu, tỷ lệ chuyển đổi, lượt hiển thị quảng cáo hoặc CPI. Phần chi phí có thể nằm ở retention, tỷ lệ người trả tiền quay lại, đánh giá, hoàn tiền, khối lượng hỗ trợ và chi phí nội dung về sau. Đặt hai nhóm chỉ số cạnh nhau trước khi kết luận một thay đổi là thành công.",
        "Khi review một thay đổi, đặt lợi ích tức thời bên cạnh phần chi phí có thể phát sinh. Lợi ích có thể là doanh thu, tỷ lệ chuyển đổi, lượt hiển thị quảng cáo hoặc CPI. Phần chi phí có thể nằm ở retention, tỷ lệ người trả tiền quay lại, đánh giá, hoàn tiền, khối lượng hỗ trợ và chi phí nội dung về sau. Đặt hai nhóm chỉ số cạnh nhau trước khi kết luận một thay đổi là thành công.",
    ),
}
for index, edit in edits.items():
    replace(doc, index, *edit)
doc.save(PATH)
print("Revised Chapter 2 opening as a self-audit.")
