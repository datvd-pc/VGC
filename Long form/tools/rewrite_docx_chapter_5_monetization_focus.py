from pathlib import Path
import re
from docx import Document

PATH = Path(r"D:\CODE\VGC\Long form\The-Art-of-Monetization-Vietnamese-Editable.docx")

def replace(document, index, expected, value):
    paragraph = document.paragraphs[index]
    compact = lambda text: re.sub(r"\s+", " ", text).strip()
    if compact(paragraph.text) != compact(expected):
        raise ValueError(f"Paragraph {index} was changed manually and was not replaced.")
    paragraph.clear()
    paragraph.add_run(value)

doc = Document(PATH)
edits = {
    127: ("5. Creative sells a feeling\nQuảng cáo bán cảm giác trước khi bán tính năng", "5. Creative sells a feeling\nCreative bán một lời hứa mà monetization phải giữ được"),
    128: ("Người chơi hiếm khi cài game vì danh sách tính năng. Họ cài vì video hứa một cảm giác: cứu vãn tình huống, lập lại trật tự, giải được điều khó, tốc độ, một cú dọn sạch đã mắt, màn lội ngược dòng hoặc bộ sưu tập ngày một lớn.", "Hãy chọn creative đã mang về nhiều lượt cài đặt nhất trong chiến dịch gần đây. Người chơi nhìn thấy tình huống gì trong ba giây đầu? Họ chờ được giải tỏa, được lập lại trật tự hay được chứng kiến một cú lội ngược dòng? Sau khi cài đặt, phiên chơi đầu có đưa họ đến cảm giác đó đủ nhanh không? Đây là monetization link của creative: lời hứa ban đầu quyết định chất lượng traffic, D1 và trần LTV mà team có thể tạo ra sau này."),
    129: ("Quảng cáo phải gọi tên công việc cảm xúc của game. Puzzle sắp xếp bán niềm vui đưa mọi thứ về trật tự. Match-3 bán chuỗi phản ứng và cảm giác tiến bộ dài hạn. Quảng cáo có thể phóng đại cảm xúc, nhưng không sống lâu nếu bán một game khác. Chi phí nằm trong funnel: CTR tốt, tỉ lệ chuyển đổi ở store yếu, D1 thấp, phiên chơi ngắn và review nói \"quảng cáo giả\".", "Một creative tốt không chỉ liệt kê feature. Nó gọi đúng công việc cảm xúc mà game làm cho người chơi. Puzzle sắp xếp có thể hứa khoảnh khắc mọi thứ trở về đúng chỗ. Match-3 có thể hứa chuỗi phản ứng và cảm giác tiến bộ qua nhiều level. Khi creative hứa một game khác, team có thể vẫn thấy CTR hoặc CPI đẹp hơn trong ngắn hạn. Phần chi phí thường xuất hiện sau đó: tỷ lệ chuyển đổi ở store listing yếu, D1 thấp, phiên chơi ngắn và đánh giá cho rằng quảng cáo không trung thực."),
    130: ("Xây creative library theo scenario. Lưu fantasy, conflict, payoff, persona, proof trong first session và cohort outcome. AI có thể tăng tốc sản xuất. AI không thể quyết định scenario có faithful với product hay không.", "Hãy xây creative library theo từng tình huống, không chỉ theo asset. Với mỗi creative, lưu lại lời hứa, xung đột, khoảnh khắc giải tỏa, nhóm người chơi hướng tới, bằng chứng trong phiên chơi đầu và kết quả theo cohort. AI có thể giúp team tạo nhiều biến thể nhanh hơn. Quyết định khó hơn vẫn thuộc về team: lời hứa này có trung thực với product không, và có đưa về nhóm người chơi mà economy của game có thể phục vụ không? Chương tiếp theo kiểm tra nơi người chơi bắt đầu xác minh lời hứa đó: store listing và lần mở game đầu."),
}
for index, edit in edits.items(): replace(doc, index, *edit)
doc.save(PATH)
print("Rewrote Chapter 5 with a monetization focus.")
