from pathlib import Path
import re

from docx import Document


PATH = Path(r"D:\CODE\VGC\Long form\The-Art-of-Monetization-Vietnamese-Editable.docx")


def replace_paragraph(document, index, expected, replacement):
    paragraph = document.paragraphs[index]
    compact = lambda value: re.sub(r"\s+", " ", value).strip()
    if compact(paragraph.text) != compact(expected):
        raise ValueError(f"Paragraph {index} was changed manually and was not replaced.")
    paragraph.clear()
    paragraph.add_run(replacement)


def set_cell(cell, value):
    paragraph = cell.paragraphs[0]
    paragraph.clear()
    paragraph.add_run(value)


def main():
    document = Document(PATH)
    edits = {
        84: (
            "Mỗi người chơi đến với một ngân sách niềm tin hữu hạn. Quảng cáo hứa sai làm giảm ngân sách đó.",
            "Tháng 3/2023, Ủy ban Thương mại Liên bang Hoa Kỳ (FTC) hoàn tất lệnh buộc Epic Games hoàn trả 245 triệu USD cho người tiêu dùng, sau các cáo buộc về dark patterns và khoản mua trong game không mong muốn. Đến tháng 12/2024, FTC cho biết đã gửi hơn 72 triệu USD tiền hoàn lại. Đây không phải là số liệu về retention, LTV hay hiệu quả kinh doanh của Fortnite. Nó là một bằng chứng công khai về giá phải trả khi luồng thanh toán khiến người chơi không còn hiểu rõ mình vừa đồng ý điều gì.",
        ),
        85: (
            "Tải chậm: Giảm niềm tin.",
            "Case này không cho phép suy ra cách Epic thiết kế economy nội bộ, cũng không cho phép quy kết một màn hình cụ thể là nguyên nhân duy nhất. Tuy vậy, nó làm rõ một điểm quan trọng: purchase flow không đứng ngoài trải nghiệm chơi. Nó là một phần của lời hứa mà game đưa ra với người chơi.",
        ),
        86: (
            "Hướng dẫn chiếm quyền điều khiển : Giảm niềm tin.",
            "Vì vậy, thay vì hỏi một màn hình có chuyển đổi tốt không, hãy hỏi thêm: người chơi có hiểu lựa chọn này không, có đủ thời gian để cân nhắc không, và liệu họ còn tin game sau khi giao dịch kết thúc không? Đó là điểm bắt đầu của trust budget.",
        ),
        87: (
            "Nút đóng khó thấy : Giảm niềm tin.",
            "Mỗi người chơi đến với một lượng niềm tin hữu hạn. Một creative hứa quá mức, thời gian tải dài, hướng dẫn giành quyền điều khiển hoặc nút đóng khó thấy đều làm lượng niềm tin đó giảm đi. Một level tạo cảm giác bất công rồi lập tức bán booster gây hại nhanh hơn, vì nó khiến người chơi nghi ngờ lý do thật của thất bại.",
        ),
        88: (
            "Một level có cảm giác không thể thắng rồi bán booster : Giảm niềm tin rất nhanh.",
            "Niềm tin không đến từ một khoảnh khắc lớn. Nó được xây qua những chi tiết bình thường: mục tiêu rõ, luật chơi dễ hiểu, thao tác đúng như lời hứa, reward được trao đúng lúc, giá bán minh bạch và hỗ trợ kịp thời khi giao dịch có lỗi.",
        ),
        89: (
            "Niềm tin được bổ sung bằng những chi tiết bình thường: luật chơi dễ đọc, thao tác đúng như lời hứa, phần thưởng được trao đúng lúc, giá rõ ràng, một lần thua dạy được điều gì đó và bộ phận hỗ trợ sửa lỗi mua hàng. Tổng hợp của chúng quyết định quảng cáo là một trao đổi công bằng hay việc mua hàng là một chiếc bẫy.",
            "Tổng hợp của các chi tiết ấy quyết định người chơi xem rewarded ad như một trao đổi công bằng hay xem in-game store như một nơi họ cần dè chừng. Framework dưới đây không đo niềm tin bằng một con số duy nhất. Nó giúp team nhìn ra những điểm đang nuôi dưỡng hoặc làm hao mòn niềm tin trong cả hành trình.",
        ),
        90: (
            "Không được đọc một mức tăng doanh thu một mình. Thêm quảng cáo xen kẽ có thể tăng ARPDAU, đồng thời làm cohort rời sớm hơn, rating giảm và chi phí kéo người chơi quay lại tăng. Bạn đã chuyển tiền về hiện tại, chưa chắc đã tạo ra giá trị.",
            "Đừng đọc một mức tăng doanh thu một mình. Thêm interstitial có thể làm ARPDAU tăng, nhưng cũng có thể khiến một cohort rời game sớm hơn, điểm đánh giá giảm và chi phí kéo người chơi quay lại cao hơn. Tiền được thu sớm hơn chưa chắc đã là giá trị được tạo ra nhiều hơn.",
        ),
        92: (
            "Với mỗi monetization surface, hỏi: nó lấy trust, bổ sung trust, hay làm cả hai?",
            "Khi audit game, hãy chọn một điểm kiếm tiền cụ thể, chẳng hạn interstitial đầu tiên hoặc offer sau khi thua. Người chơi vừa nhận được giá trị gì? Họ có hiểu điều mình sắp đổi lấy không? Sau điểm đó, chỉ số nào sẽ cho biết niềm tin được giữ lại hay đã bị đánh đổi?",
        ),
    }
    for index, edit in edits.items():
        replace_paragraph(document, index, *edit)

    table = document.tables[1]
    set_cell(table.rows[0].cells[0], "Điều nuôi dưỡng niềm tin")
    set_cell(table.rows[0].cells[1], "Điều làm hao mòn niềm tin")
    rows = [
        ("Mục tiêu rõ ràng, luật chơi dễ hiểu", "Creative hứa một trải nghiệm khác"),
        ("Thua vẫn hiểu vì sao và còn nước đi tiếp theo", "Can thiệp độ khó không được giải thích"),
        ("Rewarded ad là một trao đổi tự nguyện", "Interstitial cắt ngang lúc người chơi đang tập trung"),
        ("Giá bán và thứ nhận được được nói rõ", "Offer mơ hồ hoặc thao tác mua nhầm"),
        ("Khôi phục nhanh khi giao dịch có lỗi", "Mất reward, khôi phục yếu, hỗ trợ chậm"),
    ]
    for row, values in zip(table.rows[1:], rows):
        set_cell(row.cells[0], values[0])
        set_cell(row.cells[1], values[1])

    document.save(PATH)
    print("Rewrote Chapter 1 and its trust table.")


if __name__ == "__main__":
    main()
