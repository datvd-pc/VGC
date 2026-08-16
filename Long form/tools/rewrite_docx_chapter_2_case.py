from pathlib import Path
import re
from docx import Document

PATH = Path(r"D:\CODE\VGC\Long form\The-Art-of-Monetization-Vietnamese-Editable.docx")

def replace(document, index, expected, value):
    p = document.paragraphs[index]
    compact = lambda text: re.sub(r"\s+", " ", text).strip()
    if compact(p.text) != compact(expected):
        raise ValueError(f"Paragraph {index} was changed manually and was not replaced.")
    p.clear()
    p.add_run(value)

doc = Document(PATH)
edits = {
    93: ("2. Healthy or borrowed revenue\nDoanh thu lành mạnh hoặc đi vay từ tương lai", "2. Healthy or borrowed revenue\nDoanh thu bền vững hay doanh thu đánh đổi tương lai"),
    94: ("Monetization debt xuất hiện khi team cải thiện metric ngắn hạn bằng cách làm hại điều kiện tạo ra revenue tương lai.", "Tháng 7/2023, Sensor Tower ước tính Royal Match đạt khoảng 112 triệu USD doanh thu gộp và 14,6 triệu lượt tải, vượt Candy Crush Saga trong cùng tháng. Con số này cho thấy quy mô của một game đang vận hành tốt. Nó không cho biết ad load, economy, level design hay các quyết định nội bộ nào đã tạo ra kết quả đó."),
    95: ("Quảng cáo xen kẽ quá sớm có thể nâng doanh thu quảng cáo và hạ D1. Một bức tường độ khó có thể nâng tỉ lệ mua thêm lượt đi và tạo review gọi game là \"rigged\". Quảng cáo kịch tính quá mức có thể hạ CPI nhưng đưa sai nhóm người chơi vào funnel. Lạm phát phần thưởng có thể buộc event sau trả ngày càng nhiều mới còn cảm thấy đáng giá.", "Đó là giới hạn quan trọng của market data. Một số liệu doanh thu có thể xác nhận kết quả, nhưng không thể xác nhận con đường dẫn đến kết quả ấy. Khi một team nhìn thấy một game dẫn đầu, câu hỏi có ích không phải là “họ đã đặt gì vào game?”, mà là “hệ thống nào khiến người chơi tiếp tục ở lại đủ lâu để những lựa chọn đó có ý nghĩa?”."),
    96: ("Không lựa chọn nào tự động sai. Sai lầm là tuyên bố thắng lợi trước khi đọc hóa đơn.", "Doanh thu đánh đổi tương lai xuất hiện khi một thay đổi làm đẹp chỉ số ngắn hạn nhưng làm yếu điều kiện tạo doanh thu sau này. Interstitial quá sớm có thể tăng doanh thu quảng cáo nhưng kéo D1 xuống. Một bức tường độ khó có thể làm offer sau khi thua chuyển đổi tốt hơn, đồng thời khiến đánh giá của người chơi gọi game là bất công. Không có thay đổi nào tự động sai; điều quan trọng là team có đọc phần chi phí đi kèm hay không."),
    97: ("Immediate gain: revenue, conversion, ad impressions, CPI. Possible cost: retention, payer retention, review sentiment, refunds, support load, future content cost.", "Lợi ích tức thời có thể là doanh thu, tỷ lệ chuyển đổi, lượt hiển thị quảng cáo hoặc CPI. Phần chi phí có thể nằm ở retention, tỷ lệ người trả tiền quay lại, đánh giá, hoàn tiền, khối lượng hỗ trợ và chi phí nội dung về sau. Đặt hai nhóm chỉ số cạnh nhau trước khi kết luận một thay đổi là thành công."),
    98: ("\"Chúng tôi tăng ad load\" mô tả mechanism. \"Chúng tôi cải thiện monetization\" là kết luận cần evidence.", "“Chúng tôi tăng ad load” chỉ mô tả điều team đã làm. “Chúng tôi cải thiện monetization” là kết luận cần bằng chứng. Chương tiếp theo đi sâu hơn vào điều kiện của bằng chứng đó: một thiết kế tử tế phải cho người chơi hiểu luật chơi, lựa chọn và hệ quả của việc trả tiền."),
}
for index, edit in edits.items():
    replace(doc, index, *edit)
doc.save(PATH)
print("Rewrote Chapter 2.")
