import docx

doc_path = r'D:\CODE\VGC\Long form\The-Art-of-Monetization-Vietnamese-Editable.docx'
doc = docx.Document(doc_path)

# Verify Part I index
assert 'Part I' in doc.paragraphs[82].text, f'Paragraph 82 is {doc.paragraphs[82].text}'

# Map of paragraph replacements
p_updates = {
    2: 'THE ART OF MONETIZATION\nNghệ thuật Thiết kế Kinh tế & Kiếm tiền trong Game',
    3: 'Research Note | Ghi chú Nghiên cứu',
    4: 'Trong ngành công nghiệp game, hiếm có quyết định nào đứng vững nếu chỉ dựa vào một cơ chế chơi (mechanic) đơn lẻ, một mẫu quảng cáo (creative) bắt mắt hay một bảng dữ liệu (dashboard) bóng bẩy. Một mechanic thú vị chưa đủ tạo nên định hướng sản phẩm; một creative táo bạo không thể che lấp một nền kinh tế (economy) thiếu bền vững. Dashboard có thể chỉ ra nơi đội ngũ cần chú ý, nhưng không bao giờ thay thế được trải nghiệm thực tế khi cầm máy chơi, năng lực thiết kế màn chơi (level design), hay quyết định khó khăn nhất: dũng cảm khai tử một ý tưởng đầy hứa hẹn.',
    5: 'Tài liệu nghiên cứu này khởi nguồn từ một sự kiện do nhà phát hành (publisher) tổ chức, trước khi mở rộng thành một cuộc khảo sát sâu vào dòng game giải đố (puzzle). Thị trường luôn mở ra cơ hội, nhưng cơ hội chỉ thực sự có ý nghĩa khi đội ngũ phát triển hiểu rõ mình đang cạnh tranh ở phân khúc nào và giải quyết bài toán gì.',
    6: 'Cơ chế chơi, chủ đề và mô hình kiếm tiền của những tựa game dẫn đầu thường bị sao chép nhanh hơn tốc độ mà thị trường thực sự thấu hiểu chúng. Ẩn sau mỗi lựa chọn tưởng chừng đơn giản luôn là một tổ hợp phức tạp: tâm lý người chơi, năng lực sản xuất và logic vận hành kinh tế (economy logic). Cuốn tài liệu này được biên soạn nhằm giúp bạn phân tích sâu những tầng ẩn giấu đó, trước khi vội vã biến một tín hiệu thị trường thành quyết định đánh đổi sản phẩm.',
    7: 'Nội dung ở đây tổng hợp các giả thuyết làm game, bộ khung tư duy (frameworks), dữ liệu thực tế từ thị trường và cả những câu hỏi chưa có lời giải tuyệt đối. Tài liệu sẽ liên tục được cập nhật dựa trên phản hồi từ cộng đồng phát triển, bài học phát hành thực chiến, hành vi người chơi và dữ liệu phân tích thuần tập (cohort analysis).',
    8: 'Sự phản biện có giá trị nhất là khi chỉ ra được ranh giới khiến một kết luận không còn đúng, đưa ra các dữ liệu đối nghịch hoặc chia sẻ một công cụ ra quyết định (decision tool) đã giúp đội ngũ né tránh sai lầm đắt giá. Những góc nhìn đó chính là đòn bẩy giúp tài liệu ngày càng tiệm cận thực tế.',
    9: 'Mục tiêu cốt lõi: Thiết lập một hệ ngôn ngữ chuẩn xác và minh bạch cho các quyết định sản phẩm, rồi không ngừng hoàn thiện ngôn ngữ ấy cùng chính những người đang trực tiếp làm ra game mỗi ngày.',
    10: 'A Note to the Reader | Lời Nhắn Gửi Người Đọc',
    11: 'Làm game vốn đã khó. Kiếm tiền từ game (Monetization) lại càng phức tạp, bởi nó đòi hỏi thiết kế game (game design), kinh tế (economy), thu hút người dùng (UA), định hướng sản phẩm (product), khoa học dữ liệu (data) và vận hành trực tiếp (live ops) phải giao thoa hài hòa trong cùng một trải nghiệm duy nhất.',
    12: 'Mỗi vai trò trong dự án mang một lăng kính và thế mạnh riêng. Nhưng để đưa một tựa game đi xa, từ Founder, Product Lead, Game Designer, Data Analyst, chuyên viên UA, cho đến Publisher hay các Indie Team, tất cả đều cần một hệ quy chiếu và ngôn ngữ chung — dù mỗi vị trí nhìn nhận bài toán từ những góc độ khác nhau.',
    13: 'Mỗi framework trong cuốn sách này là một bài kiểm tra nghiêm ngặt cho sản phẩm của bạn: creative nào đáng để thử nghiệm, màn chơi nào cần tinh chỉnh lại, vị trí quảng cáo nào là hợp lý, gói ưu đãi (offer) nào thực sự có lý do tồn tại, những cặp chỉ số nào bắt buộc phải đọc song hành và thời điểm nào cần dũng cảm dừng một hướng đi.',
    14: 'Hãy mở cuốn sách này bên cạnh bản build game và bảng dữ liệu của bạn. Giá trị thực sự của tài liệu không nằm ở việc đưa ra câu trả lời có sẵn thay bạn, mà ở việc giúp đội ngũ đặt ra những câu hỏi sắc bén và chính xác hơn mỗi khi mở dự án ra xem xét.',
    15: 'Key Terms | Thuật Ngữ Cần Biết',
    16: 'Bạn không nhất thiết phải thành thạo toàn bộ thuật ngữ chuyên ngành game để đọc tài liệu này. Các thuật ngữ tiếng Anh dưới đây được giữ nguyên vì tính phổ biến trong môi trường làm việc thực tế; mỗi khi xuất hiện lần đầu trong từng phần, chúng đều được giải nghĩa theo ngữ cảnh cụ thể của cuốn sách.',
    22: 'Contents | Mục Lục',
    23: 'Hệ thống phía sau màn hình cửa hàng (The System Behind the Store)',
    24: 'Từ quảng cáo đến lần quay lại đầu tiên (From Ad to First Return)',
    25: 'Tiến trình, Áp lực và Sự công bằng (Progression, Pressure & Fairness)',
    26: 'Quảng cáo, IAP và Nền kinh tế trong game (Ads, IAP & Game Economy)',
    27: 'Tín hiệu, Quyết định và Thử nghiệm (Signals, Decisions & Experimentation)',
    28: 'Giới hạn của Dữ liệu: Dữ liệu quyết định được gì và không thể quyết định gì (What Data Can & Cannot Decide)',
    29: 'Hệ thống Vận hành Trực tiếp (Live Ops Framework)',
    30: 'Cẩm nang Thiết kế theo Thể loại (Genre-Specific Playbook)',
    31: 'Bộ Kiểm tra Tổng thể (Master Audit Checklist)',
    32: 'Nguồn Tham khảo Công khai (Public References)',
    33: 'How to Read this Playbook | Cách Sử Dụng Cuốn Sách Này',
    34: 'Đừng đọc tài liệu này như một bài blog lý thuyết. Nó chỉ phát huy tối đa giá trị khi bạn dùng nó làm công cụ giải phẫu (audit) một tựa game cụ thể.',
    35: 'Mỗi chương sẽ mổ xẻ một điểm chạm (touchpoint) cốt tử: từ creative quảng cáo, trang cửa hàng ứng dụng (store listing), trải nghiệm màn chơi đầu (FTUE), thiết kế level, khoảnh khắc thất bại (fail state), vị trí đặt quảng cáo thưởng (rewarded ads), logic ra offer, chuỗi sự kiện (events), đến bảng dữ liệu cohort, đánh giá của người dùng và các cuộc họp nội bộ.',
    36: '• Đối với game đã phát hành (Live Game): Hãy mở bản build song song khi đọc. Khi đọc về 10 màn chơi đầu, hãy tự tay chơi lại 10 màn đó. Khi đọc về rewarded ad, hãy tìm vị trí hiển thị đầu tiên và tự vấn: Người chơi đang thực sự cần gì tại khoảnh khắc này? Khi đọc về IAP, hãy mở shop và gọi tên chính xác bài toán mà từng gói nạp đang giải quyết. Khi đọc về chỉ số, hãy mở dashboard và bóc tách tín hiệu thực sự khỏi những nhiễu loạn bề nổi.',
    37: "• Đối với game ở giai đoạn Prototype: Hãy biến các chương thành những 'cổng kiểm duyệt' (quality gates) bắt buộc trước khi bước vào Soft Launch. Một tựa game chưa làm rõ được lời hứa cốt lõi, trải nghiệm 5 phút đầu, cơ chế tạo áp lực, giá trị trao đổi của quảng cáo, logic của gói bán và lý do để người chơi quay lại vào ngày mai — là tựa game chưa hề sẵn sàng để chi tiền mua người dùng (scale UA).",
    38: 'Mục tiêu của bạn không phải là đồng ý với tất cả mọi framework trong sách. Mục tiêu là gấp lại mỗi chương với một câu hỏi sắc bén hơn dành cho dự án của mình, và một cách nhìn tỉnh táo hơn để thách thức mọi câu trả lời có sẵn.',
    44: 'Monetization Starts Before the Store | Việc Kiếm Tiền Bắt Đầu Trước Màn Hình Cửa Hàng',
    45: 'Phần lớn game không chết ở màn hình cửa hàng (In-game Shop). Chúng thất bại từ rất sớm trước đó — ngay tại những mắt xích kiến tạo niềm tin khiến người chơi sẵn lòng mở ví:',
    46: '• Quảng cáo hứa hẹn một cảm xúc, nhưng 3 phút đầu vào game lại mang đến một trải nghiệm hoàn toàn lệch pha.',
    47: '• Trang Store không chứng minh được lời hứa từ Creative.',
    48: '• Phần hướng dẫn tân thủ (Tutorial) lê thê, tước đoạt quyền tự do kiểm soát của người chơi.',
    49: '• Quảng cáo xen kẽ (Interstitial) đầu tiên nhảy ra trước khi người chơi kịp quyết định xem tựa game này có đáng để họ bỏ thêm một phút nào nữa hay không.',
    50: '• Thiết kế level tạo cảm giác ức chế và bất công, rồi vội vã chìa ra một gói Booster như một liều thuốc giải vá lỗi.',
    51: '• Đội ngũ nhìn vào biểu đồ doanh thu ngắn hạn như một tín hiệu khỏe mạnh, mà không thấy tỷ lệ giữ chân (retention), điểm đánh giá (ratings), yêu cầu hoàn tiền (refunds) và niềm tin của cộng đồng đang lao dốc.',
    52: 'Màn hình cửa hàng chỉ là nơi hoàn tất một giao dịch trao đổi giá trị. Nó vĩnh viễn không thể cứu vãn một sản phẩm chưa đủ sức giữ chân người chơi.',
    53: 'Trước khi đòi hỏi tiền bạc, tựa game của bạn đã nhận được những gì từ người chơi?',
    54: '• Đầu tiên là sự chú ý.',
    55: '• Tiếp theo là cú nhấp chuột, lượt cài đặt, thời gian chờ tải, phiên trải nghiệm đầu tiên (FTUE), và lần mở lại game.',
    56: '• Dòng tiền chỉ thực sự xuất hiện khi game tích lũy đủ ngân sách niềm tin (trust budget).',
    57: 'Trong các dòng game Casual, Hybrid-casual, Puzzle và Hybrid-puzzle, Monetization là kết quả của cả một hành trình chuyển đổi:',
    58: 'Thấy Creative ➔ Nhấp chuột ➔ Trang Store ➔ Cài đặt ➔ Mở lần đầu ➔ 10 Level đầu ➔ Quay lại Ngày 1 (D1) ➔ Thói quen ➔ Chủ động xem Ads thưởng ➔ Lần nạp đầu (First IAP) ➔ Tái nạp ➔ Live Ops ➔ Giới thiệu',
    59: 'Mỗi điểm chạm đều đòi hỏi người chơi phải trao cho game một thứ:',
    60: '• Creative cần sự chú ý và tò mò.\n• Trang Store cần niềm tin ban đầu.\n• Màn hình tải game cần sự kiên nhẫn.\n• 10 màn chơi đầu cần xây đắp cảm giác thành tựu và sự tin tưởng.\n• Quảng cáo đầu tiên cần sự cho phép và đồng thuận.\n• Gói ưu đãi đầu tiên cần một lý do xứng đáng để chi trả.\n• Chuỗi sự kiện đầu tiên cần thói quen gắn bó.',
    61: 'Profit = Installs × (LTV - CPI)',
    62: 'Đây là công thức tài chính kinh điển, nhưng nó quá vĩ mô và xuất hiện quá muộn để có thể dẫn đường cho Game Designer.',
    64: 'Một công thức thực chiến (Operational Formula) hữu dụng hơn cho đội ngũ phát triển là:',
    65: 'Monetization = Nhu cầu cốt lõi × Đúng ngữ cảnh × Niềm tin tích lũy × Tốc độ phản ứng',
    66: '• Nhu cầu cốt lõi (Player Need): Cảm giác giải tỏa (relief), cơ hội thử lại, khẳng định kỹ năng, tăng tốc độ tiến trình (progression), sưu tập, sự tiện lợi, vị thế xã hội, hoặc cảm giác làm chủ tình thế.',
    67: '• Đúng ngữ cảnh (Right Context): Đề xuất ưu đãi xuất hiện chính xác vào thời điểm nhu cầu của người chơi dâng cao nhất, chứ không phải lúc Studio hay Publisher đang cần chạy KPI doanh thu.',
    68: '• Niềm tin tích lũy (Trust Budget): Cảm giác tự nhiên rằng trò chơi đối xử công bằng, minh bạch và tôn trọng thời gian/tiền bạc của họ.',
    69: '• Tốc độ phản ứng (Execution Speed): Năng lực của đội ngũ trong việc đọc nhanh dữ liệu phân tầng (ad funnel, conversion rate, cohort retention, reviews) để liên tục tối ưu vòng lặp live ops.',
    70: 'Bản đồ vận hành của cuốn sách gồm 6 trụ cột:',
    71: 'Lời hứa ⟷ Tiến trình ⟷ Áp lực ⟷ Sự đồng thuận ⟷ Giao dịch ⟷ Gắn bó dài hạn',
    72: 'Khi khuyết thiếu bất kỳ trụ cột nào, doanh thu có thể vẫn tăng vọt trong ngắn hạn — nhưng đó là cái bẫy chết người.',
    73: 'Vì sao?',
    74: 'Chỉ số IMPDAU (số lượt xem quảng cáo trung bình) có thể tăng đột biến trong khi tỷ lệ giữ chân D3 đang âm thầm sụp đổ. Một gói IAP "bẫy" người chơi có thể tạo tỷ lệ chuyển đổi cao hôm nay, nhưng phần đánh giá trên Store sẽ ngập tràn lời phàn nàn về sự bất công. Một vị trí interstitial thô bạo có thể kéo ARPDAU lên đỉnh, nhưng sẽ bóp nghẹt khả năng mở rộng quy mô UA của toàn bộ dự án.\n\n• Doanh thu lành mạnh (Healthy Revenue) tạo ra lý do để người chơi hào hứng tiếp tục cuộc hành trình sau mỗi lần xem quảng cáo hay trả phí.\n• Doanh thu vay mượn (Borrowed Revenue) vắt kiệt giá trị từ sự ức chế mà tựa game không thể bù đắp bằng chất lượng gameplay và sự công bằng.',
    75: 'Example Case: Clear Garden | Ví Dụ Phân Tích: Clear Garden',
    76: 'Hãy hình dung một dự án game Hybrid-Puzzle giả định đang bước vào giai đoạn Soft Launch với tên gọi: Clear Garden.',
    77: '• Core Loop (Vòng lặp cốt lõi): Người chơi thu dọn các vật phẩm lộn xộn trong một khu vườn hoang phế và xếp chúng vào một khay chứa giới hạn để dọn sạch không gian (tương tự cơ chế Match-3D / Grid Puzzle), từ đó tích lũy tài nguyên để phục dựng từng khu vực trong vườn (Meta-progression).\n• Creative Promise (Lời hứa từ quảng cáo): Đánh vào cảm xúc thỏa mãn khi "lập lại trật tự từ đống hỗn độn" (satisfying cleaning/organizing) và sự biến chuyển trực quan đầy cuốn hút của khu vườn.',
    78: 'Bản build đầu tiên của Clear Garden mắc phải hàng loạt "căn bệnh kinh điển":\n1. Yêu cầu quyền ATT (Tracking) ngay khi vừa mở app, trước khi người chơi kịp chạm tay vào câu đố đầu tiên.\n2. Bật Interstitial Ad ngay sau Level 2, ngắt mạch hưng phấn khi người chơi chưa kịp hiểu game.\n3. Đẩy độ khó phi lý ở Level 7 bằng cách tung ra quá nhiều biến thể vật phẩm rác, rồi lập tức "ép" người chơi mua thêm ô khay (extra slots) ngay khi vừa thất bại.\n4. Bán Starter Pack chứa một mớ tiền ảo trừu tượng nhưng không hề giải thích số tiền đó giúp giải quyết trở ngại cụ thể nào.\n5. Phát Daily Reward ồ ạt, nhưng tiến trình cải tạo khu vườn lại thiếu chiều sâu, không tạo ra bất kỳ động lực hay "móc câu tò mò" nào để người chơi mở lại game vào sáng hôm sau.',
    79: 'Nếu đây là bản build thực tế của đội ngũ bạn, những chỉ số nào sẽ gióng lên hồi chuông cảnh báo trước khi bạn lãng phí hàng ngàn USD vào việc mua thêm traffic?',
    80: 'Clear Garden là một ví dụ giả định. Nhưng những quyết định sai lầm kể trên lại là thực tế đang diễn ra hàng ngày ở vô số studio.',
    81: 'Trong suốt các chương tiếp theo của cuốn sách, chúng ta sẽ liên tục quay lại với case study Clear Garden — để chuyển hóa từng nguyên lý trừu tượng thành những giải pháp can thiệp cụ thể trên từng màn hình game.'
}

# Apply paragraph updates
for idx, new_text in p_updates.items():
    doc.paragraphs[idx].text = new_text

# Table 0 updates (Key terms)
table0_data = [
    ('Thuật ngữ', 'Ý nghĩa trong tài liệu'),
    ('Monetization', 'Mô hình và chiến lược tạo doanh thu trong game thông qua quảng cáo, giao dịch in-app và các dịch vụ bổ trợ.'),
    ('UA (User Acquisition)', 'Hoạt động thu hút người chơi mới cài đặt game, chủ yếu thông qua các chiến dịch quảng cáo trả phí (paid ads).'),
    ('IAP (In-App Purchase)', 'Giao dịch mua hàng trong ứng dụng, bao gồm gói nạp đầu (starter packs), tiền tệ in-game, vật phẩm bổ trợ hoặc gói gỡ quảng cáo (no-ads).'),
    ('Retention', 'Tỷ lệ giữ chân người chơi sau một mốc thời gian xác định (ví dụ: D1 là tỷ lệ người chơi quay lại vào ngày đầu tiên sau cài đặt, D7, D30).'),
    ('Cohort', 'Nhóm thuần tập: Tập hợp người chơi có cùng thời điểm bắt đầu hoặc chung đặc tính (ví dụ: cài game cùng ngày, đến từ cùng một mẫu creative).'),
    ('Core Loop', 'Vòng lặp cốt lõi: Chuỗi hành động chính yếu mà người chơi liên tục thực hiện và lặp lại trong suốt vòng đời trải nghiệm game.'),
    ('Live Ops', 'Vận hành trực tiếp (Live Operations): Hoạt động duy trì và làm mới game sau khi ra mắt: tổ chức chuỗi sự kiện, tung gói ưu đãi, cập nhật nội dung, gửi push notification và tinh chỉnh cấu hình từ xa (remote config).'),
    ('Creative', 'Tư liệu quảng cáo: Các định dạng nội dung (video, hình ảnh, playable ad, thông điệp) được thiết kế nhằm thu hút sự chú ý và kích thích cài đặt.'),
    ('Offer', 'Gói ưu đãi theo ngữ cảnh: Đề xuất mua vật phẩm hoặc xem quảng cáo có thưởng được kích hoạt tại đúng thời điểm và trạng thái cảm xúc của người chơi.'),
    ('Funnel', 'Phễu chuyển đổi: Chuỗi các bước tuần tự người chơi trải qua: Thấy Creative ➔ Cài đặt ➔ FTUE ➔ Chơi tiếp ➔ Quay lại ➔ Chi trả.'),
    ('LTV và CPI', 'LTV (Lifetime Value): Doanh thu trọn đời kỳ vọng từ một người chơi.\nCPI (Cost Per Install): Chi phí bình quân để có được một lượt cài đặt mới.'),
    ('ARPDAU và IMPDAU', 'ARPDAU: Doanh thu trung bình trên mỗi người chơi hoạt động hàng ngày.\nIMPDAU: Số lượt hiển thị quảng cáo trung bình trên mỗi người chơi hoạt động hàng ngày.')
]

t0 = doc.tables[0]
for r_idx, (col0, col1) in enumerate(table0_data):
    t0.rows[r_idx].cells[0].text = col0
    t0.rows[r_idx].cells[1].text = col1

doc.save(doc_path)
print('SUCCESSFULLY_UPDATED_DOCX')
